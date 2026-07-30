from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from optimize_docx_images import optimize_docx_images


ROOT = Path(__file__).resolve().parents[1]
PHOTOS = ROOT / "photos"
OUT = ROOT / "deliverables" / "J40_removed_parts_refurbishment_work_order_20260723.docx"

BLUE = RGBColor(31, 78, 121)
INK = RGBColor(32, 42, 52)
MUTED = RGBColor(90, 98, 108)
RED = RGBColor(145, 24, 24)
GREEN = RGBColor(37, 102, 63)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_labeled(doc, label, text, color=INK, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(label + " ")
    set_font(r, bold=True, color=color)
    r = p.add_run(text)
    set_font(r)
    return p


def add_image_strip(doc, filenames, max_height=3.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    count = len(filenames)
    max_width = 6.15 / count if count > 1 else 4.2
    for idx, filename in enumerate(filenames):
        path = PHOTOS / filename
        with Image.open(path) as im:
            w, h = im.size
        width = min(max_width, max_height * w / h)
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        if idx < count - 1:
            p.add_run("   ")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(9)
    r = cap.add_run("Photo evidence: " + ", ".join(Path(f).stem.split("_gp_")[0] for f in filenames))
    set_font(r, size=9, italic=True, color=MUTED)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for style_name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 12, 9, 4)):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("J40 RESTORATION • WORKSHOP CONTROL")
    set_font(r, size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Issued 23 July 2026 • Match every part before refit")
    set_font(r, size=8.5, color=MUTED)


REFURB = [
    {
        "title": "Small hinge / stop bracket",
        "images": ["20260723_013021_gp_UGFfrBkw.jpg"],
        "description": "Small steel hinge or stop bracket. Exact installed location remains to be labelled before work.",
        "instructions": [
            ("Identify", "Photograph at the installed location and mark side/orientation before dismantling."),
            ("Inspect", "Check pin play, cracks, distorted ears, elongated holes and loss of section. Fabricate by sample if unsafe."),
            ("Prepare", "Degrease, remove loose paint/rust mechanically, treat sound residual pitting, then epoxy-prime and topcoat."),
            ("Protect", "Mask the pin bore, pivot, threads and mounting faces. Lubricate only after full paint cure."),
            ("Accept", "Smooth movement without binding or excessive play; coating complete with no paint in the joint."),
        ],
    },
    {
        "title": "Door latch and lock-linkage mechanisms",
        "images": ["20260723_013044_gp_nVUdz9dQ.jpg", "20260723_013123_gp_JLEfOezg.jpg", "20260723_013303_gp_ZzbvMf2A.jpg"],
        "description": "Door latch bodies and their rods/linkages. These are separate from the failed key cylinder, which is replacement-only.",
        "instructions": [
            ("Label", "Keep left/right assemblies and every rod, spring and clip together; photograph rod routing before separation."),
            ("Clean", "Solvent-clean old grease without soaking rubber/plastic pieces; remove corrosion from non-working steel surfaces."),
            ("Inspect", "Reject cracked housings, weak or broken springs, rounded pawls, elongated pivots and latches that fail primary/secondary retention."),
            ("Finish", "Zinc-plate suitable small steel hardware or epoxy-prime and paint exterior non-working surfaces. Do not paint pawls, pivots, latch faces or threads."),
            ("Assemble", "Renew clips/bushes, lubricate lightly with latch grease and bench-cycle repeatedly before door installation."),
        ],
    },
    {
        "title": "Linkage rod / stay",
        "images": ["20260723_013050_gp_rSLoiGNw.jpg"],
        "description": "Long linkage rod or stay with eye/end fittings; installed function and side must be proven.",
        "instructions": [
            ("Identify", "Mark both endpoints and installed orientation. Do not straighten until a vehicle datum or matching sample proves the intended geometry."),
            ("Inspect", "Check bends, thread damage, worn eyes, split bushes and elongated holes."),
            ("Prepare", "Degrease, abrade loose coating/rust, treat residual pitting, epoxy-prime and topcoat the rod."),
            ("Protect", "Mask threads, spherical/eye bearing surfaces and adjustment lands."),
            ("Accept", "Correct geometry, free articulation and secure retainers after trial fit."),
        ],
    },
    {
        "title": "Glove-box lid",
        "images": ["20260723_013148_gp_AK0sjJJw.jpg"],
        "description": "Pressed-steel glove-box lid/door suitable for cosmetic and corrosion refurbishment if not distorted.",
        "instructions": [
            ("Dry fit", "Check aperture fit, hinge line, latch engagement and face alignment before final paint."),
            ("Repair", "Correct minor dents without stretching the panel; repair cracks or failed hinge fixings before coating."),
            ("Prepare", "Strip unstable coating, feather sound edges, remove corrosion, epoxy-prime and block-sand."),
            ("Paint", "Apply the approved cabin colour and sheen to visible faces and protected coating to the rear."),
            ("Accept", "Even gap, positive latch, no edge contact and no paint build-up in screw holes or hinge pivots."),
        ],
    },
    {
        "title": "Metal window channels / guides",
        "images": ["20260723_013155_gp_k26a3Rpw.jpg", "20260723_013511_gp_8BN3n86Q.jpg"],
        "description": "Steel window guide channels. The metal can be refurbished if sound; the old felt/rubber run material must be renewed.",
        "instructions": [
            ("Label", "Mark door and orientation before removing old felt or rubber."),
            ("Inspect", "Check straightness, mounting tabs, spot welds and corrosion inside the channel. Replace metal that is perforated or too thin."),
            ("Prepare", "Remove old run material, adhesive and rust; epoxy-prime and paint only the metal surfaces."),
            ("Renew", "Fit new correct-profile felt/run channel after glass and door-frame trial fit; do not paint the sliding surface."),
            ("Accept", "Glass travels freely without rattle, excessive drag, scratching or water-path obstruction."),
        ],
    },
    {
        "title": "Manual window regulator",
        "images": ["20260723_013218_gp_gz4eGcoQ.jpg"],
        "description": "Manual window regulator mechanism; refurbish only if the sector teeth, arms, shafts and rivets are sound.",
        "instructions": [
            ("Clean", "Remove dried grease and dirt without blasting bearing surfaces or packed joints."),
            ("Inspect", "Check sector teeth, crank spline, rollers, pivots, rivets, arms and end stops. Replace the regulator if teeth are stripped or arms distorted."),
            ("Finish", "Treat and repaint exposed frame/arm surfaces while masking teeth, shafts, rollers and sliding tracks."),
            ("Service", "Renew rollers/bushes as needed, apply regulator grease and bench-cycle through full travel."),
            ("Accept", "Smooth full travel, low play, no tooth skip and correct glass alignment under load."),
        ],
    },
    {
        "title": "Door pulls and interior handles",
        "images": ["20260723_013331_gp_wJ1wGBzg.jpg", "20260723_013357_gp_r5068tWQ.jpg", "20260723_013420_gp_SclaaDYg.jpg"],
        "description": "Interior pull/armrest and handle pieces. Reuse only where the substrate, fixings and shape remain sound.",
        "instructions": [
            ("Inspect", "Check cracks, stripped inserts, distorted brackets and missing captive nuts. Replace any unsafe or badly degraded substrate."),
            ("Clean", "Use plastic/vinyl-safe cleaner on trim; strip corrosion from separate steel brackets and screws."),
            ("Restore", "Repair/retrim sound pull pieces to the approved cabin finish; repaint only compatible substrates and refinish separate metal brackets."),
            ("Renew", "Fit new screws, captive nuts, spacers and rubber buffers where worn or missing."),
            ("Accept", "Secure pull load with no flex, sharp edges, loose trim or interference with latch/window operation."),
        ],
    },
    {
        "title": "Unidentified lever / bracket pieces",
        "images": ["20260723_013447_gp_mzaP0yGQ.jpg", "20260723_013456_gp_DSWTd2FA.jpg"],
        "description": "Lever, bracket and knob/shaft pieces. They remain identification-controlled and must not be altered speculatively.",
        "instructions": [
            ("Identify", "Match each piece to its installed function, location, orientation and mating hardware before work."),
            ("Inspect", "Check pivot/shaft wear, cracks, bends, damaged threads and missing bushes or clips."),
            ("Prepare", "On sound steel only, degrease, remove loose corrosion, treat residual pitting, epoxy-prime and topcoat."),
            ("Protect", "Mask shafts, threads, pivots, knob interfaces and contact faces. Straighten only against proven geometry."),
            ("Accept", "Correct fit and full movement with all bushes/clips installed; no paint-related binding."),
        ],
    },
    {
        "title": "Filler neck / outlet neck",
        "images": ["20260723_013311_gp_51vqP4uQ.jpg"],
        "description": "Cast or fabricated filler/outlet neck. Exact radiator/engine location must be confirmed before reuse.",
        "instructions": [
            ("Identify", "Confirm the circuit, mating cap or hose, pressure duty and installed orientation."),
            ("Inspect", "Descale internally; check wall pitting, cracks, flange flatness, sealing face and hose/cap seat. Replace if sealing integrity is doubtful."),
            ("Prepare", "Clean and coat the exterior only with a temperature-appropriate system after passing inspection."),
            ("Protect", "Mask all internal passages, gasket faces, cap seat, hose bead and threads; no abrasive or coating residue may remain inside."),
            ("Accept", "Flat, leak-free joint and sound cap/hose sealing under the circuit's required pressure test."),
        ],
    },
]


REPLACEMENTS = [
    ("Mechanical cable / possible handbrake wire", ["20260723_013017_gp_nBbqcMXA.jpg"], "KEEP AS PATTERN ONLY", "A rear handbrake cable is already recorded as received. Identify this loose cable by endpoints and compare length/end fittings; buy only missing sections or attachment hardware. Do not refurbish or paint."),
    ("Front seat-belt assemblies", ["20260723_013033_gp_l8ZqIJ6Q.jpg", "20260723_013103_gp_24hNTQNQ.jpg"], "REPLACE — SAFETY CRITICAL", "Old webbing, buckles and retractors are samples only. Fit new certified belts with correct anchor geometry and graded mounting hardware."),
    ("Rear tail-lamp housing / all exterior lamps", ["20260723_013210_gp_DXsn0EoQ.jpg"], "REPLACE AND EXCLUDE FROM PAINTER", "Use complete new lamp assemblies. Confirm model year, voltage, handing, brackets, gaskets, bulb holders/connectors and fittings before discarding old samples."),
    ("Window crank handle", ["20260723_013349_gp_StIZA9xA.jpg"], "REPLACE", "New reproduction handles are inexpensive. Retain the old handle only until spline, offset and clip fit are matched."),
    ("Exterior door-handle mechanism", ["20260723_013405_gp_bkbDX7gQ.jpg"], "PREFER REPLACEMENT", "This photograph is door-handle/latch hardware, not a wiper motor. Trial-fit a new reproduction handle; retain and service the old mechanism only if the replacement geometry is wrong and the original is structurally sound."),
    ("Insecure lock cylinder", ["20260723_013440_gp_bFh4nlWQ.jpg"], "REPLACE — SECURITY FAILURE", "Do not reuse a cylinder that opens with any key. Match barrel diameter, flange, clip, tailpiece and panel thickness. Prefer a complete matched lock set; use a locksmith to adapt a new cylinder if the opening is non-standard."),
    ("Plastic washer / coolant reservoirs", ["20260723_014847_gp_yhu1WZWQ.jpg", "20260723_014852_gp_aWxFuqNw.jpg"], "REPLACE", "Buy new bottles or a complete washer kit with pump, hose, nozzle and bracket. Keep originals only to confirm capacity, neck, hose and mounting positions."),
    ("Power-steering reservoir", ["20260723_014855_gp_WgIPqcIQ.jpg"], "REPLACE AFTER CIRCUIT CONFIRMATION", "Confirm the selected J60 pump/box circuit, suction/return hose IDs, cap/filter and bracket orientation before buying. Fit new compatible reservoir and hoses."),
    ("Pedal rubber pad", ["20260723_013740_gp_PDMfr5IQ.jpg"], "REPLACE", "Old rubber is a profile sample only. Fit a new pedal-pad set matched to production year and pedal shape."),
]


PURCHASES = [
    ("Matched lock set", "OEM Toyota BJ40/FJ40 cylinder lock set 69005-90316; ignition, both doors and fuel flap with matching keys. Listing states 1979–1983 fitment; verify vehicle year and barrel geometry.", "https://www.shopresponse.com/oem-toyota-land-cruiser-bj40-fj40-cylinder-lock-set-w-keys-69005-90316/"),
    ("Exterior lights", "Specter Off-Road 40-Series tail-light page; complete aftermarket pair is listed for 1/1979–1984 with connector. Use the linked catalog to select the correct production-year assemblies.", "https://www.sor.com/cat/148"),
    ("Exterior light catalog", "Cruiser Corps FJ40 exterior-light collection for model-year-specific tail, indicator, marker and license lamps. Confirm complete fittings rather than assuming every listing includes brackets/gaskets.", "https://cruisercorps.com/collections/exterior-lights/fj40"),
    ("Wiper system", "Specter Off-Road FJ40 wiper catalog covering motor, linkage, arms and blades. Match build date, motor voltage and sweep geometry before ordering.", "https://www.sor.com/cat/192"),
    ("Washer bottle kit", "Cruiser Corps universal 12 V washer kit with reservoir and motor; listing says parts required for installation are included. Confirm 12 V system and mounting space.", "https://cruisercorps.com/collections/40-series-shop-all/products/windshield-washer-kit"),
    ("Washer kit alternative", "Specter Off-Road universal washer kit, part 194-30, listed with motor, bottle, nozzle, bracket, switch, hose and fittings.", "https://www.sor.com/cat/194"),
    ("Window crank", "Cruiser Corps reproduction crank handle with retaining clip for 1975–1990 FJ40/BJ-family applications. Verify build date and spline.", "https://cruisercorps.com/products/window-regulator-crank-handle-reproduction"),
    ("Window run felt", "Cruiser Corps OEM lower window-run channel for 1975–1984 one-piece front doors; listing requires two per door. Verify door type and production year.", "https://cruisercorps.com/products/weatherstrip-front-door-lower-window-run-channel-oem"),
    ("Door handle / regulator options", "Cruiser Corps door-parts catalog with reproduction outside handles, latch sets and regulator pairs. Choose only after comparing handing, production year and mounting geometry.", "https://cruisercorps.com/collections/all-products/door-parts"),
    ("Front seat belts", "Seatbelt Planet 1979–1984 FJ40 driver/passenger retractable 3-point kit with hardware. Confirm production year, RHD anchor geometry and local compliance before purchase.", "https://www.seatbeltplanet.com/i-30498836-1979-1984-toyota-land-cruiser-fj40-driver-passenger-seat-belt-kit.html"),
    ("Pedal-pad set", "Specter Off-Road pedal catalog lists separate 1958–1978 and 1979–1984 three-pad kits. Select by vehicle production date and pedal profile.", "https://www.sor.com/cat/014d"),
]


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("J40 REMOVED PARTS")
    set_font(r, size=24, bold=True, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run("Refurbishment work order and replacement schedule")
    set_font(r, size=15, color=BLUE)

    add_labeled(doc, "Issue date", "23 July 2026")
    add_labeled(doc, "Evidence", "Google Photos intake 20260723T023624")
    add_labeled(doc, "Purpose", "Prevent replacement-grade, unsafe or inexpensive parts from entering the painter batch while giving the workshop precise instructions for the reusable items.")
    add_labeled(doc, "Painter batch", "15 photographs covering nine reusable component families. Everything else in this document is replace, retain-as-pattern, identify, or exclude.", color=GREEN)

    h = doc.add_heading("Non-negotiable controls", level=1)
    controls = [
        ("No cosmetic rescue", "Do not paint over active rust, structural thinning, cracks, unsafe wear or failed security/safety components."),
        ("Label first", "Record function, side, orientation and all mating hardware before stripping or separating parts."),
        ("Mask working areas", "No paint on threads, pivots, teeth, latch faces, bores, contacts, sealing faces, internal fluid passages or friction surfaces."),
        ("Replacement samples", "Keep excluded originals only until new parts are proven for fit, voltage, ports, handedness and mounting geometry."),
        ("Bump stops", "Their separate metal-fixture rust-treatment and preparation job remains mandatory before installation; they are not part of this loose body-hardware painter batch."),
    ]
    for label, text in controls:
        add_labeled(doc, label + ":", text)

    doc.add_page_break()
    doc.add_heading("Part A — Send for refurbishment", level=1)
    add_labeled(doc, "Scope", "Only the following photographed component families are approved for the refurbishment work order. Conditional items must pass inspection before coating.")

    for index, item in enumerate(REFURB, 1):
        if index > 1:
            doc.add_page_break()
        doc.add_heading(f"A{index}. {item['title']}", level=1)
        add_image_strip(doc, item["images"])
        add_labeled(doc, "Description", item["description"])
        add_labeled(doc, "Disposition", "REFURBISH / REPAINT IF SOUND", color=GREEN)
        doc.add_heading("Workshop instructions", level=2)
        for label, text in item["instructions"]:
            add_labeled(doc, label + ":", text)

    doc.add_page_break()
    doc.add_heading("Part B — Replace, retain as pattern, or exclude", level=1)
    add_labeled(doc, "Rule", "These parts do not enter the painter/refurbishment batch. Keep originals only until replacement fit and system compatibility are proven.", color=RED)
    for index, (title, images, status, instruction) in enumerate(REPLACEMENTS, 1):
        doc.add_heading(f"B{index}. {title}", level=2)
        add_image_strip(doc, images, max_height=2.45)
        add_labeled(doc, "Disposition", status, color=RED)
        add_labeled(doc, "Action", instruction)

    doc.add_page_break()
    doc.add_heading("Part C — Purchase links and fitment gates", level=1)
    add_labeled(doc, "Important", "Links and availability were checked on 23 July 2026. Prices, stock and international shipping can change. Do not order until the vehicle production date, system voltage, RHD/LHD relevance and sample geometry are confirmed.", color=RED)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.55)
    table.columns[1].width = Inches(4.95)
    hdr = table.rows[0].cells
    hdr[0].width = Inches(1.55)
    hdr[1].width = Inches(4.95)
    for cell, text in zip(hdr, ("Item", "Recommended source and purchase gate")):
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=10, bold=True, color=INK)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "E8EEF5")
        cell._tc.get_or_add_tcPr().append(shading)

    for item, note, url in PURCHASES:
        cells = table.add_row().cells
        cells[0].width = Inches(1.55)
        cells[1].width = Inches(4.95)
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = 1
        p = cells[0].paragraphs[0]
        r = p.add_run(item)
        set_font(r, size=9.5, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(note + " ")
        set_font(r, size=9.5)
        add_hyperlink(p, "Open purchase page", url)

    doc.add_heading("No-link procurement gates", level=2)
    add_labeled(doc, "Handbrake cable", "Do not buy another full cable blindly: one cable is already recorded as received. Match the old sample and buy only missing sections, equalizer/clevis/pins/clips or springs.")
    add_labeled(doc, "Power-steering reservoir", "No purchase link is released yet. First confirm the selected J60 pump/box, fluid circuit, suction/return hose IDs, cap/filter and bracket location; then source a new reservoir locally or from the component supplier.")
    add_labeled(doc, "Radiator/filler cap", "Already present in the project procurement plan; sample-match neck size and required pressure before buying.")

    doc.add_heading("Closeout sign-off", level=1)
    for field in ("Workshop / vendor", "Date received", "Items rejected before coating", "Items completed", "Replacement samples returned", "Checked by"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(field + ":  " + "_" * 58)
        set_font(r)

    doc.core_properties.title = "J40 Removed Parts Refurbishment Work Order"
    doc.core_properties.subject = "Photographic work order, replacement schedule and purchase links"
    doc.core_properties.author = "J40 Restoration Project"
    doc.save(OUT)
    optimize_docx_images(OUT, target_dpi=600, jpeg_quality=85)
    print(OUT)


if __name__ == "__main__":
    build()
