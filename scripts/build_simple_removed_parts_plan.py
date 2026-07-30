from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
PHOTOS = ROOT / "photos"
OUT = ROOT / "deliverables" / "J40_simple_removed_parts_plan_20260723.docx"
IMAGE_CACHE = ROOT / "deliverables" / ".simple_plan_image_cache"

BLUE = RGBColor(31, 78, 121)
GREEN = RGBColor(42, 102, 65)
RED = RGBColor(150, 32, 32)
INK = RGBColor(35, 42, 50)
MUTED = RGBColor(95, 103, 112)


def font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def margins(cell, top=80, start=110, bottom=80, end=110):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        mar.append(node)


def shade(cell, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def no_row_split(row):
    trpr = row._tr.get_or_add_trPr()
    trpr.append(OxmlElement("w:cantSplit"))


def hyperlink(paragraph, text, url):
    rid = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend((color, underline))
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    link.append(run)
    paragraph._p.append(link)


def compressed_image(filename):
    """Create a document-sized JPEG copy without changing the project photo."""
    IMAGE_CACHE.mkdir(parents=True, exist_ok=True)
    source = PHOTOS / filename
    target = IMAGE_CACHE / (Path(filename).stem + ".jpg")
    if not target.exists() or target.stat().st_mtime < source.stat().st_mtime:
        with Image.open(source) as image:
            image = image.convert("RGB")
            image.thumbnail((1200, 1200), Image.Resampling.LANCZOS)
            image.save(target, "JPEG", quality=76, optimize=True, progressive=True)
    return target


def add_picture(cell, filename, max_w=1.25, max_h=1.25):
    path = compressed_image(filename)
    with Image.open(path) as image:
        w, h = image.size
    width = min(max_w, max_h * w / h)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))


def add_header(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, before, after in (("Heading 1", 17, 15, 8), ("Heading 2", 13, 10, 5)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    hp = section.header.paragraphs[0]
    r = hp.add_run("J40 PARTS PLAN • SIMPLE VERSION")
    font(r, size=9, bold=True, color=MUTED)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("23 July 2026")
    font(r, size=8.5, color=MUTED)


def label_line(doc, label, text, color=INK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label + " ")
    font(r, bold=True, color=color)
    r = p.add_run(text)
    font(r)


def category_table(doc, rows, status_color):
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [1.45, 1.75, 4.0]
    headers = ["Photo", "Item", "What to do"]
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = Inches(widths[i])
        margins(cell)
        shade(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(headers[i])
        font(r, size=10, bold=True)
    no_row_split(table.rows[0])

    for item, images, action in rows:
        row = table.add_row()
        no_row_split(row)
        cells = row.cells
        for i, cell in enumerate(cells):
            cell.width = Inches(widths[i])
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if len(images) == 1:
            add_picture(cells[0], images[0])
        else:
            p = cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for image_name in images[:2]:
                path = compressed_image(image_name)
                with Image.open(path) as image:
                    w, h = image.size
                width = min(0.6, 1.0 * w / h)
                p.add_run().add_picture(str(path), width=Inches(width))
                p.add_run(" ")
        p = cells[1].paragraphs[0]
        r = p.add_run(item)
        font(r, size=10, bold=True, color=status_color)
        p = cells[2].paragraphs[0]
        r = p.add_run(action)
        font(r, size=10)
    return table


REFURBISH = [
    (
        "Pedal-area plunger-switch bracket",
        ["20260723_013021_gp_UGFfrBkw.jpg"],
        "Owner identifies the location as the brake/clutch pedal area. It looks like two adjustable plunger switches rather than a hydraulic junction. Label its position; test the switches and wiring; replace failed electrical parts; derust and refinish only the bracket.",
    ),
    (
        "Single stamped latch / adjuster bracket",
        ["20260723_013033_gp_l8ZqIJ6Q.jpg"],
        "Not a seatbelt. Send it with an IDENTIFY FIRST tag. It may be a bonnet/hood or other single control latch. Confirm its function and safe operation before cleaning, lubricating and refinishing reusable metal.",
    ),
    (
        "Likely front-door latch and central-locking sets",
        ["20260723_013044_gp_nVUdz9dQ.jpg", "20260723_013103_gp_24hNTQNQ.jpg"],
        "These appear to be two separate door sets: each has latch rods and an aftermarket electric lock actuator. Keep them separate as LH/RH candidates. Preserve the mechanical latches as fitment samples; replace old actuators if central locking is modernised; do not blast or paint motors, rods, pivots, catches or plated faces.",
    ),
    (
        "Parking-brake cable / linkage sample",
        ["20260723_013050_gp_rSLoiGNw.jpg"],
        "Use this to reconstruct the complete handbrake system. Send to a brake/cable specialist, not just a painter. Record length, bends and both end fittings; inspect or reproduce the cable/linkage; renew adjusters, clevises, pins and clips as needed; prove full application and free return.",
    ),
    (
        "Glove-box lid",
        ["20260723_013148_gp_AK0sjJJw.jpg"],
        "Dry-fit first. Repair dents, remove rust, epoxy-prime and paint in the chosen interior colour. Renew the hinge, latch and buffers if worn.",
    ),
    (
        "Manual window regulator",
        ["20260723_013218_gp_gz4eGcoQ.jpg"],
        "Clean out old grease. Check sector teeth, arms, rivets and rollers. Refinish exposed metal only, lubricate after the coating cures and bench-cycle through its full travel.",
    ),
    (
        "Likely windscreen-wiper motor and crank",
        ["20260723_013303_gp_ZzbvMf2A.jpg"],
        "This is a one-per-vehicle powered mechanism and its motor/gearbox/crank layout is consistent with a wiper motor, not steering. Send to an auto-electric specialist for identification and bench test. Rebuild only if worthwhile; do not blast or paint the motor, gearbox, connector, shaft or ball joint.",
    ),
    (
        "Unidentified lever pieces",
        ["20260723_013447_gp_mzaP0yGQ.jpg", "20260723_013456_gp_DSWTd2FA.jpg"],
        "Send with IDENTIFY FIRST tags. If the refurbisher confirms they are useful and sound, clean, derust and refinish the non-working surfaces. Mask pivots, bushes, shafts, holes and mating faces.",
    ),
    (
        "Likely engine oil-filler / breather tube",
        ["20260723_013311_gp_51vqP4uQ.jpg"],
        "Origin is not proven, but the large neck and small side breather are consistent with an engine filler/breather tube. Identify first. Degrease internally, inspect for deep pitting or cracks, preserve all sealing faces and coat only the sound exterior.",
    ),
    (
        "New bump-stop metal fixtures",
        ["20260722_000007_user_bump_stop_set_inventory.png"],
        "Label all four positions. Protect the new rubber from solvent, heat and abrasives. Remove loose rust from the metal plates, inspect for pitting/cracks/distortion, apply compatible rust treatment, epoxy-prime and paint the metal, then trial-fit.",
    ),
]


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    add_header(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("J40 PARTS TO SEND FOR REFURBISHMENT")
    font(r, size=24, bold=True)
    p = doc.add_paragraph()
    r = p.add_run("A simple photo guide for selecting the correct parts")
    font(r, size=14, color=BLUE)

    label_line(doc, "Purpose:", "Every photographed item in this guide should be selected, labelled and taken for refurbishment, specialist assessment or reconstruction.", GREEN)
    label_line(doc, "Not included:", "Parts planned for replacement or interior modernisation have been removed, including wiper arms/blades, lights, washer pump/bottles, failed lock, old interior pulls and padded trim, window crank and pedal rubber.", RED)
    label_line(doc, "Uncertain items:", "Send them, but attach an IDENTIFY FIRST label. The workshop should identify and assess them before blasting, dismantling or coating.")

    doc.add_heading("Simple rule for the workshop", level=1)
    label_line(doc, "Mechanical/electrical work:", "Route the pedal switches, latch sets, handbrake linkage, window regulator and wiper motor to the appropriate specialist before cosmetic work.")
    label_line(doc, "Before painting:", "Label where every part came from and check that it is not cracked, badly worn, bent or rusted too thin.")
    label_line(doc, "During painting:", "Do not paint threads, pivots, gear teeth, latch faces, seals, electrical contacts, fluid passages or mounting faces.")
    label_line(doc, "After painting:", "Fit new clips, bushes, seals or screws where needed, lubricate moving parts and test them before installation.")
    label_line(doc, "Bump stops:", "They are included because the metal fixtures need rust treatment and preparation. Protect the new rubber throughout the work.")

    doc.add_page_break()
    doc.add_heading("Select and send these parts", level=1)
    category_table(doc, REFURBISH, GREEN)

    doc.add_page_break()
    doc.add_heading("Workshop sign-off", level=1)
    label_line(doc, "Instruction:", "Record what was accepted, rejected and completed. Return every unidentified or rejected piece with its original label.")
    for field in ("Sent to", "Date", "Items rejected", "Items completed", "Checked by"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(16)
        r = p.add_run(field + ":  " + "_" * 58)
        font(r)

    doc.core_properties.title = "J40 Parts to Send for Refurbishment"
    doc.core_properties.subject = "Simple photographic refurbishment selection guide"
    doc.core_properties.author = "J40 Restoration Project"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
