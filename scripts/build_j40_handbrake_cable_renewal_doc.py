#!/usr/bin/env python3
"""Build the controlled J40 handbrake cable renewal workshop specification."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from PIL import Image, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/davidpridmore/IdeaProjects/J40")
OUT = ROOT / "docs" / "J40_Rear_Handbrake_Cable_Renewal_Specification_20260816.docx"
TMP = Path("/private/tmp/j40_handbrake_doc_assets")

PHOTO_GUIDE = [
    {
        "id": "R01", "title": "Control-end anchor",
        "path": ROOT / "photos" / "20260723_013050_gp_rSLoiGNw.jpg", "crop": (0.15, 0.02, 0.86, 0.98),
        "identify": "Short removed cable section with a swaged rectangular stop and a rectangular-loop eye. It is the strongest photo evidence for the special control-side anchorage, but the exact hand-lever attachment is not shown.",
        "action": "Bag as R01 FRONT. Match it to the hand-control lever on the vehicle and record orientation, load seat and pin/slot dimensions.",
        "confidence": "MEDIUM — on-vehicle confirmation required.",
    },
    {
        "id": "R02", "title": "Primary adjustment assembly",
        "path": ROOT / "photos" / "20260522_210729_gp_Csn9ztFA.jpg", "crop": (0.20, 0.23, 0.88, 0.88),
        "identify": "Threaded adjuster, casing abutment and support hardware from the removed custom installation. The adjacent wire is visibly frayed and is a pattern only.",
        "action": "Retrieve the complete adjuster and mating support as one labelled assembly. Measure thread diameter, pitch, hand, safe engagement and usable travel.",
        "confidence": "HIGH — multiple photo views agree.",
    },
    {
        "id": "R03", "title": "Oval compensator/equalizer",
        "path": ROOT / "photos" / "20260514_100003_gp_Vr2QI7ig.jpg", "crop": (0.02, 0.22, 0.98, 0.84),
        "identify": "Stacked oval equalizer plates through which the input and rear working cables are distributed. Hole pattern, plate order and cable load faces are unique to this installation.",
        "action": "Retrieve both plates together. Mark installed faces and cable paths; measure plate thickness, hole centres, apertures and load-seat diameters.",
        "confidence": "HIGH — installed and removed views agree.",
    },
    {
        "id": "R04", "title": "Rear reaction/equalizer bracket",
        "path": ROOT / "photos" / "20260710_103424_gp_Ab1ltlRw.jpg", "crop": (0.04, 0.06, 0.98, 0.78),
        "identify": "Heavy shaped plate and companion lever/bellcrank that react and redirect the equalized pull. The same photograph contains unrelated brake hardware; only the handbrake linkage parts are in scope.",
        "action": "Retrieve bracket and lever as a matched set with bushes/spacers in original order. Mark chassis/axle side and movement direction.",
        "confidence": "HIGH — geometry is clear in the laid-out view.",
    },
    {
        "id": "R05", "title": "Rectangular U-yoke/retainer",
        "path": ROOT / "photos" / "20260705_022559_gp_YyIEayQQ.jpg", "crop": (0.15, 0.15, 0.85, 0.78),
        "identify": "Small formed rectangular wire yoke with two open legs. This is a special linkage part, not a generic retaining clip.",
        "action": "Bag separately as R05. Photograph both faces beside a ruler and record leg spacing, wire section, overall height and fitted orientation.",
        "confidence": "HIGH — isolated component photograph.",
    },
    {
        "id": "R06", "title": "Special cable-end terminals",
        "path": ROOT / "photos" / "20260723_013050_gp_rSLoiGNw.jpg", "crop": (0.10, 0.00, 0.90, 1.00),
        "identify": "Examples include the rectangular-loop eye and swaged rectangular stop shown here, plus every clevis, eye, nipple, block and load-bearing stop on the removed cable assembly.",
        "action": "Do not discard a terminal with failed wire attached. Tag each terminal by run and end; use its functional load face as the measurement datum.",
        "confidence": "HIGH for the pictured terminals; recover all other end types.",
    },
    {
        "id": "R07", "title": "Sheath interface fittings",
        "path": ROOT / "photos" / "20260514_095826_gp_fg74oFMQ.jpg", "crop": (0.03, 0.03, 0.97, 0.88),
        "identify": "Ferrules, threaded sleeves, shoulder stops, end cups and bracket-engaging features where the outer casing reacts against the drum/backplate or support bracket.",
        "action": "Retrieve with the old sheath where possible. Label run and end; measure every shoulder, groove, thread and load-bearing abutment face.",
        "confidence": "HIGH — drum-side interface is visible.",
    },
    {
        "id": "R08", "title": "Intermediate casing locators",
        "path": ROOT / "photos" / "20260512_072730_gp_jSK3r3bg.jpg", "crop": (0.13, 0.12, 0.88, 0.88),
        "identify": "Intermediate sleeves/guides that constrain the long outer casing along the axle route and preserve working geometry through suspension movement.",
        "action": "Retrieve with the sheath if possible. Before removal, mark each locator position from the control-end functional datum and its vehicle orientation.",
        "confidence": "MEDIUM-HIGH — confirm each locator physically.",
    },
    {
        "id": "R09", "title": "Return/tension springs",
        "path": ROOT / "photos" / "20260710_103434_gp_go5dJG6w.jpg", "crop": (0.05, 0.02, 0.95, 0.92),
        "identify": "Removed spring assortment photographed with the brake parts. Only springs positively traced to cable return, lever return or equalizer stability belong to the handbrake.",
        "action": "Trace each spring to its installed position before assigning R09. Bag by location; renew any stretched, corroded, cracked or distorted spring.",
        "confidence": "MEDIUM — assortment requires location tracing.",
    },
    {
        "id": "R10", "title": "LH/RH drum interfaces",
        "path": ROOT / "photos" / "20260514_095953_gp_BXoQkXnw.jpg", "crop": (0.03, 0.04, 0.97, 0.90),
        "identify": "Rear backing-plate parking-brake lever/link and cable seating features. The replacement must preserve the standard internal Toyota drum-wire interface on both sides.",
        "action": "Retain on the vehicle or retrieve and label LH/RH. Dry-fit standard Toyota 47616-60010 wires without cutting, welding or re-swaging.",
        "confidence": "HIGH — rear drum mechanism is visible.",
    },
    {
        "id": "R11", "title": "Special mounting hardware",
        "path": ROOT / "photos" / "20260710_103424_gp_Ab1ltlRw.jpg", "crop": (0.00, 0.00, 1.00, 0.92),
        "identify": "Shaped brackets, abutments, pivots and spacers that establish cable geometry. The wheel cylinder visible here is hydraulic-brake hardware and is not part of R11.",
        "action": "Retrieve special geometry-setting parts with fitted bushes/spacers. Generic loose washers, nuts, clips and grommets remain outside this specification.",
        "confidence": "HIGH for the shaped hardware shown; reconcile against the vehicle.",
    },
]

WRONG_REFERENCE = {
    "path": ROOT / "photos" / "20260712_011332_gp_eRcLMeyg.jpg",
    "crop": (0.12, 0.02, 0.88, 0.98),
    "title": "Wrong/unverified replacement cable — quarantine",
    "body": "Clean replacement candidate photographed after purchase. It is not the known-fit removed assembly and must not supply cut lengths or terminal positions. Keep it separate and mark REJECTED — WRONG/UNVERIFIED SIZE — NOT MASTER — DO NOT INSTALL.",
}

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F6F9"
PALE_GOLD = "FFF4CE"
GOLD = "7A5A00"
PALE_RED = "FDE9E7"
RED = "9B1C1C"
GREY = "666666"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent=120) -> None:
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row) -> None:
    repeat_header(row)
    for cell in row.cells:
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(INK)


def set_repeat_table_header_dark(row) -> None:
    repeat_header(row)
    for cell in row.cells:
        set_cell_shading(cell, DARK_BLUE)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(WHITE)


def set_repeat_table_header_text(row, labels: list[str], dark=False) -> None:
    for cell, label in zip(row.cells, labels):
        cell.text = label
    if dark:
        set_repeat_table_header_dark(row)
    else:
        set_repeat_table_header(row)


def keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def keep_together(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_together = value


def heading_new_page(doc: Document, text: str, level=1):
    """Start a section on a new page without inserting a potentially blank break page."""
    heading = doc.add_heading(text, level=level)
    heading.paragraph_format.page_break_before = True
    return heading


def compact_table_text(table, size=9.5, line_spacing=1.0) -> None:
    """Tighten dense technical schedules while preserving legibility."""
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=55, start=105, bottom=55, end=105)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = line_spacing
                for run in paragraph.runs:
                    run.font.size = Pt(size)


def add_page_field(paragraph) -> None:
    paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_run(paragraph, text, bold=False, color=None, italic=False, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def add_bullet(doc, text: str, level=0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_compact_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.size = Pt(9.5)


def add_number(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_callout(doc, title: str, body: str, fill=PALE_BLUE, accent=DARK_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360], indent=0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=130, start=170, bottom=130, end=170)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title.upper(), bold=True, color=accent, size=9)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, body, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_label_value(doc, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, f"{label}: ", bold=True, color=DARK_BLUE)
    add_run(p, value)


def set_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Title", 23, INK, 0, 8),
        ("Subtitle", 11, GREY, 0, 14),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.font.bold = name not in ("Subtitle",)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(11)
        st.font.color.rgb = RGBColor.from_string(INK)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.15

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(GREY)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_together = True


def set_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        set_section(section)
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        add_run(p, "J40-HBL-RENEW-001  |  REV A", bold=True, color=DARK_BLUE, size=8)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_after = Pt(0)
        add_run(fp, "CONTROLLED WORKSHOP SPECIFICATION  •  ", color=GREY, size=8)
        add_page_field(fp)


def table_text(table, row_idx, values, bold_cols=(), colors=None) -> None:
    row = table.rows[row_idx]
    for idx, (cell, value) in enumerate(zip(row.cells, values)):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        run.bold = idx in bold_cols
        if colors and idx < len(colors) and colors[idx]:
            run.font.color.rgb = RGBColor.from_string(colors[idx])
    prevent_row_split(row)


def add_retrieval_table(doc: Document) -> None:
    rows = [
        ("R01", "Control-end anchor", "Handbrake lever/handle-side clevis, eye, pin seat or special cable anchor.", "Retrieve; label FRONT and installed orientation."),
        ("R02", "Primary adjustment assembly", "Threaded adjuster barrel/sleeve, sheath abutment, support flange/bracket and fitted jam hardware.", "Retrieve complete as an interface pattern; reuse only after inspection."),
        ("R03", "Oval compensator/equalizer", "Stacked oval plate shown in the R03 photo card, including its hole pattern and cable load faces.", "Retrieve. This unique geometry distributes the pull and is not replaced by a generic cable part."),
        ("R04", "Rear reaction/equalizer bracket", "Heavy plate and pivoting lever/bellcrank shown in the R04 photo card, including bushes/spacers and cable eye seats.", "Retrieve assembled; mark chassis/axle side and movement direction."),
        ("R05", "Rectangular U-yoke/retainer", "Loose rectangular wire yoke with two open legs shown in the R05 photo card.", "Retrieve and bag separately; photograph both faces beside a ruler."),
        ("R06", "Special cable-end terminals", "All eyes, rectangular loops, clevises, swaged blocks/nipples and load-bearing stops from every cable end.", "Retrieve even if attached to failed wire; use as terminal patterns."),
        ("R07", "Sheath interface fittings", "All ferrules, threaded sleeves, shoulder stops, end cups, transition sleeves and bracket-engaging grooves.", "Retrieve; record which end and which cable run each fitting belongs to."),
        ("R08", "Intermediate casing locators", "The fixed guide/locator sleeves and any welded or formed casing stops along the long run.", "Retrieve with the sheath if possible; copy positions from the functional datum."),
        ("R09", "Return/tension springs", "Every spring associated with cable return, lever return or equalizer stability.", "Retrieve for pattern and condition check; renew any stretched, corroded or distorted spring."),
        ("R10", "LH/RH drum interfaces", "Backing-plate parking-brake levers, links and the seating features for the standard internal drum wires.", "Retrieve/retain on vehicle; label LH and RH. Do not alter the standard interface."),
        ("R11", "Unique mounting hardware", "Any special bracket, abutment, spacer or pivot that fixes cable geometry.", "Retrieve. Generic loose washers, nuts, clips and grommets need not be specified, but keep fitted originals with the assembly for identification."),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_repeat_table_header_text(table.rows[0], ["ID", "Retained item", "What to recover", "Disposition"], dark=True)
    for item in rows:
        cells = table.add_row().cells
        for c, value in zip(cells, item):
            c.text = value
        cells[0].paragraphs[0].runs[0].bold = True
        cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(DARK_BLUE)
        prevent_row_split(table.rows[-1])
    set_table_geometry(table, [650, 1900, 3680, 3130])
    compact_table_text(table, size=9.0, line_spacing=0.95)


def add_replacement_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_repeat_table_header_text(table.rows[0], ["ID", "Replacement", "Qty", "Release basis", "Decision"], dark=True)
    data = [
        ("CBL-01", "Complete custom long hand-control-to-rear-axle cable assembly, including input run, two rear working legs and all permanent cable-end fittings/outer casing", "1", "Known-fit retained assembly measured between functional load datums under 20 N seating load", "Normally local sample-copy fabrication. Do not cut from J40 chassis dimensions alone."),
        ("CBL-02", "Standard internal rear-drum parking-brake wire", "2", "Toyota 47616-60010; one per rear wheel; unmodified standard drum interface", "Renew if either retained drum wire is frayed, bent, corroded or the terminals/guides are damaged."),
    ]
    for item in data:
        row = table.add_row()
        for cell, value in zip(row.cells, item):
            cell.text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(DARK_BLUE)
        prevent_row_split(row)
    set_table_geometry(table, [1050, 2800, 850, 2050, 2610])
    compact_table_text(table, size=9.5, line_spacing=1.0)


def add_measurement_table(doc: Document) -> None:
    measurements = [
        ("M01", "Inner effective length", "Terminal load seat A to terminal load seat B", "mm", "±1.0 mm from master"),
        ("M02", "Outer casing length", "Load-bearing casing abutment face to face", "mm", "±1.0 mm from master"),
        ("M03", "Exposed inner wire", "Casing abutment to terminal load seat at each end, cable seated", "mm", "Record each end; ±1.0 mm"),
        ("M04", "Working travel", "Full release position to full applied position at each controlled terminal", "mm", "Meet or exceed master without bottoming"),
        ("M05", "Inner wire diameter", "Three unworn straight locations; record construction/strand pattern", "mm", "Match master ±0.10 mm"),
        ("M06", "Casing OD/ID", "Three straight locations; record jacket and liner construction", "mm", "OD match ±0.25 mm; ID free-running"),
        ("M07", "Fixed feature positions", "Control-end functional datum to every ferrule, stop, guide and casing exit", "mm", "Feature order exact; positions ±1.0 mm"),
        ("M08", "Terminal geometry", "Dimension all load faces, eyes, loops, clevis jaws, nipples and stop blocks", "drawing/photo", "Exact functional interface match"),
        ("M09", "Adjuster thread/travel", "Major diameter, pitch, hand, engagement and usable adjustment", "mm/thread", "Exact interface; equal or greater safe engagement"),
        ("M10", "Route/bend envelope", "Installed route at ride height, full bump and controlled full droop", "photo/mm", "No tautness, rub or interference"),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_repeat_table_header_text(table.rows[0], ["ID", "Dimension", "Datum/method", "Record", "Acceptance"], dark=True)
    for item in measurements:
        row = table.add_row()
        for cell, value in zip(row.cells, item):
            cell.text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        prevent_row_split(row)
    set_table_geometry(table, [680, 1770, 3340, 1020, 2550])
    compact_table_text(table, size=9.0, line_spacing=0.95)


def add_signoff_table(doc: Document) -> None:
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    fields = [
        ("Master assembly tag", "", "Vehicle/axle identification", ""),
        ("Measured by", "", "Date", ""),
        ("Fabricated by", "", "Fabricator job/reference", ""),
        ("Installed/tested by", "", "Date", ""),
        ("Result", "PASS / FAIL", "Remarks", ""),
    ]
    for ridx, rowdata in enumerate(fields):
        for cidx, value in enumerate(rowdata):
            table.cell(ridx, cidx).text = value
            if cidx in (0, 2):
                set_cell_shading(table.cell(ridx, cidx), LIGHT_BLUE)
                table.cell(ridx, cidx).paragraphs[0].runs[0].bold = True
        prevent_row_split(table.rows[ridx])
    set_table_geometry(table, [1800, 2880, 1900, 2780])


def prep_image(src: Path, key: str, crop_norm=(0.0, 0.0, 1.0, 1.0)) -> Path:
    TMP.mkdir(parents=True, exist_ok=True)
    with Image.open(src) as raw:
        img = ImageOps.exif_transpose(raw).convert("RGB")
        w, h = img.size
        left, top, right, bottom = crop_norm
        crop = img.crop((int(w * left), int(h * top), int(w * right), int(h * bottom)))
        fitted = ImageOps.pad(
            crop,
            (1440, 960),
            method=Image.Resampling.LANCZOS,
            color=(248, 248, 248),
            centering=(0.5, 0.5),
        )
        out = TMP / f"retrieval_{key.lower()}.jpg"
        fitted.save(out, quality=90, optimize=True)
        return out


def set_picture_alt(inline_shape, description: str) -> None:
    inline_shape._inline.docPr.set("descr", description)


def add_figure(doc: Document, image_path: Path, caption: str, width=Inches(3.0)) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(image_path), width=width)
    set_picture_alt(picture, caption)
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.add_run(caption)


def add_photo_card(doc: Document, item: dict, image_path: Path) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [4300, 5060], indent=0)
    prevent_row_split(table.rows[0])

    image_cell, text_cell = table.rows[0].cells
    set_cell_shading(image_cell, "F8F8F8")
    set_cell_shading(text_cell, PALE_BLUE)
    set_cell_margins(image_cell, top=90, start=90, bottom=90, end=90)
    set_cell_margins(text_cell, top=110, start=140, bottom=90, end=140)

    p = image_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    picture = p.add_run().add_picture(str(image_path), width=Inches(2.85))
    set_picture_alt(
        picture,
        f"{item['id']} {item['title']}: retained handbrake component retrieval reference.",
    )

    p = text_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, f"{item['id']}  {item['title']}", bold=True, color=DARK_BLUE, size=11)
    for label, value in (
        ("Identify", item["identify"]),
        ("Retrieve", item["action"]),
        ("Evidence", item["confidence"]),
    ):
        p = text_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, f"{label}: ", bold=True, color=DARK_BLUE, size=9)
        add_run(p, value, size=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_wrong_reference_card(doc: Document, image_path: Path) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [4300, 5060], indent=0)
    prevent_row_split(table.rows[0])

    image_cell, text_cell = table.rows[0].cells
    set_cell_shading(image_cell, "F8F8F8")
    set_cell_shading(text_cell, PALE_RED)
    set_cell_margins(image_cell, top=90, start=90, bottom=90, end=90)
    set_cell_margins(text_cell, top=110, start=140, bottom=90, end=140)

    p = image_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    picture = p.add_run().add_picture(str(image_path), width=Inches(2.85))
    set_picture_alt(
        picture,
        "Wrong or unverified replacement handbrake cable retained only as a quarantined reference.",
    )

    p = text_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, WRONG_REFERENCE["title"], bold=True, color=RED, size=11)
    p = text_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, WRONG_REFERENCE["body"], size=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_styles(doc)
    set_section(doc.sections[0])

    # First-page memo masthead.
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "J40 WORKSHOP CONTROL", bold=True, color=BLUE, size=10)
    p = doc.add_paragraph(style="Title")
    p.add_run("Rear Handbrake Cable Renewal")
    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Retrieval list and fabrication specification for the retained custom/hybrid installation")

    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    meta_data = [
        ("Document", "J40-HBL-RENEW-001", "Revision", "A"),
        ("Issued", date(2026, 8, 16).isoformat(), "Status", "WORKSHOP RELEASE — MEASURE MASTER BEFORE FABRICATION"),
    ]
    for ridx, values in enumerate(meta_data):
        for cidx, value in enumerate(values):
            meta.cell(ridx, cidx).text = value
            if cidx in (0, 2):
                set_cell_shading(meta.cell(ridx, cidx), LIGHT_BLUE)
                meta.cell(ridx, cidx).paragraphs[0].runs[0].bold = True
        prevent_row_split(meta.rows[ridx])
    set_table_geometry(meta, [1350, 2600, 1050, 4360], indent=0)

    doc.add_heading("1. Release decision", level=1)
    add_callout(
        doc,
        "Required action",
        "Recover and identify all unique metal hardware from the previous installation. Renew the complete inner wire and flexible outer casing of the long assembly. Retain the old cables only as dimensional and terminal-pattern masters; do not reinstall visibly frayed wire. The two drum-entry wires must remain the standard Toyota interface.",
        fill=PALE_BLUE,
        accent=DARK_BLUE,
    )

    add_callout(
        doc,
        "Rejected reference",
        "Daraz order 243701549680938 was confirmed wrong size. Mark it “REJECTED — WRONG SIZE — NOT MASTER — DO NOT INSTALL.” Seller fitment wording such as “FJ40/BJ40” is not sufficient evidence of dimensional compatibility.",
        fill=PALE_RED,
        accent=RED,
    )

    doc.add_heading("2. System boundary", level=1)
    p = doc.add_paragraph()
    add_run(p, "Configuration. ", bold=True)
    add_run(p, "This is not the original stock 1978 J40 cable arrangement. It is a retained single-input, equalized rear system: one hand-control pull is distributed to left and right rear drum mechanisms through the recovered equalizer/yoke and brackets.")
    p = doc.add_paragraph()
    add_run(p, "Dimensional authority. ", bold=True)
    add_run(p, "The complete known-fit removed assembly and the installed functional datums control every cut length and terminal position. Standard J40 wheelbase (2,285 mm) and project rear-track reference (1,410 mm) are routing-envelope checks only; neither is a cable cut length.")
    p = doc.add_paragraph()
    add_run(p, "Procurement scope. ", bold=True)
    add_run(p, "The replacement scope is cable inner wire, flexible outer casing and permanent cable-end fittings. Generic loose washers, nuts, clips and grommets are outside the purchasing specification, although fitted originals should remain bagged with their mating parts for identification.")

    heading_new_page(doc, "3. Parts to retrieve from the old installation", level=1)
    p = doc.add_paragraph()
    p.add_run("Tag every item before separation. Use FRONT/REAR, LH/RH and installed-orientation marks; take an overview photograph and a close-up with a ruler before cutting any cable.")
    add_retrieval_table(doc)

    heading = doc.add_heading("3.1 Reuse inspection", level=2)
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(4)
    add_compact_bullet(doc, "Clean only enough to inspect. Do not grind, weld or enlarge holes before the master measurements are complete.")
    add_compact_bullet(doc, "Reject a load-bearing metal part for service if cracked, bent out of plane, deeply pitted, thinned at a load face, or if a pivot/eye is elongated enough to change travel or alignment.")
    add_compact_bullet(doc, "Keep a rejected unique part as a manufacturing pattern, clearly tagged NOT FOR REUSE.")
    add_compact_bullet(doc, "The new photos show broken/frayed strands at the equalizer and bracket area. That wire is a pattern only and must be replaced.")

    prepped = {
        item["id"]: prep_image(item["path"], item["id"], item["crop"])
        for item in PHOTO_GUIDE
    }
    wrong_prepped = prep_image(
        WRONG_REFERENCE["path"],
        "wrong_reference",
        WRONG_REFERENCE["crop"],
    )
    heading_new_page(doc, "3.2 Photo retrieval guide", level=2)
    p = doc.add_paragraph()
    add_run(p, "Review scope: ", bold=True, color=DARK_BLUE)
    add_run(
        p,
        "632 Google Photos captured from 17 April to 2 August 2026 were checked. "
        "The cards below separate the identifiable retrieval components; the source photographs are preserved unchanged in the project photo index.",
    )
    for item in PHOTO_GUIDE[:2]:
        add_photo_card(doc, item, prepped[item["id"]])

    for start in range(2, len(PHOTO_GUIDE), 2):
        heading_new_page(doc, "3.2 Photo retrieval guide (continued)", level=2)
        for item in PHOTO_GUIDE[start:start + 2]:
            add_photo_card(doc, item, prepped[item["id"]])

    heading_new_page(doc, "3.3 Quarantined reference and old-cable sizing evidence", level=2)
    add_wrong_reference_card(doc, wrong_prepped)
    add_callout(
        doc,
        "Photo screening value only",
        "The laid-out old cable appears approximately 3,280–3,290 mm overall in the tape photographs. "
        "This is a screening reference, not a released fabrication dimension: remeasure the physical master between defined load datums, and separately record casing length, exposed inner wire, feature positions, threads and working travel.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    heading_new_page(doc, "4. Cable replacement schedule", level=1)
    add_replacement_table(doc)

    doc.add_heading("4.1 CBL-01 controlled sub-runs", level=2)
    add_bullet(doc, "CBL-01-A — input run: hand control/primary adjuster to the rear equalizer reaction point.")
    add_bullet(doc, "CBL-01-B — LH working leg: equalizer to the left rear drum interface.")
    add_bullet(doc, "CBL-01-C — RH working leg: equalizer to the right rear drum interface.")
    p = doc.add_paragraph()
    p.add_run("The retained hardware determines whether these are continuous inner-wire paths, separately sheathed branches or captive subassemblies. The fabricator shall reproduce the observed architecture exactly and supply CBL-01 as one matched assembly.")

    doc.add_heading("4.2 Standard drum-entry interface", level=2)
    p = doc.add_paragraph()
    add_run(p, "Released part: ", bold=True)
    add_run(p, "Toyota 47616-60010, quantity 2 (one per rear wheel). Install unmodified and dry-fit to the actual backing-plate actuator before final assembly. If it does not seat and travel correctly, correct the vehicle-side interface or part selection—do not cut, weld or re-swage the standard wire.")

    add_callout(
        doc,
        "Photo screening values — not cut lengths",
        "The laid-out old cable appears approximately 3,280–3,290 mm overall in the tape photographs. Use this only to reject an obviously short candidate. Perspective, cable curvature and uncertain terminal datums prevent release of any fabrication length from the photographs.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    heading_new_page(doc, "5. Mandatory master measurement before fabrication", level=1)
    p = doc.add_paragraph()
    p.add_run("Lay each cable path in its natural straight condition without kinks. Apply a 20 N axial seating load only—enough to seat the terminals and remove slack, not to stretch the assembly. Measure to 1 mm using the functional load seats, not overall loose-end tips.")
    add_measurement_table(doc)

    doc.add_heading("5.1 Measurement record", level=2)
    p = doc.add_paragraph()
    p.add_run("Complete one signed record for CBL-01-A, CBL-01-B and CBL-01-C. Attach dimensioned photographs showing the datum labels, caliper readings, thread gauge and the sequence of all fixed ferrules/stops. Do not release fabrication while any measurement field is blank.")

    heading_new_page(doc, "6. Fabrication specification", level=1)
    doc.add_heading("6.1 Inner cable", level=2)
    add_bullet(doc, "Flexible, corrosion-resistant, multi-strand steel automotive parking-brake/control cable suitable for repeated pull loading and the vehicle environment.")
    add_bullet(doc, "Match the retained master’s diameter within ±0.10 mm and reproduce its flexibility/strand character; do not substitute rigid wire or a materially smaller bicycle-type cable.")
    add_bullet(doc, "Use new continuous wire for every path. No knots, twists, splices, screw-block clamps or solder-only load terminations.")
    add_bullet(doc, "The cable shop shall record the supplied cable construction and its documented tensile/load rating. The rating must be no lower than the professional shop’s approved parking-brake cable material for a comparable vehicle.")

    doc.add_heading("6.2 Outer casing", level=2)
    add_bullet(doc, "Lined spiral/helical steel control-cable casing with a weatherproof black polymer jacket, intended for automotive parking-brake service.")
    add_bullet(doc, "Match the master outside diameter within ±0.25 mm, fit every retained bracket/abutment without looseness, and provide an internal bore that permits free movement through all installed bends.")
    add_bullet(doc, "Reproduce every casing length, end ferrule, shoulder, threaded sleeve, bracket groove and intermediate locator position from the signed master record.")
    add_bullet(doc, "Seal casing ends against water and grit without reducing cable travel. Do not reuse cracked, crushed, corroded or internally rough casing.")

    heading_new_page(doc, "6.3 Terminals and retained hardware", level=2)
    add_bullet(doc, "Professionally swage/crimp all permanent terminals and ferrules. Reproduce the original load-seat geometry, orientation and terminal-to-wire datum.")
    add_bullet(doc, "Transfer an original terminal only if a competent cable fabricator confirms it is crack-free, not thinned/distorted and suitable for a new professional swage. Otherwise manufacture a new terminal from the retained pattern.")
    add_bullet(doc, "Match adjuster thread diameter, pitch, hand, engagement and usable travel exactly. Set the installed adjuster near the middle of its range after the rear shoes are correctly adjusted.")
    add_bullet(doc, "No welded cable strands, improvised cable clamps, hardware-store eye bolts, or terminal modifications that introduce sharp bending at the wire exit.")

    doc.add_heading("6.4 Installation envelope", level=2)
    add_bullet(doc, "Route as photographed/marked and support the casing at the original functional abutments. Avoid sharp bends and use no bend tighter than the master route or the cable manufacturer’s minimum radius, whichever is larger.")
    add_bullet(doc, "At ride height, controlled full bump and controlled full droop, the cable must not become taut or contact the prop shaft, exhaust, springs, tyres, sharp edges or hydraulic brake lines.")
    add_bullet(doc, "The equalizer must remain free to balance LH/RH pull without fouling a bracket or reaching a hard stop before the drums are applied.")

    heading_new_page(doc, "7. Inspection and acceptance", level=1)
    doc.add_heading("7.1 Bench release", level=2)
    for step in [
        "Compare the finished assembly against the signed master record. Confirm all effective lengths, feature order, sheath stops, terminals and threads.",
        "Cycle every inner wire through full specified travel by hand. Movement must be smooth, without scraping, kinking, birdcaging or delayed return.",
        "Have the professional cable shop proof-test every swaged termination to its documented automotive parking-brake procedure. Record load and duration; accept only with no slip, deformation, strand damage or ferrule movement.",
    ]:
        add_number(doc, step)

    doc.add_heading("7.2 Installed functional test", level=2)
    for step in [
        "Adjust the rear brake shoes correctly before using cable adjustment to remove free play.",
        "With the rear axle safely supported, operate the handbrake through 25 full apply/release cycles. Both wheels must begin applying together and release completely after every cycle.",
        "Repeat the clearance check through controlled full suspension travel. Confirm no tautness, chafing, sheath movement at abutments or equalizer interference.",
        "Perform a controlled static hold test on a suitable incline using wheel chocks and a secondary braking method. A competent mechanic must verify secure hold in both vehicle directions and full release afterward.",
        "Re-inspect terminals, brackets, pivots and adjustment after the test; record the result below. Do not road-release the vehicle after any failed criterion.",
    ]:
        add_number(doc, step)

    doc.add_heading("7.3 Sign-off", level=2)
    add_signoff_table(doc)

    heading_new_page(doc, "Appendix A — Fabricator handover checklist", level=1)
    checklist = [
        "Complete known-fit old assembly delivered, tagged and photographed before cutting.",
        "R01–R11 recovered or their vehicle-side location positively identified.",
        "CBL-01-A/B/C architecture confirmed: continuous versus separate inner wires and casing sections.",
        "M01–M10 recorded for each controlled run under 20 N seating load.",
        "Every end fitting photographed and dimensioned from its functional load seat.",
        "Wire/casing material and professional swage method agreed before fabrication.",
        "CBL-02 Toyota 47616-60010 interfaces dry-fitted at LH and RH drums if renewed.",
        "Bench proof-test record and installed acceptance results attached.",
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_repeat_table_header_text(table.rows[0], ["OK", "Handover item", "Initial/date"], dark=True)
    for item in checklist:
        row = table.add_row()
        row.cells[0].text = "☐"
        row.cells[1].text = item
        row.cells[2].text = ""
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        prevent_row_split(row)
    set_table_geometry(table, [700, 6950, 1710])

    doc.add_heading("Controlled references", level=2)
    refs = [
        "Project controlled specification: docs/j40-handbrake-line-components-exact-spec-20260816.md",
        "Project component schedule: data/manual/j40_handbrake_line_component_spec_20260816.csv",
        "Project master-measurement register: data/manual/j40_handbrake_line_master_measurements_20260816.csv",
        "Four-month Google Photos retrieval index: docs/j40-handbrake-component-photo-retrieval-index-20260816.md",
        "Separated unchanged photo copies and manifest: photos/handbrake_retrieval_components_20260816/.",
    ]
    for ref in refs:
        add_bullet(doc, ref)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "END OF CONTROLLED SPECIFICATION", bold=True, color=DARK_BLUE, size=9)

    add_header_footer(doc)
    doc.core_properties.title = "J40 Rear Handbrake Cable Renewal — Retrieval & Fabrication Specification"
    doc.core_properties.subject = "Controlled retrieval list and cable/sheath fabrication requirements"
    doc.core_properties.author = "J40 Restoration Project"
    doc.core_properties.keywords = "Toyota J40, handbrake, parking brake, cable, sheath, fabrication"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
