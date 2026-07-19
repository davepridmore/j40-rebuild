#!/usr/bin/env python3
"""Build the image-led Bilal Ganj cooling and A/C buying guide."""

from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "bilal-ganj-cooling-ac-visual-buying-guide-20260714.docx"
ASSET_DIR = ROOT / "docs" / "generated" / "bilal_ganj_cooling_ac_visual_guide_20260714"

CAP_PHOTO = ROOT / "deliverables" / "selling_site_images" / "images" / "manual_overrides" / "radiator_cap_current_car_crop_20260503.jpg"
BOTTLE_PHOTO = ROOT / "photos" / "20260512_073402_gp_P6yrwLRw.jpg"
FAN_ZONE_PHOTO = ROOT / "photos" / "20260512_100000_user_front_support_radiator_pickups_context.png"
ELECTRICAL_PANEL_PHOTO = ROOT / "photos" / "20260411_071153.jpg"
OLD_STACK_PHOTO = ROOT / "photos" / "20260420_021227_gp_iHBRfJDA.jpg"
OLD_RADIATOR_PHOTO = ROOT / "photos" / "20260503_155956_gp_P4xfMJzw.jpg"
FAN_SHROUD_PHOTO = ROOT / "photos" / "20260503_160010_gp_9F5ZH8kQ.jpg"
COMPRESSOR_PHOTO = ROOT / "photos" / "20260501_194508_gp_aSPxPLDw.jpg"
ENGINE_DRIVE_PHOTO = ROOT / "photos" / "20260501_194535_gp_mZ25Ou4A.jpg"
BOTTLES_CONTEXT_PHOTO = ROOT / "photos" / "20260503_153639_gp_ZueGlpJw.jpg"
COOLANT_HOSE_PHOTO = ROOT / "photos" / "20260529_000708_gp_v6oy8EoQ.jpg"

ARIAL = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
ARIAL_BOLD = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")

INK = "#17324D"
BLUE = "#2E74B5"
TEAL = "#147D78"
GREEN = "#2F7D32"
AMBER = "#D89000"
RED = "#B33A3A"
GRAY = "#66717D"
LIGHT = "#F3F6F8"
MID = "#D6DEE5"
WHITE = "#FFFFFF"
BLACK = "#16202A"


def fnt(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(ARIAL_BOLD if bold else ARIAL), size)


def wrapped_lines(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), trial, font=font)[2] <= width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def text_block(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    fill: str,
    max_width: int,
    line_gap: int = 8,
) -> int:
    x, y = xy
    line_height = font.size + line_gap
    for line in wrapped_lines(draw, text, font, max_width):
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height
    return y


def new_card(title: str, subtitle: str = "", size: tuple[int, int] = (1800, 720)):
    img = Image.new("RGB", size, WHITE)
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((12, 12, size[0] - 12, size[1] - 12), radius=24, outline=MID, width=4, fill=WHITE)
    draw.rectangle((12, 12, size[0] - 12, 112), fill=INK)
    title_size = 48
    subtitle_font = fnt(25)
    subtitle_width = draw.textbbox((0, 0), subtitle, font=subtitle_font)[2] if subtitle else 0
    title_limit = size[0] - 150 - subtitle_width
    while title_size > 32 and draw.textbbox((0, 0), title, font=fnt(title_size, True))[2] > title_limit:
        title_size -= 2
    draw.text((55, 34), title, font=fnt(title_size, True), fill=WHITE)
    if subtitle:
        draw.text((size[0] - 55, 48), subtitle, font=subtitle_font, fill="#DDE9F2", anchor="ra")
    return img, draw


def label(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: str = LIGHT, color: str = INK, size: int = 28):
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=MID, width=2)
    x1, y1, x2, y2 = box
    draw.text(((x1 + x2) // 2, (y1 + y2) // 2), text, font=fnt(size, True), fill=color, anchor="mm")


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = RED, width: int = 7, head: int = 20):
    draw.line((start, end), fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    for delta in (2.6, -2.6):
        p = (end[0] + head * math.cos(angle + delta), end[1] + head * math.sin(angle + delta))
        draw.line((end, p), fill=color, width=width)


def double_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], text: str, offset: tuple[int, int] = (0, -24), color: str = RED):
    arrow(draw, start, end, color=color, width=5, head=16)
    arrow(draw, end, start, color=color, width=5, head=16)
    mid = ((start[0] + end[0]) // 2 + offset[0], (start[1] + end[1]) // 2 + offset[1])
    bbox = draw.textbbox((0, 0), text, font=fnt(28, True))
    pad = 10
    draw.rounded_rectangle((mid[0] - (bbox[2] - bbox[0]) // 2 - pad, mid[1] - 21, mid[0] + (bbox[2] - bbox[0]) // 2 + pad, mid[1] + 21), radius=10, fill=WHITE)
    draw.text(mid, text, font=fnt(28, True), fill=color, anchor="mm")


def check(draw: ImageDraw.ImageDraw, xy: tuple[int, int], radius: int = 18, color: str = GREEN):
    x, y = xy
    draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=color)
    draw.line((x - 9, y, x - 2, y + 8, x + 11, y - 9), fill=WHITE, width=5, joint="curve")


def cross(draw: ImageDraw.ImageDraw, xy: tuple[int, int], radius: int = 18):
    x, y = xy
    draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=RED)
    draw.line((x - 8, y - 8, x + 8, y + 8), fill=WHITE, width=5)
    draw.line((x - 8, y + 8, x + 8, y - 8), fill=WHITE, width=5)


def save(img: Image.Image, name: str) -> Path:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    path = ASSET_DIR / name
    img.save(path, quality=95)
    return path


def paste_photo(
    img: Image.Image,
    path: Path,
    box: tuple[int, int, int, int],
    crop: tuple[int, int, int, int] | None = None,
    contain: bool = False,
    rotate: int | None = None,
) -> None:
    source = ImageOps.exif_transpose(Image.open(path)).convert("RGB")
    if rotate:
        source = source.rotate(rotate, expand=True)
    if crop:
        source = source.crop(crop)
    x1, y1, x2, y2 = box
    target_size = (x2 - x1, y2 - y1)
    if contain:
        source.thumbnail(target_size, Image.Resampling.LANCZOS)
        px = x1 + (target_size[0] - source.width) // 2
        py = y1 + (target_size[1] - source.height) // 2
        img.paste(source, (px, py))
    else:
        source = ImageOps.fit(source, target_size, method=Image.Resampling.LANCZOS)
        img.paste(source, (x1, y1))


def photo_card(
    img: Image.Image,
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    path: Path,
    title: str,
    status: str,
    note: str,
    status_color: str,
    crop: tuple[int, int, int, int] | None = None,
    rotate: int | None = None,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=20, fill=WHITE, outline=MID, width=4)
    photo_box = (x1 + 14, y1 + 14, x2 - 14, y1 + 270)
    paste_photo(img, path, photo_box, crop=crop, rotate=rotate)
    draw.rectangle(photo_box, outline=INK, width=2)
    draw.rounded_rectangle((x1 + 28, y1 + 28, x1 + 350, y1 + 82), radius=14, fill=status_color)
    draw.text((x1 + 189, y1 + 55), status, font=fnt(24, True), fill=WHITE, anchor="mm")
    draw.text((x1 + 28, y1 + 300), title, font=fnt(31, True), fill=INK)
    text_block(draw, (x1 + 28, y1 + 345), note, fnt(23, True), GRAY, x2 - x1 - 56, line_gap=5)


def draw_existing_components() -> Path:
    img, d = new_card(
        "ORIGINAL PARTS: REUSE OR TAKE AS SAMPLES",
        "actual J40 project photos",
        (1800, 1110),
    )
    photo_card(
        img,
        d,
        (55, 150, 875, 600),
        CAP_PHOTO,
        "Current radiator cap and filler neck",
        "TEMPLATE ONLY",
        "Take the cap and old radiator. Buy a NEW cap matching the new filler neck and confirmed pressure rating.",
        RED,
    )
    photo_card(
        img,
        d,
        (925, 150, 1745, 600),
        BOTTLE_PHOTO,
        "Current coolant recovery bottle",
        "REUSE IF SOUND",
        "Clean and inspect for cracks, brittle plastic, blocked nipple and a sealing cap. Otherwise buy an exact local match by sample.",
        GREEN,
        crop=(0, 1420, 1720, 3320),
    )
    photo_card(
        img,
        d,
        (55, 635, 875, 1085),
        ELECTRICAL_PANEL_PHOTO,
        "Existing relay and blade-fuse panel",
        "REUSE PANEL",
        "Confirm protected circuits for the primary electric fan and compressor clutch. Add a separate circuit only if a pusher fan is fitted.",
        BLUE,
    )
    photo_card(
        img,
        d,
        (925, 635, 1745, 1085),
        FAN_ZONE_PHOTO,
        "Radiator rear / electric-fan clearance zone",
        "PHYSICAL CHECK",
        "Confirm depth between the radiator and engine belts. The retained electric fan/shroud assembly is the full donor template.",
        AMBER,
        crop=(0, 0, 778, 900),
    )
    return save(img, "02_existing_components.png")


def draw_original_cooling_stack() -> Path:
    img, d = new_card(
        "ORIGINAL COOLING + A/C STACK BEFORE REMOVAL",
        "actual vehicle photo",
        (1800, 1110),
    )
    photo_box = (55, 145, 1745, 900)
    paste_photo(img, OLD_STACK_PHOTO, photo_box)
    d.rectangle(photo_box, outline=INK, width=4)

    label(d, (85, 170, 530, 250), "OLD ENGINE RADIATOR (REAR)", fill="#E8F1F8", color=BLUE, size=25)
    arrow(d, (530, 215), (970, 245), color=BLUE, width=7, head=22)
    label(d, (85, 720, 535, 800), "OLD A/C CONDENSER (FRONT)", fill="#E7F4F2", color=TEAL, size=25)
    arrow(d, (535, 760), (880, 580), color=TEAL, width=7, head=22)
    label(d, (1230, 170, 1715, 265), "OLD A/C HARD LINES\nROUTING SAMPLE ONLY", fill="#FFF4DD", color=AMBER, size=24)
    arrow(d, (1440, 265), (1325, 420), color=AMBER, width=7, head=22)
    label(d, (1225, 720, 1715, 815), "SIDE BRACKETS + OPENING\nGEOMETRY TEMPLATE", fill=LIGHT, color=INK, size=24)
    arrow(d, (1425, 720), (1500, 610), color=INK, width=7, head=22)

    d.text((900, 955), "KEEP THE COMPLETE OLD STACK FOR COMPARISON AT BILAL GANJ", font=fnt(30, True), fill=RED, anchor="mm")
    d.text(
        (900, 1015),
        "Copy fitting side, neck locations, hose routes and bracket geometry. Replace the aged radiator, condenser and refrigerant hoses.",
        font=fnt(24, True),
        fill=GRAY,
        anchor="mm",
    )
    return save(img, "02_original_cooling_stack.png")


def draw_original_reuse_parts() -> Path:
    img, d = new_card(
        "ORIGINAL ENGINE-SIDE PARTS TO RETAIN",
        "actual vehicle photos",
        (1800, 1110),
    )
    photo_card(
        img,
        d,
        (55, 150, 875, 600),
        FAN_SHROUD_PHOTO,
        "Electric puller fan + integrated shroud",
        "RECONDITION / REUSE",
        "Bench-test motor and direction; inspect blades, shroud, wiring and mounts. Use the complete assembly as the donor template.",
        GREEN,
        rotate=90,
    )
    photo_card(
        img,
        d,
        (925, 150, 1745, 600),
        COMPRESSOR_PHOTO,
        "Existing A/C compressor and fittings",
        "TEST THEN REUSE",
        "Check shaft, clutch bearing, clutch coil, oil condition and leaks. Copy both port/fitting sizes before hose fabrication.",
        TEAL,
        crop=(0, 0, 2300, 1100),
    )
    photo_card(
        img,
        d,
        (55, 635, 875, 1085),
        ENGINE_DRIVE_PHOTO,
        "Alternator, pulleys and belt routing",
        "TEST THEN REUSE",
        "Keep the pulley layout and old belts as templates. Fit new belts and verify alternator output with blower and electric fan running.",
        AMBER,
    )
    photo_card(
        img,
        d,
        (925, 635, 1745, 1085),
        ELECTRICAL_PANEL_PHOTO,
        "Existing relay and blade-fuse panel",
        "REUSE / ALLOCATE",
        "Confirm protected circuits for the primary fan and compressor clutch. An optional pusher fan needs its own relay and fuse.",
        BLUE,
    )
    return save(img, "03_original_reuse_parts.png")


def horizontal_photo_card(
    img: Image.Image,
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    path: Path,
    title: str,
    status: str,
    note: str,
    status_color: str,
    crop: tuple[int, int, int, int] | None = None,
    rotate: int | None = None,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=20, fill=WHITE, outline=MID, width=4)
    photo_box = (x1 + 14, y1 + 14, x1 + 440, y2 - 14)
    paste_photo(img, path, photo_box, crop=crop, rotate=rotate)
    draw.rectangle(photo_box, outline=INK, width=2)
    draw.rounded_rectangle((x1 + 465, y1 + 22, x1 + 815, y1 + 72), radius=13, fill=status_color)
    draw.text((x1 + 640, y1 + 47), status, font=fnt(22, True), fill=WHITE, anchor="mm")
    draw.text((x1 + 465, y1 + 92), title, font=fnt(29, True), fill=INK)
    text_block(draw, (x1 + 465, y1 + 137), note, fnt(22, True), GRAY, x2 - x1 - 495, line_gap=5)


def draw_original_service_samples() -> Path:
    img, d = new_card(
        "ORIGINAL COOLING PARTS TO TAKE AS SAMPLES",
        "actual removed-part photos",
        (1800, 1110),
    )
    radiator_box = (55, 150, 655, 1085)
    d.rounded_rectangle(radiator_box, radius=20, fill=WHITE, outline=MID, width=4)
    radiator_photo_box = (70, 165, 640, 860)
    paste_photo(img, OLD_RADIATOR_PHOTO, radiator_photo_box, crop=(0, 850, 1884, 3150))
    d.rectangle(radiator_photo_box, outline=INK, width=2)
    d.rounded_rectangle((85, 180, 445, 235), radius=14, fill=RED)
    d.text((265, 207), "MASTER TEMPLATE", font=fnt(23, True), fill=WHITE, anchor="mm")
    d.text((85, 890), "Removed engine radiator", font=fnt(31, True), fill=INK)
    text_block(
        d,
        (85, 935),
        "Take the complete radiator. Copy top/bottom tanks, hose necks, cap neck, overflow nipple, drain and mounting ears; do not copy the bad support leg.",
        fnt(22, True),
        GRAY,
        540,
        line_gap=5,
    )

    horizontal_photo_card(
        img,
        d,
        (700, 150, 1745, 430),
        CAP_PHOTO,
        "Radiator cap and filler neck",
        "TEMPLATE ONLY",
        "Match the neck style and confirmed pressure, but install a new cap and seal.",
        RED,
    )
    horizontal_photo_card(
        img,
        d,
        (700, 465, 1745, 745),
        BOTTLES_CONTEXT_PHOTO,
        "Coolant bottle and washer bottle",
        "DO NOT CONFUSE",
        "The cylindrical strapped bottle is coolant recovery. The shaped bottle with wiring is the windscreen washer reservoir.",
        AMBER,
        crop=(0, 1150, 1884, 3020),
    )
    horizontal_photo_card(
        img,
        d,
        (700, 780, 1745, 1060),
        COOLANT_HOSE_PHOTO,
        "Original coolant hose shape",
        "COPY SHAPE / BUY NEW",
        "Use the removed hose and metal necks for bore, bend and route. Replace all aged rubber and clamps.",
        TEAL,
    )
    return save(img, "04_original_service_samples.png")


def draw_original_fan_photo() -> Path:
    img, d = new_card(
        "ORIGINAL ELECTRIC PULLER FAN + INTEGRATED SHROUD",
        "actual removed assembly - rear / engine side of radiator",
        (1800, 980),
    )
    photo_box = (55, 145, 1260, 900)
    d.rectangle(photo_box, fill=LIGHT)
    paste_photo(img, FAN_SHROUD_PHOTO, photo_box, contain=True, rotate=90)
    d.rectangle(photo_box, outline=INK, width=4)

    label(d, (1310, 165, 1735, 260), "PRIMARY COOLING FAN", fill="#E7F4F2", color=TEAL, size=28)
    label(d, (1310, 295, 1735, 390), "RECONDITION + REUSE", fill="#EAF4E8", color=GREEN, size=28)
    label(d, (1310, 425, 1735, 540), "AIR MUST PULL THROUGH\nRADIATOR TO ENGINE", fill="#E8F1F8", color=BLUE, size=25)
    label(d, (1310, 575, 1735, 690), "MEASURE START +\nRUN CURRENT", fill="#FFF4DD", color=AMBER, size=27)
    label(d, (1310, 725, 1735, 840), "COPY COMPLETE ASSEMBLY\nIF REPLACEMENT IS NEEDED", fill=LIGHT, color=INK, size=24)
    d.text(
        (900, 940),
        "Inspect motor bearings, blades, plastic shroud, wiring, plug and mounting tabs before installation.",
        font=fnt(25, True),
        fill=RED,
        anchor="mm",
    )
    return save(img, "05_original_electric_fan.png")


def draw_original_cap_photo() -> Path:
    img, d = new_card(
        "ORIGINAL RADIATOR CAP + FILLER NECK",
        "actual removed radiator - template only",
        (1800, 800),
    )
    photo_box = (55, 145, 1135, 735)
    d.rectangle(photo_box, fill=LIGHT)
    paste_photo(img, CAP_PHOTO, photo_box, contain=True)
    d.rectangle(photo_box, outline=INK, width=4)
    label(d, (1195, 165, 1735, 260), "TAKE THIS SAMPLE", fill="#FFF0E9", color=RED, size=30)
    label(d, (1195, 300, 1735, 395), "MATCH NECK + SEAL", fill="#E8F1F8", color=BLUE, size=29)
    label(d, (1195, 435, 1735, 530), "CONFIRM PRESSURE RATING", fill="#FFF4DD", color=AMBER, size=27)
    label(d, (1195, 570, 1735, 675), "BUY A NEW CAP\nFOR THE NEW RADIATOR", fill="#EAF4E8", color=GREEN, size=27)
    return save(img, "08_original_cap.png")


def draw_stack() -> Path:
    img, d = new_card("HOW THE PARTS FIT TOGETHER", "front-to-rear side view", (1800, 980))
    y1, y2 = 210, 600
    parts = [
        (170, 250, "GRILLE", GRAY),
        (315, 365, "OPTIONAL\nPUSHER\n60-65 mm", BLUE),
        (520, 555, "CONDENSER\n21-25 mm", TEAL),
        (650, 690, "AIR GAP\n10-15 mm", AMBER),
        (790, 900, "RADIATOR\n60-64 mm", GREEN),
        (1015, 1185, "ORIGINAL ELECTRIC\nPULLER + SHROUD", BLUE),
        (1325, 1415, "ENGINE", GRAY),
    ]
    for x1, x2, name, color in parts:
        d.rounded_rectangle((x1, y1, x2, y2), radius=14, fill=color, outline=INK, width=3)
        d.multiline_text(((x1 + x2) // 2, y2 + 34), name, font=fnt(20, True), fill=INK, anchor="ma", align="center", spacing=2)
    arrow(d, (80, 460), (1500, 460), color=RED, width=9, head=30)
    d.text((1515, 460), "AIR TO ENGINE", font=fnt(34, True), fill=RED, anchor="lm")
    double_arrow(d, (520, 750), (900, 750), "BASE HARD STACK 91-104 mm", offset=(0, -28))
    label(d, (85, 835, 540, 925), "1. BUY HARD PARTS", fill="#E8F1F8", color=BLUE, size=31)
    arrow(d, (560, 880), (690, 880), color=GRAY)
    label(d, (710, 835, 1120, 925), "2. DRY-FIT ON CAR", fill="#E7F4F2", color=TEAL, size=31)
    arrow(d, (1140, 880), (1270, 880), color=GRAY)
    label(d, (1290, 835, 1715, 925), "3. MARK + CRIMP", fill="#FFF4DD", color=AMBER, size=31)
    return save(img, "01_stack.png")


def draw_radiator() -> Path:
    img, d = new_card("1. NEW HJ47 / 2H RADIATOR", "take old radiator as sample")
    x1, y1, x2, y2 = 250, 190, 1080, 585
    d.rectangle((x1, y1, x2, y2), fill="#DCE6DD", outline=INK, width=6)
    d.rectangle((x1, y1, x2, y1 + 58), fill=GREEN, outline=INK, width=4)
    d.rectangle((x1, y2 - 58, x2, y2), fill=GREEN, outline=INK, width=4)
    for x in range(x1 + 20, x2, 28):
        d.line((x, y1 + 65, x, y2 - 65), fill="#91A59A", width=2)
    for y in range(y1 + 70, y2 - 65, 20):
        d.line((x1 + 5, y, x2 - 5, y), fill="#B7C7BB", width=2)
    d.rectangle((x1 - 45, y1 + 15, x1, y2 - 15), fill="#AEB9C2", outline=INK, width=3)
    d.rectangle((x2, y1 + 15, x2 + 45, y2 - 15), fill="#AEB9C2", outline=INK, width=3)
    d.ellipse((x2 - 40, y1 - 40, x2 + 45, y1 + 20), fill=GRAY, outline=INK, width=3)
    d.rectangle((x2 - 12, y1 + 74, x2 + 90, y1 + 120), fill=GREEN, outline=INK, width=3)
    d.rectangle((x2 - 12, y2 - 125, x2 + 90, y2 - 79), fill=GREEN, outline=INK, width=3)
    double_arrow(d, (x1, 630), (x2, 630), "530 mm preferred width")
    double_arrow(d, (170, y1), (170, y2), "435 mm", offset=(-58, 0))
    d.rounded_rectangle((1220, 180, 1695, 585), radius=22, fill=LIGHT, outline=MID, width=3)
    text_block(d, (1260, 225), "CORE DEPTH", fnt(27, True), INK, 380)
    text_block(d, (1260, 270), "64 mm preferred", fnt(38, True), GREEN, 380)
    text_block(d, (1260, 340), "HOSE NECKS", fnt(27, True), INK, 380)
    text_block(d, (1260, 385), "38 mm ID - both", fnt(35, True), BLUE, 380)
    text_block(d, (1260, 455), "PATTERN", fnt(27, True), INK, 380)
    text_block(d, (1260, 500), "16400-68030", fnt(35, True), RED, 380)
    return save(img, "02_radiator.png")


def draw_condenser() -> Path:
    img, d = new_card("2. NEW PARALLEL-FLOW CONDENSER", "R134a - manifolds stay vertical")
    x1, y1, x2, y2 = 250, 195, 1025, 565
    d.rectangle((x1, y1, x2, y2), fill="#DDEBEA", outline=INK, width=5)
    for y in range(y1 + 15, y2, 16):
        d.line((x1 + 25, y, x2 - 25, y), fill="#8BB7B2", width=3)
    d.rectangle((x1, y1, x1 + 34, y2), fill=TEAL, outline=INK, width=3)
    d.rectangle((x2 - 34, y1, x2, y2), fill=TEAL, outline=INK, width=3)
    d.rectangle((x2, y1 + 65, x2 + 105, y1 + 105), fill=TEAL, outline=INK, width=3)
    d.rectangle((x2, y2 - 105, x2 + 105, y2 - 65), fill=TEAL, outline=INK, width=3)
    d.text((x2 + 125, y1 + 85), "#8 TOP\n3/4-16", font=fnt(25, True), fill=BLUE, anchor="lm")
    d.text((x2 + 125, y2 - 85), "#6 BOTTOM\n5/8-18", font=fnt(25, True), fill=TEAL, anchor="lm")
    double_arrow(d, (x1, 625), (x2, 625), "559 mm preferred / 600 mm MAX")
    double_arrow(d, (180, y1), (180, y2), "356 mm", offset=(-55, 0))
    label(d, (1400, 200, 1715, 285), "DEPTH 21 mm", fill="#E7F4F2", color=TEAL, size=28)
    label(d, (1400, 315, 1715, 400), "25 mm MAX", fill="#FFF0E9", color=RED, size=28)
    label(d, (1400, 440, 1715, 555), "NEW + CAPPED\nNO BUILT-IN DRIER", fill=LIGHT, color=INK, size=24)
    return save(img, "03_condenser.png")


def draw_fan() -> Path:
    img, d = new_card("3. OPTIONAL FRONT CONDENSER PUSHER FAN", "buy only if hot-idle testing proves it is needed")
    cx, cy, r = 620, 390, 220
    d.ellipse((cx - r, cy - r, cx + r, cy + r), fill="#E5EEF6", outline=INK, width=8)
    for angle in range(0, 360, 60):
        a = math.radians(angle)
        tip = (cx + int(r * 0.78 * math.cos(a)), cy + int(r * 0.78 * math.sin(a)))
        left = (cx + int(55 * math.cos(a + 1.5)), cy + int(55 * math.sin(a + 1.5)))
        right = (cx + int(55 * math.cos(a - 1.5)), cy + int(55 * math.sin(a - 1.5)))
        d.polygon((left, tip, right), fill=BLUE, outline=INK)
    d.ellipse((cx - 62, cy - 62, cx + 62, cy + 62), fill=INK)
    double_arrow(d, (cx - r, 655), (cx + r, 655), "14 in / 356 mm fan")
    arrow(d, (980, 385), (1550, 385), color=RED, width=12, head=35)
    d.text((1265, 330), "AIR PUSHES\nTOWARD ENGINE", font=fnt(38, True), fill=RED, anchor="mm", align="center")
    label(d, (1120, 500, 1630, 615), "OUTER SIZE <= 370 mm\nDEPTH <= 65 mm", fill=LIGHT, color=INK, size=31)
    d.text((900, 680), "PRIMARY COOLING IS THE ORIGINAL ELECTRIC PULLER FAN BEHIND THE RADIATOR", font=fnt(27, True), fill=RED, anchor="mm")
    return save(img, "04_fan.png")


def draw_drier() -> Path:
    img, d = new_card("4. RECEIVER-DRIER + TRINARY SWITCH", "buy these as a matched pair")
    x1, y1, x2, y2 = 390, 190, 690, 600
    d.rounded_rectangle((x1, y1, x2, y2), radius=70, fill="#D9E5EA", outline=INK, width=6)
    d.rectangle((x1 + 30, y1 - 35, x1 + 80, y1 + 25), fill=TEAL, outline=INK, width=3)
    d.rectangle((x2 - 80, y1 - 35, x2 - 30, y1 + 25), fill=TEAL, outline=INK, width=3)
    d.text((x1 + 55, y1 - 55), "#6 IN", font=fnt(25, True), fill=TEAL, anchor="mb")
    d.text((x2 - 55, y1 - 55), "#6 OUT", font=fnt(25, True), fill=TEAL, anchor="mb")
    d.rectangle((x2 - 55, y1 + 150, x2 + 65, y1 + 205), fill=AMBER, outline=INK, width=3)
    d.text((x2 + 85, y1 + 177), "SWITCH PORT", font=fnt(25, True), fill=INK, anchor="lm")
    d.line((x1 - 45, y1 + 80, x1 - 45, y2 - 80), fill=RED, width=5)
    arrow(d, (x1 - 45, y2 - 80), (x1 - 45, y1 + 80), color=RED, width=5)
    d.text((x1 - 75, (y1 + y2) // 2), "VERTICAL", font=fnt(29, True), fill=RED, anchor="mm")
    d.rounded_rectangle((960, 205, 1390, 485), radius=35, fill="#FFF4DD", outline=AMBER, width=5)
    for i, color in enumerate((BLACK, BLACK, AMBER, AMBER)):
        d.line((1010 + i * 100, 485, 1010 + i * 100, 575), fill=color, width=8)
    d.text((1175, 335), "TRINARY\n4 WIRES", font=fnt(46, True), fill=AMBER, anchor="mm", align="center")
    label(d, (1450, 210, 1710, 305), "LOW 22 psi", fill=LIGHT, color=INK, size=26)
    label(d, (1450, 330, 1710, 425), "FAN 240 psi", fill=LIGHT, color=INK, size=26)
    label(d, (1450, 450, 1710, 545), "HIGH 340 psi", fill=LIGHT, color=INK, size=26)
    d.text((900, 640), "DRIER MUST BE NEW, SEALED AND OPENED ONLY AT FINAL ASSEMBLY", font=fnt(30, True), fill=RED, anchor="mm")
    return save(img, "05_drier_switch.png")


def draw_coolant_hoses() -> Path:
    img, d = new_card("5. RADIATOR HOSES + JOINERS", "route after radiator is mounted")
    d.arc((160, 190, 650, 640), start=180, end=270, fill=GREEN, width=65)
    d.line((405, 600, 850, 600), fill=GREEN, width=65)
    d.line((975, 245, 1530, 245), fill=GREEN, width=65)
    d.rounded_rectangle((1190, 210, 1340, 280), radius=18, fill="#AEB9C2", outline=INK, width=3)
    for x in (1198, 1325):
        d.line((x, 213, x, 277), fill=INK, width=5)
    label(d, (180, 145, 600, 225), "2 x 90 DEG ELBOWS", fill=LIGHT, color=INK, size=28)
    label(d, (940, 330, 1565, 415), "1 m STRAIGHT HOSE", fill=LIGHT, color=INK, size=28)
    label(d, (1040, 445, 1470, 530), "2 x BEADED JOINERS", fill=LIGHT, color=INK, size=28)
    double_arrow(d, (975, 165), (1530, 165), "38 mm ID hose")
    d.text((710, 290), "8 x CLAMPS", font=fnt(34, True), fill=BLUE, anchor="mm")
    for i in range(4):
        x = 610 + i * 70
        d.ellipse((x, 330, x + 50, 400), outline=BLUE, width=7)
        d.ellipse((x, 430, x + 50, 500), outline=BLUE, width=7)
    d.text((860, 660), "LOWER HOSE MUST BE REINFORCED / ANTI-COLLAPSE", font=fnt(30, True), fill=RED, anchor="mm")
    return save(img, "06_coolant_hoses.png")


def draw_service_parts() -> Path:
    img, d = new_card("6. CAP + THERMOSTAT + DRAIN", "old cap is a sample - service parts are new")
    d.ellipse((180, 220, 540, 520), fill="#B9C2CA", outline=INK, width=6)
    d.ellipse((245, 275, 475, 465), fill=GRAY, outline=INK, width=4)
    d.text((360, 370), "CAP", font=fnt(44, True), fill=WHITE, anchor="mm")
    label(d, (155, 545, 565, 615), "16401-41021", fill=LIGHT, color=INK, size=28)
    d.text((360, 657), "MATCH NECK + PRESSURE", font=fnt(23, True), fill=RED, anchor="mm")
    d.ellipse((710, 230, 1010, 520), fill="#D7A454", outline=INK, width=6)
    d.ellipse((785, 305, 935, 455), fill=WHITE, outline=INK, width=4)
    d.text((860, 372), "82-95 C", font=fnt(32, True), fill=RED, anchor="mm")
    label(d, (655, 555, 1065, 645), "90916-03026", fill=LIGHT, color=INK, size=30)
    d.rectangle((1260, 275, 1540, 475), fill="#B8894F", outline=INK, width=6)
    d.polygon(((1540, 315), (1650, 375), (1540, 435)), fill="#B8894F", outline=INK)
    d.text((1400, 375), "DRAIN", font=fnt(38, True), fill=WHITE, anchor="mm")
    label(d, (1220, 555, 1685, 645), "96431-03873", fill=LIGHT, color=INK, size=30)
    return save(img, "07_service_parts.png")


def draw_overflow() -> Path:
    img, d = new_card("7. COOLANT RECOVERY BOTTLE", "reuse if sound - otherwise match the original")
    body = [(340, 230), (720, 230), (790, 580), (275, 580)]
    d.polygon(body, fill="#D8EEF4", outline=INK)
    d.rectangle((440, 170, 610, 240), fill=BLUE, outline=INK, width=4)
    d.line((525, 170, 525, 125), fill=GREEN, width=25)
    d.line((525, 125, 1180, 125, 1180, 420), fill=GREEN, width=20)
    d.text((530, 405), "REUSE\nOR MATCH", font=fnt(43, True), fill=BLUE, anchor="mm", align="center")
    label(d, (980, 230, 1580, 315), "OVERFLOW HOSE: 1 m", fill=LIGHT, color=INK, size=29)
    label(d, (980, 350, 1580, 435), "ID = MATCH BOTH NIPPLES", fill=LIGHT, color=INK, size=26)
    label(d, (980, 470, 1580, 555), "2 x SMALL CLAMPS", fill=LIGHT, color=INK, size=29)
    d.text((900, 655), "REPLACE IF CRACKED, BRITTLE, BLOCKED OR UNABLE TO SEAL", font=fnt(29, True), fill=RED, anchor="mm")
    return save(img, "08_overflow.png")


def draw_shroud() -> Path:
    img, d = new_card("8. PRIMARY ELECTRIC FAN + SHROUD", "reuse the complete original puller assembly if serviceable")
    x1, y1, x2, y2 = 270, 175, 1100, 615
    d.rectangle((x1, y1, x2, y2), fill="#DDE4E9", outline=INK, width=7)
    cx, cy = (x1 + x2) // 2, (y1 + y2) // 2
    d.ellipse((cx - 210, cy - 210, cx + 210, cy + 210), fill=WHITE, outline=BLUE, width=9)
    for a in range(0, 360, 60):
        ar = math.radians(a)
        d.line((cx, cy, cx + int(160 * math.cos(ar)), cy + int(160 * math.sin(ar))), fill=GRAY, width=25)
    d.ellipse((cx - 55, cy - 55, cx + 55, cy + 55), fill=INK)
    arrow(d, (cx, cy), (1210, cy), color=RED, width=9, head=28)
    double_arrow(d, (x1, 665), (x2, 665), "adapt mounts to new radiator - keep full shroud coverage")
    label(d, (1235, 175, 1680, 285), "REAR / ENGINE SIDE\nPULLER POSITION", fill="#E7F4F2", color=TEAL, size=27)
    label(d, (1235, 315, 1680, 425), "BENCH-TEST MOTOR\n+ AIR DIRECTION", fill=LIGHT, color=INK, size=27)
    label(d, (1235, 455, 1680, 565), "CLEAR OF PULLEYS,\nBELTS + ENGINE MOVEMENT", fill="#FFF0E9", color=RED, size=25)
    d.text((900, 690), "FIT ORIGINAL FIRST; COPY ITS DEPTH, PLUG AND MOUNTS IF A DONOR REPLACEMENT IS NEEDED", font=fnt(20, True), fill=RED, anchor="mm")
    return save(img, "09_shroud.png")


def draw_ac_hoses() -> Path:
    img, d = new_card("9. R134a BARRIER HOSES", "buy loose - cut after routing")
    rows = [
        (220, "#8 DISCHARGE", "13/32 in / 10 mm ID", "ABOUT 1.5 m", RED, 36),
        (385, "#6 LIQUID", "5/16 in / 8 mm ID", "ABOUT 3.0 m", TEAL, 30),
        (550, "#10 SUCTION", "1/2 in / 13 mm ID", "ABOUT 2.5 m", BLUE, 44),
    ]
    for y, name, size, length, color, width in rows:
        d.line((430, y, 1160, y), fill=color, width=width)
        d.ellipse((405, y - width // 2, 455, y + width // 2), fill=color, outline=INK, width=3)
        d.text((100, y), name, font=fnt(31, True), fill=INK, anchor="lm")
        d.text((1210, y - 22), size, font=fnt(30, True), fill=color)
        d.text((1210, y + 22), length, font=fnt(27), fill=GRAY)
    d.text((900, 660), "ADD 10% OR 200 mm SERVICE LOOP - WHICHEVER IS MORE", font=fnt(29, True), fill=AMBER, anchor="mm")
    return save(img, "10_ac_hoses.png")


def draw_ac_fittings() -> Path:
    img, d = new_card("10. A/C FITTINGS + SERVICE PORTS", "choose angles on the vehicle")
    fittings = [
        (220, "#8 FEMALE", "3/4-16", RED),
        (640, "#6 FEMALE", "5/8-18", TEAL),
        (1060, "#6 DRIER / TXV", "4 ENDS", BLUE),
    ]
    for x, name, thread, color in fittings:
        d.rectangle((x, 265, x + 220, 365), fill=color, outline=INK, width=4)
        d.polygon(((x + 220, 250), (x + 320, 315), (x + 220, 380)), fill="#BBC4CC", outline=INK)
        d.line((x + 320, 315, x + 320, 470), fill=color, width=36)
        d.text((x + 160, 200), name, font=fnt(29, True), fill=INK, anchor="mm")
        d.text((x + 160, 520), thread, font=fnt(31, True), fill=color, anchor="mm")
    d.rectangle((1490, 230, 1580, 470), fill="#C9D4DD", outline=INK, width=4)
    d.rectangle((1460, 205, 1610, 255), fill=BLUE, outline=INK, width=3)
    d.text((1535, 520), "HIGH + LOW\nSERVICE PORTS", font=fnt(27, True), fill=BLUE, anchor="mm", align="center")
    d.text((900, 645), "BUY STRAIGHT / 45 / 90 DEG ONLY AFTER MOCK-UP", font=fnt(31, True), fill=RED, anchor="mm")
    return save(img, "11_ac_fittings.png")


def draw_crimp_workflow() -> Path:
    img, d = new_card("A/C HOSE CRIMPING - DO NOT BUY THE TOOL", "Bilal Ganj A/C shop should crimp", (1800, 1100))
    steps = [
        (75, 175, 420, 765, "1", "FIT PARTS", "Mount compressor, condenser, drier and evaporator."),
        (500, 175, 845, 765, "2", "ROUTE HOSE", "Push fittings into loose hose. Do not crimp yet."),
        (925, 175, 1270, 765, "3", "MARK IT", "Mark cut point and draw one straight clock line."),
        (1350, 175, 1695, 765, "4", "SHOP CRIMP", "Remove without twisting. Shop crimps with correct dies."),
    ]
    for x1, y1, x2, y2, number, title, body in steps:
        d.rounded_rectangle((x1, y1, x2, y2), radius=28, fill=LIGHT, outline=MID, width=4)
        d.ellipse((x1 + 25, y1 + 25, x1 + 105, y1 + 105), fill=BLUE)
        d.text((x1 + 65, y1 + 65), number, font=fnt(42, True), fill=WHITE, anchor="mm")
        d.text(((x1 + x2) // 2, y1 + 145), title, font=fnt(34, True), fill=INK, anchor="mm")
        if number == "1":
            d.rectangle((x1 + 90, y1 + 240, x2 - 90, y1 + 430), fill="#DDEBEA", outline=INK, width=5)
            d.rectangle((x2 - 115, y1 + 275, x2 - 45, y1 + 320), fill=TEAL, outline=INK)
        elif number == "2":
            d.arc((x1 + 65, y1 + 245, x2 - 65, y1 + 480), start=190, end=350, fill=RED, width=36)
            d.rectangle((x1 + 65, y1 + 380, x1 + 125, y1 + 440), fill=GRAY, outline=INK)
            d.rectangle((x2 - 125, y1 + 380, x2 - 65, y1 + 440), fill=GRAY, outline=INK)
        elif number == "3":
            d.line((x1 + 65, y1 + 365, x2 - 65, y1 + 365), fill=RED, width=44)
            d.line((x1 + 180, y1 + 320, x1 + 180, y1 + 410), fill=BLACK, width=8)
            d.line((x1 + 235, y1 + 320, x1 + 235, y1 + 410), fill=BLACK, width=8)
            d.text(((x1 + x2) // 2, y1 + 465), "CUT + CLOCK MARKS", font=fnt(23, True), fill=RED, anchor="mm")
        else:
            d.rectangle((x1 + 85, y1 + 245, x2 - 85, y1 + 405), fill="#C9D4DD", outline=INK, width=5)
            d.ellipse((x1 + 130, y1 + 280, x2 - 130, y1 + 390), fill=INK)
            d.text(((x1 + x2) // 2, y1 + 465), "CORRECT DIES", font=fnt(23, True), fill=BLUE, anchor="mm")
        text_block(d, (x1 + 35, y1 + 505), body, fnt(22, True), INK, x2 - x1 - 70, line_gap=5)
    label(d, (120, 835, 1680, 945), "BEST METHOD: MARKED HOSES GO BACK TO THE A/C SHOP FOR CRIMPING", fill="#E7F4F2", color=TEAL, size=35)
    d.text((900, 1015), "Puray system ko pehle fit karo. Hose mark karo. Crimp baad mein karwao.", font=fnt(30, True), fill=RED, anchor="mm")
    return save(img, "12_crimp_workflow.png")


def draw_mounts() -> Path:
    img, d = new_card("11. STEEL FRAME + RUBBER MOUNTS", "radiator and condenser stay separate")
    d.rectangle((300, 170, 400, 620), fill=INK)
    d.rectangle((1160, 170, 1260, 620), fill=INK)
    d.rectangle((400, 230, 1160, 570), fill="#DCE6DD", outline=GREEN, width=8)
    d.rectangle((350, 280, 1210, 520), outline=TEAL, width=10)
    d.rectangle((475, 570, 600, 615), fill=GRAY, outline=INK)
    d.rectangle((960, 570, 1085, 615), fill=GRAY, outline=INK)
    double_arrow(d, (300, 655), (1260, 655), "opening 640-670 mm - confirm on car")
    label(d, (1315, 180, 1695, 270), "UPRIGHTS\n50 x 50 x 4 mm", fill=LIGHT, color=INK, size=28)
    label(d, (1315, 300, 1695, 390), "CROSSBAR\n25 x 25 x 3 mm", fill=LIGHT, color=INK, size=28)
    label(d, (1315, 420, 1695, 510), "RADIATOR: M8\nA/C PARTS: M6", fill=LIGHT, color=INK, size=28)
    label(d, (1315, 540, 1695, 630), "RUBBER ISOLATORS", fill="#FFF4DD", color=AMBER, size=28)
    return save(img, "13_mounts.png")


def draw_electrical() -> Path:
    img, d = new_card("12. ELECTRIC FAN + A/C CIRCUITS", "reuse the existing relay/fuse panel after testing")
    boxes = [
        (85, 180, 555, 440, "PRIMARY PULLER FAN", "OWN RELAY + FUSE", BLUE),
        (665, 180, 1135, 440, "COMPRESSOR CLUTCH", "30-40 A RELAY\n10-15 A FUSE", TEAL),
        (1245, 180, 1715, 440, "OPTIONAL PUSHER", "SEPARATE RELAY + FUSE\nONLY IF FITTED", AMBER),
    ]
    for x1, y1, x2, y2, title, body, color in boxes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=25, fill=LIGHT, outline=color, width=6)
        d.text(((x1 + x2) // 2, y1 + 60), title, font=fnt(27, True), fill=color, anchor="mm")
        d.text(((x1 + x2) // 2, y1 + 160), body, font=fnt(31, True), fill=INK, anchor="mm", align="center")
    label(d, (85, 485, 615, 590), "MEASURE FAN START + RUN CURRENT", fill="#E8F1F8", color=BLUE, size=26)
    label(d, (650, 485, 1180, 590), "SIZE FUSE + CABLE FROM TEST", fill=LIGHT, color=INK, size=26)
    label(d, (1215, 485, 1715, 590), "THERMAL + A/C FAN REQUEST", fill="#E7F4F2", color=TEAL, size=25)
    d.text((900, 635), "TRINARY SWITCH REQUESTS THE PRIMARY FAN; THERMAL CONTROL MUST ALSO RUN IT WITHOUT A/C", font=fnt(25, True), fill=RED, anchor="mm")
    d.text((900, 685), "PRELIMINARY RELAY CLASS: 30-40 A - FINAL FUSE, CABLE AND CONNECTORS FOLLOW THE MOTOR RATING AND MEASURED CURRENT", font=fnt(21, True), fill=GREEN, anchor="mm")
    return save(img, "14_electrical.png")


def draw_belts() -> Path:
    img, d = new_card(
        "13. ORIGINAL BELTS + PULLEY ROUTING",
        "actual converted engine - use this photo and the removed belts as templates",
        (1800, 1250),
    )
    photo_box = (405, 145, 1395, 1160)
    paste_photo(img, ENGINE_DRIVE_PHOTO, photo_box, crop=(0, 500, 1884, 3000))
    d.rectangle(photo_box, outline=INK, width=4)

    label(d, (45, 185, 355, 300), "A/C COMPRESSOR\n+ BELT", fill="#FFF0E9", color=RED, size=27)
    arrow(d, (355, 242), (1260, 470), color=RED, width=7, head=22)
    label(d, (1445, 250, 1755, 365), "ALTERNATOR\n+ BELT", fill="#E8F1F8", color=BLUE, size=28)
    arrow(d, (1445, 308), (1220, 680), color=BLUE, width=7, head=22)
    label(d, (45, 730, 355, 845), "LOWER DRIVE +\nTENSIONER", fill="#FFF4DD", color=AMBER, size=27)
    arrow(d, (355, 785), (690, 930), color=AMBER, width=7, head=22)
    label(d, (1445, 725, 1755, 840), "MATCH EACH\nPULLEY GROOVE", fill=LIGHT, color=INK, size=27)
    label(d, (1445, 885, 1755, 1000), "MEASURE AFTER\nFINAL ADJUSTMENT", fill="#EAF4E8", color=GREEN, size=27)
    d.text((900, 1205), "BUY NEW BELTS ONLY AFTER THE ACCESSORIES ARE MOUNTED AND THE ADJUSTERS ARE MID-TRAVEL", font=fnt(25, True), fill=RED, anchor="mm")
    return save(img, "15_belts.png")


def draw_fluids() -> Path:
    img, d = new_card("14. FLUIDS + FINAL TESTS", "buy only after installation is ready")
    for x, color, text in ((180, GREEN, "16 L\n50/50\nCOOLANT"), (540, BLUE, "20 L\nDISTILLED\nWATER"), (900, TEAL, "R134a\nCHARGE BY\nWEIGHT")):
        d.rounded_rectangle((x, 220, x + 270, 545), radius=30, fill=color, outline=INK, width=5)
        d.rectangle((x + 75, 165, x + 195, 235), fill=color, outline=INK, width=4)
        d.text((x + 135, 385), text, font=fnt(34, True), fill=WHITE, anchor="mm", align="center")
    d.rounded_rectangle((1280, 180, 1700, 585), radius=30, fill=LIGHT, outline=MID, width=4)
    checks = ["RADIATOR PRESSURE", "DRY NITROGEN LEAK", "VACUUM >= 30 MIN", "CHARGE + CURRENT TEST"]
    for i, text in enumerate(checks):
        y = 245 + i * 85
        check(d, (1330, y))
        d.text((1370, y), text, font=fnt(25, True), fill=INK, anchor="lm")
    return save(img, "16_fluids_tests.png")


def draw_checklist() -> Path:
    img, d = new_card("BILAL GANJ - SIMPLE BUYING CHECKLIST", "show this page to the seller", (1800, 1080))
    left = [
        "OLD RADIATOR + CAP WITH YOU",
        "ORIGINAL ELECTRIC FAN + SHROUD",
        "CURRENT BOTTLE WITH YOU",
        "NEW HJ47 / 2H RADIATOR",
        "NEW CONDENSER",
        "NEW SEALED DRIER + SWITCH",
    ]
    right = [
        "PRIMARY FAN BENCH-TESTED",
        "OPTIONAL PUSHER ONLY IF NEEDED",
        "#6 / #8 / #10 HOSES ROUTED",
        "CUT + CLOCK MARKS DRAWN",
        "A/C SHOP CRIMPS HOSES",
        "PRESSURE / VACUUM TEST DONE",
    ]
    for col, items in enumerate((left, right)):
        x = 120 + col * 850
        for i, item in enumerate(items):
            y = 205 + i * 120
            d.rounded_rectangle((x, y, x + 735, y + 82), radius=18, fill=LIGHT, outline=MID, width=3)
            d.rectangle((x + 22, y + 21, x + 62, y + 61), fill=WHITE, outline=INK, width=3)
            d.text((x + 85, y + 41), item, font=fnt(25, True), fill=INK, anchor="lm")
    label(d, (140, 935, 1660, 1025), "NO USED DRIER - NO THROUGH-FIN TIES - NO HOSE CRIMP BEFORE MOCK-UP", fill="#FFF0E9", color=RED, size=31)
    return save(img, "17_checklist.png")


def set_cell_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_repeatable_font(run, name: str, size: float, bold: bool = False, color: str = "000000"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_picture(doc: Document, path: Path, width: float, alt_text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text)
    return p


def add_part(doc: Document, heading: str, image: Path, instruction: str, full: bool = False):
    p = doc.add_paragraph(heading, style="Heading 2")
    keep_with_next(p)
    add_picture(doc, image, 6.5, heading)


def configure_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = header.add_run("J40 VISUAL BUYING GUIDE  |  BILAL GANJ")
    set_repeatable_font(hr, "Calibri", 8.5, True, "66717D")
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("J40 cooling + A/C  |  page ")
    set_repeatable_font(fr, "Calibri", 8.5, False, "66717D")
    add_page_field(footer)
    return doc


def build_doc(images: dict[str, Path]) -> None:
    doc = configure_doc()

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    r = kicker.add_run("VISUAL MARKET BUYING PACK")
    set_repeatable_font(r, "Calibri", 11, True, "D89000")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    r = title.add_run("J40 Radiator + A/C Parts")
    set_repeatable_font(r, "Calibri", 28, True, "17324D")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("For buying new parts and reusing original templates in Bilal Ganj")
    set_repeatable_font(r, "Calibri", 12.5, False, "66717D")

    callout = doc.add_paragraph()
    callout.paragraph_format.left_indent = Inches(0.15)
    callout.paragraph_format.right_indent = Inches(0.15)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(10)
    set_cell_shading(callout, "E7F4F2")
    r = callout.add_run("CRIMPING DECISION: Do not buy the tool for this one job. Fit and mark the loose hoses on the car, then have the A/C shop crimp them with the correct dies.")
    set_repeatable_font(r, "Calibri", 11, True, "145E59")

    add_picture(doc, images["stack"], 6.5, "Front-to-rear cooling stack dimensions and buying sequence")
    roman = doc.add_paragraph()
    roman.paragraph_format.space_before = Pt(5)
    roman.paragraph_format.space_after = Pt(0)
    r = roman.add_run("Purana radiator sample saath le kar jao. Pehle parts gaari par fit karo. Hose mark karo. Crimp baad mein karwao.")
    set_repeatable_font(r, "Calibri", 10.5, True, "B33A3A")

    doc.add_page_break()
    add_part(doc, "Original cooling stack before removal", images["original_stack"], "Use the original stack to identify component order, fitting side, hose routing and mounting geometry.")
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(0)
    set_cell_shading(note, "FFF4DD")
    r = note.add_run("The photographed original is an electric puller fan with an integrated shroud. Recondition and reuse it behind the radiator as the primary fan. Add a separate front pusher only if hot-idle cooling and A/C testing proves it is necessary.")
    set_repeatable_font(r, "Calibri", 10.5, True, "8A5A00")

    doc.add_page_break()
    add_part(doc, "Original engine-side parts to retain", images["original_reuse"], "Recondition and reuse these parts only after the checks printed beside each photograph.")

    doc.add_page_break()
    add_part(doc, "Original electric puller fan and integrated shroud", images["original_fan"], "RECONDITION / REUSE: this is the primary radiator fan and the donor template for any exact replacement.")
    add_part(doc, "Primary electric fan installation", images["shroud"], "FIT BEHIND RADIATOR: preserve shroud coverage, airflow direction and clearance from the engine drive.")

    doc.add_page_break()
    add_part(doc, "Original cooling parts to carry as samples", images["original_samples"], "Carry the actual radiator, cap, coolant bottle and hose samples when matching replacement parts.")

    doc.add_page_break()
    add_part(doc, "Radiator", images["radiator"], "BUY: one new HJ47/2H copper-brass radiator. Old sample controls neck direction and mounts.")
    add_part(doc, "A/C condenser", images["condenser"], "BUY: one new parallel-flow condenser. The body and both fittings must stay inside the front opening.")

    doc.add_page_break()
    add_part(doc, "Optional front condenser pusher fan", images["fan"], "CONDITIONAL PURCHASE: add only if hot-idle tests show inadequate airflow with the primary puller fan operating correctly.")
    add_part(doc, "Receiver-drier and pressure switch", images["drier"], "BUY: one new sealed drier and one matched trinary switch. Never buy an open display drier.")

    doc.add_page_break()
    add_part(doc, "Radiator hoses and joiners", images["coolant_hoses"], "BUY AFTER DRY-FIT: hose angles and lengths must follow the installed radiator.")
    add_part(doc, "Recovery bottle", images["overflow"], "REUSE IF SOUND: clean and inspect the current bottle. Otherwise buy an exact local match using it as the sample.")

    doc.add_page_break()
    add_part(doc, "Radiator service parts", images["service_parts"], "BUY NEW: cap, thermostat/gaskets and drain. Use the old cap only to identify the neck style; confirm the new radiator's pressure requirement.")
    add_part(doc, "Original radiator cap and filler neck", images["original_cap"], "TEMPLATE ONLY: take this original with the radiator, but install a new correctly rated cap.")

    doc.add_page_break()
    add_part(doc, "A/C barrier hoses", images["ac_hoses"], "BUY LOOSE: #8, #6 and #10 R134a barrier hose. Do not cut by estimate alone.")
    add_part(doc, "A/C fittings and service ports", images["ac_fittings"], "BUY AFTER MOCK-UP: choose straight, 45-degree or 90-degree fittings on the vehicle.")

    doc.add_page_break()
    add_part(doc, "Hose marking and professional crimping", images["crimp"], "TAKE TO A/C SHOP: marked hoses, fittings and clock lines. Ask the shop to inspect every crimp.")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    set_cell_shading(p, "FFF0E9")
    r = p.add_run("Why not buy the tool? A proper hydraulic kit must match the hose construction, ferrules and dies. One wrong crimp can leak refrigerant, admit moisture or fail under pressure. Professional crimping is cheaper and easier to verify for a one-off build.")
    set_repeatable_font(r, "Calibri", 11, True, "8B2F2F")

    doc.add_page_break()
    add_part(doc, "Steel frame and rubber mounts", images["mounts"], "MAKE: radiator, condenser and fan must have separate removable mounts.")
    add_part(doc, "Relays, fuses and wiring", images["electrical"], "REUSE PANEL: protect the primary fan and clutch separately. Measure the original fan current before final fuse and cable sizing; add a third protected circuit only if a pusher is installed.")

    doc.add_page_break()
    add_part(doc, "Original belts and pulley routing", images["belts"], "BUY NEW AFTER MOCK-UP: use the actual route, old belts, pulley groove width and mid-travel adjustment length.")

    doc.add_page_break()
    add_part(doc, "Coolant, refrigerant and tests", images["fluids"], "COMMISSION: radiator pressure test, dry-nitrogen leak test, 30-minute vacuum and charge by weight.")

    doc.add_page_break()
    add_part(doc, "Final market checklist", images["checklist"], "Tick each box. Do not pay for the main parts until their measurements are photographed.")
    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(10)
    source.paragraph_format.space_after = Pt(0)
    r = source.add_run("Pattern references only: Toyota 16400-68030 radiator; 03262-VUC/CNFP1422 condenser class; Mastercool 61-25642 trinary-switch function. New purchases are local; the original electric puller fan/shroud, bottle and electrical panel are reuse candidates subject to inspection and testing.")
    set_repeatable_font(r, "Calibri", 9, False, "66717D")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main() -> None:
    images = {
        "stack": draw_stack(),
        "original_stack": draw_original_cooling_stack(),
        "original_reuse": draw_original_reuse_parts(),
        "original_samples": draw_original_service_samples(),
        "original_fan": draw_original_fan_photo(),
        "original_cap": draw_original_cap_photo(),
        "radiator": draw_radiator(),
        "condenser": draw_condenser(),
        "fan": draw_fan(),
        "drier": draw_drier(),
        "coolant_hoses": draw_coolant_hoses(),
        "service_parts": draw_service_parts(),
        "overflow": draw_overflow(),
        "shroud": draw_shroud(),
        "ac_hoses": draw_ac_hoses(),
        "ac_fittings": draw_ac_fittings(),
        "crimp": draw_crimp_workflow(),
        "mounts": draw_mounts(),
        "electrical": draw_electrical(),
        "belts": draw_belts(),
        "fluids": draw_fluids(),
        "checklist": draw_checklist(),
    }
    build_doc(images)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
