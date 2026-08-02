from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET = ROOT / "data/manual/fabrication/front_cooling_stack_rev_c/work_document_assets"
ASSET.mkdir(parents=True, exist_ok=True)
OUT = ROOT / "docs/J40-existing-radiator-refurbishment-and-turbo-fan-specification-rev-i.docx"
R0_PHOTO = ROOT / "photos/20260802_173643_gp_tJNrLg8A.jpg"
STONE_GUARD_PHOTO = ROOT / "photos/20260802_173652_gp_XFyn0ruQ.jpg"
R0_REFERENCE = ASSET / "rev_i_r01_actual_removed_radiator_reference.jpg"
STONE_GUARD_REFERENCE = ASSET / "rev_i_r02_actual_stone_guard_reference.jpg"

NAVY = "17324D"
BLUE = "2E74B5"
CYAN = "3A9FBF"
GREEN = "3E7B60"
GOLD = "C28B28"
RED = "B84842"
INK = "1F2B35"
MUTED = "64727D"
LIGHT = "E8EEF4"
PALE = "F5F7F9"
WHITE = "FFFFFF"
LINE = "B8C5CE"


def colour(value):
    return "#" + value


def font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrapped(draw, xy, text, face, fill, max_width, spacing=7):
    words = text.split()
    lines = []
    line = ""
    for word in words:
        trial = f"{line} {word}".strip()
        if draw.textbbox((0, 0), trial, font=face)[2] <= max_width or not line:
            line = trial
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    draw.multiline_text(xy, "\n".join(lines), font=face, fill=fill, spacing=spacing)


def title(draw, code, heading, subtitle):
    draw.rectangle((0, 0, 1800, 118), fill=colour(NAVY))
    draw.text((48, 25), f"{code}  {heading}", font=font(34, True), fill="white")
    draw.text((48, 78), subtitle, font=font(19), fill="#DCE8F1")
    draw.text((1510, 36), "REV I · 02 AUG 2026", font=font(19, True), fill="#DCE8F1")


def footer(draw, code):
    draw.line((45, 1142, 1755, 1142), fill=colour(LINE), width=2)
    draw.text((48, 1153), f"J40 COOLING PACK · {code} · ALL SHOP DIMENSIONS IN mm", font=font(17, True), fill=colour(MUTED))
    draw.text((1270, 1153), "FIT RELEASE = PHYSICAL R0 SAMPLE + 1:1 FIXTURE", font=font(17, True), fill=colour(RED))


def arrow(draw, start, end, fill=CYAN, width=8):
    draw.line((start, end), fill=colour(fill), width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 20
    draw.polygon(
        [
            end,
            (end[0] - size * math.cos(angle - 0.5), end[1] - size * math.sin(angle - 0.5)),
            (end[0] - size * math.cos(angle + 0.5), end[1] - size * math.sin(angle + 0.5)),
        ],
        fill=colour(fill),
    )


def fan(draw, centre, diameter, label, fill=GOLD):
    cx, cy = centre
    r = diameter / 2
    draw.ellipse((cx-r, cy-r, cx+r, cy+r), fill="#FBF4E4", outline=colour(fill), width=7)
    for deg in range(0, 360, 60):
        a = math.radians(deg)
        ex = cx + r * 0.70 * math.cos(a)
        ey = cy + r * 0.70 * math.sin(a)
        draw.line((cx, cy, ex, ey), fill=colour(fill), width=max(14, int(diameter / 15)))
    hub = diameter * 0.12
    draw.ellipse((cx-hub, cy-hub, cx+hub, cy+hub), fill=colour(fill))
    box = draw.textbbox((0, 0), label, font=font(19, True))
    draw.rectangle((cx-(box[2]-box[0])/2-9, cy+r-35, cx+(box[2]-box[0])/2+9, cy+r+2), fill="white")
    draw.text((cx-(box[2]-box[0])/2, cy+r-31), label, font=font(19, True), fill=colour(fill))


def dim_h(draw, x1, x2, y, ref_y, label):
    draw.line((x1, ref_y, x1, y), fill=colour(RED), width=3)
    draw.line((x2, ref_y, x2, y), fill=colour(RED), width=3)
    draw.line((x1, y, x2, y), fill=colour(RED), width=4)
    draw.polygon([(x1, y), (x1+17, y-9), (x1+17, y+9)], fill=colour(RED))
    draw.polygon([(x2, y), (x2-17, y-9), (x2-17, y+9)], fill=colour(RED))
    box = draw.textbbox((0, 0), label, font=font(20, True))
    cx = (x1 + x2) / 2
    draw.rectangle((cx-(box[2]-box[0])/2-8, y-18, cx+(box[2]-box[0])/2+8, y+18), fill="white")
    draw.text((cx-(box[2]-box[0])/2, y-14), label, font=font(20, True), fill=colour(RED))


def dim_v(draw, y1, y2, x, ref_x, label):
    draw.line((ref_x, y1, x, y1), fill=colour(RED), width=3)
    draw.line((ref_x, y2, x, y2), fill=colour(RED), width=3)
    draw.line((x, y1, x, y2), fill=colour(RED), width=4)
    draw.polygon([(x, y1), (x-9, y1+17), (x+9, y1+17)], fill=colour(RED))
    draw.polygon([(x, y2), (x-9, y2-17), (x+9, y2-17)], fill=colour(RED))
    draw.text((x+12, (y1+y2)/2-12), label, font=font(20, True), fill=colour(RED))


def note(draw, box, heading, body, accent=BLUE):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=15, fill=colour(PALE), outline=colour(accent), width=3)
    draw.rectangle((x1, y1, x1+12, y2), fill=colour(accent))
    draw.text((x1+28, y1+17), heading, font=font(22, True), fill=colour(accent))
    wrapped(draw, (x1+28, y1+52), body, font(18), colour(INK), x2-x1-52, 5)


def build_d11():
    im = Image.new("RGB", (1800, 1200), "white")
    d = ImageDraw.Draw(im)
    title(d, "D11", "R0 ORIGINAL RADIATOR — RECORE RELEASE", "Copy the physical sample; photographs do not release fabrication dimensions.")

    # Simplified front view of the photographed copper/brass radiator.
    x1, y1, x2, y2 = 355, 245, 1240, 915
    d.rounded_rectangle((x1, y1, x2, y2), radius=22, fill="#B98B62", outline="#513D2C", width=9)
    d.rectangle((x1+38, y1+90, x2-38, y2-77), fill="#423932", outline="#2B2622", width=5)
    for x in range(x1+48, x2-45, 16):
        d.line((x, y1+98, x, y2-84), fill="#9B7454", width=2)
    d.rounded_rectangle((x1+60, y1-24, x2-60, y1+92), radius=20, fill="#4A423B", outline="#251F1A", width=7)
    d.rounded_rectangle((x1+55, y2-77, x2-55, y2+24), radius=18, fill="#4A423B", outline="#251F1A", width=7)
    # Filler, top/bottom coolant necks and ears.
    d.rectangle((x1+312, y1-72, x1+400, y1-20), fill="#9B7752", outline="#4B3828", width=5)
    d.ellipse((x1+305, y1-91, x1+407, y1-47), fill="#C3B197", outline="#4B3828", width=5)
    d.rectangle((x2-5, y1+78, x2+128, y1+136), fill="#8C6748", outline="#4B3828", width=6)
    d.rectangle((x1-118, y2-142, x1+5, y2-82), fill="#8C6748", outline="#4B3828", width=6)
    for yy in (y1+28, y2-110):
        d.rounded_rectangle((x1-82, yy, x1-3, yy+58), radius=10, fill="#4A423B", outline="#251F1A", width=5)
        d.ellipse((x1-54, yy+18, x1-34, yy+38), fill="white", outline="#251F1A", width=3)
        d.rounded_rectangle((x2+3, yy, x2+82, yy+58), radius=10, fill="#4A423B", outline="#251F1A", width=5)
        d.ellipse((x2+34, yy+18, x2+54, yy+38), fill="white", outline="#251F1A", width=3)

    dim_h(d, x1, x2, 1010, y2, "R0-W — MEASURE COMPLETE SAMPLE")
    dim_v(d, y1, y2, 215, x1, "R0-H — MEASURE")
    d.line((798, 210, 798, 960), fill=colour(CYAN), width=3)
    d.text((812, 192), "vehicle C0", font=font(18, True), fill=colour(CYAN))

    note(d, (1310, 186, 1740, 378), "KEEP / COPY EXACTLY", "Sound top and bottom tanks, both large coolant neck positions and angles, filler/cap seat, overflow, drain, side rails, upper ears and lower locators.", GREEN)
    note(d, (1310, 407, 1740, 600), "REPLACE", "Old core. Fit a new heavy-duty high-efficiency copper/brass core on the original header/tank footprint. Four-row is preferred only if depth and tests pass.", RED)
    note(d, (1310, 629, 1740, 822), "FIXTURE CONTROL", "Record T0-L/T0-R upper-hole coordinates and L0-L/L0-R lower saddles in a rigid 1:1 jig. Upper ears bolt to the existing top holes; lower EPDM saddles carry the weight.", BLUE)
    note(d, (1310, 851, 1740, 1055), "DO NOT GUESS", "Measure R0-W/H/D, active core RC-W/H/D, every neck OD/centre/angle, cap type, drain and mounting centres before ordering the core.", GOLD)
    footer(d, "R0 SAMPLE CONTROL")
    path = ASSET / "rev_i_d11_refurbished_radiator_release.png"
    im.save(path, optimize=True)
    return path


def build_d12():
    im = Image.new("RGB", (1800, 1200), "white")
    d = ImageDraw.Draw(im)
    title(d, "D12", "COMPLETE COOLING PACK — SIDE DEPTH & AIRFLOW", "All three heat exchangers and all three fans; front is left, engine is right.")

    top, bottom = 278, 845
    layers = [
        (135, 175, "STONE\nGUARD", "10–15 gap", MUTED),
        (208, 330, "2 × 248\nPUSHERS", "55 max", GOLD),
        (350, 475, "CHARGE-AIR\nCORE", "50 target", BLUE),
        (500, 565, "A/C\nCONDENSER", "sample depth", GREEN),
        (592, 752, "R0 RECORED\nRADIATOR", "R0-D measured", RED),
        (775, 930, "SEALED\nSHROUD", "full face", NAVY),
    ]
    for x1, x2, label, depth, accent in layers:
        d.rounded_rectangle((x1, top, x2, bottom), radius=12, fill=colour(PALE), outline=colour(accent), width=5)
        bbox = d.multiline_textbbox((0, 0), label, font=font(22, True), align="center")
        d.multiline_text(((x1+x2-(bbox[2]-bbox[0]))/2, top+200), label, font=font(22, True), fill=colour(accent), align="center", spacing=5)
        bbox2 = d.textbbox((0, 0), depth, font=font(17, True))
        d.text(((x1+x2-(bbox2[2]-bbox2[0]))/2, bottom+13), depth, font=font(17, True), fill=colour(accent))

    fan(d, (1070, 562), 380, "FAN 3 · 2H MECHANICAL PULLER", NAVY)
    d.rectangle((1260, 430, 1710, 700), fill="#E9EDF0", outline=colour(NAVY), width=5)
    d.text((1390, 535), "ENGINE", font=font(38, True), fill=colour(NAVY))
    d.text((1320, 590), "fan belt / pulley plane", font=font(20), fill=colour(MUTED))

    for y in (200, 235, 900, 935):
        arrow(d, (80, y), (1695, y), CYAN, 7)
    d.text((530, 148), "AIRFLOW: GRILLE → ENGINE", font=font(28, True), fill=colour(CYAN))

    dim_h(d, 135, 930, 1020, bottom, "FRONT PACK: TARGET ≤225 INSTALLED — VERIFY ACTUAL")
    dim_h(d, 930, 1260, 1090, 700, "RADIATOR-TO-BLADE ≥20; 25–30 PREFERRED")
    note(d, (1250, 752, 1725, 915), "CLEARANCE GATES", "Rear blade radial clearance to shroud ≥15. Nothing may touch at engine-rock, belt change, hose movement or full fan sweep.", RED)
    footer(d, "STACK DEPTH")
    path = ASSET / "rev_i_d12_refurbished_pack_dimensions.png"
    im.save(path, optimize=True)
    return path


def build_d13():
    im = Image.new("RGB", (1800, 1200), "white")
    d = ImageDraw.Draw(im)
    title(d, "D13", "CENTRED FRONT FANS & CONTROL LOGIC", "Two matched Toyota/Denso pushers share the stack; Fan 2 is the turbo/charge-air assist.")

    cx, cy = 900, 495
    d.line((cx, 150, cx, 920), fill=colour(CYAN), width=5)
    d.text((cx+14, 150), "VEHICLE C0", font=font(20, True), fill=colour(CYAN))
    d.rounded_rectangle((310, 190, 1490, 825), radius=24, fill="#F8FAFB", outline=colour(NAVY), width=6)
    fan(d, (634, cy), 440, "FAN 1 · RADIATOR / A-C", GOLD)
    fan(d, (1166, cy), 440, "FAN 2 · TURBO / CHARGE-AIR", GOLD)
    dim_h(d, 634, 900, 850, cy, "C0 −133")
    dim_h(d, 900, 1166, 890, cy, "C0 +133")
    dim_h(d, 634, 1166, 930, cy, "266 CENTRE-TO-CENTRE")

    note(d, (55, 968, 505, 1115), "MECHANICAL FIT", "Nominal blade diameter 248. Clear ring ≥258. Fan-group midpoint to C0 ±2 and both hubs at equal height ±2.", NAVY)
    note(d, (675, 968, 1125, 1115), "NORMAL CONTROL", "Fan 1: A/C demand or coolant lead. Fan 2: high IAT / boost lead. On high coolant, high IAT, high A/C pressure or high load: BOTH ON.", BLUE)
    note(d, (1295, 968, 1745, 1115), "FAILSAFE", "One relay and fuse branch per motor. Manual override and sensor fault command BOTH electric fans ON. Fan 3 remains engine driven.", RED)
    footer(d, "FAN CENTRELINE")
    path = ASSET / "rev_i_d13_fan_control_and_centreline.png"
    im.save(path, optimize=True)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def add_table(doc, headers, rows, widths=None, size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for idx, heading in enumerate(headers):
        set_cell_shading(hdr.cells[idx], NAVY)
        p = hdr.cells[idx].paragraphs[0]
        run = p.add_run(heading)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(size)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(value))
            run.font.size = Pt(size)
            run.font.color.rgb = RGBColor.from_string(INK)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_lead and text.startswith(bold_lead):
        run = p.add_run(bold_lead)
        run.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2.5)
        p.add_run(item)


def add_callout(doc, heading, body, fill="FFF4DD", border=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "10")
        tag.set(qn("w:color"), border)
        borders.append(tag)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(heading + "\n")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(border)
    p.add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(doc, name, width=6.7, caption=None):
    path = ASSET / name
    if not path.exists():
        return
    add_image_path(doc, path, width, caption)


def add_image_path(doc, path, width=6.7, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].italic = True
        c.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
        c.runs[0].font.size = Pt(8.5)


def prepare_reference_photo(source, output):
    """Create a document-sized copy while retaining the untouched picker import."""
    with Image.open(source) as original:
        image = ImageOps.exif_transpose(original).convert("RGB")
        image.thumbnail((1600, 1600), Image.Resampling.LANCZOS)
        image.save(output, "JPEG", quality=84, optimize=True, progressive=True)
    return output


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.addnext(field)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    for level, size, col in ((1, 17, NAVY), (2, 12.5, BLUE), (3, 10.5, GREEN)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(col)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "J40 · ORIGINAL-RADIATOR REFURBISHMENT + TURBO-FAN COOLING PACK · REV I"
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    page_number(section.footer.paragraphs[0])


def build_docx():
    prepare_reference_photo(R0_PHOTO, R0_REFERENCE)
    prepare_reference_photo(STONE_GUARD_PHOTO, STONE_GUARD_REFERENCE)

    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("J40 ORIGINAL-RADIATOR\nREFURBISHMENT + TURBO-FAN PACK")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = doc.add_paragraph("FABRICATOR SPECIFICATION · REV I · 02 AUGUST 2026")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True
    p2.runs[0].font.size = Pt(11)
    p2.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

    add_callout(
        doc,
        "START HERE — TAKE THE ACTUAL RADIATOR TO THE SHOP",
        "The photographed copper/brass engine radiator is R0, the master fit sample. Do not build from photo dimensions. Recore R0 and preserve—or exactly copy—its tanks, two coolant necks, filler/cap seat, overflow, drain, side rails, upper mounting ears and lower locators. The old core is not to be reused.",
        "FFF0EE",
        RED,
    )
    add_image(doc, "rev_i_ph01_refurbished_cooling_pack_assembled.png", 6.7, "Opaque assembled concept. Final metal dimensions come from R0 and the chassis fixture.")
    add_para(doc, "Purpose: a simple, locally repairable cooling pack for Pakistan use, including 50 °C ambient operation and the planned 150 bhp / 112 kW turbo ceiling.")
    add_para(doc, "Air order: grille → removable stone guard → two centred electric pushers → charge-air cooler → A/C condenser → recored R0 radiator → sealed shroud → 2H mechanical puller → engine.")

    doc.add_page_break()
    add_heading(doc, "Physical donor samples — identify these parts first")
    add_image_path(
        doc,
        R0_REFERENCE,
        6.75,
        "R01 — Actual removed copper/brass engine radiator R0. This physical part is the fit and interface master; measure it before dismantling.",
    )
    add_image_path(
        doc,
        STONE_GUARD_REFERENCE,
        6.75,
        "R02 — Actual removable expanded-mesh stone guard. Refurbish or reproduce it inside the measured opening with at least 70% open area.",
    )

    doc.add_page_break()
    add_heading(doc, "1. What is being built")
    add_table(doc, ["Item", "Count", "Job"], [
        ("R0 original engine radiator", "1", "Dismantle, inspect and fit a new heavy-duty copper/brass core; retain exact vehicle interfaces."),
        ("Charge-air cooler", "1", "Slim front core for the turbo; two 57 mm beaded air ports only."),
        ("A/C condenser", "1", "Inspect/identify existing thin aluminium sample; reuse only after pressure/flush/fit pass, otherwise fit checked replacement."),
        ("Front electric pushers", "2", "Matched Toyota/Denso 248 mm units, horizontally centred. Fan 2 is the turbo/charge-air assist."),
        ("Rear mechanical puller", "1", "Retained Toyota 2H engine-driven fan behind a full-face shroud; pulls air toward engine."),
        ("Stone guard", "1", "Refurbish/reproduce the photographed removable expanded-mesh guard, ≥70% open area."),
    ], [1.5, 0.55, 4.65])
    add_callout(doc, "FAN COUNT IS THREE", "Two front electric fans plus one rear engine-driven mechanical fan. Fan 2 is the turbo/charge-air-assist role; it is not a fourth fan. Both electric fans push through the shared stack.", "EAF4FA", BLUE)

    add_heading(doc, "2. Dimension and fit-release sheet")
    add_table(doc, ["Code", "Measure / requirement", "Release value"], [
        ("R0-W/H/D", "Complete radiator envelope including tanks, seams, ears, filler, necks and drain", "MEASURE PHYSICAL R0"),
        ("RC-W/H/D", "Active core face and depth", "MEASURE R0; new core matches header/tank footprint"),
        ("T0-L / T0-R", "Upper chassis-hole 3D coordinates and centre-to-centre", "COPY IN RIGID 1:1 FIXTURE"),
        ("L0-L / L0-R", "Lower saddle/tab coordinates and span", "COPY IN SAME FIXTURE"),
        ("N1 / N2", "Coolant neck OD, centre and angle; also filler, overflow and drain", "COPY R0 EXACTLY"),
        ("W0", "Clear opening at top, middle and bottom", "MEASURE CHASSIS; every component stays inside"),
        ("F1", "Front fan hubs", "C0−133 and C0+133; 266 C-C; midpoint and height ±2"),
        ("F2", "Rear radiator-to-blade static clearance", "≥20; 25–30 preferred"),
        ("F3", "Rear blade radial clearance inside shroud", "≥15 throughout sweep"),
        ("P0", "Front stack installed depth", "Target ≤225; prove against actual fore/aft space"),
    ], [0.75, 3.4, 2.55])
    add_para(doc, "The previous 530 × 435 × 64 mm radiator concept is not a fit-release dimension. R0 and the 1:1 chassis fixture now control fit.", "The previous")
    add_image(doc, "rev_i_d11_refurbished_radiator_release.png", 6.75)

    doc.add_page_break()
    add_heading(doc, "3. Radiator-shop instructions")
    add_bullets(doc, [
        "Tag R0 before dismantling. Photograph and measure every neck, cap seat, drain, bracket and locating feature.",
        "Strip paint and scale; dismantle and chemically clean. Pressure- and flow-check tanks, seams, necks, side rails and ears.",
        "Reject any pitted, thinned or cracked tank/neck/rail. If an original piece is unsafe, reproduce its external geometry from the physical sample before discarding it.",
        "Install a new heavy-duty high-efficiency copper/brass core on the original header/tank footprint. Four-row construction is preferred only if R0-D, airflow and clearance gates pass; thermal test performance controls acceptance.",
        "Use thin fin-safe black coating only. Do not bury fins or solder joints under thick paint.",
        "Pressure-test at the verified system/shop test pressure. The cap rating must match the actual cap neck and engine system; do not guess.",
        "Flow-test, hot-recirculation test and leak-test after final soldering. Record the results on the release sheet.",
    ])
    add_heading(doc, "4. Mounting")
    add_bullets(doc, [
        "Use the existing two upright side rails. Upper radiator ears must visibly bolt through the existing holes at the top returns using M8 hardware, sleeves and rubber isolation.",
        "Lower EPDM saddles carry radiator weight. Upper bolts locate and retain; they must not crush or twist the tanks.",
        "Keep at least 5 mm hard clearance to metal throughout the radiator, guard, cores and fan frames.",
        "No crossbar, fan frame, fuse box or relay box may widen the pack beyond W0. Electrical boxes use independent high/rear brackets and rain hoods.",
        "Stone guard stays independently removable 10–15 mm ahead of the front fan/shroud face and must not rub the fans.",
    ])
    add_image(doc, "rev_i_d12_refurbished_pack_dimensions.png", 6.75)

    doc.add_page_break()
    add_heading(doc, "5. Fans, turbo function and wiring")
    add_table(doc, ["Fan", "Physical part", "Lead trigger", "Direction"], [
        ("Fan 1", "Toyota/Denso 248 mm front pusher", "A/C demand or coolant-temperature lead", "Grille → engine"),
        ("Fan 2", "Matching Toyota/Denso 248 mm front pusher", "High intake-air temperature / boost lead", "Grille → engine"),
        ("Fan 3", "Toyota 2H engine-driven mechanical puller", "Engine speed / fan coupling as fitted", "Pulls grille → engine"),
    ], [0.75, 2.1, 2.55, 1.15])
    add_bullets(doc, [
        "Use two matched Prado 120 / GX470 Toyota/Denso 248 mm fan assemblies. Nominal hub centres are C0−133 mm and C0+133 mm, exactly 266 mm apart; group midpoint and hub height tolerance ±2 mm.",
        "Provide at least a 258 mm clear ring for each nominal 248 mm blade. Confirm the actual donor frame, plug, tab and wire-bend envelope before steel cutting.",
        "Fan 2 is called the turbo fan because its electrical lead trigger is charge-air temperature/boost. It cools the charge-air core by pushing ambient air through the shared stack; it does not blow directly onto the turbocharger housing.",
        "On high coolant temperature, high IAT, high A/C pressure, high load, sensor fault or manual override, both electric fans run together.",
        "Use one sealed relay and one closed MIDI-fuse branch per electric motor. Earth each motor with equal-size cable to a clean chassis/negative point. Put no exposed fuses in the airflow or rain path.",
    ])
    add_image(doc, "rev_i_d13_fan_control_and_centreline.png", 6.75)

    doc.add_page_break()
    add_heading(doc, "6. Standard donor parts — buy only after sample check")
    add_table(doc, ["Part", "Preferred source / reference", "Rule"], [
        ("Front fan assemblies ×2", "Prado 120 / GX470 Toyota/Denso 248 mm; 88590-60040 / -60050 / -60051 / -60060", "Buy a matching pair with frames, plugs and pigtails; verify actual envelope."),
        ("Fan motor", "Toyota/Denso 88550-12160 reference", "Use matched working motors; confirm connector and rotation."),
        ("Fan blade", "Toyota 88453-60010 reference", "Confirm 248 mm blade and correct pusher rotation."),
        ("Rear mechanical blade", "Toyota 16361-68030 / 16361-68031 candidate", "Use only after hub, sweep and blade-count check against engine."),
        ("Rear shroud", "Toyota 16711-47040 candidate", "Use only if it seals R0 face and passes clearances; otherwise fabricate full-face shroud."),
        ("A/C condenser", "Existing sample after identification/test, or common 14 × 22 in R134a parallel-flow", "Bench-measure fittings, drier, frame and W0 before purchase."),
        ("Receiver/drier", "Toyota 88471-34010 candidate or common #6", "Port and refrigerant compatibility must be checked."),
        ("Relays / fuses", "Toyota 90987-02027 or sealed ISO 40 A; closed MIDI fuse box", "One relay and fuse branch per motor; rate from measured current."),
    ], [1.35, 2.7, 2.85], 8.2)

    add_heading(doc, "7. 50 °C thermal release")
    add_bullets(doc, [
        "Radiator duty: ≥115 kW continuous heat rejection and ≥130 kW for 10 minutes at 50 °C ambient, or equivalent supplier-certified/tested performance on the released core.",
        "Charge-air core: target 500 × 180 × 50 mm only after W0/P0 check; ≥15 kW duty, two 57 mm beaded ports, route pressure drop ≤10 kPa, and post-cooler IAT ≤80 °C at the agreed full-load test.",
        "Two-front-fan installed airflow target ≥3,000 m³/h aggregate; 3,300 m³/h preferred. Accept installed measurements, not free-air catalogue numbers alone.",
        "Run hot-idle soak with A/C on, repeated low-speed pulls, and sustained loaded road testing. No boiling, coolant loss, hose collapse, fan/shroud contact, electrical overheating or progressive temperature rise.",
        "Record coolant in/out temperatures, IAT before/after cooler, ambient, fan states, voltage and current. A 50 °C claim is not released until the completed vehicle passes the agreed test or equivalent instrumented simulation.",
    ])

    add_heading(doc, "8. Fabricator acceptance checklist")
    add_table(doc, ["Check", "Pass requirement", "Initial / date"], [
        ("R0 measured before dismantling", "All R0/RC/T0/L0/N1/N2 fields recorded", "__________"),
        ("Core and tanks", "New core; tanks/rails/necks sound or exact-copy replacements", "__________"),
        ("Top-hole attachment", "Both upper ears bolt to existing top holes without strain", "__________"),
        ("Lower support", "Both lower EPDM saddles carry weight and cannot walk out", "__________"),
        ("Fan centring", "C0±133; 266 C-C; midpoint/height ±2", "__________"),
        ("Rear fan clearance", "≥20 axial (25–30 preferred), ≥15 radial, all conditions", "__________"),
        ("Stack width/depth", "Inside W0; target front pack ≤225; no electrical side growth", "__________"),
        ("Pressure/flow/electrical", "No leaks; flow pass; current/voltage/fuse/relay records complete", "__________"),
        ("50 °C thermal validation", "Recorded full-system pass at agreed test condition", "__________"),
    ], [1.75, 3.75, 1.35], 8.4)

    # Visual appendix. Images are included when present so the builder can be rerun
    # after the photoreal design set has been generated.
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Appendix A — Visual assembly reference")
    add_image(doc, "rev_i_ph02_refurbished_cooling_pack_exploded.png", 6.75, "Exploded component relationship. Do not dimension from the rendering.")
    add_image(doc, "rev_i_ph03_refurbished_pack_installed.png", 6.75, "Opaque installed concept: upper ears attach to the existing top holes; rear fan is behind the radiator.")
    add_image(doc, "rev_i_ph04_refurbished_pack_engine_side.png", 6.75, "Engine-side reference: full-face shroud and rear mechanical puller remain visible and serviceable.")

    props = doc.core_properties
    props.title = "J40 Original-Radiator Refurbishment and Turbo-Fan Cooling Pack — Rev I"
    props.subject = "Pakistan radiator-fabricator specification"
    props.author = "J40 Rebuild Project"
    props.keywords = "J40, radiator, recore, turbo fan, Prado, Denso, Pakistan, 50 C"
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    for builder in (build_d11, build_d12, build_d13):
        print(builder())
    print(build_docx())
