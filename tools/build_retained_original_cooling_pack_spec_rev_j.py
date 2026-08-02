from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont, ImageOps
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET = ROOT / "data/manual/fabrication/front_cooling_stack_rev_c/work_document_assets"
ASSET.mkdir(parents=True, exist_ok=True)
OUT = ROOT / "docs/J40-retained-original-radiator-and-fans-cooling-pack-specification-rev-j.docx"

R0_PHOTO = ROOT / "photos/20260802_173643_gp_tJNrLg8A.jpg"
GUARD_PHOTO = ROOT / "photos/20260802_173652_gp_XFyn0ruQ.jpg"
FAN_PHOTO = ROOT / "photos/20260317_235229.jpg"
R0_REFERENCE = ASSET / "rev_j_r01_actual_removed_radiator_reference.jpg"
GUARD_REFERENCE = ASSET / "rev_j_r02_actual_stone_guard_reference.jpg"
FAN_REFERENCE = ASSET / "rev_j_r03_actual_large_fan_context.jpg"

PH01 = ASSET / "rev_j_ph01_retained_cooling_pack_assembled.jpg"
PH02 = ASSET / "rev_j_ph02_retained_cooling_pack_exploded.jpg"
PH03 = ASSET / "rev_j_ph03_retained_pack_installed.jpg"
PH04 = ASSET / "rev_j_ph04_retained_pack_engine_side.jpg"

D14 = ASSET / "rev_j_d14_reuse_release.png"
D15 = ASSET / "rev_j_d15_matched_fans_split_auxiliary_plane.png"
D16 = ASSET / "rev_j_d16_stack_depth_and_airflow.png"

NAVY = "17324D"
BLUE = "2E74B5"
CYAN = "3A9FBF"
GREEN = "3E7B60"
GOLD = "C28B28"
RED = "B84842"
INK = "1F2B35"
MUTED = "64727D"
LIGHT = "E8EEF5"
PALE = "F5F7F9"
WHITE = "FFFFFF"
LINE = "B8C5CE"
BLACK = "111820"


def colour(value):
    return "#" + value


def pil_font(size, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def wrap_lines(draw, text, face, max_width):
    lines = []
    for paragraph in text.split("\n"):
        words = paragraph.split()
        line = ""
        for word in words:
            trial = f"{line} {word}".strip()
            if not line or draw.textbbox((0, 0), trial, font=face)[2] <= max_width:
                line = trial
            else:
                lines.append(line)
                line = word
        lines.append(line)
    return lines


def draw_wrapped(draw, box, text, face, fill=colour(INK), spacing=6, align="left"):
    x0, y0, x1, _ = box
    lines = wrap_lines(draw, text, face, x1 - x0)
    draw.multiline_text((x0, y0), "\n".join(lines), font=face, fill=fill, spacing=spacing, align=align)


def diagram_header(draw, code, title, subtitle):
    draw.rectangle((0, 0, 1800, 125), fill=colour(NAVY))
    draw.text((48, 23), f"{code}  {title}", font=pil_font(35, True), fill="white")
    draw.text((48, 78), subtitle, font=pil_font(20), fill=colour(LIGHT))


def diagram_footer(draw, text):
    draw.rectangle((0, 1138, 1800, 1200), fill=colour(NAVY))
    draw.text((48, 1155), text, font=pil_font(18, True), fill="white")


def rounded_box(draw, box, heading, body, fill=PALE, border=LINE, heading_colour=NAVY):
    draw.rounded_rectangle(box, radius=18, fill=colour(fill), outline=colour(border), width=4)
    x0, y0, x1, _ = box
    draw.text((x0 + 22, y0 + 18), heading, font=pil_font(24, True), fill=colour(heading_colour))
    draw_wrapped(draw, (x0 + 22, y0 + 58, x1 - 22, box[3] - 16), body, pil_font(19), colour(INK), 5)


def arrow(draw, start, end, fill=BLUE, width=8):
    draw.line((start, end), fill=colour(fill), width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 22
    for delta in (2.55, -2.55):
        point = (
            end[0] + length * math.cos(angle + delta),
            end[1] + length * math.sin(angle + delta),
        )
        draw.line((end, point), fill=colour(fill), width=width)


def dimension(draw, start, end, label, offset=0, fill=NAVY):
    sx, sy = start
    ex, ey = end
    if sy == ey:
        sy += offset
        ey += offset
        draw.line((sx, sy - 12, sx, sy + 12), fill=colour(fill), width=3)
        draw.line((ex, ey - 12, ex, ey + 12), fill=colour(fill), width=3)
        draw.line((sx, sy, ex, ey), fill=colour(fill), width=3)
        draw.polygon([(sx, sy), (sx + 16, sy - 7), (sx + 16, sy + 7)], fill=colour(fill))
        draw.polygon([(ex, ey), (ex - 16, ey - 7), (ex - 16, ey + 7)], fill=colour(fill))
        tw = draw.textbbox((0, 0), label, font=pil_font(19, True))[2]
        draw.rectangle(((sx + ex - tw) / 2 - 8, sy - 16, (sx + ex + tw) / 2 + 8, sy + 17), fill="white")
        draw.text(((sx + ex - tw) / 2, sy - 13), label, font=pil_font(19, True), fill=colour(fill))


def save_diagram(image, path):
    image.save(path, "PNG", optimize=True)
    return path


def build_d14():
    image = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(image)
    diagram_header(draw, "D14", "RETAIN / TEST / REPLACE RELEASE", "Use the actual vehicle parts as masters; buy only after a measured failure or match decision.")

    columns = [
        (70, "R0 ENGINE RADIATOR", "Clean, inspect, pressure-test and flow-test. Record full envelope, core, tanks, necks, ears and lower locators.", "PASS: retain/refurbish", "FAIL: repair or recore while preserving R0 external geometry"),
        (610, "FRONT ELECTRIC FAN", "Photograph, measure complete module, verify pusher rotation, start/run current and bench airflow. Find a true Toyota/Denso twin.", "MATCH + PASS: install existing fan plus twin", "NO MATCH / FAIL: install matched pair; keep old fan as spare"),
        (1150, "A/C CONDENSER", "Identify, pressure-test, flush, inspect fins and prove ports/drier fit in one half of the shallow auxiliary plane.", "PASS: retain existing condenser", "FAIL: measured replacement in the same released zone"),
    ]
    for x, heading, body, good, bad in columns:
        rounded_box(draw, (x, 190, x + 480, 430), heading, body, "F5F7F9", LINE)
        arrow(draw, (x + 240, 430), (x + 240, 500), BLUE)
        draw.rounded_rectangle((x + 70, 500, x + 410, 565), radius=22, fill=colour(GOLD))
        draw.text((x + 188, 516), "TEST", font=pil_font(25, True), fill="white")
        arrow(draw, (x + 240, 565), (x + 130, 660), GREEN)
        arrow(draw, (x + 240, 565), (x + 350, 660), RED)
        rounded_box(draw, (x, 660, x + 225, 880), "PASS", good, "EAF4EF", GREEN, GREEN)
        rounded_box(draw, (x + 255, 660, x + 480, 880), "FAIL", bad, "FFF0EE", RED, RED)

    draw.rounded_rectangle((70, 930, 1690, 1085), radius=20, fill=colour("FFF4DD"), outline=colour(GOLD), width=4)
    draw.text((95, 955), "NON-NEGOTIABLE RELEASE", font=pil_font(25, True), fill=colour(GOLD))
    draw_wrapped(
        draw,
        (95, 995, 1665, 1065),
        "Two matching equal-envelope front pushers only. One CAC only. R0 remains the sole engine-coolant radiator. The CAC and condenser share one side-by-side plane. Both top M8 bolts use the existing horizontal-return holes.",
        pil_font(22, True),
    )
    diagram_footer(draw, "REV J · NO CLEAN-SHEET RADIATOR · NO UNEQUAL FRONT FAN PAIR · NO SERIAL CAC/CONDENSER STACK")
    return save_diagram(image, D14)


def build_d15():
    image = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(image)
    diagram_header(draw, "D15", "FRONT VIEW — MATCHED FANS + SPLIT AUXILIARY PLANE", "All release dimensions come from the accepted complete modules and the 1:1 vehicle fixture.")

    # Usable carrier and core zones
    left, top, right, bottom = 150, 215, 1285, 925
    draw.rounded_rectangle((left, top, right, bottom), radius=24, fill=colour("EEF2F5"), outline=colour(NAVY), width=7)
    draw.rectangle((left + 55, top + 70, 715, bottom - 65), fill=colour("D8EEF4"), outline=colour(CYAN), width=5)
    draw.rectangle((715, top + 70, right - 55, bottom - 65), fill=colour("DDE2E5"), outline=colour(MUTED), width=5)
    draw.text((285, top + 92), "CAC ZONE", font=pil_font(27, True), fill=colour(CYAN))
    draw.text((895, top + 92), "A/C CONDENSER ZONE", font=pil_font(24, True), fill=colour(MUTED))
    draw.line((715, top + 70, 715, bottom - 65), fill=colour(GOLD), width=12)

    c1, c2, radius = (445, 585), (990, 585), 205
    for cx, cy in (c1, c2):
        draw.ellipse((cx - radius, cy - radius, cx + radius, cy + radius), fill=colour("1A2229"), outline=colour(BLACK), width=8)
        for angle in range(0, 360, 45):
            a = math.radians(angle)
            p1 = (cx + 45 * math.cos(a), cy + 45 * math.sin(a))
            p2 = (cx + 175 * math.cos(a + 0.22), cy + 175 * math.sin(a + 0.22))
            draw.line((p1, p2), fill=colour("77828B"), width=22)
        draw.ellipse((cx - 43, cy - 43, cx + 43, cy + 43), fill=colour("0F151A"), outline=colour(LINE), width=4)
    draw.line((717, top + 5, 717, bottom + 15), fill=colour(RED), width=4)
    draw.text((680, top - 42), "C0", font=pil_font(25, True), fill=colour(RED))

    dimension(draw, (left, bottom + 65), (right, bottom + 65), "Wc — measured usable width")
    dimension(draw, (c1[0] - radius, top + 15), (c1[0] + radius, top + 15), "F")
    dimension(draw, (c1[0] + radius, bottom + 20), (c2[0] - radius, bottom + 20), "g")
    dimension(draw, (left, bottom - 10), (c1[0] - radius, bottom - 10), "M")
    dimension(draw, (c2[0] + radius, bottom - 10), (right, bottom - 10), "M")

    # Top mount inset
    draw.rounded_rectangle((1340, 210, 1730, 785), radius=20, fill=colour(PALE), outline=colour(LINE), width=4)
    draw.text((1363, 235), "TOP MOUNT — BOTH SIDES", font=pil_font(22, True), fill=colour(NAVY))
    draw.rectangle((1395, 375, 1665, 420), fill=colour("9BA5AD"), outline=colour(MUTED), width=4)
    draw.rectangle((1395, 420, 1455, 720), fill=colour("9BA5AD"), outline=colour(MUTED), width=4)
    draw.ellipse((1570, 386, 1604, 420), fill="white", outline=colour(MUTED), width=4)
    draw.rectangle((1578, 335, 1596, 555), fill=colour("B8BEC4"), outline=colour(BLACK), width=3)
    draw.ellipse((1558, 320, 1616, 348), fill=colour("A6ADB3"), outline=colour(BLACK), width=3)
    draw.ellipse((1553, 421, 1621, 474), fill=colour(BLACK), outline=colour(BLACK), width=3)
    draw.rectangle((1550, 474, 1624, 650), fill=colour(BLACK), outline=colour(BLACK), width=4)
    draw_wrapped(draw, (1365, 735, 1705, 775), "Vertical M8 through exact existing horizontal-plate hole", pil_font(18, True), colour(RED), 4)
    arrow(draw, (1675, 360), (1603, 360), RED, 5)
    draw.text((1365, 805), "Lower saddles carry weight.\nTop bolts locate and restrain.", font=pil_font(20, True), fill=colour(GREEN))

    draw.rounded_rectangle((1340, 900, 1730, 1085), radius=18, fill=colour("FFF4DD"), outline=colour(GOLD), width=4)
    draw.text((1365, 925), "CENTRING RULE", font=pil_font(23, True), fill=colour(GOLD))
    draw.text((1365, 972), "M = (Wc − 2F − g) / 2", font=pil_font(23, True), fill=colour(INK))
    draw_wrapped(draw, (1365, 1012, 1705, 1070), "M > 0; equal margins ±2 mm; group midpoint on C0 ±2 mm; hubs level ±2 mm.", pil_font(17))
    diagram_footer(draw, "REV J · LEFT/RIGHT CORE ORIENTATION IS RELEASED ONLY AFTER PIPE, DRIER AND SERVICE-ROUTE MOCK-UP")
    return save_diagram(image, D15)


def build_d16():
    image = Image.new("RGB", (1800, 1200), "white")
    draw = ImageDraw.Draw(image)
    diagram_header(draw, "D16", "AIRFLOW + STACK DEPTH", "The two auxiliary cores share one depth plane; R0 is the only full-width engine-coolant radiator.")

    y0, y1 = 340, 820
    layers = [
        (105, 150, "STONE\nGUARD", MUTED),
        (235, 335, "MATCHED\nPUSHERS", BLACK),
        (425, 555, "ONE SPLIT\nAUX PLANE", CYAN),
        (675, 820, "R0 ENGINE\nRADIATOR", GOLD),
        (920, 1020, "SEALED\nSHROUD", BLACK),
        (1105, 1300, "LARGE 2H\nPULLER", GREEN),
        (1460, 1680, "ENGINE", NAVY),
    ]
    for x0, x1, label, fill in layers:
        draw.rounded_rectangle((x0, y0, x1, y1), radius=15, fill=colour(fill), outline=colour(NAVY), width=4)
        lines = label.split("\n")
        for idx, line in enumerate(lines):
            box = draw.textbbox((0, 0), line, font=pil_font(20, True))
            tx = (x0 + x1 - (box[2] - box[0])) / 2
            draw.text((tx, 545 + idx * 34), line, font=pil_font(20, True), fill="white")

    for sy in (420, 575, 730):
        arrow(draw, (45, sy), (1725, sy), BLUE, 7)
    draw.text((70, 250), "AMBIENT / GRILLE SIDE", font=pil_font(24, True), fill=colour(BLUE))
    draw.text((1450, 250), "ENGINE SIDE", font=pil_font(24, True), fill=colour(NAVY))

    draw.rounded_rectangle((390, 175, 590, 300), radius=15, fill=colour("D8EEF4"), outline=colour(CYAN), width=4)
    draw.text((433, 198), "CAC", font=pil_font(26, True), fill=colour(CYAN))
    draw.rounded_rectangle((605, 175, 805, 300), radius=15, fill=colour("DDE2E5"), outline=colour(MUTED), width=4)
    draw.text((627, 198), "CONDENSER", font=pil_font(22, True), fill=colour(MUTED))
    draw.text((405, 263), "SIDE-BY-SIDE — SAME DEPTH", font=pil_font(16, True), fill=colour(INK))

    dimension(draw, (105, 900), (1300, 900), "P0 — measure complete installed front/rear envelope")
    rounded_box(
        draw,
        (100, 970, 825, 1090),
        "WHY THREE HEAT EXCHANGERS?",
        "R0 carries engine coolant. The condenser carries R134a refrigerant. The CAC carries compressed intake air. None can perform another circuit's job.",
        "F5F7F9",
        LINE,
    )
    rounded_box(
        draw,
        (900, 970, 1700, 1090),
        "DO NOT BUILD IN SERIES",
        "Do not place CAC in front of the condenser. Their shared shallow carrier is split left/right, then both discharge into the full-width R0 behind.",
        "FFF0EE",
        RED,
        RED,
    )
    diagram_footer(draw, "REV J · EXACTLY THREE FANS TOTAL · ALL AIR MOVES GRILLE → ENGINE")
    return save_diagram(image, D16)


def prepare_reference_photo(source, output):
    if not source.exists():
        return None
    with Image.open(source) as original:
        image = ImageOps.exif_transpose(original).convert("RGB")
        image.thumbnail((1600, 1400), Image.Resampling.LANCZOS)
        image.save(output, "JPEG", quality=84, optimize=True, progressive=True)
    return output


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:cantSplit")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def set_table_width(table, dxa=9360, indent=120):
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(dxa))
    width.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths=None, font_size=8.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    header = table.rows[0]
    repeat_header(header)
    prevent_row_split(header)
    for index, heading in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(heading)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    for values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for index, value in enumerate(values):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            run.font.name = "Calibri"
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor.from_string(INK)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_page_heading(doc, text, level=1):
    paragraph = add_heading(doc, text, level)
    paragraph.paragraph_format.page_break_before = True
    return paragraph


def add_para(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        run = paragraph.add_run(bold_lead)
        run.bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_numbered(doc, items):
    for index, item in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(f"{index}. ").bold = True
        paragraph.add_run(item)


def add_callout(doc, heading, body, fill="FFF4DD", border=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 160, 130, 160)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "10")
        tag.set(qn("w:color"), border)
        borders.append(tag)
    tc_pr.append(borders)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(heading + "\n")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(border)
    paragraph.add_run(body)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.3


def add_image_path(doc, path, width=6.35, caption=None):
    if not path or not Path(path).exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(6)
        run = cap.runs[0]
        run.italic = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Page ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    settings = {
        1: (16, 18, 10, NAVY),
        2: (13, 14, 7, BLUE),
        3: (12, 10, 5, GREEN),
    }
    for level, (size, before, after, tone) in settings.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(tone)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "J40 · RETAINED-ORIGINAL COOLING PACK · REV J"
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(section.footer.paragraphs[0])


def build_docx():
    prepare_reference_photo(R0_PHOTO, R0_REFERENCE)
    prepare_reference_photo(GUARD_PHOTO, GUARD_REFERENCE)
    prepare_reference_photo(FAN_PHOTO, FAN_REFERENCE)

    document = Document()
    configure(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    run = title.add_run("J40 RETAINED-ORIGINAL\nCOOLING PACK")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor.from_string(NAVY)
    issue = document.add_paragraph("FABRICATOR SPECIFICATION · REV J · 02 AUGUST 2026")
    issue.alignment = WD_ALIGN_PARAGRAPH.CENTER
    issue.runs[0].bold = True
    issue.runs[0].font.size = Pt(11)
    issue.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

    add_callout(
        document,
        "THE RELEASED IDEA",
        "Refurbish the actual R0 engine radiator. Fit exactly two matching, equal-size electric pushers centrally at the front. Put one CAC and the existing serviceable A/C condenser side-by-side in one shallow plane behind them. Keep the large 2H mechanical puller behind R0. Use both intended horizontal top holes.",
        "EAF4FA",
        BLUE,
    )
    add_image_path(document, PH01, 6.35, "PH01 — Opaque assembled proposal. Appearance and relationships only; do not scale this image.")
    add_para(document, "Duty: Toyota 2H with conservative future turbo plan, maximum 150 bhp / 112 kW, R134a A/C and instrumented 50 °C ambient acceptance.")
    add_para(document, "Karigar ke liye: purana bara radiator engine ke liye rahe ga. Aagay do aik-jaisay pankhay beech mein. Un ke peechay turbo cooler aur A/C condenser side-by-side. Dono asal upar walay suraakh istemal karo.")

    document.add_page_break()
    add_heading(document, "Start with these actual parts")
    add_image_path(document, R0_REFERENCE, 6.35, "R01 — Actual removed R0 radiator. Its physical geometry is the master for every radiator interface.")
    add_image_path(document, GUARD_REFERENCE, 6.35, "R02 — Actual removable stone guard. Straighten/repair and keep separately removable.")
    add_image_path(document, FAN_REFERENCE, 6.35, "R03 — Existing installed context for the retained large 2H mechanical fan. Inspect fan and shroud before reuse.")
    add_callout(
        document,
        "FRONT-FAN HOLD",
        "The archive has no reliable isolated photograph or dimensions for the existing front electric fan. Photograph, measure and test that complete module before buying its twin. If it cannot be truly matched, install a matched Toyota/Denso pair and keep the old fan as a spare.",
        "FFF4DD",
        GOLD,
    )

    add_page_heading(document, "1. What is actually in the pack")
    add_table(document, ["Circuit", "Physical item", "Rev J treatment"], [
        ("Engine coolant", "Actual full-width R0 copper/brass radiator", "Retain/refurbish. Recore only if condition, pressure/flow or thermal testing requires it."),
        ("R134a refrigerant", "Existing A/C condenser", "Reuse only after identification, pressure, flush, fin and fit pass; otherwise measured replacement."),
        ("Turbo charge air", "One compact CAC/intercooler", "New, pressure-tested core with measured tanks/ports and ≤10 kPa route pressure-drop target."),
    ], [1.1, 2.15, 3.15])
    add_para(document, "The three heat exchangers are not duplicates. They contain different fluids and serve different circuits. R0 cannot condense refrigerant or cool turbo charge air.")
    add_image_path(document, D16, 5.75, "D16 — Air order and side-by-side auxiliary-plane rule.")
    add_callout(
        document,
        "AIR ORDER",
        "Stone guard → two matching front pushers → one shallow CAC/condenser split plane → full-width R0 → sealed rear shroud → retained large mechanical puller → engine.",
        "EAF4FA",
        BLUE,
    )

    add_page_heading(document, "2. Retain first; replace only after evidence")
    add_image_path(document, D14, 6.35, "D14 — Retain/test/replace release logic.")
    add_table(document, ["Item", "Mandatory inspection / test", "Release"], [
        ("R0 radiator", "Clean; tank, seam, neck, rail and ear inspection; pressure and flow test; record R0-W/H/D and active core.", "Retain/refurbish. Recore only after a failed gate; preserve external geometry."),
        ("Large rear fan/shroud", "Hub/bearing, blade cracks/runout, sweep, shroud sealing, axial/radial clearance and engine rock.", "Retain/recondition; minimum sealed adapter only if needed."),
        ("Existing front electric fan", "Complete F×H×D, blade sweep, label/plug, tabs, rotation, start/run current and bench test.", "One half of installed pair only if a true equal-envelope twin is found and pair passes."),
        ("Existing condenser", "Identify; pressure/flush; fin, port, drier and side-by-side fit checks.", "Reuse if passed; otherwise measured replacement."),
        ("Stone guard + electrics", "Guard straightness/open mesh; closed fuse covers, cable bends, service access.", "Reuse inside current pack silhouette; no side annex."),
    ], [1.3, 3.15, 2.2])

    add_page_heading(document, "3. Dimensions — measure before fabrication")
    add_image_path(document, D15, 6.35, "D15 — Two matching front modules centred on C0, with the exact top-hole mounting detail.")
    add_callout(
        document,
        "CENTRING KEY",
        "M = (Wc − 2F − g) / 2. Release only if M is positive, the left and right margins agree within 2 mm, the complete pair midpoint is on C0 within 2 mm, and both hubs are level within 2 mm.",
        "EAF4FA",
        BLUE,
    )
    add_page_heading(document, "Dimension release sheet", level=2)
    add_table(document, ["Code", "Measure / requirement", "Release value"], [
        ("W0", "Clear chassis opening at top, middle and bottom, including latch/bonnet/steering/service obstructions.", "MEASURE vehicle"),
        ("P0", "Available grille-to-engine depth with guard, full fan frames, auxiliary carrier, R0, shroud and engine sweep.", "MEASURE vehicle"),
        ("R0-W/H/D", "Complete R0 envelope including tanks, seams, cap, necks, drain, rails, ears and locators.", "MEASURE R0"),
        ("Wc", "Usable front fan carrier/opening width between released obstructions.", "MEASURE fixture"),
        ("F", "Outside width of each accepted complete matching fan module.", "MEASURE both"),
        ("g", "Clear gap between complete modules.", "10–15 mm preferred"),
        ("M", "Equal outer margin: (Wc − 2F − g) / 2.", ">0; left/right equal ±2 mm"),
        ("C0", "Vehicle centreline and complete fan-pair midpoint.", "Coincident ±2 mm"),
        ("B0", "Rigid 1:1 template of two original top holes: C-C, diameter, edge distance, plate thickness and height.", "COPY exact vehicle holes"),
    ], [0.75, 4.1, 1.7], 8.4)
    add_para(document, "No fixed fan diameter, hub coordinate or part number is released from an AI image. Both hubs must be level within 2 mm. Confirm guards, tabs, plugs, wire bends and fan sweep before steel cutting.")

    add_page_heading(document, "4. Mount the radiator in the intended top holes")
    add_image_path(document, PH03, 6.35, "PH03 — Installed proposal. Both bolt heads visibly occupy the existing holes in the inward horizontal top-return plates.")
    add_numbered(document, [
        "Each grey upright ends in an inward-facing horizontal top-return plate with its existing circular hole.",
        "A short black removable R0-side ear rises from the matching radiator side rail and sits directly below that plate.",
        "One vertical M8 bolt and large washer sit on top of each grey plate and pass through that exact existing hole.",
        "The bolt continues through a sleeved EPDM/rubber bush and the black radiator ear, with washer and locking nut below.",
        "Do not drill, slot, ream or substitute a bolt through an upright side face. Do not leave a visual or mechanical air gap.",
        "Two lower rubber saddles carry the radiator weight. The top bolts locate and restrain without hard-clamping tank movement.",
    ])
    add_callout(
        document,
        "TOP-MOUNT ACCEPTANCE",
        "Two bolts total. Both original horizontal top holes occupied. B0 template agrees with the vehicle. R0 sits naturally in the lower saddles and neither tank nor rail is pulled into alignment.",
        "FFF0EE",
        RED,
    )

    add_page_heading(document, "5. Front fans and shallow auxiliary carrier")
    add_table(document, ["Requirement", "Fabricator instruction"], [
        ("Front pair", "Exactly two matching, equal-size complete Toyota/Denso-family pushers. Unequal fans are not permitted."),
        ("Reuse route", "Use the tested existing electric fan plus a genuine matching twin only if complete envelopes, airflow direction, electrical load and condition match."),
        ("Fallback route", "If no true twin or any gate fails, install a measured matched pair and keep the old electric fan as a spare."),
        ("Centring", "Complete group midpoint on C0 ±2 mm; equal outer margins ±2 mm; hubs at the same height ±2 mm."),
        ("Installed airflow", "≥3,000 m³/h aggregate through finished guard + split auxiliary plane + R0 at 13.5 V; ≥3,300 preferred."),
        ("Auxiliary carrier", "CAC and condenser side-by-side in one depth plane, separated by a sealed divider; each behind one fan zone."),
        ("Orientation", "Choose CAC-left or CAC-right only after turbo pipe, A/C port/drier, bonnet latch and service-envelope mock-up."),
        ("Width/service", "All cores, ports, brackets and closed electrical boxes stay inside W0; each part removable without cutting."),
    ], [1.45, 5.05])
    add_image_path(document, PH02, 6.35, "PH02 — Wide component split-out: guard, matched front pair, side-by-side auxiliary plane, R0, rear shroud and large fan.")

    add_page_heading(document, "6. Electrical and controls")
    add_table(document, ["Fan", "Normal lead request", "Fail-safe request"], [
        ("Front fan — A/C zone", "A/C pressure/demand leads.", "High coolant, high IAT, high A/C pressure, high load, fault or manual emergency: BOTH front fans."),
        ("Front fan — CAC zone", "Boost/high IAT leads.", "Same shared BOTH-fans request."),
        ("Rear mechanical fan", "Engine driven as fitted.", "Not a substitute for either electric branch and not counted in electrical airflow test."),
    ], [1.55, 2.25, 2.7])
    add_bullets(document, [
        "Both electric fans push grille-to-engine. Verify rotation and airflow on the finished vehicle.",
        "Keep one independently fused relay branch per motor in the existing covered relay/MIDI carrier arrangement.",
        "Size relay, MIDI/fuse, cable, connector and earth from measured start/run current and voltage-drop results.",
        "Keep covers closed, cable exits down/rear and service access clear. Do not hang electrical weight on a core or radiator.",
        "Final thresholds and engine/A/C interlocks belong to the electrical calibration release.",
    ])
    add_heading(document, "7. Performance targets for Pakistan duty")
    add_table(document, ["Item", "Acceptance target"], [
        ("R0 heat rejection", "≥115 kW continuous; ≥130 kW for 10 minutes at the defined 50 °C duty."),
        ("Front-pair installed airflow", "≥3,000 m³/h aggregate; ≥3,300 m³/h preferred, through the complete installed stack at 13.5 V."),
        ("CAC thermal result", "Post-CAC IAT ≤80 °C at the agreed 150 bhp / boost test point."),
        ("CAC pressure loss", "Complete charge route ≤10 kPa at that point."),
        ("Derate", "No cooling-system-caused boost or engine-load derate inside the accepted 150 bhp envelope."),
    ], [2.0, 4.5])

    add_page_heading(document, "8. Minimum fabrication and purchase")
    add_table(document, ["Reuse / buy", "Qty", "Item and rule"], [
        ("REUSE", "1", "Actual R0 radiator; test first, preserve its complete external geometry."),
        ("REUSE", "1", "Large rear 2H mechanical fan and shroud; inspect/recondition."),
        ("REUSE IF PASS", "1", "Existing front electric fan as one half of the pair; otherwise service spare."),
        ("REUSE IF PASS", "1", "Existing A/C condenser."),
        ("REUSE", "1", "Stone guard and existing covered relay/MIDI carrier arrangement."),
        ("BUY PREFERRED", "1", "True matching Toyota/Denso-family twin for accepted existing fan, with plug/pigtail."),
        ("BUY FALLBACK", "2", "Physically matched Toyota/Denso-family complete fan pair if the existing fan cannot be matched/accepted."),
        ("BUY", "1", "Compact CAC after W0/P0 and routing release; include pressure-test evidence."),
        ("BUY / FABRICATE", "1 set", "Centred fan carrier; split auxiliary carrier/divider; seals, edging, P-clips and isolators."),
        ("BUY", "2 assemblies", "Two M8 top-locator assemblies total; one per exact original horizontal top hole."),
        ("BUY", "2", "Lower radiator rubber saddles matched to R0."),
        ("BUY AS MEASURED", "2 branches", "Relays, MIDI/fuses, cable, connectors, earths and sleeving sized from measured currents."),
    ], [1.2, 0.8, 4.5], 8.2)
    add_callout(
        document,
        "WHAT NOT TO BUY",
        "Do not buy a second engine radiator. Do not buy two extra auxiliary radiators. Do not buy a guessed fan diameter or fixed Prado part number. Prado/Land Cruiser/Toyota breaker stock is only the practical sample-matching pool.",
        "FFF4DD",
        GOLD,
    )

    add_page_heading(document, "9. Fabrication and test sequence")
    add_numbered(document, [
        "Photograph/tag R0, guard, rear fan/shroud, existing electric fan, condenser and existing covered electrical arrangement.",
        "Record W0, P0, R0-W/H/D, Wc, F, g, C0, B0, lower saddles, engine-fan sweep and every pipe/cable/service envelope.",
        "Clean and test R0. Refurbish it; recore only if a recorded failure earns that work. Test the condenser and existing fans.",
        "Find a true fan twin or release a matched pair. Bench-check rotation, start/run current and condition.",
        "Mock up the opaque full pack 1:1 with grille/front panel, bonnet/latch, guard, fan frames, auxiliary cores/ports, R0, rear shroud/fan, hoses, drier, charge pipes and closed fuse/relay carrier.",
        "Transfer B0 only to the two removable R0-side ears. Fit lower saddles; insert the two vertical M8 bolts through the exact horizontal top-return holes without forcing alignment.",
        "Fabricate the minimum removable fan and split auxiliary carriers. Seal edge bypass and the central divider. Do not pass ties through fins or load one core from another.",
        "Measure installed front-pair airflow at 13.5 V through the finished stack. Correct and repeat until ≥3,000 m³/h.",
        "Pressure/flow/electrical check, then instrument hot-idle A/C and loaded 50 °C testing or documented equivalent simulation plus hottest available road test.",
    ])
    add_page_heading(document, "10. Acceptance record")
    add_table(document, ["Gate", "Pass evidence", "Initial/date"], [
        ("R0 identity/condition", "Geometry record + pressure/flow result; recore reason if applicable.", "__________"),
        ("Matching front pair", "Two equal envelopes; rotation/current/condition records; centred formula complete.", "__________"),
        ("Side-by-side auxiliary plane", "CAC and condenser in one depth plane; divider/edge seals; ports/service proven.", "__________"),
        ("Top attachment", "Both exact horizontal holes occupied by vertical M8 assemblies; B0 matches; no strain.", "__________"),
        ("Rear fan clearance", "Full sweep and engine-rock clearance; shroud sealed; no contact.", "__________"),
        ("Installed airflow", "≥3,000 m³/h aggregate at 13.5 V through complete stack.", "__________"),
        ("50 °C thermal duty", "Ambient, coolant in/out, IAT pre/post, boost, EGT, commands, voltage/current logged.", "__________"),
    ], [1.65, 3.95, 0.9], 8.2)
    add_para(document, "Reject progressive temperature rise, boiling/coolant loss, hose collapse, fan/shroud contact, electrical overheating, uncontrolled IAT, excessive pressure drop or cooling-system boost/load derate. Correct and repeat.")
    add_callout(
        document,
        "FINAL SHOP RULE",
        "The physical R0, existing parts, B0/C0/W0/P0 measurements, deterministic drawings, 1:1 fixture and signed test record control manufacture. No render or nominal target guarantees fit or 50 °C performance.",
        "FFF0EE",
        RED,
    )

    add_page_heading(document, "Appendix A — Opaque visual cross-check")
    add_image_path(document, PH04, 6.35, "PH04 — Engine-side proposal: retained large 2H mechanical puller immediately behind R0; both upper top-return holes use vertical bolts.")
    add_image_path(document, PH03, 6.35, "PH03 repeated for the front-side top-hole attachment. AI images do not release dimensions.")

    properties = document.core_properties
    properties.title = "J40 Retained-Original Cooling Pack — Rev J"
    properties.subject = "Pakistan radiator-fabricator specification"
    properties.author = "J40 Rebuild Project"
    properties.keywords = "J40, retained radiator, matched fans, side-by-side CAC condenser, Toyota Denso, Pakistan, 50 C"
    document.save(OUT)
    return OUT


if __name__ == "__main__":
    for builder in (build_d14, build_d15, build_d16):
        print(builder())
    print(build_docx())
