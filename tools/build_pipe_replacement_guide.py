from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "J40-pipe-and-hose-replacement-location-guide-20260723.docx"
PHOTOS = ROOT / "photos"

IMAGES = {
    "overview": PHOTOS / "20260723_055107_gp_e0sQOMzg.jpg",
    "large": PHOTOS / "20260723_055100_gp_X5l447Ow.jpg",
    "overflow": PHOTOS / "20260723_055053_gp_Nbrk41Tw.jpg",
    "long_small": PHOTOS / "20260723_055044_gp_kapiWhpQ.jpg",
    "dn10_pair": PHOTOS / "20260723_055049_gp_gDhxAOhg.jpg",
    "long_medium": PHOTOS / "20260723_055040_gp_KIqpzvmw.jpg",
    "dn10_overview": PHOTOS / "20260723_055032_gp_WF9QPyuA.jpg",
    "dn10_marking": PHOTOS / "20260723_054008_gp_fJeietsQ.jpg",
    "engine_before_front": PHOTOS / "20260317_235150.jpg",
    "engine_before_overview": PHOTOS / "20260317_235229.jpg",
    "engine_front_joiner": PHOTOS / "20260430_215957_gp_2iBbUagw.jpg",
}

BLUE = RGBColor(46, 116, 181)
NAVY = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 98, 108)
PALE = "E8EEF5"
WARN = "FFF4CE"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    set_font(p.add_run(text), size=9, italic=True, color=MUTED)


def add_image(doc, key, width, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    run.add_picture(str(IMAGES[key]), width=Inches(width))
    add_caption(doc, caption)


def add_status_box(doc, label, text, fill=PALE):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [1700, 7660])
    table.rows[0].cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table.rows[0].cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(table.rows[0].cells[0], fill)
    p = table.rows[0].cells[0].paragraphs[0]
    set_font(p.add_run(label), bold=True, color=NAVY)
    p = table.rows[0].cells[1].paragraphs[0]
    set_font(p.add_run(text))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, NAVY, 10, 5),
):
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = color
    st.font.bold = True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("J40 / 2H pipe and hose installation guide"), size=9, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("Prepared 23 July 2026 | Dry-fit before final clamping"), size=8.5, color=MUTED)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(3)
set_font(p.add_run("PIPE AND HOSE REPLACEMENT GUIDE"), size=23, bold=True, color=NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
set_font(p.add_run("Toyota Land Cruiser J40 / HJ47 with 2H diesel"), size=14, color=MUTED)

add_status_box(
    doc,
    "Purpose",
    "Identify where each new hose belongs, retain the original metal coolant joiners where serviceable, and prevent a hose from being installed merely because its diameter appears close.",
)

doc.add_heading("Installation rule", level=1)
add_bullet(doc, "Keep each old hose beside its replacement until both ends, bend direction and installed route are confirmed.")
add_bullet(doc, "Reuse the two original metal coolant pipes after internal cleaning, corrosion inspection and pressure testing. They are joining pieces, not scrap.")
add_bullet(doc, "Use new smooth-band or constant-tension clamps on coolant hoses. Do not rely on the old perforated worm-drive clamps where they are rusty or distorted.")
add_bullet(doc, "Do not trim a molded hose until the radiator, engine and retained metal joiner are in their final positions.")

doc.add_heading("Quick location map", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
for i, text in enumerate(("Replacement group", "Where it installs", "Release position")):
    set_cell_shading(table.rows[0].cells[i], PALE)
    set_font(table.rows[0].cells[i].paragraphs[0].add_run(text), bold=True, color=NAVY)
rows = [
    ("Large molded hoses + metal joiners", "Radiator upper/lower outlets to the 2H thermostat outlet and water-pump inlet", "Dry-fit; likely correct"),
    ("Small clear/black hose", "Radiator filler-neck overflow nipple to reserve bottle", "Replace clear old hose"),
    ("Medium hoses", "Heater feed/return or breather route; identify by old end diameter and fitted nipples", "Label before fitting"),
    ("ARCHOR DN10 hose", "Likely brake-booster vacuum and/or low-pressure oil route", "Conditional; not coolant or diesel-approved by marking"),
]
for vals in rows:
    cells = table.add_row().cells
    for i, value in enumerate(vals):
        set_font(cells[i].paragraphs[0].add_run(value), size=9.5)
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
set_table_widths(table, [2200, 4300, 2860])

doc.add_page_break()
doc.add_heading("1. Main radiator coolant assemblies", level=1)
p = doc.add_paragraph()
set_font(p.add_run("Where installed. "), bold=True)
set_font(p.add_run("Across the front of the 2H engine, between the radiator and the two large engine coolant necks. The upper circuit connects the thermostat/water outlet to the radiator upper inlet. The lower circuit connects the radiator lower outlet to the water-pump/lower inlet."))

add_image(doc, "large", 2.65, "New molded hose beside its old counterpart. The new part may appear slimmer but the bend and inside diameter control fit.")

doc.add_heading("Before dismantling — location reference", level=2)
add_image(doc, "engine_before_front", 2.45, "Looking down at the front of the installed engine: the upper radiator hose is the large black elbow at lower centre, joining the thermostat/water outlet to the radiator top tank. The brown small-bore hose crossing the radiator top is the overflow line.")
add_image(doc, "engine_front_joiner", 2.35, "Front of engine after the radiator was removed: this close view records one large coolant route and its original metal elbow. Rebuild this same three-piece geometry with the corresponding new rubber ends; the open end at the top continues toward the radiator.")

doc.add_heading("How to assemble", level=2)
add_bullet(doc, "Build the upper and lower routes using the new molded rubber pieces and the matching original metal joining pipe shown on page 1.")
add_bullet(doc, "Orient each metal elbow exactly as it sat in the old assembly; swapping the two joiners can aim a hose toward the fan or belts.")
add_bullet(doc, "Push every hose fully over its neck or metal pipe, with the clamp behind the retention bead or raised end—not on the lip.")
add_bullet(doc, "Dry-fit with the radiator and fan installed. Reject any arrangement that twists, flattens or places a hose against a pulley, belt, fan, radiator edge or body bracket.")

add_status_box(doc, "Decision", "The new molded coolant hoses are reasonable replacement candidates. Their apparent thinner wall is not proof of lower quality, but no readable coolant standard is shown. Accept them only after diameter, shape and installed clearance checks.")

doc.add_heading("Metal joiner acceptance", level=2)
add_bullet(doc, "Reuse if the tube is structurally solid, free of deep pitting or pinholes, internally unrestricted and able to hold the cooling-system pressure test.")
add_bullet(doc, "Clean internally and externally, remove loose scale, preserve or restore the hose-retention beads, then protect the exterior after the successful dry-fit.")

doc.add_heading("2. Radiator overflow hose", level=1)
add_image(doc, "overflow", 2.45, "New small black hose beside the aged transparent overflow hose.")
add_image(doc, "engine_before_overview", 3.9, "Pre-dismantling overview: the old translucent brown overflow hose starts at the radiator filler neck/cap at the front centre, crosses the radiator top, and ends at the pale reserve bottle at the front-left corner of the photograph.")
p = doc.add_paragraph()
set_font(p.add_run("Where installed. "), bold=True)
set_font(p.add_run("From the small overflow nipple at the radiator filler neck/cap seat to the nipple on the coolant reserve/overflow bottle."))
add_bullet(doc, "Route continuously downhill where practical, without a low loop that can trap coolant.")
add_bullet(doc, "Keep it clear of the fan, belts and exhaust heat; reuse the original clip locations.")
add_bullet(doc, "Confirm the new hose grips both small nipples without splitting or requiring excessive clamp force." )
add_status_box(doc, "Decision", "The black replacement is preferable to the hardened, discoloured transparent hose if its inside diameter matches and it is coolant-compatible EPDM.")

doc.add_heading("3. Heater, breather and other medium hoses", level=1)
add_image(doc, "long_small", 2.3, "New and old small/medium hose comparison. Match by both end diameter and the original fitted route.")
add_image(doc, "long_medium", 2.3, "Another long-hose comparison. Similar length alone does not identify its circuit.")
p = doc.add_paragraph()
set_font(p.add_run("Likely installation points. "), bold=True)
set_font(p.add_run("The two heater hoses run between the engine heater-feed/return nipples and the heater-core connections at the firewall. A larger oil-resistant hose may instead serve the crankcase breather/oil-mist route."))
add_status_box(doc, "Photo limit", "The pre-dismantling engine photographs show the firewall area but do not trace both ends of these medium hoses clearly enough for a reliable one-to-one assignment. Identify them from the old hose endpoints and nipple diameters before fitting.", WARN)
add_bullet(doc, "Label the old sample at both ends before removal: engine, firewall, breather or other component.")
add_bullet(doc, "Heater hose must be hot-coolant-rated (SAE J20R3 or equivalent). Breather hose must be oil-mist-resistant.")
add_bullet(doc, "Do not substitute one circuit's hose for another simply because the outside diameter looks similar.")
add_status_box(doc, "Decision", "The photographed shapes and lengths are plausible, but these unmarked hoses cannot be finally assigned or approved from the flat-lay photographs alone.", WARN)

doc.add_heading("4. ARCHOR DN10 reinforced hose", level=1)
add_image(doc, "dn10_pair", 2.45, "New ARCHOR hose beside the old fabric-braided hose. The smaller outside diameter is not automatically a weakness.")
add_image(doc, "dn10_marking", 3.8, "Readable marking: DIN EN 854 / SAE 100R6 / ISO 4079, DN10 (3/8 in), working pressure 406 psi / 28 bar.")

p = doc.add_paragraph()
set_font(p.add_run("What the marking proves. "), bold=True)
set_font(p.add_run("This is a textile-reinforced low-pressure hydraulic hose with a nominal 10 mm bore and 28 bar working-pressure rating. Modern reinforcement can legitimately give it a smaller outside diameter than the old hose."))

p = doc.add_paragraph()
set_font(p.add_run("Most likely installation. "), bold=True)
set_font(p.add_run("The principal candidate is the long brake-booster vacuum route from the 2H vacuum pump/check-valve plumbing to the brake booster at the firewall. A second length may be intended for a low-pressure oil/vacuum-pump return route, but it must be identified against the labelled old hose and its fitted endpoints."))

add_status_box(doc, "Photo limit", "The available pre-dismantling views do not show the complete DN10 hose from source to destination. Do not copy the path of an adjacent air-conditioning, fuel or coolant hose; confirm the vacuum-pump/check-valve endpoint and the brake-booster endpoint directly on the vehicle.", WARN)

add_status_box(doc, "Suitable", "Potentially suitable for brake-booster vacuum or a low-pressure oil circuit if the manufacturer confirms continuous-vacuum service, oil compatibility and the required engine-bay temperature range.")
add_status_box(doc, "Not released", "Do not use it for radiator/heater coolant: no SAE J20 coolant rating is shown. Do not use it for diesel feed/return: no SAE J30 or DIN 73379 fuel rating is shown.", WARN)

doc.add_heading("Final checks before starting the engine", level=1)
add_bullet(doc, "Photograph and label every completed route before the radiator shroud or other parts hide it.")
add_bullet(doc, "Pressure-test the cooling system cold; inspect every metal-pipe joint and clamp for seepage.")
add_bullet(doc, "Start the engine, warm it fully and recheck for hose softening, collapse, ballooning, abrasion and fan/belt contact.")
add_bullet(doc, "For the brake-booster vacuum hose, confirm strong brake assist and verify that the hose does not flatten when the engine is running.")
add_bullet(doc, "After the first full heat cycle, allow the engine to cool and recheck clamp seating and coolant level.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
