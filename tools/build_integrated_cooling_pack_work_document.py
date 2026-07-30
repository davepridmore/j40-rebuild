from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "J40-integrated-cooling-pack-work-document-rev-b.docx"
ASSET = ROOT / "data" / "manual" / "fabrication" / "front_cooling_stack_rev_a" / "work_document_assets"
ASSET.mkdir(parents=True, exist_ok=True)

NAVY = "17324D"; BLUE = "2C6E9B"; CYAN = "4BA3C7"; RED = "C9534B"
GOLD = "C9952E"; GREEN = "4E7D61"; LIGHT = "EAF0F4"; PALE = "F5F7F9"
INK = "1E2933"; MUTED = "62717D"; WHITE = "FFFFFF"; LINE = "B7C3CC"

def font(size, bold=False):
    paths = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for p in paths:
        if Path(p).exists(): return ImageFont.truetype(p, size)
    return ImageFont.load_default()

def text_center(draw, box, text, fnt, fill=INK, spacing=5):
    x1,y1,x2,y2=box
    bb=draw.multiline_textbbox((0,0), text, font=fnt, align="center", spacing=spacing)
    draw.multiline_text(((x1+x2-(bb[2]-bb[0]))/2,(y1+y2-(bb[3]-bb[1]))/2),text,font=fnt,fill="#"+fill,align="center",spacing=spacing)

def rounded(draw, box, fill, outline=LINE, radius=20, width=3):
    draw.rounded_rectangle(box, radius=radius, fill="#"+fill, outline="#"+outline, width=width)

def arrow(draw, a, b, color=BLUE, width=10):
    draw.line([a,b], fill="#"+color, width=width)
    import math
    ang=math.atan2(b[1]-a[1],b[0]-a[0]); s=22
    pts=[b,(b[0]-s*math.cos(ang-.55),b[1]-s*math.sin(ang-.55)),(b[0]-s*math.cos(ang+.55),b[1]-s*math.sin(ang+.55))]
    draw.polygon(pts, fill="#"+color)

def dim_h(draw, x1, x2, y, ref_y, label, color=RED):
    draw.line((x1,ref_y,x1,y),fill="#"+color,width=3)
    draw.line((x2,ref_y,x2,y),fill="#"+color,width=3)
    draw.line((x1,y,x2,y),fill="#"+color,width=4)
    draw.polygon([(x1,y),(x1+18,y-9),(x1+18,y+9)],fill="#"+color)
    draw.polygon([(x2,y),(x2-18,y-9),(x2-18,y+9)],fill="#"+color)
    bb=draw.textbbox((0,0),label,font=font(23,True)); tw=bb[2]-bb[0]
    draw.rectangle((x1+(x2-x1-tw)/2-10,y-18,x1+(x2-x1+tw)/2+10,y+18),fill="white")
    draw.text((x1+(x2-x1-tw)/2,y-15),label,font=font(23,True),fill="#"+color)

def dim_v(draw, y1, y2, x, ref_x, label, color=RED):
    draw.line((ref_x,y1,x,y1),fill="#"+color,width=3)
    draw.line((ref_x,y2,x,y2),fill="#"+color,width=3)
    draw.line((x,y1,x,y2),fill="#"+color,width=4)
    draw.polygon([(x,y1),(x-9,y1+18),(x+9,y1+18)],fill="#"+color)
    draw.polygon([(x,y2),(x-9,y2-18),(x+9,y2-18)],fill="#"+color)
    imtxt=Image.new("RGBA",(420,45),(255,255,255,235)); td=ImageDraw.Draw(imtxt); td.text((8,5),label,font=font(22,True),fill="#"+color)
    imtxt=imtxt.rotate(90,expand=True); draw._image.paste(imtxt,(int(x-22),int((y1+y2-imtxt.height)/2)),imtxt)

def save_front():
    im=Image.new("RGB",(1800,1100),"white"); d=ImageDraw.Draw(im)
    d.text((70,50),"FRONT ELEVATION — LOOKING THROUGH THE GRILLE",font=font(42,True),fill="#"+NAVY)
    d.text((70,105),"Component bodies shown; mounting tabs and tank seams remain site-fit.",font=font(25),fill="#"+MUTED)
    # uprights
    rounded(d,(155,205,255,970),NAVY,NAVY,12); rounded(d,(1545,205,1645,970),NAVY,NAVY,12)
    d.text((166,505),"RIGHT\nUPRIGHT",font=font(24,True),fill="white",spacing=7)
    d.text((1556,505),"LEFT\nUPRIGHT",font=font(24,True),fill="white",spacing=7)
    # radiator rear outline
    rounded(d,(330,265,1470,925),LIGHT,BLUE,24,6)
    text_center(d,(420,790,1380,900),"ENGINE RADIATOR\n530 × 435 × 64 mm core",font(32,True),BLUE)
    # condenser
    rounded(d,(370,315,1430,790),"DDECF3",CYAN,20,6)
    text_center(d,(480,360,1320,500),"A/C CONDENSER\n559 × 356 × 21 mm nominal",font(30,True),NAVY)
    # intercooler lower band
    rounded(d,(450,650,1350,905),"E4EEE8",GREEN,20,6)
    text_center(d,(520,700,1280,850),"CHARGE-AIR INTERCOOLER\n500 × 180 × 60 mm core\n57 mm beaded outlets",font(29,True),GREEN)
    # fan upper
    d.ellipse((695,300,1105,710),fill="#F3E8D2",outline="#"+GOLD,width=8)
    d.ellipse((840,445,960,565),fill="#"+GOLD)
    for ang in range(0,360,60):
        import math
        x=900+150*math.cos(math.radians(ang)); y=505+150*math.sin(math.radians(ang))
        d.line((900,505,x,y),fill="#"+GOLD,width=38)
    text_center(d,(650,210,1150,305),"SLIM 12-IN A/C PUSHER FAN\nupper band; frame-mounted",font(27,True),GOLD)
    # drier
    rounded(d,(1480,400,1530,715),"EEE7D7",GOLD,14,5)
    d.text((1350,735),"Receiver-drier\n(vertical, outside airflow)",font=font(22,True),fill="#"+GOLD,align="center")
    im.save(ASSET/"front_elevation.png",quality=95)

def save_side():
    im=Image.new("RGB",(1800,1050),"white"); d=ImageDraw.Draw(im)
    d.text((70,50),"SIDE SECTION — AIRFLOW AND REQUIRED GAPS",font=font(42,True),fill="#"+NAVY)
    d.text((70,105),"Nominal lower-band stack depth: 170 mm; allow 180–190 mm installed.",font=font(25),fill="#"+MUTED)
    # grille & engine
    rounded(d,(120,250,190,900),"D8DEE3",MUTED,8); d.text((100,920),"GRILLE",font=font(24,True),fill="#"+MUTED)
    rounded(d,(1560,270,1690,880),"E8E3DD",INK,18); d.text((1573,900),"ENGINE",font=font(24,True),fill="#"+INK)
    # blocks
    rounded(d,(330,535,560,865),"E4EEE8",GREEN,12,5); text_center(d,(335,555,555,835),"INTERCOOLER\n60 mm\nlower band",font(28,True),GREEN)
    rounded(d,(690,300,800,865),"DDECF3",CYAN,12,5); text_center(d,(685,350,805,810),"A/C\nCONDENSER\n21 mm",font(26,True),NAVY)
    rounded(d,(930,250,1190,900),LIGHT,BLUE,12,5); text_center(d,(940,320,1180,830),"ENGINE\nRADIATOR\n64 mm",font(29,True),BLUE)
    # fan/shroud
    rounded(d,(1215,310,1510,840),PALE,MUTED,18,4); d.ellipse((1270,410,1455,595),outline="#"+INK,width=8)
    text_center(d,(1230,630,1495,790),"REAR SHROUD\nengine-driven fan\n≥20 mm static\n25–30 preferred",font(24,True),INK)
    # gaps/dims
    arrow(d,(560,760),(685,760),MUTED,5); arrow(d,(685,760),(560,760),MUTED,5)
    text_center(d,(555,680,690,740),"10 mm clear",font(20,True),MUTED)
    arrow(d,(800,760),(925,760),MUTED,5); arrow(d,(925,760),(800,760),MUTED,5)
    text_center(d,(795,680,930,740),"15 mm clear",font(20,True),MUTED)
    # airflow
    for y in (380,500,620): arrow(d,(210,y),(1510,y),CYAN,8)
    d.text((220,310),"RAM / FAN AIRFLOW",font=font(24,True),fill="#"+CYAN)
    # pusher fan upper band
    d.ellipse((315,245,565,495),outline="#"+GOLD,width=7); d.text((315,200),"PUSHER FAN ABOVE I/C",font=font(22,True),fill="#"+GOLD)
    im.save(ASSET/"side_stack.png",quality=95)

def save_flows():
    im=Image.new("RGB",(1800,1100),"white"); d=ImageDraw.Draw(im)
    d.text((70,50),"THREE SEPARATE CIRCUITS — NEVER PLUMB THEM TOGETHER",font=font(40,True),fill="#"+NAVY)
    sections=[
      (190,"ENGINE COOLANT",BLUE,["Engine hot outlet","Radiator inlet","Radiator core","Engine return"]),
      (505,"TURBO CHARGE AIR",GREEN,["Air filter","Turbo compressor","Intercooler","Intake plenum"]),
      (820,"R134a AIR CONDITIONING",GOLD,["Compressor","Condenser","Receiver-drier","TXV / evaporator","Compressor suction"])]
    for y,title,col,items in sections:
        d.text((80,y),title,font=font(27,True),fill="#"+col)
        n=len(items); start=80; gap=25; w=(1640-gap*(n-1))/n
        boxes=[]
        for i,item in enumerate(items):
            x=start+i*(w+gap); box=(x,y+55,x+w,y+190); rounded(d,box,"F7F9FA",col,16,4); text_center(d,box,item,font(23,True),col); boxes.append(box)
        for a,b in zip(boxes,boxes[1:]): arrow(d,(a[2]+5,(a[1]+a[3])/2),(b[0]-7,(b[1]+b[3])/2),col,7)
    d.text((80,1020),"Receiver-drier is A/C hardware mounted beside the condenser—not a fourth heat-exchanger core.",font=font(25,True),fill="#"+INK)
    im.save(ASSET/"system_flows.png",quality=95)

def save_dimensioned_front():
    im=Image.new("RGB",(1800,1320),"white"); d=ImageDraw.Draw(im)
    d.text((60,35),"DIMENSIONED FRONT ELEVATION — CORE BODIES",font=font(40,True),fill="#"+NAVY)
    d.text((60,85),"All values mm. Red = manufacture target. Purple M-dimensions = measure on vehicle before release.",font=font(23),fill="#"+MUTED)
    # Scale: 1 mm = 1.75 px; center x=900; common top baseline chosen only for envelope comparison.
    sx=1.75; cx=900
    rw,rh=530*sx,435*sx; rx1,ry1=cx-rw/2,355; rx2,ry2=cx+rw/2,355+rh
    cw,ch=559*sx,356*sx; cx1,cy1=cx-cw/2,410; cx2,cy2=cx+cw/2,410+ch
    iw,ih=500*sx,180*sx; ix1,iy1=cx-iw/2,760; ix2,iy2=cx+iw/2,760+ih
    # uprights / inside opening
    ux1=340; ux2=1460
    rounded(d,(260,300,340,1170),NAVY,NAVY,10); rounded(d,(1460,300,1540,1170),NAVY,NAVY,10)
    text_center(d,(260,545,340,930),"STRUCTURAL\nUPRIGHT",font(20,True),WHITE)
    text_center(d,(1460,545,1540,930),"STRUCTURAL\nUPRIGHT",font(20,True),WHITE)
    rounded(d,(rx1,ry1,rx2,ry2),LIGHT,BLUE,15,6); text_center(d,(rx1+80,ry1+70,rx2-80,ry1+220),"RADIATOR CORE\n530 W × 435 H × 64 D",font(29,True),BLUE)
    rounded(d,(cx1,cy1,cx2,cy2),"DDECF3",CYAN,15,6); text_center(d,(cx1+100,cy1+245,cx2-100,cy1+370),"A/C CONDENSER BODY\n559 W × 356 H × 21 D",font(27,True),NAVY)
    rounded(d,(ix1,iy1,ix2,iy2),"E4EEE8",GREEN,15,6); text_center(d,(ix1+70,iy1+45,ix2-70,iy2-45),"INTERCOOLER CORE\n500 W × 180 H × 60 D",font(28,True),GREEN)
    # width dims
    dim_h(d,rx1,rx2,285,ry1,"530 radiator core width")
    dim_h(d,cx1,cx2,220,cy1,"559 condenser body width")
    dim_h(d,ix1,ix2,1190,iy2,"500 intercooler core width")
    dim_h(d,ux1,ux2,1260,1170,"M1 clear inside width: measure top / middle / bottom",color="7B4AA8")
    # height dims
    dim_v(d,ry1,ry2,1660,rx2,"435 radiator core height")
    dim_v(d,cy1,cy2,1575,cx2,"356 condenser body height")
    dim_v(d,iy1,iy2,220,ix1,"180 I/C core height")
    # release notes
    rounded(d,(60,150,500,255),"F3ECF8","7B4AA8",12,3)
    text_center(d,(70,155,490,250),"M1 RELEASE\n≥540 for 530 radiator core\n≥569 for 559 condenser + 5/side",font(21,True),"7B4AA8")
    d.text((60,1280),"Mounting ears, tanks, manifolds and fittings must remain inside the verified M1 envelope or use forward adapter geometry without rubbing uprights.",font=font(20,True),fill="#"+INK)
    im.save(ASSET/"dimensioned_front_elevation.png",quality=95)

def save_dimensioned_side():
    im=Image.new("RGB",(1800,1240),"white"); d=ImageDraw.Draw(im)
    d.text((60,35),"DIMENSIONED SIDE SECTION — LOWER AIRFLOW BAND",font=font(40,True),fill="#"+NAVY)
    d.text((60,85),"All values mm. Component thicknesses + clear gaps form the controlling front-to-rear stack.",font=font(23),fill="#"+MUTED)
    y1,y2=350,930; x=220; scale=5.0
    parts=[("INTERCOOLER",60,GREEN,"E4EEE8"),("CLEAR",10,MUTED,"FFFFFF"),("CONDENSER",21,CYAN,"DDECF3"),("CLEAR",15,MUTED,"FFFFFF"),("RADIATOR",64,BLUE,LIGHT)]
    bounds=[]
    for name,w,col,fill in parts:
        x2=x+w*scale; bounds.append((name,x,x2,w,col))
        if name=="CLEAR":
            d.rectangle((x,y1,x2,y2),fill="white",outline="#"+MUTED,width=3)
            text_center(d,(x,y1+120,x2,y2-120),f"{w}\nCLEAR",font(20,True),MUTED)
        else:
            rounded(d,(x,y1,x2,y2),fill,col,10,5); text_center(d,(x+5,y1+80,x2-5,y2-80),f"{name}\n{w} DEEP",font(25,True),col)
        dim_h(d,x,x2,1010,y2,f"{w}")
        x=x2
    stack_start=bounds[0][1]; stack_end=bounds[-1][2]
    dim_h(d,stack_start,stack_end,1125,y2,"170 nominal core/gap stack")
    dim_h(d,stack_start,stack_end+100,1190,y2,"Allow 180–190 installed incl. seams/brackets",color="7B4AA8")
    # grille and fan clearance
    rounded(d,(80,300,145,980),"D8DEE3",MUTED,8); d.text((73,995),"GRILLE",font=font(20,True),fill="#"+MUTED)
    rx=stack_end
    rounded(d,(rx+80,390,rx+300,890),PALE,MUTED,12,4); text_center(d,(rx+90,500,rx+290,790),"SHROUD +\nENGINE FAN",font(23,True),INK)
    rounded(d,(rx+470,330,rx+600,950),"E8E3DD",INK,12,4); text_center(d,(rx+475,550,rx+595,760),"ENGINE",font(23,True),INK)
    dim_h(d,rx,rx+80,270,y1,"M5 ≥20 static; 25–30 preferred",color="7B4AA8")
    arrow(d,(155,500),(rx+50,500),CYAN,8); arrow(d,(155,680),(rx+50,680),CYAN,8)
    rounded(d,(70,140,710,250),"F3ECF8","7B4AA8",12,3)
    text_center(d,(80,145,700,245),"M4 RELEASE: usable lower-band depth ≥180\nplus 5–10 fabrication tolerance; stop if actual space <170",font(22,True),"7B4AA8")
    rounded(d,(970,140,1720,250),PALE,MUTED,12,3)
    text_center(d,(980,145,1710,245),"M3: measure grille-to-radiator plane separately\nat upper fan band and lower intercooler band",font(22,True),MUTED)
    im.save(ASSET/"dimensioned_side_section.png",quality=95)

def save_mounting_detail():
    im=Image.new("RGB",(1800,1280),"white"); d=ImageDraw.Draw(im)
    d.text((60,35),"DIMENSIONED MOUNTING DETAILS — REV B",font=font(40,True),fill="#"+NAVY)
    d.text((60,85),"All values mm. Verify handed geometry on the vehicle before cutting the mirrored upright.",font=font(23),fill="#"+MUTED)
    # Upright developed profile schematic
    d.text((80,155),"A — MIRRORED STRUCTURAL UPRIGHT / SADDLE",font=font(27,True),fill="#"+NAVY)
    x=260; top=245; bot=925
    d.line((x,top,x,bot),fill="#"+NAVY,width=34); d.line((x,top,x+210,top),fill="#"+NAVY,width=34)
    d.line((x,bot,x+250,bot),fill="#"+NAVY,width=34); d.line((x+250,bot,x+250,1110),fill="#"+NAVY,width=34)
    dim_v(d,top,bot,150,x,"410 upright height")
    dim_h(d,x,x+210,175,top,"58 top return")
    dim_h(d,x,x+250,985,bot,"70 chassis bridge")
    dim_v(d,bot,1110,590,x+250,"80 outer saddle leg")
    dim_h(d,x-17,x+17,1165,1110,"48 main face (section)")
    # Adapter tab
    d.text((780,155),"B — REMOVABLE ADAPTER TAB",font=font(27,True),fill="#"+NAVY)
    ax1,ay1,ax2,ay2=890,285,1240,775
    rounded(d,(ax1,ay1,ax2,ay2),LIGHT,BLUE,12,5)
    rounded(d,(1038,390,1092,670),WHITE,BLUE,26,4)
    dim_h(d,ax1,ax2,230,ay1,"50 wide")
    dim_v(d,ay1,ay2,1305,ax2,"70 high")
    dim_h(d,1038,1092,835,670,"9 slot width")
    dim_v(d,390,670,1370,1092,"20 slot length")
    d.text((830,880),"• 4 mm mild steel\n• M8 class 8.8 bolt\n• Slot vertical\n• Max simple offset M7 = 20\n• Box spacer above 20",font=font(23,True),fill="#"+INK,spacing=12)
    d.text((650,1080),"UPRIGHT MATERIAL: 4 mm mild-steel formed angle. Match existing; fabricate left piece as a handed mirror.",font=font(22,True),fill="#"+INK)
    im.save(ASSET/"dimensioned_mounting_details.png",quality=95)

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn("w:shd"))
    if shd is None: shd=OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"),fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar=OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m,v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=tcMar.find(qn("w:"+m))
        if node is None: node=OxmlElement("w:"+m); tcMar.append(node)
        node.set(qn("w:w"),str(v)); node.set(qn("w:type"),"dxa")

def set_table_widths(table, widths):
    table.autofit=False
    tblPr=table._tbl.tblPr
    tblW=tblPr.find(qn("w:tblW")); tblW.set(qn("w:w"),str(sum(widths))); tblW.set(qn("w:type"),"dxa")
    ind=OxmlElement("w:tblInd"); ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa"); tblPr.append(ind)
    grid=table._tbl.tblGrid
    for c in list(grid): grid.remove(c)
    for w in widths:
        c=OxmlElement("w:gridCol"); c.set(qn("w:w"),str(w)); grid.append(c)
    for row in table.rows:
        for cell,w in zip(row.cells,widths):
            tcW=cell._tc.get_or_add_tcPr().find(qn("w:tcW")); tcW.set(qn("w:w"),str(w)); tcW.set(qn("w:type"),"dxa")
            set_cell_margins(cell); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER

def set_run(run,size=11,bold=False,color=INK,italic=False):
    run.font.name="Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),"Calibri"); run._element.rPr.rFonts.set(qn("w:hAnsi"),"Calibri")
    run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.color.rgb=RGBColor.from_string(color)

def add_p(doc,text,style=None,bold_lead=None,after=6):
    p=doc.add_paragraph(style=style); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.2
    if bold_lead and text.startswith(bold_lead):
        r=p.add_run(bold_lead); set_run(r,bold=True); r=p.add_run(text[len(bold_lead):]); set_run(r)
    else: r=p.add_run(text); set_run(r)
    return p

def add_bullets(doc,items):
    for item in items: add_p(doc,item,"List Bullet",after=4)

def add_heading(doc,text,level=1):
    p=doc.add_paragraph(text,style=f"Heading {level}"); return p

def add_table(doc,headers,rows,widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.style="Table Grid"
    for i,h in enumerate(headers):
        set_cell_shading(t.rows[0].cells[i],LIGHT); p=t.rows[0].cells[i].paragraphs[0]; r=p.add_run(h); set_run(r,10,bold=True,color=NAVY); p.paragraph_format.space_after=Pt(0)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            p=cells[i].paragraphs[0]; r=p.add_run(str(val)); set_run(r,9.5); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.08
    set_table_widths(t,widths); return t

def page_field(p):
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p.add_run("J40 Cooling Pack  |  Rev B  |  "); set_run(r,9,color=MUTED)
    fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); p._p.append(fld)

def build():
    save_front(); save_side(); save_flows(); save_dimensioned_front(); save_dimensioned_side(); save_mounting_detail()
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Inches(.72); sec.bottom_margin=Inches(.72); sec.left_margin=Inches(.78); sec.right_margin=Inches(.78)
    sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)
    # styles: compact_reference_guide, explicit tokens
    st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(11); st.font.color.rgb=RGBColor.from_string(INK)
    st.paragraph_format.space_after=Pt(6); st.paragraph_format.line_spacing=1.25
    for n,size,before,after,col in ((1,16,18,10,BLUE),(2,13,14,7,BLUE),(3,12,10,5,NAVY)):
        s=doc.styles[f"Heading {n}"]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
    for name in ("List Bullet","List Number"):
        s=doc.styles[name]; s.font.name="Calibri"; s.font.size=Pt(11); s.paragraph_format.left_indent=Inches(.375); s.paragraph_format.first_line_indent=Inches(-.188); s.paragraph_format.space_after=Pt(4); s.paragraph_format.line_spacing=1.25
    hp=sec.header.paragraphs[0]; hp.text="J40 FRONT COOLING SYSTEM  •  FABRICATOR WORK DOCUMENT"; hp.alignment=WD_ALIGN_PARAGRAPH.LEFT; set_run(hp.runs[0],9,bold=True,color=MUTED)
    page_field(sec.footer.paragraphs[0])

    # cover
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(62); p.paragraph_format.space_after=Pt(12)
    r=p.add_run("FABRICATION & INSTALLATION GUIDE"); set_run(r,11,bold=True,color=GOLD)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8); r=p.add_run("J40 Integrated Cooling Pack"); set_run(r,28,bold=True,color=NAVY)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(26); r=p.add_run("Engine radiator • Turbo intercooler • A/C condenser • Receiver-drier • Fans"); set_run(r,14,color=MUTED)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(18); r=p.add_run("REV B  |  23 JULY 2026  |  MOCK-UP AND QUOTATION RELEASE"); set_run(r,10,bold=True,color=BLUE)
    add_table(doc,["STATUS","STRUCTURAL DECISION","FINAL-CORE RELEASE"],[
        ("Mock-up approved","Retain existing upright; weld a mirrored upright on the left","HOLD until M1–M7 are recorded")],[1700,4300,3360])
    doc.add_paragraph().paragraph_format.space_after=Pt(12)
    p=doc.add_paragraph(); set_cell=False
    r=p.add_run("DESIGN INTENT\n"); set_run(r,11,bold=True,color=NAVY)
    r=p.add_run("Build one modular front cooling package. Keep all pressure circuits separate and every component individually removable. No exchanger carries another exchanger’s weight."); set_run(r,12)
    p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(18); p.paragraph_format.line_spacing=1.25
    doc.add_picture(str(ASSET/"side_stack.png"),width=Inches(6.75))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc,"1. Assembly at a glance",1)
    add_p(doc,"The front pack has three heat-exchanger cores and two nearby A/C/fan auxiliaries. The receiver-drier is not a fourth radiator. The turbo itself remains on the engine; only its charge-air intercooler is in the front stack.")
    doc.add_picture(str(ASSET/"front_elevation.png"),width=Inches(6.75)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc,"What must be independently removable",2)
    add_bullets(doc,["Engine radiator, rear shroud and their rubber-isolated saddles/tabs.","A/C condenser, receiver-drier and slim pusher fan.","Turbo charge-air intercooler and its four isolated mounting tabs.","Side rails and any crossrails needed for component removal or service."])

    doc.add_page_break(); add_heading(doc,"2. Dimensioned fabrication sheets",1)
    add_p(doc,"Dimension authority: red dimensions are the Rev B manufacture targets. Purple M-dimensions are mandatory vehicle measurements and override assumptions. Do not scale the printed drawing; use the written dimensions.")
    doc.add_picture(str(ASSET/"dimensioned_front_elevation.png"),width=Inches(6.75)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc,["FRONT-ELEVATION RELEASE","REQUIREMENT"],[
        ("Radiator","530 × 435 core; overall width ≤ M1 − 10"),
        ("Condenser","559 × 356 body nominal; verify M1 ≥569 or use proven forward adapter geometry"),
        ("Intercooler","500 × 180 core; overall width ≤ M1 − 10"),
        ("Vertical placement","Site-fit from M2, M6, bonnet/latch and service-removal checks")],[2700,6660])

    doc.add_page_break(); add_heading(doc,"2A. Dimensioned side section",1)
    doc.add_picture(str(ASSET/"dimensioned_side_section.png"),width=Inches(6.75)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_p(doc,"Controlling arithmetic: 60 intercooler + 10 clear + 21 condenser + 15 clear + 64 radiator = 170 mm nominal. Allow 180–190 mm for seams, brackets and imperfect planes. The fan and shroud sit behind the radiator and are controlled separately by M5.")

    doc.add_page_break(); add_heading(doc,"2B. Dimensioned mounting details",1)
    doc.add_picture(str(ASSET/"dimensioned_mounting_details.png"),width=Inches(6.75)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc,["DETAIL","DIMENSION / MATERIAL"],[
        ("Removable side rail","30 × 3 flat OR 25 × 25 × 3 angle; cut length M2 − 10; keep 5 top and bottom clearance"),
        ("Radiator lower saddle","3–4 plate with 5 EPDM pad; upper tabs restrain only"),
        ("Condenser tabs","3 plate, rubber washers and M6 bolts"),
        ("Intercooler tabs","4 plate, M8 isolators; 2–3 thermal-movement slot at one upper point")],[2600,6760])
    add_p(doc,"The upright dimensions reproduce the measured existing-side bracket basis. Mirror and verify the left part against the chassis before welding. Slots belong in removable adapter tabs, never in structural uprights.")

    doc.add_page_break(); add_heading(doc,"3. Front-to-rear package",1)
    doc.add_picture(str(ASSET/"side_stack.png"),width=Inches(6.75)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_table(doc,["ORDER","COMPONENT","DESIGN ENVELOPE / RULE"],[
        ("1 — front upper","Slim A/C pusher fan","12 in nominal; ≤330 mm diameter × 65 mm deep; upper band"),
        ("1 — front lower","Charge-air intercooler","500 W × 180 H × 60 D mm; 57 mm OD outlets"),
        ("2","Clear air gap","10 mm minimum between intercooler and condenser"),
        ("3","R134a condenser","559 W × 356 H × 21 D mm nominal; ≤600 × 370 × 25"),
        ("4","Clear air gap","15 mm target; 10 mm absolute minimum to radiator"),
        ("5 — rear","Engine radiator","530 W × 435 H × 64 D mm core; high-efficiency four-row pattern"),
        ("6 — engine side","Rear shroud and mechanical fan","≥20 mm static tip clearance; 25–30 mm preferred")],[1450,2600,5310])

    doc.add_page_break(); add_heading(doc,"4. Separate circuits and connections",1)
    doc.add_picture(str(ASSET/"system_flows.png"),width=Inches(6.75)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc,"Interface holds",2)
    add_bullets(doc,["Copy radiator neck side, angle and centre position from the old radiator and installed hose mock-up; target neck OD is 38 mm.","Do not finalize intercooler outlet direction until turbo compressor outlet and intake-plenum inlet are marked on the vehicle.","Confirm condenser #8 inlet and #6 outlet threads before crimping hoses; mount the receiver-drier vertically after the condenser outlet.","Do not open the new receiver-drier to atmosphere until final A/C assembly and charging preparation."])

    doc.add_page_break(); add_heading(doc,"5. Structure and mounting",1)
    add_p(doc,"Fixed structural decision: retain the measured right-side formed upright and fabricate a handed/mirrored upright on the left. Dry-fit the complete pack before final welding or drilling.",bold_lead="Fixed structural decision:")
    add_table(doc,["ITEM","FABRICATION REQUIREMENT"],[
        ("Structural uprights","410 mm high; 48 mm main face; 4 mm mild steel; mirror existing 58 mm top return, 70 mm bridge and 80 mm saddle leg."),
        ("Removable side rails","30 × 3 mm flat or 25 × 25 × 3 mm angle; final length M2 − 10 mm."),
        ("Rail attachment","Two M8 class 8.8 bolts per upright; four total. Put 9 × 20 mm vertical slots in adapter tabs—not structural uprights."),
        ("Radiator","Weight on two 3–4 mm lower saddles with 5 mm EPDM pads. Upper tabs restrain only."),
        ("Condenser","Four independent 3 mm tabs, rubber washers and M6 bolts."),
        ("Intercooler","Four independent 4 mm tabs and M8 isolators; one upper point gets 2–3 mm horizontal thermal-movement slot."),
        ("Fan","Removable hoop/crossrail. Never use through-core plastic rods."),
        ("Receiver-drier","Rubber-lined removable vertical clamp outside the primary airflow.")],[2100,7260])
    add_heading(doc,"Absolute no-go conditions",2)
    add_bullets(doc,["No welding, drilling or hard clamping through tanks, tubes, cores or fins.","No radiator carrying the condenser, intercooler or fan.","No condenser carrying the fan or intercooler.","No forced bolts that twist a core or pull the radiator toward the engine fan.","No permanent crossrail across the radiator core face."])

    doc.add_page_break(); add_heading(doc,"6. Mandatory vehicle measurements — M1 to M7",1)
    add_p(doc,"Record these with the second upright tacked in place and the grille/front panel, bonnet latch, engine fan and body position represented. Final core manufacture remains on hold until all seven are complete.")
    rows=[
      ("M1","Minimum clear inside width","Top / middle / bottom; use smallest","≥540 for 530 core; ≥569 for unmodified 559 condenser + clearance"),
      ("M2","Clear vertical opening","Lower pad plane to bonnet/latch obstruction","Radiator overall height + 10 mm"),
      ("M3","Grille to radiator front plane","Upper fan band and lower I/C band","Must accept chosen front components"),
      ("M4","Usable lower-band stack depth","Grille/guard to radiator rear-component datum","≥180 mm plus 5–10 mm tolerance"),
      ("M5","Radiator rear face to fan tips","Closest blade through full rotation","≥20 mm; 25–30 preferred"),
      ("M6","Lowest safe intercooler edge","Relative to protected frame/bumper line","≥25 mm above lowest protected line"),
      ("M7","Upright face to side-rail offset","All four corners","0–20 simple tab; boxed spacer above 20")]
    add_table(doc,["ID","MEASUREMENT","METHOD","PASS / RELEASE CRITERION"],rows,[650,2400,2860,3450])
    add_heading(doc,"Also capture",2)
    add_bullets(doc,["Coolant-neck centres, OD and angles; cap, overflow and drain positions.","Condenser port side and actual threads.","Turbo compressor outlet and intake-plenum inlet locations.","Bonnet closure, steering/body clearance and service-removal path.","A ruler/tape visible in every measurement photograph."])

    doc.add_page_break(); add_heading(doc,"7. Fabricator sequence",1)
    steps=[
      "Duplicate the existing upright as a left-hand mirror; tack it and remeasure the opening.",
      "Make full-size radiator, condenser, intercooler and fan envelopes from cardboard or plywood.",
      "Set the radiator plane from rear fan clearance and coolant-hose sweep.",
      "Position the condenser with its clear air gap, then place the intercooler only in the lower frontal band.",
      "Place the pusher fan principally in the upper band and verify grille/bonnet-latch clearance.",
      "Mark removable rails, tabs, saddles and the vertical receiver-drier clamp from the real components.",
      "Confirm charge-pipe and A/C hose routes before outlet welding or hose crimping.",
      "Record M1–M7, photograph the mock-up, and obtain owner release for final cores.",
      "Bench pressure/flow-test the radiator; leak-test the intercooler at 20 psi; obtain A/C-shop condenser evidence.",
      "Final-fit, corrosion-protect, install and validate at warm idle with A/C on."]
    for s in steps: add_p(doc,s,"List Number",after=6)
    add_heading(doc,"Decision rules if depth is short",2)
    add_bullets(doc,["Protect radiator size and rear fan clearance first.","Reduce intercooler thickness from 60 to 50 mm only with a reputable high-efficiency core.","Move or split the pusher fan before reducing radiator area.","If lower-band usable depth is under 170 mm, stop and review a remote/water-to-air charge cooler."])

    doc.add_page_break(); add_heading(doc,"8. Inspection and handover record",1)
    checks=[
      ("☐","M1–M7 sheet complete with ruler photographs"),("☐","Front and side full-size mock-up photographs"),("☐","Grille, bonnet/latch and fan represented during tack-fit"),("☐","As-built sketch: cores, tanks, ports, holes and offsets"),("☐","Radiator pressure and flow result recorded"),("☐","Intercooler 20 psi leak-test result recorded"),("☐","Condenser/A/C pressure-test evidence recorded"),("☐","Bare-metal bracket and weld photographs"),("☐","No core carries another component"),("☐","Fan airflow direction confirmed with paper-strip test"),("☐","Warm-idle and A/C-on validation complete"),("☐","All exchangers removable without cutting welded parts")]
    add_table(doc,["CHECK","ACCEPTANCE ITEM"],checks,[900,8460])
    add_heading(doc,"Release note",2)
    add_p(doc,"This document releases mock-up, quotation and measurement capture. It does not release final core manufacture until M1–M7 and the actual component interfaces have been approved.")
    add_heading(doc,"Controlled project references",2)
    add_bullets(doc,["docs/j40-integrated-cooling-pack-fabricator-handoff-20260717.md","data/manual/fabrication/front_cooling_stack_rev_a/integrated_cooling_pack_dimensions_rev_b.csv","docs/engine-radiator-recore-release-20260529.md","docs/2h-turbo-suitability-and-options-20260717.md","docs/ac-hvac-workstream.md"])

    doc.core_properties.title="J40 Integrated Cooling Pack — Fabrication & Installation Guide — Rev B"
    doc.core_properties.subject="Radiator, turbo intercooler and A/C front cooling package"
    doc.core_properties.author="J40 Project"
    doc.save(OUT)
    print(OUT)

if __name__ == "__main__": build()
