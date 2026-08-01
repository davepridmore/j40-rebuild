from pathlib import Path
import math

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "J40-integrated-cooling-pack-fabricator-specification-rev-c.docx"
SOURCE = ROOT / "docs" / "J40-integrated-cooling-pack-fabricator-specification-rev-c.md"
ASSET = (
    ROOT
    / "data"
    / "manual"
    / "fabrication"
    / "front_cooling_stack_rev_c"
    / "work_document_assets"
)
ASSET.mkdir(parents=True, exist_ok=True)

NAVY = "17324D"
BLUE = "2E74B5"
CYAN = "4BA3C7"
GREEN = "4E7D61"
GOLD = "C9952E"
RED = "C9534B"
PURPLE = "7B4AA8"
INK = "1E2933"
MUTED = "62717D"
LIGHT = "E8EEF5"
PALE = "F5F7F9"
WHITE = "FFFFFF"
LINE = "B7C3CC"


def image_font(size, bold=False):
    paths = [
        (
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
            if bold
            else "/System/Library/Fonts/Supplemental/Arial.ttf"
        ),
        (
            "/Library/Fonts/Arial Bold.ttf"
            if bold
            else "/Library/Fonts/Arial.ttf"
        ),
    ]
    for path in paths:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def hex_colour(value):
    return "#" + value


def rounded(draw, box, fill, outline=LINE, radius=18, width=3):
    draw.rounded_rectangle(
        box,
        radius=radius,
        fill=hex_colour(fill),
        outline=hex_colour(outline),
        width=width,
    )


def centred_text(draw, box, text, font, fill=INK, spacing=5):
    x1, y1, x2, y2 = box
    bounds = draw.multiline_textbbox(
        (0, 0), text, font=font, align="center", spacing=spacing
    )
    width = bounds[2] - bounds[0]
    height = bounds[3] - bounds[1]
    draw.multiline_text(
        ((x1 + x2 - width) / 2, (y1 + y2 - height) / 2),
        text,
        font=font,
        fill=hex_colour(fill),
        align="center",
        spacing=spacing,
    )


def arrow(draw, start, end, colour=CYAN, width=8):
    draw.line((start, end), fill=hex_colour(colour), width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 20
    points = [
        end,
        (
            end[0] - size * math.cos(angle - 0.55),
            end[1] - size * math.sin(angle - 0.55),
        ),
        (
            end[0] - size * math.cos(angle + 0.55),
            end[1] - size * math.sin(angle + 0.55),
        ),
    ]
    draw.polygon(points, fill=hex_colour(colour))


def horizontal_dimension(draw, x1, x2, y, reference_y, label, colour=RED):
    draw.line((x1, reference_y, x1, y), fill=hex_colour(colour), width=3)
    draw.line((x2, reference_y, x2, y), fill=hex_colour(colour), width=3)
    draw.line((x1, y, x2, y), fill=hex_colour(colour), width=4)
    draw.polygon(
        [(x1, y), (x1 + 18, y - 9), (x1 + 18, y + 9)],
        fill=hex_colour(colour),
    )
    draw.polygon(
        [(x2, y), (x2 - 18, y - 9), (x2 - 18, y + 9)],
        fill=hex_colour(colour),
    )
    font = image_font(22, True)
    bounds = draw.textbbox((0, 0), label, font=font)
    text_width = bounds[2] - bounds[0]
    centre = x1 + (x2 - x1) / 2
    draw.rectangle(
        (centre - text_width / 2 - 10, y - 18, centre + text_width / 2 + 10, y + 18),
        fill="white",
    )
    draw.text(
        (centre - text_width / 2, y - 15),
        label,
        font=font,
        fill=hex_colour(colour),
    )


def vertical_dimension(draw, y1, y2, x, reference_x, label, colour=RED):
    draw.line((reference_x, y1, x, y1), fill=hex_colour(colour), width=3)
    draw.line((reference_x, y2, x, y2), fill=hex_colour(colour), width=3)
    draw.line((x, y1, x, y2), fill=hex_colour(colour), width=4)
    draw.polygon(
        [(x, y1), (x - 9, y1 + 18), (x + 9, y1 + 18)],
        fill=hex_colour(colour),
    )
    draw.polygon(
        [(x, y2), (x - 9, y2 - 18), (x + 9, y2 - 18)],
        fill=hex_colour(colour),
    )
    layer = Image.new("RGBA", (430, 44), (255, 255, 255, 235))
    layer_draw = ImageDraw.Draw(layer)
    layer_draw.text(
        (7, 4), label, font=image_font(21, True), fill=hex_colour(colour)
    )
    layer = layer.rotate(90, expand=True)
    draw._image.paste(
        layer,
        (int(x - 22), int((y1 + y2 - layer.height) / 2)),
        layer,
    )


def draw_fan(draw, box, label):
    x1, y1, x2, y2 = box
    draw.ellipse(box, fill="#F6EEDB", outline=hex_colour(GOLD), width=7)
    cx = (x1 + x2) / 2
    cy = (y1 + y2) / 2
    radius = min(x2 - x1, y2 - y1) * 0.35
    for angle in range(0, 360, 60):
        ex = cx + radius * math.cos(math.radians(angle))
        ey = cy + radius * math.sin(math.radians(angle))
        draw.line((cx, cy, ex, ey), fill=hex_colour(GOLD), width=28)
    hub = min(x2 - x1, y2 - y1) * 0.13
    draw.ellipse(
        (cx - hub, cy - hub, cx + hub, cy + hub),
        fill=hex_colour(GOLD),
    )
    centred_text(
        draw,
        (x1, y2 - 52, x2, y2 - 7),
        label,
        image_font(20, True),
        GOLD,
    )


def save_front_layout():
    image = Image.new("RGB", (1800, 1280), "white")
    draw = ImageDraw.Draw(image)
    draw.text(
        (55, 35),
        "REV C FRONT ELEVATION — DUAL-FAN / LOWER-INTERCOOLER LAYOUT",
        font=image_font(37, True),
        fill=hex_colour(NAVY),
    )
    draw.text(
        (55, 84),
        "All values mm. Outside fan envelope includes frames, tabs, plugs, guards and wire bends.",
        font=image_font(22),
        fill=hex_colour(MUTED),
    )

    scale = 1.58
    centre_x = 900
    radiator_width = 530 * scale
    radiator_height = 435 * scale
    radiator_left = centre_x - radiator_width / 2
    radiator_top = 270
    radiator_right = centre_x + radiator_width / 2
    radiator_bottom = radiator_top + radiator_height

    condenser_width = 559 * scale
    condenser_height = 356 * scale
    condenser_left = centre_x - condenser_width / 2
    condenser_top = radiator_top + 18
    condenser_right = centre_x + condenser_width / 2
    condenser_bottom = condenser_top + condenser_height

    intercooler_width = 500 * scale
    intercooler_height = 180 * scale
    intercooler_left = centre_x - intercooler_width / 2
    intercooler_bottom = radiator_bottom
    intercooler_top = intercooler_bottom - intercooler_height
    intercooler_right = centre_x + intercooler_width / 2

    fan_group_width = 515 * scale
    fan_group_height = 245 * scale
    fan_left = centre_x - fan_group_width / 2
    fan_top = radiator_top
    fan_right = centre_x + fan_group_width / 2
    fan_bottom = fan_top + fan_group_height

    upright_left = radiator_left - 110
    upright_right = radiator_right + 110
    rounded(
        draw,
        (upright_left - 45, 230, upright_left + 45, 1060),
        NAVY,
        NAVY,
        10,
        4,
    )
    rounded(
        draw,
        (upright_right - 45, 230, upright_right + 45, 1060),
        NAVY,
        NAVY,
        10,
        4,
    )
    centred_text(
        draw,
        (upright_left - 43, 475, upright_left + 43, 790),
        "RIGHT\nUPRIGHT",
        image_font(20, True),
        WHITE,
    )
    centred_text(
        draw,
        (upright_right - 43, 475, upright_right + 43, 790),
        "LEFT\nUPRIGHT",
        image_font(20, True),
        WHITE,
    )

    rounded(
        draw,
        (radiator_left, radiator_top, radiator_right, radiator_bottom),
        LIGHT,
        BLUE,
        18,
        6,
    )
    centred_text(
        draw,
        (
            radiator_left + 120,
            radiator_bottom - 115,
            radiator_right - 120,
            radiator_bottom - 25,
        ),
        "RADIATOR CORE 530 W × 435 H",
        image_font(26, True),
        BLUE,
    )

    rounded(
        draw,
        (condenser_left, condenser_top, condenser_right, condenser_bottom),
        "DDECF3",
        CYAN,
        16,
        5,
    )
    centred_text(
        draw,
        (
            condenser_left + 110,
            condenser_top + 25,
            condenser_right - 110,
            condenser_top + 85,
        ),
        "CONDENSER 559 W × 356 H",
        image_font(24, True),
        NAVY,
    )

    fan_gap = 12
    fan_box_width = (fan_group_width - fan_gap) / 2
    draw_fan(
        draw,
        (fan_left, fan_top, fan_left + fan_box_width, fan_bottom),
        "9 in PUSHER 1",
    )
    draw_fan(
        draw,
        (fan_left + fan_box_width + fan_gap, fan_top, fan_right, fan_bottom),
        "9 in PUSHER 2",
    )

    rounded(
        draw,
        (
            intercooler_left,
            intercooler_top,
            intercooler_right,
            intercooler_bottom,
        ),
        "E4EEE8",
        GREEN,
        16,
        6,
    )
    centred_text(
        draw,
        (
            intercooler_left + 80,
            intercooler_top + 55,
            intercooler_right - 80,
            intercooler_bottom - 55,
        ),
        "INTERCOOLER CORE\n500 W × 180 H × 60 D",
        image_font(26, True),
        GREEN,
    )

    horizontal_dimension(
        draw,
        radiator_left,
        radiator_right,
        205,
        radiator_top,
        "530 radiator core",
    )
    horizontal_dimension(
        draw,
        condenser_left,
        condenser_right,
        145,
        condenser_top,
        "559 condenser body",
        CYAN,
    )
    horizontal_dimension(
        draw,
        fan_left,
        fan_right,
        1110,
        fan_bottom,
        "515 released max / 520 mock-up only",
        GOLD,
    )
    horizontal_dimension(
        draw,
        intercooler_left,
        intercooler_right,
        1175,
        intercooler_bottom,
        "500 intercooler core",
        GREEN,
    )
    vertical_dimension(
        draw,
        radiator_top,
        radiator_bottom,
        1640,
        radiator_right,
        "435 radiator core height",
    )
    vertical_dimension(
        draw,
        fan_top,
        fan_bottom,
        1515,
        fan_right,
        "245 released max / 250 mock-up only",
        GOLD,
    )
    vertical_dimension(
        draw,
        intercooler_top,
        intercooler_bottom,
        215,
        intercooler_left,
        "180 intercooler core",
        GREEN,
    )

    rounded(draw, (55, 1060, 560, 1240), "F3ECF8", PURPLE, 14, 4)
    centred_text(
        draw,
        (65, 1070, 550, 1230),
        "M1 WIDTH GATE\n≥540 for radiator core\n≥569 for condenser + 5/side\nactual tanks, ports and ears must fit",
        image_font(20, True),
        PURPLE,
    )
    rounded(draw, (1230, 1070, 1740, 1230), "FDF2F1", RED, 14, 4)
    centred_text(
        draw,
        (1240, 1080, 1730, 1220),
        "NO OVERLAP\nFan frame, plug, guard and wire\nmust remain above the complete\nintercooler envelope.",
        image_font(20, True),
        RED,
    )

    image.save(ASSET / "rev_c_front_layout.png", quality=95)


def side_band(
    draw,
    top,
    title,
    components,
    total_label,
    gate_label,
    gate_colour,
):
    draw.text((70, top), title, font=image_font(27, True), fill=hex_colour(NAVY))
    y1 = top + 75
    y2 = y1 + 285
    x = 230
    scale = 5.0
    start = x
    for name, depth, colour, fill in components:
        x2 = x + depth * scale
        if name == "CLEAR":
            draw.rectangle(
                (x, y1, x2, y2),
                fill="white",
                outline=hex_colour(MUTED),
                width=3,
            )
            centred_text(
                draw,
                (x, y1 + 20, x2, y2 - 20),
                f"{depth}\nCLEAR",
                image_font(18, True),
                MUTED,
            )
        else:
            rounded(draw, (x, y1, x2, y2), fill, colour, 9, 4)
            centred_text(
                draw,
                (x + 4, y1 + 25, x2 - 4, y2 - 25),
                f"{name}\n{depth} D",
                image_font(21, True),
                colour,
            )
        horizontal_dimension(draw, x, x2, y2 + 55, y2, str(depth), colour)
        x = x2
    horizontal_dimension(
        draw, start, x, y2 + 125, y2, total_label, gate_colour
    )
    rounded(draw, (1230, y1 + 15, 1725, y2 - 15), PALE, gate_colour, 14, 4)
    centred_text(
        draw,
        (1240, y1 + 25, 1715, y2 - 25),
        gate_label,
        image_font(21, True),
        gate_colour,
    )
    arrow(draw, (80, y1 + 90), (215, y1 + 90), CYAN, 7)
    arrow(draw, (80, y1 + 190), (215, y1 + 190), CYAN, 7)
    draw.text(
        (75, y1 + 230),
        "GRILLE",
        font=image_font(18, True),
        fill=hex_colour(MUTED),
    )


def save_side_geometry():
    image = Image.new("RGB", (1800, 1330), "white")
    draw = ImageDraw.Draw(image)
    draw.text(
        (55, 35),
        "REV C SIDE GEOMETRY — UPPER AND LOWER AIRFLOW BANDS",
        font=image_font(38, True),
        fill=hex_colour(NAVY),
    )
    draw.text(
        (55, 85),
        "Front/grille is left. Engine is right. Written dimensions control; do not scale this sheet.",
        font=image_font(22),
        fill=hex_colour(MUTED),
    )

    side_band(
        draw,
        150,
        "A — UPPER BAND: ELECTRIC PUSHERS + CONDENSER + RADIATOR",
        [
            ("2-FAN\nASSEMBLY", 55, GOLD, "F6EEDB"),
            ("CLEAR", 5, MUTED, WHITE),
            ("CONDENSER", 21, CYAN, "DDECF3"),
            ("CLEAR", 15, MUTED, WHITE),
            ("RADIATOR", 64, BLUE, LIGHT),
        ],
        "160 component/gap depth incl. radiator",
        "M3-U: radiator FRONT face to\nclosest fixed obstruction.\n55 + 5 + 21 + 15 + 5 build\n= 101 minimum.",
        PURPLE,
    )

    side_band(
        draw,
        735,
        "B — LOWER BAND: INTERCOOLER + CONDENSER + RADIATOR",
        [
            ("INTERCOOLER", 60, GREEN, "E4EEE8"),
            ("CLEAR", 10, MUTED, WHITE),
            ("CONDENSER", 21, CYAN, "DDECF3"),
            ("CLEAR", 15, MUTED, WHITE),
            ("RADIATOR", 64, BLUE, LIGHT),
        ],
        "170 nominal lower stack",
        "M3-L: radiator FRONT face to\nclosest fixed obstruction:\n106 + 10 build = 116 minimum.\nM4 ≥180; 190 preferred.\n170–179 not released; <170 STOP.",
        RED,
    )

    rounded(draw, (65, 1250, 1730, 1310), "FDF2F1", RED, 12, 3)
    centred_text(
        draw,
        (75, 1255, 1720, 1305),
        "Behind radiator: full close shroud + original engine-driven puller fan. M5 ≥20 static; 25–30 preferred.",
        image_font(21, True),
        RED,
    )
    image.save(ASSET / "rev_c_side_geometry.png", quality=95)


def save_fan_wiring():
    image = Image.new("RGB", (1800, 1020), "white")
    draw = ImageDraw.Draw(image)
    draw.text(
        (55, 35),
        "DUAL CONDENSER FAN — MOUNTING AND WIRING CONCEPT",
        font=image_font(38, True),
        fill=hex_colour(NAVY),
    )
    draw.text(
        (55, 85),
        "Final fuse values follow measured run/start current and wire capacity.",
        font=image_font(22),
        fill=hex_colour(MUTED),
    )

    rounded(draw, (80, 230, 300, 430), LIGHT, BLUE, 16, 4)
    centred_text(
        draw,
        (90, 240, 290, 420),
        "BATTERY /\nCHARGING\nSYSTEM\n12 V",
        image_font(25, True),
        BLUE,
    )
    rounded(draw, (390, 170, 610, 310), "FDF2F1", RED, 14, 4)
    rounded(draw, (390, 350, 610, 490), "FDF2F1", RED, 14, 4)
    centred_text(
        draw,
        (400, 180, 600, 300),
        "FUSE 1\nfrom measured\nFan 1 current",
        image_font(21, True),
        RED,
    )
    centred_text(
        draw,
        (400, 360, 600, 480),
        "FUSE 2\nfrom measured\nFan 2 current",
        image_font(21, True),
        RED,
    )
    rounded(draw, (710, 170, 930, 310), PALE, GOLD, 14, 4)
    rounded(draw, (710, 350, 930, 490), PALE, GOLD, 14, 4)
    centred_text(
        draw,
        (720, 180, 920, 300),
        "SEALED\nRELAY 1",
        image_font(24, True),
        GOLD,
    )
    centred_text(
        draw,
        (720, 360, 920, 480),
        "SEALED\nRELAY 2",
        image_font(24, True),
        GOLD,
    )
    rounded(draw, (1040, 170, 1280, 310), "F6EEDB", GOLD, 14, 4)
    rounded(draw, (1040, 350, 1280, 490), "F6EEDB", GOLD, 14, 4)
    centred_text(
        draw,
        (1050, 180, 1270, 300),
        "9 in\nPUSHER 1",
        image_font(24, True),
        GOLD,
    )
    centred_text(
        draw,
        (1050, 360, 1270, 480),
        "9 in\nPUSHER 2",
        image_font(24, True),
        GOLD,
    )
    rounded(draw, (1410, 245, 1705, 425), LIGHT, GREEN, 14, 4)
    centred_text(
        draw,
        (1420, 255, 1695, 415),
        "CLEAN GROUND STUD\nsame-capacity ground\nNo sheet-metal screw",
        image_font(21, True),
        GREEN,
    )

    for y in (240, 420):
        arrow(draw, (305, 330), (380, y), RED, 6)
        arrow(draw, (615, y), (700, y), RED, 6)
        arrow(draw, (935, y), (1030, y), GOLD, 6)
        arrow(draw, (1285, y), (1400, 330), GREEN, 6)

    rounded(draw, (590, 600, 1110, 770), "F3ECF8", PURPLE, 14, 4)
    centred_text(
        draw,
        (600, 610, 1100, 760),
        "COMMON LOW-CURRENT TRIGGER\nBoth fans ON with A/C clutch command.\nTrinary also protects compressor and\nmay request fans at high pressure.",
        image_font(22, True),
        PURPLE,
    )
    arrow(draw, (720, 600), (820, 500), PURPLE, 6)
    arrow(draw, (980, 600), (820, 500), PURPLE, 6)

    rounded(draw, (80, 845, 1720, 970), PALE, MUTED, 14, 3)
    centred_text(
        draw,
        (95, 855, 1705, 960),
        "6 mm² common feed/ground unless testing proves 4 mm² adequate; ≥4 mm² each branch.\n"
        "Weatherproof holders/connectors. Loaded motor voltage: within 0.5 V of battery.",
        image_font(19, True),
        INK,
    )
    image.save(ASSET / "rev_c_fan_wiring.png", quality=95)


SHEET_W = 2400
SHEET_H = 1700
SHEET_FOOTER_Y = 1550


def wrap_drawing_text(draw, text, font, max_width):
    lines = []
    for paragraph in str(text).split("\n"):
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        line = words[0]
        for word in words[1:]:
            candidate = line + " " + word
            if draw.textlength(candidate, font=font) <= max_width:
                line = candidate
            else:
                lines.append(line)
                line = word
        lines.append(line)
    return "\n".join(lines)


def dashed_line(draw, points, fill, width=4, dash=18, gap=10):
    (x1, y1), (x2, y2) = points
    length = math.hypot(x2 - x1, y2 - y1)
    if length == 0:
        return
    ux = (x2 - x1) / length
    uy = (y2 - y1) / length
    distance = 0
    while distance < length:
        end = min(distance + dash, length)
        draw.line(
            (
                x1 + ux * distance,
                y1 + uy * distance,
                x1 + ux * end,
                y1 + uy * end,
            ),
            fill=hex_colour(fill),
            width=width,
        )
        distance += dash + gap


def dashed_rectangle(draw, box, colour=PURPLE, width=4, dash=18, gap=10):
    x1, y1, x2, y2 = box
    dashed_line(draw, ((x1, y1), (x2, y1)), colour, width, dash, gap)
    dashed_line(draw, ((x2, y1), (x2, y2)), colour, width, dash, gap)
    dashed_line(draw, ((x2, y2), (x1, y2)), colour, width, dash, gap)
    dashed_line(draw, ((x1, y2), (x1, y1)), colour, width, dash, gap)


def drawing_panel(draw, box, title, body="", fill=PALE, outline=LINE, title_colour=NAVY):
    rounded(draw, box, fill, outline, 12, 3)
    x1, y1, x2, _ = box
    draw.text(
        (x1 + 18, y1 + 14),
        title,
        font=image_font(25, True),
        fill=hex_colour(title_colour),
    )
    if body:
        font = image_font(20)
        wrapped = wrap_drawing_text(draw, body, font, x2 - x1 - 36)
        draw.multiline_text(
            (x1 + 18, y1 + 51),
            wrapped,
            font=font,
            fill=hex_colour(INK),
            spacing=5,
        )


def drawing_note(draw, box, title, body, colour=PURPLE, fill="F3ECF8"):
    rounded(draw, box, fill, colour, 12, 4)
    x1, y1, x2, _ = box
    draw.text(
        (x1 + 16, y1 + 12),
        title,
        font=image_font(22, True),
        fill=hex_colour(colour),
    )
    font = image_font(18)
    wrapped = wrap_drawing_text(draw, body, font, x2 - x1 - 32)
    draw.multiline_text(
        (x1 + 16, y1 + 46),
        wrapped,
        font=font,
        fill=hex_colour(INK),
        spacing=4,
    )


def dimension_h(draw, x1, x2, y, ref_y, label, colour=RED, font_size=22):
    draw.line((x1, ref_y, x1, y), fill=hex_colour(colour), width=3)
    draw.line((x2, ref_y, x2, y), fill=hex_colour(colour), width=3)
    draw.line((x1, y, x2, y), fill=hex_colour(colour), width=4)
    draw.polygon(
        [(x1, y), (x1 + 20, y - 10), (x1 + 20, y + 10)],
        fill=hex_colour(colour),
    )
    draw.polygon(
        [(x2, y), (x2 - 20, y - 10), (x2 - 20, y + 10)],
        fill=hex_colour(colour),
    )
    font = image_font(font_size, True)
    bounds = draw.textbbox((0, 0), label, font=font)
    width = bounds[2] - bounds[0]
    centre = (x1 + x2) / 2
    draw.rounded_rectangle(
        (centre - width / 2 - 12, y - font_size, centre + width / 2 + 12, y + font_size),
        radius=7,
        fill="white",
    )
    draw.text(
        (centre - width / 2, y - font_size + 2),
        label,
        font=font,
        fill=hex_colour(colour),
    )


def dimension_v(draw, y1, y2, x, ref_x, label, colour=RED, label_side="right"):
    draw.line((ref_x, y1, x, y1), fill=hex_colour(colour), width=3)
    draw.line((ref_x, y2, x, y2), fill=hex_colour(colour), width=3)
    draw.line((x, y1, x, y2), fill=hex_colour(colour), width=4)
    draw.polygon(
        [(x, y1), (x - 10, y1 + 20), (x + 10, y1 + 20)],
        fill=hex_colour(colour),
    )
    draw.polygon(
        [(x, y2), (x - 10, y2 - 20), (x + 10, y2 - 20)],
        fill=hex_colour(colour),
    )
    font = image_font(21, True)
    wrapped = wrap_drawing_text(draw, label, font, 245)
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=3)
    width = bbox[2] - bbox[0]
    height = bbox[3] - bbox[1]
    tx = x + 16 if label_side == "right" else x - width - 18
    ty = (y1 + y2 - height) / 2
    draw.rounded_rectangle(
        (tx - 7, ty - 6, tx + width + 7, ty + height + 6),
        radius=6,
        fill="white",
    )
    draw.multiline_text(
        (tx, ty),
        wrapped,
        font=font,
        fill=hex_colour(colour),
        spacing=3,
    )


def leader_note(draw, anchor, box, title, body, colour=PURPLE):
    x1, y1, x2, y2 = box
    edge = (x1, (y1 + y2) / 2) if anchor[0] < x1 else (x2, (y1 + y2) / 2)
    draw.line((anchor, edge), fill=hex_colour(colour), width=4)
    draw.ellipse(
        (anchor[0] - 7, anchor[1] - 7, anchor[0] + 7, anchor[1] + 7),
        fill=hex_colour(colour),
    )
    drawing_note(draw, box, title, body, colour)


def numbered_bubble(draw, centre, number, colour=PURPLE):
    cx, cy = centre
    draw.ellipse(
        (cx - 23, cy - 23, cx + 23, cy + 23),
        fill=hex_colour(colour),
        outline="white",
        width=3,
    )
    label = str(number)
    font = image_font(23, True)
    bounds = draw.textbbox((0, 0), label, font=font)
    draw.text(
        (cx - (bounds[2] - bounds[0]) / 2, cy - (bounds[3] - bounds[1]) / 2 - 2),
        label,
        font=font,
        fill="white",
    )


def new_drawing_sheet(code, title, subtitle):
    image = Image.new("RGB", (SHEET_W, SHEET_H), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((24, 24, SHEET_W - 24, SHEET_H - 24), outline=hex_colour(NAVY), width=5)
    draw.rectangle((24, 24, SHEET_W - 24, 158), fill=hex_colour(NAVY))
    draw.text(
        (52, 43),
        f"{code}  |  {title}",
        font=image_font(36, True),
        fill="white",
    )
    draw.text(
        (54, 101),
        subtitle,
        font=image_font(21),
        fill="#DDE8F2",
    )
    rounded(draw, (1980, 48, 2342, 135), "FDF2F1", RED, 12, 3)
    centred_text(
        draw,
        (1990, 53, 2332, 130),
        "FINAL CORE\nMANUFACTURE: HOLD",
        image_font(20, True),
        RED,
    )
    return image, draw


def finish_drawing_sheet(draw, code, description):
    draw.line((24, SHEET_FOOTER_Y, SHEET_W - 24, SHEET_FOOTER_Y), fill=hex_colour(NAVY), width=4)
    draw.rectangle((24, SHEET_FOOTER_Y, 420, SHEET_H - 24), fill="#FDF2F1")
    draw.rectangle((420, SHEET_FOOTER_Y, 820, SHEET_H - 24), fill="#F3ECF8")
    draw.text((45, 1571), "RED = FIXED / RELEASE LIMIT", font=image_font(18, True), fill=hex_colour(RED))
    draw.text((441, 1571), "PURPLE = FIELD MEASURE", font=image_font(18, True), fill=hex_colour(PURPLE))
    draw.text((45, 1608), "ALL DIMENSIONS mm", font=image_font(18, True), fill=hex_colour(NAVY))
    draw.text((441, 1608), "DO NOT SCALE DRAWING", font=image_font(18, True), fill=hex_colour(NAVY))
    draw.text(
        (850, 1570),
        description,
        font=image_font(19, True),
        fill=hex_colour(INK),
    )
    draw.text(
        (850, 1610),
        "Unlabelled geometry and every tank / neck / port / bracket position: copy actual parts and measure on vehicle.",
        font=image_font(17),
        fill=hex_colour(MUTED),
    )
    draw.text((2050, 1570), "REV C.1", font=image_font(21, True), fill=hex_colour(BLUE))
    draw.text((2050, 1608), "30 JUL 2026", font=image_font(18, True), fill=hex_colour(INK))
    draw.text((2250, 1570), code, font=image_font(27, True), fill=hex_colour(NAVY))


def save_d01_complete_stack():
    image, draw = new_drawing_sheet(
        "D01",
        "COMPLETE COOLING PACK — FRONT ELEVATION",
        "Nominal component bodies and released dual-fan envelope; physical tanks, ears, plugs and obstructions still control fit.",
    )
    cx = 1210
    scale = 1.95
    rad = (cx - 530 * scale / 2, 335, cx + 530 * scale / 2, 335 + 435 * scale)
    cond = (cx - 559 * scale / 2, 370, cx + 559 * scale / 2, 370 + 356 * scale)
    ic = (cx - 500 * scale / 2, rad[3] - 180 * scale, cx + 500 * scale / 2, rad[3])
    fans = (cx - 515 * scale / 2, rad[1], cx + 515 * scale / 2, rad[1] + 245 * scale)
    upright_l = rad[0] - 125
    upright_r = rad[2] + 125

    rounded(draw, (upright_l - 34, 270, upright_l + 34, 1305), NAVY, NAVY, 8, 3)
    rounded(draw, (upright_r - 34, 270, upright_r + 34, 1305), NAVY, NAVY, 8, 3)
    draw.text((upright_l - 26, 1230), "R", font=image_font(23, True), fill="white")
    draw.text((upright_r - 26, 1230), "L", font=image_font(23, True), fill="white")

    rounded(draw, rad, LIGHT, BLUE, 15, 6)
    for y in range(int(rad[1] + 34), int(rad[3] - 30), 23):
        draw.line((rad[0] + 20, y, rad[2] - 20, y), fill="#B9CEDF", width=2)
    centred_text(draw, (rad[0] + 40, rad[3] - 88, rad[2] - 40, rad[3] - 24), "RADIATOR ACTIVE-CORE BASIS 530 × 435 × 64 — COMPLETE R0 TBD", image_font(21, True), BLUE)

    draw.rectangle(cond, fill="#DDECF3", outline=hex_colour(CYAN), width=6)
    for x in range(int(cond[0] + 24), int(cond[2] - 20), 25):
        draw.line((x, cond[1] + 20, x, cond[3] - 20), fill="#A8D0E2", width=2)
    draw.text((cond[0] + 25, cond[1] + 20), "CONDENSER BODY 559 × 356 × 21", font=image_font(24, True), fill=hex_colour(NAVY))

    rounded(draw, ic, "E4EEE8", GREEN, 13, 6)
    for y in range(int(ic[1] + 22), int(ic[3] - 16), 18):
        draw.line((ic[0] + 18, y, ic[2] - 18, y), fill="#BAD4C2", width=2)
    centred_text(draw, (ic[0] + 20, ic[1] + 80, ic[2] - 20, ic[3] - 50), "INTERCOOLER CORE\n500 × 180 × 60\n57 OD BEADED OUTLETS", image_font(24, True), GREEN)

    gap = 14
    fan_width = (fans[2] - fans[0] - gap) / 2
    draw_fan(draw, (fans[0], fans[1], fans[0] + fan_width, fans[3]), "9 in PUSHER 1")
    draw_fan(draw, (fans[0] + fan_width + gap, fans[1], fans[2], fans[3]), "9 in PUSHER 2")

    dimension_h(draw, rad[0], rad[2], 294, rad[1], "530 RADIATOR CORE", RED)
    dimension_h(draw, cond[0], cond[2], 228, cond[1], "559 CONDENSER BODY", RED)
    dimension_h(draw, fans[0], fans[2], 1278, fans[3], "≤515 COMPLETE DUAL-FAN ENVELOPE", RED)
    dimension_h(draw, ic[0], ic[2], 1342, ic[3], "500 INTERCOOLER CORE", RED)
    dimension_v(draw, rad[1], rad[3], 2185, rad[2], "435 RADIATOR CORE", RED, "left")
    dimension_v(draw, fans[1], fans[3], 2070, fans[2], "≤245 FAN PACKAGE", RED, "right")
    dimension_v(draw, ic[1], ic[3], 245, ic[0], "180 I/C CORE", RED, "left")

    drawing_note(
        draw,
        (55, 205, 425, 455),
        "M1 WIDTH / CHORAI",
        "Measure clear opening top, middle and bottom. ≥540 for 530 core. ≥569 for condenser plus 5 each side. Complete tanks, ears and ports must also fit.",
        PURPLE,
    )
    drawing_note(
        draw,
        (55, 980, 425, 1248),
        "M2 + LOWER PROTECTION",
        "M2 = actual radiator overall height + 10 to bonnet/latch. M6 = intercooler lower edge ≥25 above protected frame/bumper line.",
        PURPLE,
    )
    drawing_note(
        draw,
        (1975, 1070, 2340, 1348),
        "NO OVERLAP / NO RUB",
        "Fan frame, guard, plug and wire bend stay above the complete intercooler envelope. F5 ≥5 under vibration and load; 10 target.",
        RED,
        "FDF2F1",
    )
    finish_drawing_sheet(draw, "D01", "COMPLETE STACK FRONT ELEVATION")
    image.save(ASSET / "rev_c_d01_complete_stack.png", quality=96)


def save_d02_radiator_assembly():
    image, draw = new_drawing_sheet(
        "D02",
        "RADIATOR ASSEMBLY — ORTHOGRAPHIC & CONTROLLED INTERFACES",
        "Core size is fixed. Dashed tanks, necks, cap, drain and brackets are placeholders only: copy the original radiator.",
    )
    drawing_panel(draw, (45, 185, 1420, 1280), "FRONT ELEVATION — CORE SIZE CONTROLS", fill="F8FAFC")
    core = (245, 350, 1160, 1100)
    rounded(draw, core, LIGHT, BLUE, 12, 6)
    for y in range(380, 1080, 22):
        draw.line((265, y, 1140, y), fill="#B9CEDF", width=2)
    dashed_rectangle(draw, (190, 300, 1215, 1150), PURPLE, 5)
    draw.rounded_rectangle((190, 300, 1215, 350), radius=10, outline=hex_colour(PURPLE), width=5)
    draw.rounded_rectangle((190, 1100, 1215, 1150), radius=10, outline=hex_colour(PURPLE), width=5)
    draw.ellipse((310, 265, 415, 345), outline=hex_colour(PURPLE), width=5)
    draw.line((350, 265, 350, 230), fill=hex_colour(PURPLE), width=5)
    draw.ellipse((1040, 1135, 1105, 1175), outline=hex_colour(PURPLE), width=5)
    dashed_rectangle(draw, (135, 425, 190, 520), PURPLE, 4)
    dashed_rectangle(draw, (1215, 895, 1270, 990), PURPLE, 4)
    dimension_h(draw, core[0], core[2], 285, core[1], "530 CORE WIDTH", RED)
    dimension_v(draw, core[1], core[3], 1295, core[2], "435 CORE HEIGHT", RED, "left")
    centred_text(draw, (core[0] + 80, 640, core[2] - 80, 810), "4-ROW HIGH-EFFICIENCY\nCOPPER / BRASS CORE\n530 W × 435 H × 64 D", image_font(31, True), BLUE)
    leader_note(draw, (350, 285), (50, 1190, 510, 1495), "R01 FILLER / CAP", "Copy cap-seat type, cap rating, filler height and overflow from serviceable original. Do not fit a random higher-pressure cap.", PURPLE)
    leader_note(draw, (1242, 940), (535, 1190, 990, 1495), "R02 MOUNTING", "Copy actual mounting plane, bracket centres, hole sizes and isolator stack. Weight sits on two lower saddles; upper tabs restrain only.", PURPLE)
    leader_note(draw, (1070, 1150), (1015, 1190, 1420, 1495), "R03 DRAIN", "Accessible low point. Position and thread copied from sample or approved service fitting.", PURPLE)

    drawing_panel(draw, (1460, 185, 2350, 740), "SIDE ELEVATION — DEPTH", fill="F8FAFC")
    sx, sy = 1710, 325
    draw.rectangle((sx, sy, sx + 210, sy + 315), fill=hex_colour(LIGHT), outline=hex_colour(BLUE), width=6)
    dashed_rectangle(draw, (sx - 45, sy - 35, sx + 255, sy + 350), PURPLE, 5)
    dimension_h(draw, sx, sx + 210, 675, sy + 315, "64 CORE DEPTH", RED)
    draw.text((1510, 235), "FRONT", font=image_font(21, True), fill=hex_colour(CYAN))
    arrow(draw, (1560, 310), (1680, 310), CYAN, 7)
    leader_note(draw, (sx - 20, sy + 55), (1955, 300, 2315, 500), "R04 NECKS", "38 OD target only. Side, angle, centres, insertion and bead are copied from sample/hose mock-up.", PURPLE)
    leader_note(draw, (sx + 225, sy + 255), (1955, 515, 2315, 700), "R05 TANK DEPTH", "Record complete tank, seam, bracket and hose-bend envelope. Core depth alone is not the vehicle-fit envelope.", PURPLE)

    drawing_panel(draw, (1460, 775, 2350, 1495), "TOP VIEW + INTERFACE REGISTER", fill="F8FAFC")
    draw.rectangle((1610, 865, 2170, 1035), fill=hex_colour(LIGHT), outline=hex_colour(BLUE), width=6)
    dashed_rectangle(draw, (1560, 825, 2220, 1075), PURPLE, 5)
    dimension_v(draw, 865, 1035, 2270, 2170, "64 CORE", RED, "left")
    rows = [
        ("R06", "TOP/BOTTOM TANK SHAPE", "COPY ORIGINAL"),
        ("R07", "INLET / OUTLET SIDE + ANGLE", "FIELD MEASURE"),
        ("R08", "OVERFLOW / DRAIN / BRACKETS", "FIELD MEASURE"),
        ("M5", "REAR FACE TO FAN", "≥20; 25–30 PREF"),
    ]
    y = 1125
    for code, item, control in rows:
        draw.line((1510, y, 2310, y), fill=hex_colour(LINE), width=2)
        draw.text((1520, y + 10), code, font=image_font(19, True), fill=hex_colour(PURPLE if code.startswith("R") else RED))
        draw.text((1605, y + 10), item, font=image_font(17, True), fill=hex_colour(INK))
        draw.text((2055, y + 10), control, font=image_font(17, True), fill=hex_colour(PURPLE if "FIELD" in control or "COPY" in control else RED))
        y += 78
    finish_drawing_sheet(draw, "D02", "RADIATOR ORTHOGRAPHIC / INTERFACE CONTROL")
    image.save(ASSET / "rev_c_d02_radiator_assembly.png", quality=96)


def save_d03_radiator_components():
    image, draw = new_drawing_sheet(
        "D03",
        "RADIATOR — EXPLODED COMPONENT & FABRICATION PARTS",
        "Schematic assembly order only. Drawn neck, port and bracket positions are not dimensions and must not be copied.",
    )
    drawing_panel(draw, (45, 185, 1665, 1495), "EXPLODED ASSEMBLY — NOT TO SCALE", fill="F8FAFC")
    cx = 830
    top_y = 250
    # Top tank and fittings.
    dashed_rectangle(draw, (500, top_y, 1160, top_y + 105), PURPLE, 5)
    draw.arc((560, top_y - 20, 690, top_y + 90), 180, 360, fill=hex_colour(PURPLE), width=5)
    draw.ellipse((805, top_y - 45, 900, top_y + 25), outline=hex_colour(PURPLE), width=5)
    numbered_bubble(draw, (465, top_y + 35), 1)
    numbered_bubble(draw, (855, top_y - 35), 2)
    # Header plate.
    draw.rectangle((450, 405, 1210, 450), fill="#E1E7EC", outline=hex_colour(NAVY), width=4)
    numbered_bubble(draw, (420, 425), 3, NAVY)
    # Core.
    core = (485, 535, 1175, 1080)
    rounded(draw, core, LIGHT, BLUE, 10, 6)
    for y in range(560, 1060, 20):
        draw.line((505, y, 1155, y), fill="#B9CEDF", width=2)
    centred_text(draw, (540, 720, 1120, 900), "4-ROW COPPER / BRASS ACTIVE CORE\n530 × 435 × 64 THERMAL BASIS ONLY\nCOMPLETE R0 IS FIELD-MEASURED", image_font(24, True), BLUE)
    numbered_bubble(draw, (450, 790), 4, BLUE)
    # Bottom header and tank.
    draw.rectangle((450, 1165, 1210, 1210), fill="#E1E7EC", outline=hex_colour(NAVY), width=4)
    dashed_rectangle(draw, (500, 1280, 1160, 1385), PURPLE, 5)
    draw.ellipse((1025, 1365, 1090, 1410), outline=hex_colour(PURPLE), width=5)
    numbered_bubble(draw, (420, 1188), 5, NAVY)
    numbered_bubble(draw, (465, 1335), 6)
    numbered_bubble(draw, (1060, 1400), 7)
    # Side plates and mounts.
    draw.rectangle((365, 535, 405, 1080), fill="#D9E0E6", outline=hex_colour(NAVY), width=4)
    draw.rectangle((1255, 535, 1295, 1080), fill="#D9E0E6", outline=hex_colour(NAVY), width=4)
    dashed_rectangle(draw, (310, 640, 365, 750), PURPLE, 4)
    dashed_rectangle(draw, (1295, 870, 1350, 980), PURPLE, 4)
    numbered_bubble(draw, (385, 1115), 8, NAVY)
    numbered_bubble(draw, (1340, 925), 9)
    # Explode guides.
    for x in (520, 830, 1140):
        dashed_line(draw, ((x, 360), (x, 405)), MUTED, 3, 10, 7)
        dashed_line(draw, ((x, 450), (x, 535)), MUTED, 3, 10, 7)
        dashed_line(draw, ((x, 1080), (x, 1165)), MUTED, 3, 10, 7)
        dashed_line(draw, ((x, 1210), (x, 1280)), MUTED, 3, 10, 7)

    drawing_panel(draw, (1705, 185, 2350, 1050), "PARTS / KARIGAR LIST", fill="F8FAFC")
    parts = [
        ("1", "Top tank", "Reuse/repair or copy sample"),
        ("2", "Filler + cap + overflow", "Copy serviceable original"),
        ("3", "Top header plate", "Match core/tank method"),
        ("4", "4-row active-core matrix", "530 × 435 × 64 thermal basis; complete R0 TBD"),
        ("5", "Bottom header plate", "Match core/tank method"),
        ("6", "Bottom tank", "Reuse/repair or copy sample"),
        ("7", "Drain assembly", "Accessible; sample controls"),
        ("8", "Side support plates", "No load into fins/tubes"),
        ("9", "Mounting brackets", "Transfer actual centres"),
        ("10", "Inlet/outlet necks", "38 OD target; sample controls"),
        ("11", "Lower saddle + EPDM", "3–4 steel + 5 pad"),
        ("12", "Close shroud", "Removable, full perimeter"),
    ]
    y = 245
    for n, name, control in parts:
        draw.ellipse((1725, y, 1765, y + 40), fill=hex_colour(PURPLE if n not in ("4", "11") else RED))
        draw.text((1737, y + 5), n, font=image_font(17, True), fill="white")
        draw.text((1780, y), name, font=image_font(19, True), fill=hex_colour(INK))
        draw.text((1780, y + 25), control, font=image_font(16), fill=hex_colour(MUTED))
        y += 65

    drawing_note(
        draw,
        (1705, 1080, 2350, 1495),
        "FABRICATION CONTROL",
        "Pressure parts: radiator specialist only. Braze/solder process, header engagement and tube sealing follow the selected core system. Bench pressure/flow test and record result. No welding/drilling into a core or tank. No through-core ties. Keep final core away from grinding and structural welding.",
        RED,
        "FDF2F1",
    )
    finish_drawing_sheet(draw, "D03", "RADIATOR EXPLODED PARTS / ASSEMBLY ORDER")
    image.save(ASSET / "rev_c_d03_radiator_components.png", quality=96)


# Legacy Rev C definition retained for historical reproducibility.  The later
# Rev G definition with the same name is the active generator.
def save_d04_component_dimensions():
    image, draw = new_drawing_sheet(
        "D04",
        "FRONT COMPONENTS — INDIVIDUAL DIMENSION SHEET",
        "Nominal bodies and released fan-package limits. Actual manifolds, tanks, ports, ears, plugs and wire bends are measured.",
    )
    panels = [(45, 185, 785, 1495), (830, 185, 1570, 1495), (1615, 185, 2350, 1495)]
    drawing_panel(draw, panels[0], "A — R134a CONDENSER", fill="F8FAFC")
    c = (155, 370, 675, 1040)
    draw.rectangle(c, fill="#DDECF3", outline=hex_colour(CYAN), width=6)
    for x in range(175, 660, 22):
        draw.line((x, 390, x, 1020), fill="#A8D0E2", width=2)
    dimension_h(draw, c[0], c[2], 325, c[1], "559 BODY", RED)
    dimension_v(draw, c[1], c[3], 720, c[2], "356 BODY", RED, "left")
    draw.rectangle((610, 350, 650, 1060), outline=hex_colour(PURPLE), width=5)
    draw.ellipse((645, 430, 710, 475), outline=hex_colour(PURPLE), width=5)
    draw.ellipse((645, 900, 710, 945), outline=hex_colour(PURPLE), width=5)
    drawing_note(draw, (85, 1110, 745, 1455), "COMPLETE ENVELOPE / PORTS", "Body: 559 W × 356 H × 21 D nominal parallel-flow. Same-side top #8 inlet and lower #6 outlet are routing assumptions only; confirm actual fitting type and orientation. Four independent 3 mm tabs, M6 with rubber washers. No fan or core load.", PURPLE)

    drawing_panel(draw, panels[1], "B — INTERCOOLER", fill="F8FAFC")
    i = (920, 590, 1480, 895)
    rounded(draw, i, "E4EEE8", GREEN, 12, 6)
    for y in range(610, 880, 18):
        draw.line((940, y, 1460, y), fill="#BAD4C2", width=2)
    draw.ellipse((845, 670, 935, 815), outline=hex_colour(PURPLE), width=6)
    draw.ellipse((1465, 670, 1555, 815), outline=hex_colour(PURPLE), width=6)
    dimension_h(draw, i[0], i[2], 535, i[1], "500 CORE", RED)
    dimension_v(draw, i[1], i[3], 1515, i[2], "180 CORE", RED, "left")
    dimension_h(draw, 845, 935, 985, 815, "57 OD", RED, 19)
    drawing_note(draw, (870, 1110, 1530, 1455), "TANKS / OUTLETS / TEST", "Core: 500 W × 180 H × 60 D. 57 OD / 2.25 in beaded outlets. Actual tank shape, outlet side, angle and complete hose sweep are FIELD MEASURE. Four independent 4 mm tabs, M8 isolators; one upper 2–3 horizontal movement slot. Leak test 20 psi.", PURPLE)

    drawing_panel(draw, panels[2], "C — DUAL 9 in PUSHERS", fill="F8FAFC")
    f = (1680, 445, 2290, 855)
    gap = 12
    fw = (f[2] - f[0] - gap) / 2
    draw_fan(draw, (f[0], f[1], f[0] + fw, f[3]), "FAN 1")
    draw_fan(draw, (f[0] + fw + gap, f[1], f[2], f[3]), "FAN 2")
    dimension_h(draw, f[0], f[2], 395, f[1], "≤515 COMPLETE", RED)
    dimension_v(draw, f[1], f[3], 2320, f[2], "≤245 COMPLETE", RED, "left")
    draw.rectangle((1800, 930, 2180, 1010), fill="#F6EEDB", outline=hex_colour(GOLD), width=5)
    dimension_v(draw, 930, 1010, 2240, 2180, "≤55 DEPTH", RED, "left")
    drawing_note(draw, (1655, 1110, 2310, 1455), "COMPLETE FAN PACKAGE", "Two matching 12 V, 9 in nominal pushers. Envelope ≤515 W × 245 H × 55 D includes frames, tabs, guards, plugs and wire bends. Documented combined free-air rating ≥2,200 m³/h at about 13 V. Separate fuse and sealed relay for each fan; every fuse part remains inside a closed gasketed MIDI box.", RED, "FDF2F1")
    finish_drawing_sheet(draw, "D04", "CONDENSER / INTERCOOLER / DUAL-FAN DIMENSIONS")
    image.save(ASSET / "rev_c_d04_component_dimensions.png", quality=96)


def save_d05_mounting_shroud():
    image, draw = new_drawing_sheet(
        "D05",
        "MOUNTING, LOWER SADDLE & MECHANICAL-FAN SHROUD",
        "All exchangers and the electric-fan frame mount independently and remain removable; no through-core ties.",
    )
    drawing_panel(draw, (45, 185, 790, 1495), "A — FORMED UPRIGHT + RAIL", fill="F8FAFC")
    ux = 245
    draw.rectangle((ux, 390, ux + 95, 1130), fill="#D9E0E6", outline=hex_colour(NAVY), width=6)
    draw.rectangle((ux, 315, ux + 250, 390), fill="#D9E0E6", outline=hex_colour(NAVY), width=6)
    draw.rectangle((ux - 45, 1130, ux + 310, 1210), fill="#D9E0E6", outline=hex_colour(NAVY), width=6)
    draw.rectangle((ux + 310, 1130, ux + 390, 1290), fill="#D9E0E6", outline=hex_colour(NAVY), width=6)
    dimension_h(draw, ux, ux + 95, 350, 390, "48 FACE", RED, 18)
    dimension_v(draw, 390, 1130, 145, ux, "410 UPRIGHT", RED, "right")
    dimension_h(draw, ux, ux + 250, 270, 315, "58 TOP RETURN", RED, 18)
    dimension_h(draw, ux - 45, ux + 310, 1255, 1210, "70 CHASSIS BRIDGE", RED, 18)
    dimension_v(draw, 1130, 1290, 695, ux + 390, "80 OUTER LEG", RED, "left")
    draw.rectangle((455, 540, 665, 610), fill="#E8EEF5", outline=hex_colour(BLUE), width=5)
    draw.rounded_rectangle((520, 553, 550, 598), radius=12, outline=hex_colour(PURPLE), width=5)
    draw.rounded_rectangle((600, 553, 630, 598), radius=12, outline=hex_colour(PURPLE), width=5)
    leader_note(draw, (575, 575), (75, 1320, 750, 1460), "REMOVABLE RAIL / ADAPTER", "30 × 3 flat or 25 × 25 × 3 angle. Rail length = M2 − 10. Two M8 class 8.8 bolts per upright. Adapter slot 9 × 20 vertical. M7 offset 0–20 simple tab; boxed spacer above 20.", PURPLE)

    drawing_panel(draw, (835, 185, 1580, 1495), "B — LOWER SADDLE DETAIL", fill="F8FAFC")
    draw.rectangle((1000, 440, 1400, 1040), fill=hex_colour(LIGHT), outline=hex_colour(BLUE), width=6)
    draw.rectangle((950, 1040, 1450, 1090), fill="#6E7F8C", outline=hex_colour(NAVY), width=5)
    draw.rectangle((975, 1015, 1425, 1040), fill="#6FA77C", outline=hex_colour(GREEN), width=3)
    draw.line((950, 1090, 950, 1180), fill=hex_colour(NAVY), width=16)
    draw.line((1450, 1090, 1450, 1180), fill=hex_colour(NAVY), width=16)
    dimension_v(draw, 1015, 1040, 1485, 1425, "5 EPDM", RED, "left")
    dimension_v(draw, 1040, 1090, 1535, 1450, "3–4 STEEL", RED, "left")
    draw.ellipse((920, 350, 980, 410), fill="#6FA77C")
    draw.line((950, 410, 950, 1015), fill=hex_colour(PURPLE), width=5)
    drawing_note(draw, (875, 1215, 1540, 1460), "LOAD PATH", "Radiator weight sits on two lower saddles with 5 mm EPDM. M8 fasteners use large washers and rubber bushes. Upper tabs restrain fore/aft and side movement only; do not hang radiator weight from upper ears.", RED, "FDF2F1")

    drawing_panel(draw, (1625, 185, 2350, 1495), "C — CLOSE SHROUD / PULLER", fill="F8FAFC")
    draw.rectangle((1710, 410, 1770, 1110), fill=hex_colour(LIGHT), outline=hex_colour(BLUE), width=6)
    draw.rectangle((1770, 440, 2120, 1080), outline=hex_colour(NAVY), width=7)
    draw.ellipse((1820, 575, 2200, 955), outline=hex_colour(RED), width=8)
    draw.ellipse((1900, 655, 2120, 875), fill="#F6EEDB", outline=hex_colour(GOLD), width=7)
    draw.line((2010, 765, 2285, 765), fill=hex_colour(GOLD), width=22)
    dimension_h(draw, 1710, 1820, 1170, 1110, "≥20 STATIC; 25–30 PREF", RED, 18)
    dimension_h(draw, 1900, 2120, 1020, 955, "35–50% BLADE DEPTH", RED, 18)
    draw.text((1740, 300), "RADIATOR", font=image_font(20, True), fill=hex_colour(BLUE))
    draw.text((1840, 490), "FULL-PERIMETER\nREMOVABLE SHROUD", font=image_font(20, True), fill=hex_colour(NAVY))
    draw.text((1940, 725), "ORIGINAL\nPULLER", font=image_font(19, True), fill=hex_colour(GOLD))
    drawing_note(draw, (1665, 1215, 2310, 1460), "F1–F3 + M5", "Measure actual fan OD, swept circle, axial blade depth and centre X/Y through full hand rotation. Shroud aperture clears swept circle plus engine movement by ≥15 radial. Seal shroud perimeter. Reject damaged or weld-repaired blades.", PURPLE)
    finish_drawing_sheet(draw, "D05", "STRUCTURAL MOUNTS / SADDLE / CLOSE SHROUD")
    image.save(ASSET / "rev_c_d05_mounting_shroud.png", quality=96)


def draw_depth_band_v2(draw, top, title, components, front_gate, full_gate, gate_body):
    draw.text((80, top), title, font=image_font(27, True), fill=hex_colour(NAVY))
    y1 = top + 65
    y2 = y1 + 245
    start = 390
    x = start
    scale = 5.2
    for name, depth, colour, fill in components:
        x2 = x + depth * scale
        if name == "CLEAR":
            draw.rectangle((x, y1, x2, y2), fill="white", outline=hex_colour(MUTED), width=3)
        else:
            rounded(draw, (x, y1, x2, y2), fill, colour, 8, 4)
        if depth <= 15:
            short_name = "B" if name == "BUILD" else "C"
            centred_text(
                draw,
                (x - 2, y1 + 80, x2 + 2, y2 - 80),
                f"{short_name}{depth}",
                image_font(13, True),
                colour,
            )
        else:
            centred_text(
                draw,
                (x + 2, y1 + 25, x2 - 2, y2 - 20),
                f"{name}\n{depth}",
                image_font(19, True),
                colour,
            )
        dimension_h(draw, x, x2, y2 + 42, y2, str(depth), colour, 18)
        x = x2
    rad_front = x - 64 * scale
    dimension_h(draw, start, rad_front, y2 + 94, y2, front_gate, PURPLE, 19)
    dimension_h(draw, start, x, y2 + 144, y2, full_gate, RED, 19)
    drawing_note(draw, (1640, y1 - 5, 2320, y2 + 145), "RELEASE GATE", gate_body, PURPLE)
    arrow(draw, (80, y1 + 80), (350, y1 + 80), CYAN, 8)
    arrow(draw, (80, y1 + 170), (350, y1 + 170), CYAN, 8)
    draw.text((84, y2 - 35), "GRILLE / FRONT", font=image_font(18, True), fill=hex_colour(MUTED))


def save_d06_side_geometry():
    image, draw = new_drawing_sheet(
        "D06",
        "COMPLETE SIDE SECTION — DEPTH BUDGET & CLEARANCE GATES",
        "Front/grille is left; engine is right. Measure to the nearest fixed obstruction, not an ideal grille plane.",
    )
    draw_depth_band_v2(
        draw,
        205,
        "A — UPPER BAND: PUSHERS + CONDENSER + RADIATOR",
        [
            ("BUILD", 5, MUTED, WHITE),
            ("DUAL\nFANS", 55, GOLD, "F6EEDB"),
            ("CLEAR", 5, MUTED, WHITE),
            ("COND", 21, CYAN, "DDECF3"),
            ("CLEAR", 15, MUTED, WHITE),
            ("RAD", 64, BLUE, LIGHT),
        ],
        "M3-U ≥101 TO RADIATOR FRONT",
        "165 INSTALLED ENVELOPE incl. 5 BUILD",
        "M3-U = 5 build/plane + 55 fan + 5 clear + 21 condenser + 15 clear = 101 to radiator front. Check both fan centres, connectors and guards. If using grille plane, subtract rearward guard/obstruction projection.",
    )
    draw_depth_band_v2(
        draw,
        820,
        "B — LOWER BAND: INTERCOOLER + CONDENSER + RADIATOR",
        [
            ("BUILD", 10, MUTED, WHITE),
            ("I/C", 60, GREEN, "E4EEE8"),
            ("CLEAR", 10, MUTED, WHITE),
            ("COND", 21, CYAN, "DDECF3"),
            ("CLEAR", 15, MUTED, WHITE),
            ("RAD", 64, BLUE, LIGHT),
        ],
        "M3-L ≥116 TO RADIATOR FRONT",
        "M4 180 INSTALLED PASS; 190 PREF",
        "Base lower stack = 170. Add 10 build/plane allowance: M3-L = 116 to radiator front and M4 = 180 to radiator rear. Measure at seams/outlets and several lower points. 170–179 NOT RELEASED; <170 STOP.",
    )
    # Rear shroud / mechanical fan cue.
    draw.rectangle((1325, 1265, 1385, 1460), fill="#E8EEF5", outline=hex_colour(BLUE), width=5)
    draw.rectangle((1385, 1280, 1565, 1445), outline=hex_colour(NAVY), width=5)
    draw.ellipse((1500, 1300, 1655, 1455), outline=hex_colour(GOLD), width=7)
    dimension_h(draw, 1385, 1500, 1500, 1445, "M5 ≥20; 25–30 PREF", RED, 17)
    draw.text((1030, 1350), "RADIATOR REAR", font=image_font(19, True), fill=hex_colour(BLUE))
    draw.text((1395, 1290), "CLOSE\nSHROUD", font=image_font(17, True), fill=hex_colour(NAVY))
    draw.text((1515, 1355), "PULLER", font=image_font(17, True), fill=hex_colour(GOLD))
    finish_drawing_sheet(draw, "D06", "SIDE-SECTION PACKAGING / M3-U / M3-L / M4 / M5")
    image.save(ASSET / "rev_c_d06_side_geometry.png", quality=96)


# Legacy Rev C definition retained for historical reproducibility.  The later
# Rev G definition with the same name is the active generator.
def save_d07_fan_wiring():
    image, draw = new_drawing_sheet(
        "D07",
        "DUAL PUSHER MODULE — MOUNTING, AIRFLOW & WIRING",
        "Two matching 9 in pushers are mandatory for the condenser; each motor has its own protected power branch, with all fuse parts inside one closed gasketed MIDI box.",
    )
    drawing_panel(draw, (45, 185, 980, 1495), "A — REMOVABLE UPPER FAN FRAME", fill="F8FAFC")
    frame = (135, 350, 890, 900)
    draw.rectangle(frame, outline=hex_colour(NAVY), width=9)
    draw_fan(draw, (170, 390, 515, 835), "9 in PUSHER 1")
    draw_fan(draw, (525, 390, 870, 835), "9 in PUSHER 2")
    draw.rectangle((110, 320, 250, 355), fill="#D9E0E6", outline=hex_colour(NAVY), width=4)
    draw.rectangle((775, 320, 915, 355), fill="#D9E0E6", outline=hex_colour(NAVY), width=4)
    draw.rectangle((110, 895, 250, 930), fill="#D9E0E6", outline=hex_colour(NAVY), width=4)
    draw.rectangle((775, 895, 915, 930), fill="#D9E0E6", outline=hex_colour(NAVY), width=4)
    dimension_h(draw, frame[0], frame[2], 285, frame[1], "≤515 COMPLETE", RED)
    dimension_v(draw, frame[1], frame[3], 930, frame[2], "≤245 COMPLETE", RED, "left")
    draw.text((125, 970), "AIRFLOW", font=image_font(21, True), fill=hex_colour(CYAN))
    for y in (1000, 1050, 1100):
        arrow(draw, (220, y), (815, y), CYAN, 8)
    drawing_note(draw, (85, 1160, 940, 1455), "MOUNTING / FIT", "Independent removable upper hoop/crossrails; never through-core plastic ties and never hang motors on condenser fins/tubes. Airflow: grille → fans → condenser → radiator → engine. F4 includes guards, plugs, tabs and wire bends. F5 ≥5 no-rub to intercooler; 10 target.", PURPLE)

    drawing_panel(draw, (1025, 185, 2350, 1495), "B — TWO INDEPENDENT PROTECTED CIRCUITS", fill="F8FAFC")
    rounded(draw, (1080, 400, 1290, 590), LIGHT, BLUE, 14, 4)
    centred_text(draw, (1090, 410, 1280, 580), "BATTERY /\nCHARGING\n12 V", image_font(24, True), BLUE)
    ys = (350, 680)
    for index, y in enumerate(ys, start=1):
        rounded(draw, (1390, y, 1580, y + 135), "FDF2F1", RED, 12, 4)
        centred_text(draw, (1400, y + 8, 1570, y + 127), f"MIDI BOX\nFUSE {index}\nMEASURED", image_font(18, True), RED)
        rounded(draw, (1690, y, 1875, y + 135), "F6EEDB", GOLD, 12, 4)
        centred_text(draw, (1700, y + 8, 1865, y + 127), f"SEALED\nRELAY {index}", image_font(20, True), GOLD)
        rounded(draw, (1985, y, 2185, y + 135), "F6EEDB", GOLD, 12, 4)
        centred_text(draw, (1995, y + 8, 2175, y + 127), f"9 in\nFAN {index}", image_font(21, True), GOLD)
        arrow(draw, (1295, 495), (1380, y + 67), RED, 6)
        arrow(draw, (1585, y + 67), (1680, y + 67), RED, 6)
        arrow(draw, (1880, y + 67), (1975, y + 67), GOLD, 6)
        arrow(draw, (2190, y + 67), (2260, 930), GREEN, 6)
    rounded(draw, (2170, 900, 2315, 1060), "E4EEE8", GREEN, 12, 4)
    centred_text(draw, (2180, 910, 2305, 1050), "CLEAN\nGROUND\nSTUD", image_font(19, True), GREEN)
    rounded(draw, (1420, 960, 2010, 1135), "F3ECF8", PURPLE, 12, 4)
    centred_text(draw, (1435, 970, 1995, 1125), "COMMON LOW-CURRENT TRIGGER\nBoth fans ON with A/C clutch command.\nTrinary/high-pressure may also request fans.", image_font(21, True), PURPLE)
    arrow(draw, (1580, 960), (1780, 815), PURPLE, 6)
    arrow(draw, (1850, 960), (1780, 485), PURPLE, 6)
    drawing_note(draw, (1080, 1190, 2305, 1455), "WIRE / TEST", "6 mm² copper common feed and ground unless loaded test proves 4 mm² adequate; ≥4 mm² each motor branch. Every fuse element, holder, bus and live fuse stud is inside the closed gasketed MIDI box; relays are in their separate covered box. Use sealed connectors, glands, loom and strain relief. Fuse from measured running/start current and conductor capacity. Each running motor must be within 0.5 V of loaded battery charging voltage.", RED, "FDF2F1")
    finish_drawing_sheet(draw, "D07", "DUAL-FAN MODULE / AIRFLOW / TWO-BRANCH WIRING")
    image.save(ASSET / "rev_c_d07_fan_wiring.png", quality=96)


# Legacy Rev C dispatcher; the later Rev G dispatcher is the active definition.
def save_dimensioned_drawing_set():
    save_d01_complete_stack()
    save_d02_radiator_assembly()
    save_d03_radiator_components()
    save_d04_component_dimensions()
    save_d05_mounting_shroud()
    save_d06_side_geometry()
    save_d07_fan_wiring()


DRAWING_SHEETS = [
    (
        "D01",
        "Complete cooling-pack front elevation",
        "rev_c_d01_complete_stack.png",
        "Radiator, condenser, intercooler, two 9 in pushers and M1/M2/F5 packaging.",
    ),
    (
        "D02",
        "Radiator assembly orthographic",
        "rev_c_d02_radiator_assembly.png",
        "530 × 435 × 64 core plus controlled sample/vehicle interfaces R01–R08.",
    ),
    (
        "D03",
        "Radiator exploded parts",
        "rev_c_d03_radiator_components.png",
        "Core, tanks, headers, filler, drain, side plates, brackets, saddles and shroud.",
    ),
    (
        "D04",
        "Individual component dimensions",
        "rev_c_d04_component_dimensions.png",
        "Condenser, intercooler and complete dual-fan module requirements.",
    ),
    (
        "D05",
        "Mounting, saddle and shroud details",
        "rev_c_d05_mounting_shroud.png",
        "Transferred upright, removable rail, lower saddle and mechanical-puller shroud.",
    ),
    (
        "D06",
        "Complete side-section depth budget",
        "rev_c_d06_side_geometry.png",
        "M3-U, M3-L, M4 and M5 clearance arithmetic and stop gates.",
    ),
    (
        "D07",
        "Dual-pusher mounting and wiring",
        "rev_c_d07_fan_wiring.png",
        "Removable upper module, airflow direction and two protected motor branches.",
    ),
]


def set_run(run, size=11, bold=False, colour=INK, italic=False):
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:cs"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(colour)


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    properties.append(indent)
    grid = table._tbl.tblGrid
    for item in list(grid):
        grid.remove(item)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(document, headers, rows, widths, font_size=9.3):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], LIGHT)
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(str(header))
        set_run(run, 9.5, True, NAVY)
    set_repeat_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(str(value))
            set_run(run, font_size)
    set_table_geometry(table, widths)
    return table


def add_paragraph(
    document,
    text,
    style=None,
    bold_lead=None,
    after=6,
    italic=False,
):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        first = paragraph.add_run(bold_lead)
        set_run(first, bold=True)
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run(rest, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_run(run, italic=italic)
    return paragraph


def add_bullets(document, items):
    for item in items:
        add_paragraph(document, item, "List Bullet", after=4)


def add_numbers(document, items):
    # Use literal numbering so every separately controlled sequence restarts at 1.
    # Word's built-in List Number style otherwise continues the previous list
    # after intervening headings/tables in some LibreOffice renderers.
    for index, item in enumerate(items, 1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.30)
        paragraph.paragraph_format.first_line_indent = Inches(-0.30)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(f"{index}.  {item}")
        set_run(run)


def add_heading(document, text, level=1):
    return document.add_paragraph(text, style=f"Heading {level}")


def add_heading_new_page(document, text, level=1):
    """Start a section without inserting a blank break-only paragraph.

    A break-only paragraph can be pushed onto the next page when the preceding
    content exactly fills a page, producing an unintended blank sheet in the
    LibreOffice render.  Applying the break to the heading keeps pagination
    deterministic.
    """
    paragraph = add_heading(document, text, level)
    paragraph.paragraph_format.page_break_before = True
    return paragraph


def add_callout(document, title, body, fill="FDF2F1", colour=RED):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    # Keep the title and body together.  Without this, LibreOffice can split
    # the single callout row across pages and leave an orphaned banner.
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    first = paragraph.add_run(title + "\n")
    set_run(first, 10.5, True, colour)
    second = paragraph.add_run(body)
    set_run(second, 10.5, False, INK)
    set_table_geometry(table, [9360])
    return table


def add_picture(document, path, width=6.35):
    document.add_picture(str(path), width=Inches(width))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.paragraphs[-1].paragraph_format.space_after = Pt(4)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("J40 Cooling Pack  |  Rev C.1  |  ")
    set_run(run, 9, False, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for level, size, before, after, colour in (
        (1, 16, 18, 10, BLUE),
        (2, 13, 14, 7, BLUE),
        (3, 12, 10, 5, NAVY),
    ):
        style = document.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "J40 FRONT COOLING SYSTEM  •  LOCAL FABRICATOR SPECIFICATION"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.runs[0], 9, True, MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_landscape_drawing_appendix(document):
    add_heading(document, "Appendix A. Full-size dimensioned drawing sheets", 1)
    add_paragraph(
        document,
        "These seven sheets form the Rev C.1 drawing register. Red dimensions are fixed nominal requirements "
        "or release limits. Purple dimensions and dashed geometry must be measured from the vehicle or copied "
        "from the actual component. Written dimensions control; do not scale the images.",
    )
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.40)
    section.right_margin = Inches(0.40)
    section.header_distance = Inches(0.15)
    section.footer_distance = Inches(0.15)
    for index, (code, title, filename, _) in enumerate(DRAWING_SHEETS):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"{code} — {title}")
        set_run(run, 10, True, NAVY)
        document.add_picture(str(ASSET / filename), width=Inches(9.8))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.paragraphs[-1].paragraph_format.space_after = Pt(1)
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(0)
        run = caption.add_run(
            "Controlled drawing: written dimensions govern; red = fixed/release limit; "
            "purple = field measure / owner approval."
        )
        set_run(run, 8, False, MUTED, True)
        if index < len(DRAWING_SHEETS) - 1:
            document.add_page_break()


def build_document_legacy_reference_do_not_call():
    """Historical pre-Rev-F document builder retained only for source traceability."""
    save_dimensioned_drawing_set()

    document = Document()
    configure_document(document)

    # Cover.
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(48)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run("FABRICATOR SPECIFICATION")
    set_run(run, 11, True, GOLD)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run("J40 Integrated Radiator\nand Front Cooling Pack")
    set_run(run, 27, True, NAVY)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run(
        "2H diesel • manual gearbox basis • 5–7 psi future turbo • R134a A/C"
    )
    set_run(run, 13, False, MUTED)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(16)
    run = paragraph.add_run("REV C.1  |  30 JULY 2026  |  ALL DIMENSIONS IN mm")
    set_run(run, 10, True, BLUE)
    add_callout(
        document,
        "RELEASE STATUS — MOCK-UP / MEASUREMENT ONLY",
        "Final radiator, condenser and intercooler manufacture is HOLD until M1–M7 and F1–F5 pass, "
        "the complete actual parts are mock-fitted, and the owner releases manufacture in writing.",
    )
    add_paragraph(
        document,
        "SHOP RULE / KARIGAR KE LIYE: Pehle gaari par naap lo. Phir cardboard ya plywood dummy fit karo. "
        "Owner ki written approval ke baad hi final core banao.",
        after=12,
        italic=True,
    )
    add_table(
        document,
        ["VEHICLE", "FIXED COOLING DECISION", "REV C CHANGE"],
        [
            (
                "1978 J40\nToyota 2H",
                "530 × 435 × 64 four-row copper/brass radiator; original mechanical fan + close shroud",
                "Two compact 9 in condenser pushers replace the colliding single 12 in layout",
            )
        ],
        [1700, 3860, 3800],
        9.7,
    )
    # Shop summary.
    document.add_page_break()
    add_heading(document, "1. One-page shop instruction", 1)
    add_paragraph(
        document,
        "The engine radiator, A/C condenser and turbo intercooler are three separate pressure circuits. "
        "Each component has its own removable, rubber-isolated mounts. No core carries another part."
    )
    add_picture(document, ASSET / "rev_c_d01_complete_stack.png", 5.10)
    add_table(
        document,
        ["PART", "BASELINE SIZE / DECISION"],
        [
            ("Engine radiator core", "530 W × 435 H × 64 D; high-efficiency 4-row copper/brass"),
            ("A/C condenser", "559 W × 356 H × 21 D nominal R134a parallel-flow"),
            ("Intercooler core", "500 W × 180 H × 60 D; 57 OD / 2.25 in beaded outlets"),
            (
                "Front electric fans",
                "2 matching 12 V, 9 in nominal pushers; released max ≤515 × 245 × 55; "
                "≤520 × 250 × 55 is mock-up/owner-approval only",
            ),
            (
                "Main radiator fan",
                "retain/rebuild existing engine-driven puller; full close shroud required",
            ),
        ],
        [2700, 6660],
    )
    add_callout(
        document,
        "NEVER",
        "No welding/drilling a core or tank. No through-core plastic ties. No fan load on the condenser. "
        "Do not delete the mechanical fan. Do not guess necks, cap pressure, port threads or fan current.",
    )
    add_heading(document, "1.1 Rev C.1 drawing register", 2)
    add_table(
        document,
        ["SHEET", "DRAWING", "FABRICATOR USE"],
        [(code, title, use) for code, title, _, use in DRAWING_SHEETS],
        [900, 3200, 5260],
        8.7,
    )
    add_callout(
        document,
        "DRAWING LEGEND",
        "RED = fixed nominal dimension or release limit. PURPLE = measure on vehicle, copy actual component, "
        "or obtain written approval. Every unlabelled tank, neck, port, bracket and obstruction position is "
        "field-measured. Written dimensions control; never scale an image.",
        "F3ECF8",
        PURPLE,
    )

    # Status and radiator.
    document.add_page_break()
    add_heading(document, "2. Latest component basis and release status", 1)
    add_table(
        document,
        ["ITEM", "CURRENT BASIS", "SHOP PROOF REQUIRED"],
        [
            (
                "Engine / gearbox",
                "2H diesel; H55F / 5-speed manual candidate",
                "No automatic-transmission oil cooler. Re-open if gearbox changes.",
            ),
            (
                "Original radiator",
                "Physical sample is interface authority",
                "Inspect tanks, headers, necks, cap seat, overflow, drain and brackets; quote recore first.",
            ),
            (
                "Mechanical fan",
                "Present in engine photograph",
                "Inspect fan/hub/water pump/mounts; record complete rotating envelope.",
            ),
            (
                "Condenser",
                "559 × 356 × 21 nominal design basis",
                "Tape and photograph actual body, seams, ports, ears and complete envelope.",
            ),
            (
                "Intercooler",
                "planned, not final-fabrication released",
                "Mark turbo and intake connections on vehicle before outlet welding.",
            ),
            (
                "Electric fans",
                "Rev C dual compact arrangement",
                "Actual models must pass full envelope, airflow, current and mock-up checks.",
            ),
        ],
        [1800, 3000, 4560],
        8.8,
    )
    add_heading(document, "3. Engine radiator specification", 1)
    add_paragraph(
        document,
        "Preferred first option: recore the original copper/brass radiator if tanks, filler neck, header "
        "plates, drain, overflow and brackets are sound. Otherwise make a new copper/brass unit from the sample."
    )
    add_picture(document, ASSET / "rev_c_d02_radiator_assembly.png", 4.20)
    add_table(
        document,
        ["PARAMETER", "REQUIREMENT"],
        [
            ("Core", "530 W × 435 H × 64 D target; high-efficiency 4-row copper/brass"),
            (
                "Thermal duty",
                "Toyota 2H diesel plus planned conservative 5–7 psi turbo heat load",
            ),
            (
                "Necks",
                "38 OD target; sample/hose mock-up controls side, angle, centres, insertion and bead",
            ),
            (
                "Filler / cap",
                "retain correct seat, overflow and verified original/service cap rating; never fit random higher pressure",
            ),
            ("Drain", "accessible at low point"),
            ("Oil cooler", "none on present manual-gearbox basis"),
            (
                "Finish",
                "clean/flush; thin black radiator coating only; no heavy paint blocking fins",
            ),
            (
                "Test",
                "verified system/cap test pressure, ≥5 min, no leak/sweat/loss; flow-test and record",
            ),
        ],
        [2300, 7060],
        9.0,
    )
    add_callout(
        document,
        "ALTERNATIVE SIZE HOLD",
        "A 540 W × 435 H × about 60 D core is not released unless M1 is at least 550 for that core and the "
        "complete tank/bracket/condenser/removal envelopes are also proven.",
        "FFF8E8",
        GOLD,
    )
    add_picture(document, ASSET / "rev_c_d03_radiator_components.png", 6.35)

    # Fan specification.
    document.add_page_break()
    add_heading(document, "4. All required fans and airflow", 1)
    add_table(
        document,
        ["COMPONENT", "FAN REQUIREMENT", "REV C DECISION"],
        [
            (
                "Engine radiator",
                "main idle/low-speed airflow",
                "original engine-driven puller + full close shroud",
            ),
            (
                "A/C condenser",
                "forced hot-idle airflow",
                "2 matching compact 12 V pusher fans",
            ),
            (
                "Intercooler",
                "no dedicated fan in this stack",
                "ram air + airflow drawn by main fan; extra fan would add blockage/depth",
            ),
            (
                "Cabin evaporator",
                "cabin blower required",
                "A/C installer scope, outside radiator-shop metalwork",
            ),
        ],
        [1900, 2900, 4560],
        9.0,
    )
    add_heading(document, "4.1 Engine-driven puller fan and close shroud", 2)
    add_bullets(
        document,
        [
            "Reject cracked, loose, distorted, contact-marked or weld-repaired blades.",
            "Check hub, pulley, water-pump bearing, fan run-out and engine mounts before setting the radiator plane.",
            "Record fan OD, axial fore/aft blade depth, centre X/Y and closest blade plane through a full hand rotation.",
            "Fit a rigid removable full-face shroud sealed around the radiator perimeter.",
            "Blade insertion into shroud opening: 35–50% of axial fore/aft blade depth.",
            "Shroud aperture must clear the rotating sweep plus checked engine-movement envelope by ≥15 radial.",
            "Radiator rear face to nearest fan point: ≥20 static; 25–30 preferred.",
        ],
    )
    add_heading(document, "4.2 Dual A/C pusher fans", 2)
    add_table(
        document,
        ["PARAMETER", "REQUIRED"],
        [
            ("Quantity / size", "2 matching × 9 in nominal; 12 V pusher"),
            (
                "Released maximum envelope",
                "≤515 W × 245 H × 55 D, including all tabs, guards, plugs and wire bends",
            ),
            (
                "Mock-up-only exception",
                "≤520 W × 250 H × 55 D; placement drawing and owner approval required before purchase",
            ),
            (
                "Location",
                "side-by-side above complete I/C envelope; ≥5 no-rub under vibration/load, 10 target",
            ),
            ("Direction", "grille → condenser → radiator → engine"),
            (
                "Airflow screen",
                "credible documented combined free-air rating ≥2,200 m³/h at about 13 V",
            ),
            (
                "Build / mounts",
                "sealed weather-resistant automotive motors; independent removable upper frame",
            ),
        ],
        [2500, 6860],
        9.0,
    )
    add_paragraph(
        document,
        "Record make/model, published curve, measured running current and measured startup current for each fan. "
        "A market-stall CFM claim without model data is not acceptance evidence."
    )
    add_callout(
        document,
        "12-INCH FALLBACK",
        "Not the Rev C baseline. Reconsider only after owner-approved fit and hot-idle A/C proof.",
        "FFF8E8",
        GOLD,
    )

    # Wiring and geometry.
    wiring_heading = add_heading(document, "5. Dual-fan wiring", 1)
    wiring_heading.paragraph_format.page_break_before = True
    add_picture(document, ASSET / "rev_c_d07_fan_wiring.png", 6.35)
    add_bullets(
        document,
        [
            "Two separately fused branches and two sealed relays.",
            "Both fans switch ON whenever the A/C clutch is commanded ON; trinary/high-pressure control may also request them.",
            "The clutch/switch circuit triggers relays only and never carries fan-motor current; A/C technician verifies logic/protection.",
            "Fuse each branch from measured run/start current and wire capacity; a typical result may be 20–30 A, but measurements control.",
            "Use 6 mm² copper common feed/ground unless a voltage-drop test proves 4 mm² adequate; use ≥4 mm² for each branch.",
            "Use a clean proper ground stud, weatherproof holders/connectors, loom, grommets and strain relief.",
            "Both motors must be within 0.5 V of loaded battery charging voltage while running.",
        ],
    )
    add_heading(document, "6. Stack geometry and depth gates", 1)
    add_picture(document, ASSET / "rev_c_d06_side_geometry.png", 6.35)
    add_table(
        document,
        ["CHECK", "ARITHMETIC / RELEASE"],
        [
            (
                "Upper band M3-U",
                "55 fan + 5 clear + 21 condenser + 15 clear + 5 build = 101 from radiator front to closest fixed obstruction",
            ),
            (
                "Lower band M3-L",
                "60 I/C + 10 clear + 21 condenser + 15 clear = 106; require 116 to closest fixed obstruction",
            ),
            (
                "Lower full M4",
                "106 front components/gaps + 64 radiator = 170 nominal; ≥180 PASS; 170–179 not released; <170 STOP",
            ),
        ],
        [2400, 6960],
        9.2,
    )

    # Condenser, IC, structure.
    document.add_page_break()
    add_heading(document, "7. Condenser, receiver-drier and intercooler", 1)
    add_picture(document, ASSET / "rev_c_d04_component_dimensions.png", 5.15)
    add_table(
        document,
        ["PART", "COMPLETE SPECIFICATION"],
        [
            (
                "R134a condenser",
                "parallel-flow; 559 W × 356 H × 21 D nominal. Same-side vertical manifolds: top #8 inlet, lower #6 "
                "outlet, subject to physical confirmation before crimping. Four independent 3 mm tabs, rubber washers, "
                "M6. A 600-wide candidate needs a proven opening/forward mount; if between uprights, M1 ≥610. No fan/core load.",
            ),
            (
                "Receiver-drier",
                "vertical in rubber-lined removable clamp outside primary airflow; keep sealed until final A/C assembly, evacuation and charge.",
            ),
            (
                "Intercooler",
                "500 W × 180 H × 60 D core; 57 OD / 2.25 in beaded outlets; four independent 4 mm tabs with M8 "
                "isolators; one upper 2–3 horizontal movement slot; 20 psi leak test.",
            ),
            (
                "Intercooler fallback",
                "50-deep reputable high-efficiency core if depth is marginal; do not go below 450 W × 160 H × 50 D "
                "without a full thermal-system review.",
            ),
        ],
        [1900, 7460],
        8.9,
    )
    add_heading(document, "8. Structure and independent mounting", 1)
    add_paragraph(
        document,
        "Retain the existing formed upright and the newly added handed/mirrored mate. Inspect attachment zones, "
        "weld penetration, cracking, distortion and squareness before paint. The installed sample controls final geometry."
    )
    add_picture(document, ASSET / "rev_c_d05_mounting_shroud.png", 6.35)
    add_table(
        document,
        ["ITEM", "FABRICATION REQUIREMENT"],
        [
            (
                "Structural basis",
                "4 mm formed mild steel: 48 main face, 410 upright, 58 top return, 70 chassis bridge, 80 outer leg; "
                "installed parts control—do not substitute generic 50 × 50 angle",
            ),
            (
                "Removable rails",
                "30 × 3 flat or 25 × 25 × 3 angle; cut to M2 − 10 for 5 top/bottom clearance; 2 × M8 class 8.8 bolts/upright",
            ),
            (
                "Adapter",
                "9 × 20 vertical slots in removable tabs only; simple offset 0–20, boxed spacer above 20",
            ),
            (
                "Radiator",
                "weight on two 3–4 mm lower saddles with 5 mm EPDM pads; M8 fasteners, large washers/rubber bushes; upper tabs restrain only",
            ),
            ("Fan assembly", "independent removable upper hoop/crossrails; never plastic through-core ties"),
            (
                "Serviceability",
                "each exchanger/fan separately removable; no permanent crossrail blocks core face or removal",
            ),
        ],
        [2300, 7060],
        9.0,
    )
    add_callout(
        document,
        "WORKSHOP PROTECTION",
        "Deburr and corrosion-protect after bare-metal inspection. Keep final cores away during welding/grinding. "
        "No forced bolt alignment and no metal edge against a tank, pipe, hose, wire or fin.",
        "FFF8E8",
        GOLD,
    )

    # Measurements.
    document.add_page_break()
    add_heading(document, "9. Mandatory fit gates — all must PASS", 1)
    add_paragraph(
        document,
        "Measure with both uprights represented, body settled on mounts, grille/front panel and bonnet latch fitted "
        "or accurately represented, and the original engine fan present. Put a tape/ruler in every evidence photograph."
    )
    measurements = [
        (
            "M1",
            "minimum clear width, top/mid/bottom",
            "≥540 for 530 core; ≥569 for condenser + 5/side; complete tanks/ports/ears fit",
        ),
        (
            "M2",
            "lower saddle plane to bonnet/latch obstruction",
            "actual radiator overall H + 10 and vertical package fits",
        ),
        (
            "M3-U",
            "radiator front to closest fixed obstruction, fan band",
            "≥101; if using grille plane, subtract measured rearward obstruction projection",
        ),
        (
            "M3-L",
            "radiator front to closest fixed obstruction, lower band",
            "≥116; if using grille plane, subtract measured rearward obstruction projection",
        ),
        (
            "M4",
            "closest fixed front obstruction to radiator rear, lower band",
            "≥180; 190 preferred; 170–179 not released; <170 STOP",
        ),
        (
            "M5",
            "radiator rear to nearest mechanical-fan point",
            "≥20 static; 25–30 preferred through rotation/movement",
        ),
        (
            "M6",
            "lowest intercooler edge vs protected line",
            "≥25 above protected frame/bumper line",
        ),
        (
            "M7",
            "upright face to rail offset at 4 corners",
            "0–20 simple tab; boxed spacer above 20",
        ),
        (
            "F1",
            "mechanical fan OD / complete swept circle",
            "record; shroud aperture based on actual sweep",
        ),
        (
            "F2",
            "axial blade depth / shroud insertion",
            "35–50% inside; sweep + engine movement has ≥15 radial clearance",
        ),
        (
            "F3",
            "mechanical fan centre X/Y",
            "record on as-built drawing",
        ),
        (
            "F4",
            "complete dual-pusher outside envelope",
            "released max ≤515 × 245 × 55; ≤520 × 250 × 55 mock-up/owner-approval only",
        ),
        (
            "F5",
            "pusher to I/C / edge clearance under vibration/load",
            "≥5 no-rub everywhere; 10 target at frame, plugs and wiring",
        ),
    ]
    add_table(
        document,
        ["ID", "MEASURE", "PASS / RELEASE CRITERION"],
        measurements,
        [750, 3780, 4830],
        8.6,
    )
    add_heading(document, "Also record", 2)
    add_bullets(
        document,
        [
            "Radiator overall tank/bracket size, neck centres/OD/angles, cap rating, overflow and drain.",
            "Condenser complete body/seam/manifold/port/bracket envelope.",
            "Intercooler complete tank/outlet envelope and charge-pipe route.",
            "Bonnet, latch, grille, winch, bumper/guard, steering/body, hose and service-removal clearances.",
        ],
    )

    # Sequence and commissioning.
    document.add_page_break()
    add_heading(document, "10. Required fabrication sequence", 1)
    add_numbers(
        document,
        [
            "Inspect both installed uprights for weld quality, distortion and squareness; alter/repair only for an identified defect with owner approval.",
            "Measure M1–M7 and the original mechanical fan F1–F3.",
            "Obtain the radiator sample, actual condenser, candidate fans and representative intercooler envelope.",
            "Make full-size dummies including tanks, ears, plugs, hose bends, fan guards and cable exits.",
            "Set radiator plane from mechanical-fan clearance and coolant-hose sweep.",
            "Position condenser with 15 target rear gap; place intercooler low with 10 front gap.",
            "Place both pushers side-by-side above the complete intercooler envelope.",
            "Close bonnet and fit/represent grille, latch, winch, bumper and guard; prove removal paths.",
            "Photograph all gates, prepare an as-measured sketch and obtain owner written release.",
            "Manufacture/recore; bench-test radiator, intercooler and both fans; retain records.",
            "Install, wire and complete static, hot-idle, A/C and road commissioning.",
        ],
    )
    add_heading(document, "11. Commissioning acceptance", 1)
    add_table(
        document,
        ["STAGE", "ACCEPTANCE"],
        [
            (
                "Static",
                "isolators present; no forced alignment/contact; independent removal; full fan rotation and movement "
                "clear; both pushers blow toward engine; shroud perimeter sealed",
            ),
            (
                "Warm idle",
                "no leak/wobble/contact; thermostat operation established; no continuous temperature rise or boiling",
            ),
            (
                "A/C hot idle",
                "both pushers run whenever clutch is commanded; technician records high-/low-side pressures, ambient and vent temperature",
            ),
            (
                "Electrical",
                "correct fan voltage drop; no abnormal current, relay heating, connector heating or wiring movement",
            ),
            (
                "Road / recheck",
                "low- and normal-speed road test; repeat hot-idle A/C check; no vibration, rubbing or fin damage",
            ),
        ],
        [1700, 7660],
        8.9,
    )
    add_callout(
        document,
        "IDLE AIRFLOW HOLD",
        "If the condenser lower portion has inadequate hot-idle airflow behind the intercooler, stop and revise "
        "fan ducting/placement. Do not delete the intercooler or shrink the radiator without a complete review.",
    )

    # Handover.
    document.add_page_break()
    add_heading(document, "12. Handover record and signatures", 1)
    records = [
        ("[ ]", "M1–M7 and F1–F5 complete", ""),
        ("[ ]", "Radiator make/core/tank overall size", ""),
        ("[ ]", "Radiator test pressure / time / result", ""),
        ("[ ]", "Verified cap rating and source", ""),
        ("[ ]", "Condenser make/model/actual envelope", ""),
        ("[ ]", "Intercooler envelope and 20 psi test", ""),
        ("[ ]", "Fan 1 model/flow/run/start current", ""),
        ("[ ]", "Fan 2 model/flow/run/start current", ""),
        ("[ ]", "Ambient temperature / test conditions", ""),
        ("[ ]", "Charging voltage / each fan run current / voltage drop", ""),
        ("[ ]", "A/C high/low pressures / vent temperature", ""),
        ("[ ]", "Warm-idle and road-test result", ""),
        ("[ ]", "As-built drawing and photographs", ""),
    ]
    add_table(
        document,
        ["CHECK", "RECORD", "RESULT / EVIDENCE"],
        records,
        [800, 4200, 4360],
        8.9,
    )
    add_paragraph(document, "\nFabricator name/signature: ______________________________________________")
    add_paragraph(document, "A/C electrician/technician: _____________________________________________")
    add_paragraph(document, "Owner final-core release: ______________________________________________")
    add_paragraph(document, "Date: ______________________________")
    add_heading(document, "Controlled references", 2)
    add_bullets(
        document,
        [
            "docs/J40-integrated-cooling-pack-fabricator-specification-rev-c.md",
            "docs/j40-integrated-cooling-pack-fabricator-handoff-20260717.md",
            "docs/radiator-workstream.md",
            "docs/engine-radiator-recore-release-20260529.md",
            "docs/project-progress-update-20260729.md",
            "data/manual/fabrication/front_cooling_stack_rev_a/integrated_cooling_pack_dimensions_rev_b.csv",
            "SPAL Axial Fans General Catalogue — reference performance benchmark only; no brand mandate.",
        ],
    )
    add_landscape_drawing_appendix(document)
    document.core_properties.title = (
        "J40 Integrated Radiator and Front Cooling Pack — Fabricator Specification — Rev C.1"
    )
    document.core_properties.subject = (
        "Local-fabricator specification for engine radiator, condenser, intercooler and required fans"
    )
    document.core_properties.author = "J40 Project"
    document.core_properties.comments = f"Controlled text source: {SOURCE.name}"
    document.save(OUT)
    print(OUT)


# ---------------------------------------------------------------------------
# Rev G overrides
#
# The older Rev C implementations above are retained for traceability.  The
# controlled outputs and stable URLs keep their historical "rev_c" filenames,
# but the functions below replace the prior architecture with the Rev G
# Toyota-donor-led, best-effort all-front, three-core architecture.
# ---------------------------------------------------------------------------


def finish_drawing_sheet(draw, code, description):
    draw.line((24, SHEET_FOOTER_Y, SHEET_W - 24, SHEET_FOOTER_Y), fill=hex_colour(NAVY), width=4)
    draw.rectangle((24, SHEET_FOOTER_Y, 420, SHEET_H - 24), fill="#FDF2F1")
    draw.rectangle((420, SHEET_FOOTER_Y, 820, SHEET_H - 24), fill="#F3ECF8")
    draw.text((45, 1571), "RED = FIXED / RELEASE LIMIT", font=image_font(18, True), fill=hex_colour(RED))
    draw.text((441, 1571), "PURPLE = FIELD MEASURE", font=image_font(18, True), fill=hex_colour(PURPLE))
    draw.text((45, 1608), "ALL DIMENSIONS mm", font=image_font(18, True), fill=hex_colour(NAVY))
    draw.text((441, 1608), "DO NOT SCALE DRAWING", font=image_font(18, True), fill=hex_colour(NAVY))
    draw.text((850, 1570), description, font=image_font(19, True), fill=hex_colour(INK))
    draw.text(
        (850, 1610),
        "Performance values and written dimensions govern. Unlabelled interfaces: measure actual vehicle/components.",
        font=image_font(17),
        fill=hex_colour(MUTED),
    )
    draw.text((2050, 1570), "REV G", font=image_font(21, True), fill=hex_colour(BLUE))
    draw.text((2050, 1608), "01 AUG 2026", font=image_font(18, True), fill=hex_colour(INK))
    draw.text((2250, 1570), code, font=image_font(27, True), fill=hex_colour(NAVY))


def _air_arrows(draw, x1, x2, ys, colour=CYAN):
    for y in ys:
        arrow(draw, (x1, y), (x2, y), colour, 7)


def save_d01_complete_stack():
    image, draw = new_drawing_sheet(
        "D01",
        "COMPLETE REV G COOLING SYSTEM — TOYOTA-DONOR-LED ALL-FRONT STACK",
        "Best-effort 50°C/A/C design basis only; no 50°C or no-derate claim until the complete vehicle passes test.",
    )
    drawing_panel(draw, (45, 185, 2355, 1485), "A — BETWEEN-UPRIGHTS ALL-FRONT STACK", fill="F8FAFC")
    fans, ic, cond, rad, shroud = (80, 385, 430, 1260), (480, 335, 775, 1310), (875, 365, 1165, 1280), (1275, 315, 1700, 1330), (1815, 375, 2210, 1280)
    rounded(draw, fans, "F6EEDB", GOLD, 12, 5)
    centred_text(draw, (90, 420, 420, 500), "FRONT ELEVATION\nCENTRED / LEVEL", image_font(15, True), GOLD)
    draw_fan(draw, (95, 670, 250, 825), "248 mm")
    draw_fan(draw, (260, 670, 415, 825), "248 mm")
    centred_text(draw, (90, 885, 420, 1045), "MATCHED\nTOYOTA/DENSO\n266 mm C-C", image_font(18, True), GOLD)
    rounded(draw, ic, "E4EEE8", GREEN, 12, 6)
    centred_text(draw, ic, "SLIM HORIZONTAL\nCHARGE-AIR COOLER\n\n≥15 kW TARGET\nIAT ≤80°C TARGET\nROUTE Δp ≤10 kPa\nTEST-GATED", image_font(22, True), GREEN)
    draw.rectangle(cond, outline=hex_colour(CYAN), width=7)
    centred_text(draw, cond, "14 × 22\nR134a CONDENSER\nMEASURED PART", image_font(22, True), CYAN)
    rounded(draw, rad, LIGHT, BLUE, 12, 6)
    centred_text(draw, rad, "MEASURED-FIT\nENGINE RADIATOR\n\nBETWEEN UPRIGHTS\nPERFORMANCE / FIT\nTEST-GATED", image_font(25, True), BLUE)
    rounded(draw, shroud, PALE, NAVY, 10, 5)
    draw_fan(draw, (1870, 545, 2150, 1080), "MECHANICAL\nPULLER")
    centred_text(draw, (1825, 1100, 2200, 1230), "SEALED FULL-FACE SHROUD\n≥9,000 m³/h INSTALLED TARGET\n@ 1,500 rpm", image_font(16, True), RED)
    _air_arrows(draw, 70, 2290, (270, 1390), CYAN)
    draw.text((70, 230), "GRILLE → TWIN PUSHERS → CHARGE COOLER → GAP → 14×22 CONDENSER → GAP → RADIATOR → PULLER → ENGINE", font=image_font(17, True), fill=hex_colour(CYAN))
    centred_text(draw, (760, 315, 1510, 365), "2 ELECTRIC FRONT + 1 ENGINE-DRIVEN REAR = 3 FANS TOTAL", image_font(17, True), RED)
    dimension_h(draw, fans[0], rad[2], 300, fans[1], "M3: FRONT MODULE THROUGH RADIATOR FRONT — MEASURED FIT", PURPLE, 18)
    dimension_h(draw, ic[2], cond[0], 1350, ic[3], "≥10 SERVICE / AIR GAP", RED, 17)
    dimension_h(draw, cond[2], rad[0], 1400, cond[3], "15 TARGET / 10 ABS. MIN GAP", RED, 17)
    dimension_v(draw, rad[1], rad[3], 1745, rad[2], "M1/M2 — BETWEEN UPRIGHTS", PURPLE, "left")
    drawing_note(
        draw,
        (70, 1430, 2295, 1465),
        "REV G RELEASE RULE",
        "Use Prado/GX donor motors, blades, plugs and pigtails only in the custom centred shroud. All exchangers stay between uprights and are rubber-isolated/removable. No Prado radiator/shroud direct-fit assumption, side pack, extra fan or added width. Exactly three fans; 50°C/no-derate remains test-gated.",
        PURPLE,
    )
    finish_drawing_sheet(draw, "D01", "ALL-FRONT THREE-CORE STACK / INSTALLED-PERFORMANCE TARGETS")
    image.save(ASSET / "rev_c_d01_complete_stack.png", quality=96)


def save_d02_radiator_assembly():
    image, draw = new_drawing_sheet(
        "D02",
        "ENGINE RADIATOR — PERFORMANCE-CONTROLLED ASSEMBLY",
        "Row count and material do not release the radiator; supplier/bench thermal and pressure-drop evidence is mandatory.",
    )
    drawing_panel(draw, (45, 185, 1510, 1490), "A — FRONT / SIDE ORTHOGRAPHIC", fill="F8FAFC")
    core = (300, 390, 1180, 1240)
    rounded(draw, core, LIGHT, BLUE, 10, 6)
    for y in range(425, 1210, 28):
        draw.line((330, y, 1150, y), fill="#9DB5C7", width=2)
    top_tank = (260, 300, 1220, 430)
    bottom_tank = (260, 1210, 1220, 1340)
    rounded(draw, top_tank, "E9D6B4", GOLD, 12, 5)
    rounded(draw, bottom_tank, "E9D6B4", GOLD, 12, 5)
    draw.ellipse((1015, 240, 1110, 335), outline=hex_colour(GOLD), width=5)
    draw.rectangle((1120, 330, 1280, 410), outline=hex_colour(PURPLE), width=5)
    draw.rectangle((180, 1245, 285, 1315), outline=hex_colour(PURPLE), width=5)
    centred_text(
        draw,
        (370, 600, 1110, 1020),
        "530 × 435 ACTIVE FACE\n0.231 m² GROSS\n≥115 kW CONTINUOUS\n≥130 kW FOR 10 min\n50°C GRILLE AMBIENT + A/C",
        image_font(30, True),
        BLUE,
    )
    dimension_h(draw, core[0], core[2], 470, core[1], "530 ACTIVE CORE NOMINAL", RED)
    dimension_v(draw, core[1], core[3], 1325, core[2], "435 ACTIVE CORE NOMINAL", RED, "left")
    side = (1350, 390, 1450, 1240)
    rounded(draw, side, LIGHT, BLUE, 8, 5)
    dimension_h(draw, side[0], side[2], 1320, side[3], "64 ACTIVE CORE NOMINAL", RED, 18)
    leader_note(draw, (1055, 285), (745, 205, 1110, 360), "R01 FILLER / CAP", "Copy verified 2H filler seat, overflow and cap. Never raise cap pressure.", PURPLE)
    leader_note(draw, (1200, 360), (1130, 205, 1490, 360), "R02 UPPER NECK", "38 OD target; actual hose, centre and angle control.", PURPLE)
    leader_note(draw, (220, 1280), (55, 1190, 360, 1460), "R03 LOWER NECK", "Flow-sized formed hose; prevent suction collapse.", PURPLE)
    drawing_note(
        draw,
        (1545, 185, 2355, 780),
        "B — THERMAL REPORT MUST STATE",
        "• radiator air-in after operating condenser\n• coolant in/out temperatures\n• 50/50 mixture and coolant flow\n• coolant-side pressure loss\n• installed airflow / face-velocity grid\n• air-side pressure loss\n• voltage, fan/engine speed and duration\n• ≥60 min stable continuous condition\n• 130 kW/10 min from stable condition",
        RED,
        "FDF2F1",
    )
    drawing_note(
        draw,
        (1545, 820, 2355, 1480),
        "C — TANKS / CIRCUIT / TEST",
        "Copy actual sample for tanks, brackets, necks, filler, overflow and drain. Verify thermostat, bypass, high-point bleed, lower-hose anti-collapse and coolant flow. Pressure-test at verified system test pressure ≥5 min; flow-test uniformity. No AT oil cooler on current manual basis. Thin fin-safe coating only.",
        PURPLE,
    )
    finish_drawing_sheet(draw, "D02", "RADIATOR ASSEMBLY / PERFORMANCE + FIELD-COPIED INTERFACES")
    image.save(ASSET / "rev_c_d02_radiator_assembly.png", quality=96)


def save_d03_radiator_components():
    image, draw = new_drawing_sheet(
        "D03",
        "RADIATOR, SHROUD AND DUCT — COMPONENT / ASSEMBLY BREAKDOWN",
        "Every pressure part and structural part is separately testable, isolated and removable.",
    )
    parts = [
        ("1", "TOP TANK + VERIFIED FILLER", (80, 250, 570, 470), GOLD, "E9D6B4"),
        ("2", "PERFORMANCE-QUALIFIED CORE", (690, 250, 1230, 620), BLUE, LIGHT),
        ("3", "BOTTOM TANK + DRAIN", (1350, 250, 1840, 470), GOLD, "E9D6B4"),
        ("4", "SIDE PLATES / REMOVABLE RAILS", (80, 760, 570, 1010), NAVY, PALE),
        ("5", "LOWER SADDLES + EPDM", (690, 760, 1230, 1010), GREEN, "E4EEE8"),
        ("6", "FULL-FACE SEALED SHROUD", (1350, 700, 1840, 1080), NAVY, PALE),
        ("7", "ENGINE-DRIVEN PULLER FAN", (1900, 480, 2320, 960), GOLD, "F6EEDB"),
    ]
    for number, label, box, colour, fill in parts:
        rounded(draw, box, fill, colour, 12, 5)
        numbered_bubble(draw, (box[0] + 35, box[1] + 35), number, colour)
        centred_text(draw, (box[0] + 50, box[1] + 20, box[2] - 25, box[3] - 15), label, image_font(23, True), colour)
    arrow(draw, (570, 360), (680, 420), MUTED, 6)
    arrow(draw, (1235, 420), (1340, 360), MUTED, 6)
    arrow(draw, (570, 880), (680, 880), MUTED, 6)
    arrow(draw, (1235, 880), (1340, 880), MUTED, 6)
    arrow(draw, (1845, 850), (1890, 760), MUTED, 6)
    drawing_note(
        draw,
        (75, 1160, 1120, 1470),
        "ASSEMBLY ORDER",
        "Inspect sample → approve thermal curve → make dummies → tack removable rails/saddles → dry-fit tanks/core → fit condenser and fan module → fit shroud/fan → pressure/flow test → instrumented 50°C test.",
        PURPLE,
    )
    drawing_note(
        draw,
        (1180, 1160, 2325, 1470),
        "NEVER",
        "No through-core ties. No welding/drilling a tank, header, tube, core or fin. No forced bolt alignment. Upper tabs restrain only; lower EPDM saddles carry weight. Keep final core away from grinding/welding.",
        RED,
        "FDF2F1",
    )
    finish_drawing_sheet(draw, "D03", "RADIATOR PARTS / ASSEMBLY / SHOP PROHIBITIONS")
    image.save(ASSET / "rev_c_d03_radiator_components.png", quality=96)


def save_d04_component_dimensions():
    image, draw = new_drawing_sheet(
        "D04",
        "ALL-FRONT COMPONENTS — PUSHER MODULE, CHARGE COOLER, CONDENSER AND RADIATOR",
        "Actual outside envelopes control. The whole three-core stack stays between uprights; performance is installed-condition, not free-air.",
    )
    drawing_panel(draw, (45, 185, 770, 1475), "A — R134a CONDENSER", fill="F8FAFC")
    cond = (145, 390, 665, 1110)
    draw.rectangle(cond, outline=hex_colour(CYAN), width=7)
    for y in range(410, 1090, 20):
        draw.line((165, y, 645, y), fill="#9BCBD8", width=2)
    dimension_h(draw, cond[0], cond[2], 330, cond[1], "559 BODY NOMINAL", RED)
    dimension_v(draw, cond[1], cond[3], 710, cond[2], "356 BODY NOMINAL", RED, "left")
    centred_text(
        draw,
        (195, 485, 535, 610),
        "R134a\nPARALLEL-FLOW",
        image_font(23, True),
        CYAN,
    )
    drawing_note(
        draw,
        (90, 1170, 720, 1435),
        "DEPTH / MOUNT / GAPS",
        "21 body depth nominal; actual seams, manifolds, ports and ears field measure. "
        "4 independent rubber-isolated tabs. Carries no fan or core load. "
        "10–15 mm service/air gaps to adjacent cores; each adjacent gap is field-verified before release.",
        PURPLE,
    )

    drawing_panel(draw, (815, 185, 1570, 1475), "B — CENTRED TOYOTA FAN CANDIDATES", fill="F8FAFC")
    # Candidate arrangement: two Toyota/Denso 248 mm blades/motors symmetric
    # about the measured active-fin centreline in one custom sealed shroud.
    # Service parts are mounted behind an upright or above/rear of the carrier;
    # nothing is allowed to add front-view width or displace either fan.
    module = (850, 470, 1535, 845)
    rounded(draw, module, "F6EEDB", GOLD, 12, 6)
    centre_x = (module[0] + module[2]) // 2
    for y in range(module[1] - 25, module[3] + 45, 22):
        draw.line((centre_x, y, centre_x, min(y + 12, module[3] + 35)), fill=hex_colour(PURPLE), width=3)
    fan_d = 248
    ring_d = 258
    fan_gap = 18
    left_ring = (centre_x - fan_gap // 2 - ring_d, 530, centre_x - fan_gap // 2, 530 + ring_d)
    right_ring = (centre_x + fan_gap // 2, 530, centre_x + fan_gap // 2 + ring_d, 530 + ring_d)
    for ring in (left_ring, right_ring):
        draw.ellipse(ring, outline=hex_colour(RED), width=5)
    left_cx = (left_ring[0] + left_ring[2]) // 2
    right_cx = (right_ring[0] + right_ring[2]) // 2
    fan_y = (left_ring[1] + left_ring[3]) // 2
    left_fan = (left_cx - fan_d // 2, fan_y - fan_d // 2, left_cx + fan_d // 2, fan_y + fan_d // 2)
    right_fan = (right_cx - fan_d // 2, fan_y - fan_d // 2, right_cx + fan_d // 2, fan_y + fan_d // 2)
    draw_fan(draw, left_fan, "248 mm")
    draw_fan(draw, right_fan, "248 mm")
    centred_text(
        draw,
        (875, 275, 1510, 335),
        "SERVICE PARTS MOUNT REAR/UPPER ON SEPARATE BRACKETS — SEE D09",
        image_font(13, True),
        BLUE,
    )
    dimension_h(draw, module[0], module[2], 430, module[1], "F1 COMPLETE W — MUST FIT W0/M1", PURPLE)
    dimension_h(
        draw,
        left_cx,
        right_cx,
        875,
        left_fan[1],
        "266 C-C CAND.",
        RED,
        17,
    )
    dimension_v(draw, module[1], module[3], 1555, module[2], "F1 COMPLETE H — MUST FIT M2", PURPLE, "left")
    centred_text(draw, (900, 825, 1490, 855), "ACTIVE-FIN C/L — BOTH MOTOR CENTRES AT SAME Y", image_font(13, True), PURPLE)
    drawing_note(
        draw,
        (840, 920, 1545, 1165),
        "STANDARD TOYOTA CANDIDATE — HOLD",
        "Matched Prado 120 / GX470 family motors/blades: 248 mm swept circles, ≥258 ring openings, "
        "266 motor C-C (±133 from active-fin C/L), 524 ring-group width. One CUSTOM close sealed shroud. "
        "Centres: x = W_active/2 ±133, same y. Complete donor shrouds are not assumed to fit. Measure first.",
        PURPLE,
        "F3ECF8",
    )
    drawing_note(
        draw,
        (840, 1185, 1545, 1435),
        "F1 FIT + F3 INSTALLED DUTY",
        "No side service tower: drier behind upright; Relay Rev D and MIDI Rev D use separate staggered structural brackets per D09/E2. "
        "PASS ≥3,000 m³/h at 13.5 V through the complete three-core stack; ≥3,300 procurement target. "
        "Published 96 W / 8 A / 2,400 rpm equivalent data is identification only—not a pressure-flow pass.",
        RED,
        "FDF2F1",
    )

    drawing_panel(draw, (1615, 185, 2355, 1475), "C — SLIM HORIZONTAL CHARGE-AIR COOLER", fill="F8FAFC")
    ic = (1700, 410, 2270, 1080)
    rounded(draw, ic, "E4EEE8", GREEN, 12, 6)
    for y in range(450, 1050, 24):
        draw.line((1730, y, 2240, y), fill="#A6C8AF", width=2)
    centred_text(draw, (1750, 560, 2190, 930), "SLIM HORIZONTAL\nCHARGE-AIR COOLER\n\nBETWEEN UPRIGHTS\n≥15 kW TARGET\nIAT ≤80°C TARGET\nROUTE Δp ≤10 kPa\nALL TEST-GATED", image_font(23, True), GREEN)
    dimension_h(draw, ic[0], ic[2], 350, ic[1], "500 BODY NOMINAL", RED, 17)
    dimension_v(draw, ic[1], ic[3], 2300, ic[2], "180 BODY NOMINAL", RED, "left")
    drawing_note(draw, (1660, 1165, 2310, 1435), "REV G COMPONENT RULE", "50 body depth nominal with 57 mm beaded outlets. No side pack, side fan, service tower or extra width. Charge pipes use measured adapters/supports only; the cooler is removable and isolated. 50°C capability remains a complete-stack vehicle-test claim.", RED, "FDF2F1")
    finish_drawing_sheet(draw, "D04", "ALL-FRONT COMPONENT ENVELOPES / INSTALLED-FLOW ACCEPTANCE")
    image.save(ASSET / "rev_c_d04_component_dimensions.png", quality=96)


def save_d05_mounting_shroud():
    image, draw = new_drawing_sheet(
        "D05",
        "INDEPENDENT MOUNTS, FULL SHROUD AND SERVICE ACCESS",
        "Existing formed uprights are retained only after weld, alignment and strength inspection.",
    )
    drawing_panel(draw, (45, 185, 1120, 1480), "A — MAIN PACK SUPPORT", fill="F8FAFC")
    draw.rectangle((180, 300, 260, 1290), fill="#D9E0E6", outline=hex_colour(NAVY), width=5)
    draw.rectangle((900, 300, 980, 1290), fill="#D9E0E6", outline=hex_colour(NAVY), width=5)
    # The photographed chassis uprights finish in inward horizontal returns with
    # one existing vertical hole per side.  Show the positive through-bolt
    # pickup explicitly; the earlier generic rail could be misread as merely
    # sitting close to those returns.
    draw.rectangle((220, 285, 470, 350), fill="#C8D0D7", outline=hex_colour(NAVY), width=5)
    draw.rectangle((690, 285, 940, 350), fill="#C8D0D7", outline=hex_colour(NAVY), width=5)
    for bolt_x in (420, 740):
        draw.ellipse((bolt_x - 24, 270, bolt_x + 24, 298), fill="#63717D", outline=hex_colour(NAVY), width=3)
        draw.rectangle((bolt_x - 8, 294, bolt_x + 8, 394), fill="#63717D", outline=hex_colour(NAVY), width=2)
        draw.rectangle((bolt_x - 54, 350, bolt_x + 54, 375), fill="#6B8F71", outline=hex_colour(GREEN), width=3)
        draw.rectangle((bolt_x - 72, 375, bolt_x + 72, 403), fill="#263746", outline=hex_colour(NAVY), width=3)
        draw.polygon(
            [
                (bolt_x - 24, 405),
                (bolt_x + 24, 405),
                (bolt_x + 31, 420),
                (bolt_x + 18, 438),
                (bolt_x - 18, 438),
                (bolt_x - 31, 420),
            ],
            fill="#63717D",
            outline=hex_colour(NAVY),
        )
    # The radiator has two removable structural side rails.  Each rail carries
    # one short upper ear directly beneath the existing return hole; there is
    # deliberately no transverse top carrier.
    draw.rectangle((330, 430, 830, 1185), outline=hex_colour(BLUE), width=7)
    draw.rectangle((295, 400, 335, 1205), fill="#263746", outline=hex_colour(NAVY), width=4)
    draw.rectangle((825, 400, 865, 1205), fill="#263746", outline=hex_colour(NAVY), width=4)
    draw.rectangle((310, 375, 455, 405), fill="#263746", outline=hex_colour(NAVY), width=4)
    draw.rectangle((705, 375, 850, 405), fill="#263746", outline=hex_colour(NAVY), width=4)
    rounded(draw, (250, 1180, 380, 1290), "E4EEE8", GREEN, 10, 4)
    rounded(draw, (780, 1180, 910, 1290), "E4EEE8", GREEN, 10, 4)
    centred_text(draw, (360, 760, 800, 940), "RADIATOR WEIGHT ON\nLOWER EPDM SADDLES\nDIRECT SIDE EARS RESTRAIN", image_font(23, True), BLUE)
    leader_note(draw, (220, 420), (65, 245, 520, 365), "UPRIGHTS", "Inspect welds, distortion, squareness and attachment before paint.", PURPLE)
    leader_note(
        draw,
        (740, 330),
        (555, 205, 1090, 300),
        "2 × EXISTING TOP HOLES",
        "One positive through-bolt per side. Transfer the actual holes; do not drill or slot the chassis returns.",
        RED,
    )
    leader_note(draw, (315, 1200), (70, 1310, 540, 1450), "LOWER SADDLE", "Under the same side rail as the top ear; 3–4 mm steel + 5 mm EPDM; no tank contact.", GREEN)
    leader_note(draw, (420, 390), (555, 305, 1090, 430), "DIRECT SIDE EAR", "Short removable radiator-rail ear lands under the original hole. No transverse top crossbar.", BLUE)

    drawing_panel(draw, (1160, 185, 2355, 930), "B — REAR SHROUD / ENGINE FAN", fill="F8FAFC")
    shroud = (1260, 325, 2070, 820)
    rounded(draw, shroud, PALE, NAVY, 10, 6)
    draw_fan(draw, (1450, 385, 1900, 765), "ENGINE PULLER")
    dimension_h(draw, 1390, 1450, 865, 820, "M5 ≥20; 25–30 PREF", RED, 18)
    drawing_note(draw, (2090, 325, 2315, 820), "FAN FIT", "35–50% blade depth in opening. ≥15 radial through engine movement. Seal full radiator perimeter.", RED, "FDF2F1")

    drawing_panel(draw, (1160, 970, 2355, 1480), "C — SERVICE / ISOLATION RULE", fill="F8FAFC")
    drawing_note(
        draw,
        (1210, 1060, 2300, 1435),
        "REMOVE SEPARATELY",
        "Radiator, condenser, slim charge-air cooler, front pusher module, drier, mechanical shroud and the two electrical boxes on their own brackets each remove separately without cutting. Use rubber isolation, accessible fasteners/captive nuts and protected edges. Prove bonnet, hose, wiring, lid/cover sweeps and tool paths on the full-size dummy.",
        PURPLE,
    )
    finish_drawing_sheet(draw, "D05", "MOUNTS / LOWER SADDLES / SHROUD / SERVICEABILITY")
    image.save(ASSET / "rev_c_d05_mounting_shroud.png", quality=96)


def save_d06_side_geometry():
    image, draw = new_drawing_sheet(
        "D06",
        "ALL-FRONT SIDE SECTION — THREE-CORE DEPTH BUDGET",
        "Front/grille is left and engine is right. Actual component bodies, service gaps and nearest obstruction control M3–M5.",
    )
    drawing_panel(draw, (45, 185, 2355, 900), "A — REV G ALL-FRONT PACK / M3–M5", fill="F8FAFC")
    x = 230
    y1, y2 = 350, 670
    items = [
        ("FRONT\nCLEAR", 95, MUTED, WHITE, "ACTUAL"),
        ("TWIN\nPUSHERS", 220, GOLD, "F6EEDB", "55 MAX"),
        ("FAN/IC\nGAP", 70, MUTED, WHITE, "5 MIN\n10 PREF"),
        ("CHARGE\nCOOLER", 180, GREEN, "E4EEE8", "50"),
        ("GAP 1", 85, MUTED, WHITE, "10 MIN"),
        ("14×22\nCONDENSER", 145, CYAN, "DDECF3", "21"),
        ("GAP 2", 90, MUTED, WHITE, "15 PREF\n10 MIN"),
        ("RADIATOR", 270, BLUE, LIGHT, "64 CORE"),
        ("SHROUD", 220, NAVY, PALE, "ACTUAL"),
        ("M5", 135, RED, WHITE, "≥20\n25–30"),
        ("ENGINE\nFAN", 260, GOLD, "F6EEDB", "PULLER"),
    ]
    starts = {}
    for name, width, colour, fill, label in items:
        starts[name] = (x, x + width)
        rounded(draw, (x, y1, x + width, y2), fill, colour, 7, 4)
        centred_text(draw, (x + 4, y1 + 25, x + width - 4, y2 - 25), f"{name}\n{label}", image_font(18, True), colour)
        x += width
    _air_arrows(draw, 70, 2280, (285,), CYAN)
    dimension_h(draw, starts["TWIN\nPUSHERS"][0], starts["RADIATOR"][1], 745, y2, "NOMINAL COMPONENT PACK = 215 ABSOLUTE / 225 PREFERRED — EXTRAS NOT INCLUDED", RED, 15)
    dimension_h(draw, starts["FRONT\nCLEAR"][0], starts["RADIATOR"][1], 805, y2, "M3 AVAILABLE MUST EXCEED COMPLETE REAL PACK + GUARDS / PLUGS / CLEARANCE", RED, 16)
    drawing_note(draw, (1510, 720, 2300, 870), "GAP / DEPTH RULE", "215 = 55+5+50+10+21+10+64. Preferred 225 uses 10 fan/IC and 15 condenser/radiator. Both exclude screen, plugs, seams, brackets, pipes, tools and vibration space. Actual M3 and the full-size dummy control.", RED, "FDF2F1")
    drawing_panel(draw, (45, 950, 2355, 1480), "B — RELEASE / SERVICE / AIR-SEAL RULES", fill="F8FAFC")
    blocks = [
        (110, 1060, 600, 1400, "ALL EXCHANGERS\nBETWEEN UPRIGHTS\nRUBBER-ISOLATED\nREMOVABLE", GREEN, "E4EEE8"),
        (720, 1060, 1210, 1400, "PUSHER SHROUD\nSEALED TO\nCHARGE COOLER", GOLD, "F6EEDB"),
        (1330, 1060, 1820, 1400, "PERIMETER SEALS\nNO BYPASS AROUND\nTHREE CORES", CYAN, "DDECF3"),
        (1940, 1060, 2300, 1400, "MECHANICAL\nFULL-FACE\nPULLER SHROUD", NAVY, PALE),
    ]
    for x1, yy1, x2, yy2, label, colour, fill in blocks:
        rounded(draw, (x1, yy1, x2, yy2), fill, colour, 9, 5)
        centred_text(draw, (x1 + 10, yy1 + 10, x2 - 10, yy2 - 10), label, image_font(20, True), colour)
    _air_arrows(draw, 70, 2310, (1005, 1450), CYAN)
    centred_text(draw, (280, 1430, 2200, 1470), "EXACTLY 3 FANS — NO EXTRA ELECTRIC/SIDE FAN / WIDTH — D08 B0/H0-L/H0-R/W0/P0/R0 CONTROLS", image_font(16, True), PURPLE)
    finish_drawing_sheet(draw, "D06", "ALL-FRONT DEPTH GATES / GAPS / REMOVABILITY / SEALING")
    image.save(ASSET / "rev_c_d06_side_geometry.png", quality=96)


def save_d07_fan_wiring():
    image, draw = new_drawing_sheet(
        "D07",
        "ELECTRIC FAN POWER, CONTROL, FAULT INDICATION AND AIRFLOW",
        "Every motor branch is separately protected; final fuse/wire values follow measured current and voltage drop.",
    )
    drawing_panel(draw, (45, 185, 2355, 1050), "A — POWER DISTRIBUTION", fill="F8FAFC")
    draw.text(
        (390, 250),
        "TWO CLOSED ELECTRICAL BOXES ON SEPARATE STAGGERED BRACKETS — D09 / E2 / E2a",
        font=image_font(19, True),
        fill=hex_colour(BLUE),
    )
    rounded(draw, (90, 370, 330, 900), LIGHT, BLUE, 12, 5)
    centred_text(draw, (100, 380, 320, 890), "BATTERY-SIDE\nMASTER CUTOFF\n+ PROTECTED FEED\n13.5 V TEST", image_font(20, True), BLUE)
    rounded(draw, (430, 315, 750, 835), "FFF7F7", RED, 14, 5)
    rounded(draw, (780, 315, 1090, 835), "FFF9EC", GOLD, 14, 5)
    centred_text(
        draw,
        (445, 705, 735, 815),
        "CLOSED GASKETED MIDI BOX\nALL FUSES / HOLDERS / LIVE STUDS INSIDE\nLID SHUT + LATCHED IN USE",
        image_font(13, True),
        RED,
    )
    centred_text(
        draw,
        (795, 705, 1075, 815),
        "COVERED RELAY BOX\nNO EXPOSED LIVE TERMINALS\nCOVER SHUT IN USE",
        image_font(13, True),
        GOLD,
    )
    branches = [
        (450, "FRONT MOTOR 1", "A/C + TRINARY"),
        (650, "FRONT MOTOR 2*", "A/C + TRINARY"),
    ]
    for y, motor, trigger in branches:
        rounded(draw, (470, y - 100, 690, y + 35), "FDF2F1", RED, 10, 4)
        centred_text(draw, (480, y - 92, 680, y + 27), "FUSE\nFROM MEASURED I", image_font(17, True), RED)
        rounded(draw, (820, y - 100, 1030, y + 35), "F6EEDB", GOLD, 10, 4)
        relay_label = "90987-02027 CAND.\nOR SEALED ISO 40 A" if motor.startswith("FRONT") else "SEALED\nRELAY"
        centred_text(draw, (830, y - 92, 1020, y + 27), relay_label, image_font(14, True), GOLD)
        rounded(draw, (1170, y - 100, 1470, y + 35), "F6EEDB", GOLD, 10, 4)
        centred_text(draw, (1180, y - 92, 1460, y + 27), motor, image_font(18, True), GOLD)
        rounded(draw, (800, y + 45, 1050, y + 88), "F3ECF8", PURPLE, 8, 3)
        centred_text(draw, (808, y + 48, 1042, y + 85), trigger, image_font(13, True), PURPLE)
        arrow(draw, (335, 635), (460, y - 32), RED, 6)
        arrow(draw, (700, y - 32), (810, y - 32), RED, 6)
        arrow(draw, (1040, y - 32), (1160, y - 32), GOLD, 6)
        arrow(draw, (925, y + 44), (925, y + 36), PURPLE, 4)
    rounded(draw, (2050, 370, 2290, 900), "E4EEE8", GREEN, 12, 5)
    centred_text(draw, (2060, 380, 2280, 890), "CLEAN\nGROUND\nSTUD\nEQUAL-CAPACITY\nRETURN", image_font(20, True), GREEN)
    for y, _, _ in branches:
        arrow(draw, (1480, y - 32), (2040, y - 32), GREEN, 5)
    arrow(draw, (590, 840), (590, 900), RED, 5)
    arrow(draw, (935, 840), (935, 900), GOLD, 5)
    centred_text(draw, (420, 895, 1110, 935), "DOWNWARD / REARWARD CABLE EXITS THROUGH GLANDS — DRIP LOOPS BELOW BOXES", image_font(14, True), NAVY)
    draw.text((90, 965), "Rev G uses two front motors, each with its own protected branch. No exposed fuse, holder or live fuse stud is permitted.", font=image_font(17), fill=hex_colour(MUTED))

    drawing_panel(draw, (45, 1100, 1120, 1480), "B — AIRFLOW POLARITY", fill="F8FAFC")
    draw.text((90, 1180), "MAIN:", font=image_font(20, True), fill=hex_colour(NAVY))
    arrow(draw, (220, 1200), (1000, 1200), CYAN, 8)
    draw.text((220, 1225), "GRILLE → TWIN PUSHERS → CHARGE COOLER → CONDENSER → RADIATOR → ENGINE", font=image_font(15, True), fill=hex_colour(CYAN))
    draw.text((90, 1320), "PULLER:", font=image_font(20, True), fill=hex_colour(NAVY))
    arrow(draw, (220, 1340), (1000, 1340), GREEN, 8)
    draw.text((220, 1365), "MECHANICAL PULLER: RADIATOR → SEALED SHROUD → ENGINE", font=image_font(17, True), fill=hex_colour(GREEN))
    drawing_panel(draw, (1160, 1100, 2355, 1480), "C — LOADED ELECTRICAL TEST", fill="F8FAFC")
    drawing_note(draw, (1210, 1170, 2300, 1435), "PASS", "Hot idle with A/C, cabin blower, lights and both pushers ON. Each motor within 0.5 V of battery/alternator unless supplier tighter. Record run/start current, voltage, alternator output and connector/relay temperature. Provide fault indication for each front-fan branch.", RED, "FDF2F1")
    finish_drawing_sheet(draw, "D07", "CLOSED FUSE/RELAY BOXES / FAN POWER / CONTROL / HOT-IDLE TEST")
    image.save(ASSET / "rev_c_d07_fan_wiring.png", quality=96)


def save_d08_existing_top_hole_mount():
    image, draw = new_drawing_sheet(
        "D08",
        "RADIATOR TO EXISTING SIDE POSTS — DIRECT TOP-EAR ATTACHMENT",
        "The radiator side rails and lower saddles are set from the real bracket holes. No transverse top carrier and no chassis-hole alteration.",
    )

    drawing_panel(draw, (45, 185, 1190, 1480), "A — BRACKET-FIRST FRONT ELEVATION / COMPLETE LOAD PATH", fill="F8FAFC")

    # Include the two factual top-return crops so the fabricator can recognise
    # the exact existing holes being controlled by this sheet.
    source_photo = ROOT / "photos" / "20260722_000001_user_second_radiator_arm_welded_front_structure.png"

    def paste_crop(crop_box, target_box):
        if not source_photo.exists():
            return
        with Image.open(source_photo) as source:
            crop = source.convert("RGB").crop(crop_box)
        tx1, ty1, tx2, ty2 = target_box
        max_w, max_h = tx2 - tx1, ty2 - ty1
        scale = min(max_w / crop.width, max_h / crop.height)
        resized = crop.resize(
            (max(1, int(crop.width * scale)), max(1, int(crop.height * scale))),
            Image.Resampling.LANCZOS,
        )
        px = int(tx1 + (max_w - resized.width) / 2)
        py = int(ty1 + (max_h - resized.height) / 2)
        image.paste(resized, (px, py))
        draw.rectangle((tx1, ty1, tx2, ty2), outline=hex_colour(NAVY), width=4)

    paste_crop((250, 40, 950, 460), (80, 255, 585, 500))
    paste_crop((2000, 40, 2900, 460), (650, 255, 1155, 500))
    centred_text(draw, (80, 505, 585, 540), "PHOTO: LEFT EXISTING TOP HOLE", image_font(16, True), PURPLE)
    centred_text(draw, (650, 505, 1155, 540), "PHOTO: RIGHT EXISTING TOP HOLE", image_font(16, True), PURPLE)

    # Grey is existing chassis metal. Black is one removable structural rail
    # on each radiator side. Each short ear lands under one existing hole.
    draw.rectangle((120, 625, 210, 1325), fill="#C8D0D7", outline=hex_colour(NAVY), width=5)
    draw.rectangle((1040, 625, 1130, 1325), fill="#C8D0D7", outline=hex_colour(NAVY), width=5)
    draw.rectangle((165, 585, 500, 650), fill="#C8D0D7", outline=hex_colour(NAVY), width=5)
    draw.rectangle((740, 585, 1085, 650), fill="#C8D0D7", outline=hex_colour(NAVY), width=5)
    bolt_centres = (440, 800)
    for bolt_x in bolt_centres:
        draw.ellipse((bolt_x - 27, 565, bolt_x + 27, 598), fill="#63717D", outline=hex_colour(NAVY), width=3)
        draw.rectangle((bolt_x - 8, 594, bolt_x + 8, 735), fill="#63717D", outline=hex_colour(NAVY), width=2)
        draw.rectangle((bolt_x - 60, 650, bolt_x + 60, 680), fill="#6B8F71", outline=hex_colour(GREEN), width=3)
        draw.polygon(
            [
                (bolt_x - 25, 720),
                (bolt_x + 25, 720),
                (bolt_x + 34, 736),
                (bolt_x + 20, 756),
                (bolt_x - 20, 756),
                (bolt_x - 34, 736),
            ],
            fill="#63717D",
            outline=hex_colour(NAVY),
        )
    # Short direct ears and vertical radiator side rails. No cross-car bar.
    draw.rectangle((330, 680, 460, 720), fill="#263746", outline=hex_colour(NAVY), width=5)
    draw.rectangle((780, 680, 910, 720), fill="#263746", outline=hex_colour(NAVY), width=5)
    draw.rectangle((315, 705, 355, 1220), fill="#263746", outline=hex_colour(NAVY), width=5)
    draw.rectangle((885, 705, 925, 1220), fill="#263746", outline=hex_colour(NAVY), width=5)
    draw.rectangle((355, 755, 885, 1190), fill="#E8EEF5", outline=hex_colour(BLUE), width=7)
    centred_text(
        draw,
        (390, 835, 850, 1070),
        "530 × 435 ACTIVE CORE BASIS\n\nCOMPLETE R0 ENVELOPE\nCOPIED FROM SAMPLE + FIXTURE",
        image_font(23, True),
        BLUE,
    )
    for saddle_x in (335, 905):
        rounded(draw, (saddle_x - 65, 1195, saddle_x + 65, 1310), "E4EEE8", GREEN, 10, 4)
    draw.rectangle((245, 1300, 995, 1345), fill="#C8D0D7", outline=hex_colour(NAVY), width=5)
    dimension_h(draw, bolt_centres[0], bolt_centres[1], 555, 585, "B0 — ACTUAL TOP-HOLE C-C", PURPLE, 18)
    dimension_h(draw, 210, 1040, 1345, 1345, "W0 — MINIMUM CLEAR WIDTH BETWEEN POST INNER FACES", PURPLE, 16)
    dimension_v(draw, 617, 1300, 1150, 1130, "H0-L / H0-R — HOLE PLANE TO SADDLE SEAT", PURPLE, "left")
    drawing_note(
        draw,
        (80, 1360, 1155, 1465),
        "BRACKET BASIS + FIT RULE",
        "Existing bracket: 48 face / 410 upright / 58 inward return; verify B0/H0-L/H0-R/W0/P0. Set saddles first. Both bolts enter fully by hand with ≥5 mm rail/post clearance and zero pull. The 410 upright is not radiator height.",
        RED,
        "FDF2F1",
    )

    drawing_panel(draw, (1230, 185, 2355, 1010), "B — SECTION THROUGH ONE DIRECT RADIATOR EAR", fill="F8FAFC")
    bolt_x = 1780
    # M8 head and large washer above the existing return.
    draw.polygon(
        [(1735, 285), (1825, 285), (1850, 315), (1825, 345), (1735, 345), (1710, 315)],
        fill="#63717D",
        outline=hex_colour(NAVY),
    )
    draw.rectangle((1685, 345, 1875, 365), fill="#8E9AA3", outline=hex_colour(NAVY), width=3)
    draw.rectangle((bolt_x - 10, 340, bolt_x + 10, 680), fill="#63717D", outline=hex_colour(NAVY), width=2)
    # Existing chassis return and its retained hole.
    draw.rectangle((1390, 365, 2200, 430), fill="#C8D0D7", outline=hex_colour(NAVY), width=5)
    draw.rectangle((bolt_x - 15, 365, bolt_x + 15, 430), fill="#FFFFFF", outline=hex_colour(PURPLE), width=3)
    # Controlled isolation gap with crush sleeve around the bolt.
    draw.rectangle((1630, 430, 1930, 482), fill="#6B8F71", outline=hex_colour(GREEN), width=4)
    draw.rectangle((bolt_x - 18, 430, bolt_x + 18, 482), fill="#8E9AA3", outline=hex_colour(NAVY), width=3)
    # One short 4 mm ear is bolted to/welded on the removable radiator side
    # rail. A small independent transfer plate is allowed only if the sampled
    # radiator ear cannot reach the hole without changing the chassis.
    draw.rectangle((1510, 482, 2050, 525), fill="#263746", outline=hex_colour(NAVY), width=5)
    draw.rectangle((1988, 520, 2050, 790), fill="#263746", outline=hex_colour(NAVY), width=5)
    draw.polygon([(1885, 525), (1988, 525), (1988, 630)], fill="#52606A", outline=hex_colour(NAVY))
    draw.rectangle((1940, 720, 2110, 790), fill="#E8EEF5", outline=hex_colour(BLUE), width=5)
    # Washer and nut under the direct side ear.
    draw.rectangle((1695, 525, 1865, 548), fill="#8E9AA3", outline=hex_colour(NAVY), width=3)
    draw.polygon(
        [(1735, 548), (1825, 548), (1850, 578), (1825, 608), (1735, 608), (1710, 578)],
        fill="#63717D",
        outline=hex_colour(NAVY),
    )
    draw.line((1345, 398, 1605, 398), fill=hex_colour(PURPLE), width=4)
    draw.text((1285, 370), "EXISTING CHASSIS\nTOP RETURN — KEEP", font=image_font(18, True), fill=hex_colour(PURPLE))
    leader_note(
        draw,
        (1800, 315),
        (1915, 240, 2315, 355),
        "POSITIVE FIXING",
        "M8 × 1.25 class 8.8 provisional; large washer above, prevailing-torque nut below.",
        RED,
    )
    leader_note(
        draw,
        (1900, 455),
        (1990, 385, 2320, 520),
        "ISOLATOR",
        "6 mm EPDM, 60–65 Shore A, with hard steel crush sleeve; no uncontrolled rubber squash.",
        GREEN,
    )
    leader_note(
        draw,
        (2030, 520),
        (2040, 550, 2320, 735),
        "EAR-L / EAR-R",
        "4 mm ear on the removable side rail. Separate transfer plate only after sample fit. No crossbar.",
        BLUE,
    )
    dimension_h(draw, 1510, 2050, 850, 790, "EAR LENGTH FROM 1:1 TEMPLATE — NOT CATALOGUE", PURPLE, 16)
    drawing_note(
        draw,
        (1280, 885, 2305, 975),
        "HOLE RULE",
        "Round ear/transfer-plate hole only: Ø9.0 +0.2 if the measured chassis hole safely accepts M8. Otherwise revise bolt, ear hole and sleeve together. No slot or reaming in chassis.",
        RED,
        "FDF2F1",
    )

    drawing_panel(draw, (1230, 1050, 2355, 1480), "C — FIELD TRANSFER / RELEASE CHECK", fill="F8FAFC")
    drawing_note(
        draw,
        (1275, 1120, 1765, 1435),
        "B0 / H0-L / H0-R / W0 / P0 / R0",
        "B0 hole C-C + rigid 1:1 template\nH0-L and H0-R: each hole plane to saddle seat\nW0 inside clear width top/mid/low\nP0 return/post plane to radiator rail\nR0 complete radiator W×H×D, ear and locator pattern\nAlso record hole Ø/edge/tab/bonnet clearance",
        PURPLE,
        "F3ECF8",
    )
    drawing_note(
        draw,
        (1795, 1120, 2310, 1435),
        "KARIGAR FIT SEQUENCE",
        "1 Copy holes with one rigid 1:1 template.\n2 Measure complete radiator and old ear pattern.\n3 Set two lower saddles under side rails.\n4 Tack EAR-L/R directly under returns.\n5 Refit: both bolts hand-enter with zero pull.\n6 Check bonnet/cap/hose/tool/fan clearances.\n7 Finish-weld off vehicle; coat after inspection.",
        GREEN,
        "E4EEE8",
    )
    finish_drawing_sheet(draw, "D08", "SIDE POSTS / DIRECT RADIATOR EARS / BRACKET-FIRST DIMENSION CONTROL")
    image.save(ASSET / "rev_d_d08_existing_top_hole_mount.png", quality=96)


def save_d09_independent_electrical_mounts():
    image, draw = new_drawing_sheet(
        "D09",
        "CLOSED FUSE + RELAY ENCLOSURES — TWO INDEPENDENT STAGGERED MOUNTS",
        "Each donor enclosure keeps its own footprint and service sweep. Neither box, bracket or hood may load the radiator or widen the front pack.",
    )

    drawing_panel(draw, (45, 185, 1540, 1030), "A — FRONT VIEW / NESTED WIDTH, STAGGERED HEIGHT + DEPTH — NOT TO SCALE", fill="F8FAFC")
    active = (190, 750, 1390, 855)
    draw.rectangle(active, fill="#E8EEF5", outline=hex_colour(BLUE), width=6)
    for y in range(active[1] + 20, active[3] - 10, 22):
        draw.line((active[0] + 20, y, active[2] - 20, y), fill="#B7C8D5", width=2)
    centred_text(draw, active, "ACTIVE FIN / SEALED AIRFLOW APERTURE\nKEEP COMPLETELY CLEAR", image_font(26, True), BLUE)
    draw.rectangle((145, 290, 205, 980), fill="#C8D0D7", outline=hex_colour(NAVY), width=5)
    draw.rectangle((1375, 290, 1435, 980), fill="#C8D0D7", outline=hex_colour(NAVY), width=5)
    # In front elevation the smaller MIDI box is wholly inside the Relay box
    # width.  The real parts are vertically and/or depth staggered on separate
    # local brackets; there is no shared plate or cross-car carrier.
    draw.line((790, 245, 790, 855), fill="#8A9BA8", width=3)
    # Two deliberately separate local bracket stubs.  They terminate at
    # different non-radiator structural pickup symbols so the front view
    # cannot be mistaken for one shared cross-car carrier.
    draw.line((420, 405, 350, 405, 350, 445), fill=hex_colour(NAVY), width=10)
    draw.ellipse((334, 429, 366, 461), fill="#FFFFFF", outline=hex_colour(NAVY), width=6)
    draw.line((1010, 620, 1080, 620, 1080, 665), fill=hex_colour(NAVY), width=10)
    draw.ellipse((1064, 649, 1096, 681), fill="#FFFFFF", outline=hex_colour(NAVY), width=6)
    draw.polygon([(405, 235), (1175, 235), (1140, 270), (440, 270)], fill="#B8C1C8", outline=hex_colour(NAVY))
    draw.polygon([(555, 465), (1025, 465), (995, 500), (585, 500)], fill="#B8C1C8", outline=hex_colour(NAVY))
    relay_plate = (420, 270, 1160, 500)
    midi_plate = (570, 500, 1010, 700)
    rounded(draw, relay_plate, "D9E1E8", NAVY, 12, 6)
    rounded(draw, midi_plate, "D9E1E8", NAVY, 12, 6)
    relay = (445, 295, 1135, 475)
    rounded(draw, relay, "F6EEDB", GOLD, 10, 4)
    centred_text(draw, relay, "COVERED RELAY REV D BOX\n360 × 245 × 3 BASE / 300 × 197 × 3 INSULATOR\nCOVER SHUT IN OPERATION", image_font(15, True), GOLD)
    midi = (590, 520, 990, 680)
    rounded(draw, midi, "F3ECF8", PURPLE, 10, 4)
    centred_text(draw, midi, "CLOSED GASKETED MIDI REV D\n210 × 165 × 65 BODY\n230 × 185 LID ENVELOPE\nALL FUSES + LIVE STUDS INSIDE", image_font(13, True), PURPLE)
    centred_text(draw, (405, 240, 1175, 268), "RELAY: OWN HOOD + UPPER/FORWARD LOCAL BRACKET", image_font(12, True), BLUE)
    centred_text(draw, (555, 470, 1025, 498), "MIDI: OWN HOOD + LOWER/REAR LOCAL BRACKET", image_font(11, True), PURPLE)
    centred_text(draw, (230, 445, 455, 470), "PICKUP A — NOT RADIATOR", image_font(9, True), NAVY)
    centred_text(draw, (1000, 665, 1230, 690), "PICKUP B — NOT RADIATOR", image_font(9, True), NAVY)
    draw.line((420, 250, 420, 730), fill=hex_colour(RED), width=3)
    draw.line((1160, 250, 1160, 730), fill=hex_colour(RED), width=3)
    centred_text(draw, (380, 705, 1200, 740), "CENTRED ON PACK C/L — PROJECTED WIDTH ≤ RELAY 360 mm ENVELOPE", image_font(12, True), RED)
    arrow(draw, (675, 680), (675, 730), PURPLE, 4)
    arrow(draw, (905, 475), (905, 730), RED, 4)
    centred_text(draw, (230, 725, 1350, 752), "CABLES EXIT DOWN / REAR — BOTH BOXES AND HOODS STAY INSIDE W0 / UPRIGHTS", image_font(12, True), RED)
    drawing_note(
        draw,
        (90, 875, 1490, 1022),
        "NO EXTRA SIDE WIDTH — NEVER SIDE-BY-SIDE ACROSS THE CAR",
        "Keep both complete assemblies inside W0 and the existing uprights. Stagger MIDI behind or below Relay so its front projection stays inside the Relay 360 mm width. Use two local brackets only: no common plate, cross-car carrier or radiator load. Prove both positions at E2/E2a.",
        RED,
        "FDF2F1",
    )

    drawing_panel(draw, (1580, 185, 2355, 1030), "B — CONTROLLED COMPONENT ENVELOPES", fill="F8FAFC")
    drawing_note(
        draw,
        (1625, 255, 2310, 555),
        "COVERED RELAY REV D — REUSE",
        "360 W × 245 H × 3 aluminium base\n300 W × 197 H × 3 insulating sheet\nSeparate closed cover; no exposed live terminals\n\nCover stays shut in operation. Mock its opening, terminals, boots, tools, drip protection and heat clearance with the actual assembly.",
        GOLD,
        "F6EEDB",
    )
    drawing_note(
        draw,
        (1625, 585, 2310, 995),
        "CLOSED MIDI REV D — REUSE",
        "210 W × 165 H × 65 body\n230 W × 185 H lid envelope\n140 × 85 × 5 insulating subplate\n1 input + 5 output glands; 28 mm doubled-cable exit\n\nEvery fuse, holder, bus and live fuse stud is inside. Lid stays shut/latched in operation; mock lid swing and fuse withdrawal.",
        PURPLE,
        "F3ECF8",
    )

    drawing_panel(draw, (45, 1070, 2355, 1480), "C — TWO LOCAL MOUNTS / WEATHER / SERVICE / E2 + E2a RELEASE", fill="F8FAFC")
    blocks = [
        (90, 1160, 540, 1415, "BATTERY-SIDE\nMASTER CUTOFF\nSTAYS AT BATTERY", BLUE, LIGHT),
        (650, 1160, 1050, 1415, "2 LOCAL BRACKETS\n3 mm 5052-H32 Al\n20 mm FOLDED RETURNS\nFOLLOW REAL BOXES", RED, "FDF2F1"),
        (1160, 1160, 1560, 1415, "2 HOODS 1.2–1.5 Al\n25 mm TURNED EDGES\nDOWN/REAR GLANDS\nBAFFLED DRAINS", GREEN, "E4EEE8"),
        (1670, 1160, 2310, 1415, "E2/E2a PASS EACH BOX\nACTUAL LID/COVER SWEEP\n≥25 CABLE/COVER CLEARANCE\nHOSE-SPRAY + DUST: DRY/CLEAN INSIDE\nNO IP CLAIM UNLESS RATED/TESTED", PURPLE, "F3ECF8"),
    ]
    for x1, y1, x2, y2, label, colour, fill in blocks:
        rounded(draw, (x1, y1, x2, y2), fill, colour, 10, 5)
        centred_text(draw, (x1 + 10, y1 + 10, x2 - 10, y2 - 10), label, image_font(17, True), colour)
    arrow(draw, (550, 1288), (640, 1288), RED, 6)
    arrow(draw, (1060, 1288), (1150, 1288), GREEN, 6)
    arrow(draw, (1570, 1288), (1660, 1288), PURPLE, 6)

    finish_drawing_sheet(draw, "D09", "TWO CLOSED BOXES / INDEPENDENT BRACKETS / E2 + E2a FIT AND WEATHER GATES")
    image.save(ASSET / "rev_g_d09_independent_electrical_mounts.png", quality=96)


def save_d10_procurement_fit_control():
    image, draw = new_drawing_sheet(
        "D10",
        "REV G PROCUREMENT DIMENSIONS AND PHYSICAL-FIT RELEASE",
        "Fixed nominal part bodies support sourcing; the actual complete parts, vehicle measurements and full-size dummy control cutting.",
    )

    drawing_panel(draw, (45, 185, 1125, 780), "A — PURCHASE / SAMPLE DIMENSIONS", fill="F8FAFC")
    rows = [
        ("RADIATOR", "530 × 435 × 64 active core; R0 outside TBD", BLUE, LIGHT),
        ("CONDENSER", "559 W × 356 H × 21 D body", CYAN, "DDECF3"),
        ("CHARGE COOLER", "500 W × 180 H × 50 D; 57 outlets", GREEN, "E4EEE8"),
        ("FRONT FAN PARTS", "2 × 248 sweep; ≥258 rings; 266 C-C", GOLD, "F6EEDB"),
        ("ELECTRICAL BOXES", "Relay 360×245 base; MIDI 210×165×65", PURPLE, "F3ECF8"),
    ]
    for index, (label, value, colour, fill) in enumerate(rows):
        y1 = 245 + index * 82
        rounded(draw, (90, y1, 1080, y1 + 62), fill, colour, 9, 4)
        draw.text((115, y1 + 14), label, font=image_font(15, True), fill=hex_colour(colour))
        draw.text((390, y1 + 14), value, font=image_font(15, True), fill=hex_colour(INK))
    drawing_note(
        draw,
        (90, 660, 1080, 760),
        "DONOR BOUNDARY",
        "Prado 120 / GX470-family parts are physical fan candidates only. Buy two matching measured motors/blades/plugs; do not assume one donor supplies a pair, or that a Prado radiator/shroud fits.",
        RED,
        "FDF2F1",
    )

    drawing_panel(draw, (1160, 185, 2355, 780), "B — CENTRED FRONT GEOMETRY", fill="F8FAFC")
    cond_left, cond_right = 1280, 2230
    cond_top, cond_bottom = 315, 650
    draw.rectangle((cond_left, cond_top, cond_right, cond_bottom), fill="#DDECF3", outline=hex_colour(CYAN), width=6)
    centre_x = (cond_left + cond_right) / 2
    draw.line((centre_x, 260, centre_x, 710), fill=hex_colour(PURPLE), width=4)
    scale = (cond_right - cond_left) / 559
    radiator_width = 530 * scale
    dashed_rectangle(
        draw,
        (centre_x - radiator_width / 2, cond_top + 10, centre_x + radiator_width / 2, cond_bottom - 10),
        BLUE,
        4,
    )
    ring_d = 258 * scale
    first_centre = cond_left + 146.5 * scale
    second_centre = cond_left + 412.5 * scale
    fan_y = (cond_top + cond_bottom) / 2
    for index, fan_x in enumerate((first_centre, second_centre), start=1):
        draw_fan(
            draw,
            (fan_x - ring_d / 2, fan_y - ring_d / 2, fan_x + ring_d / 2, fan_y + ring_d / 2),
            f"FAN {index}",
        )
    dimension_h(draw, cond_left, cond_right, 275, cond_top, "559 CONDENSER BODY", RED, 17)
    dimension_h(draw, first_centre, second_centre, 690, cond_bottom, "266 MOTOR C-C", RED, 17)
    centred_text(draw, (1270, 705, 2240, 750), "CENTRES FROM CONDENSER LEFT: x = 146.5 / 412.5; y = 178", image_font(16, True), PURPLE)
    centred_text(draw, (1290, 615, 2215, 646), "530 RADIATOR ACTIVE-WIDTH DATUM (DASHED); 524 RING GROUP; 17.5 SIDE BANDS", image_font(13, True), NAVY)

    drawing_panel(draw, (45, 825, 1440, 1480), "C — FRONT-TO-REAR COMPONENT BUDGET", fill="F8FAFC")
    x = 100
    items = [
        ("FANS", 55, 165, GOLD, "F6EEDB"),
        ("GAP", 5, 75, MUTED, WHITE),
        ("IC", 50, 150, GREEN, "E4EEE8"),
        ("GAP", 10, 80, MUTED, WHITE),
        ("COND", 21, 120, CYAN, "DDECF3"),
        ("GAP", 10, 80, MUTED, WHITE),
        ("RAD", 64, 180, BLUE, LIGHT),
    ]
    for label, nominal, width, colour, fill in items:
        rounded(draw, (x, 980, x + width, 1220), fill, colour, 7, 4)
        centred_text(draw, (x + 3, 995, x + width - 3, 1205), f"{label}\n{nominal}", image_font(17, True), colour)
        x += width
    dimension_h(draw, 100, x, 1265, 1220, "215 ABSOLUTE COMPONENT STACK", RED, 17)
    drawing_note(
        draw,
        (90, 1320, 1395, 1445),
        "PREFERRED / COMPLETE DEPTH",
        "225 preferred uses 10 fan/IC gap and 15 condenser/radiator gap. Both figures exclude stone screen, connectors, seams, brackets, pipes, tools and vibration clearance. M3 and the complete dummy must exceed the real finished pack.",
        RED,
        "FDF2F1",
    )

    drawing_panel(draw, (1480, 825, 2355, 1480), "D — DO NOT CUT UNTIL THESE PASS", fill="F8FAFC")
    drawing_note(
        draw,
        (1525, 880, 2310, 1045),
        "W0 / R0 / M3",
        "W0 is the smallest clear post width. R0 is the COMPLETE radiator W×H×D, ears and lower locators—not only the 530×435×64 active core. M3 fits the real depth and service clearances.",
        PURPLE,
        "F3ECF8",
    )
    drawing_note(
        draw,
        (1525, 1070, 2310, 1195),
        "B0 / H0-L / H0-R / DIRECT SIDE EARS",
        "Copy both original inward top-return holes with one rigid 1:1 template. Match H0-L and H0-R with lower saddles under the same radiator side rails. Both bolts hand-enter; no crossbar or chassis alteration.",
        GREEN,
        "E4EEE8",
    )
    drawing_note(
        draw,
        (1525, 1205, 2310, 1325),
        "E2 / E2a — TWO BOXES",
        "Mount the actual closed boxes independently on two local structural brackets and stagger vertically and/or in depth. Prove each lid/cover sweep and ≥25 cable clearance. No shared tray or common plate. Interiors stay dry/clean.",
        PURPLE,
        "F3ECF8",
    )
    drawing_note(
        draw,
        (1525, 1345, 2310, 1455),
        "FINAL RELEASE HOLD",
        "Do not buy final radiator/core or drill/coat until B0/H0-L/H0-R/W0/P0/R0, BRKT, M1–M6 and E2/E2a are signed. AI images are not evidence.",
        RED,
        "FDF2F1",
    )
    finish_drawing_sheet(draw, "D10", "PROCUREMENT DIMENSIONS / CENTRED FANS / DEPTH BUDGET / FIT RELEASE")
    image.save(ASSET / "rev_g_d10_procurement_fit_control.png", quality=96)


def save_dimensioned_drawing_set():
    save_d01_complete_stack()
    save_d02_radiator_assembly()
    save_d03_radiator_components()
    save_d04_component_dimensions()
    save_d05_mounting_shroud()
    save_d06_side_geometry()
    save_d07_fan_wiring()
    save_d08_existing_top_hole_mount()
    save_d09_independent_electrical_mounts()
    save_d10_procurement_fit_control()


DRAWING_SHEETS = [
    ("D01", "Complete Rev G all-front stack", "rev_c_d01_complete_stack.png", "Three-core all-front airflow path and installed-performance targets."),
    ("D02", "Performance-controlled radiator", "rev_c_d02_radiator_assembly.png", "Thermal duty, core selection and field-copied interfaces."),
    ("D03", "Radiator component breakdown", "rev_c_d03_radiator_components.png", "Pressure parts, mounts, shroud and assembly sequence."),
    ("D04", "All-front component envelopes", "rev_c_d04_component_dimensions.png", "Pusher, charge-air cooler, condenser and radiator fit/flow controls."),
    ("D05", "Mounting, shroud and service access", "rev_c_d05_mounting_shroud.png", "Independent mounts, saddles and mechanical-fan geometry."),
    ("D06", "All-front depth and gap gates", "rev_c_d06_side_geometry.png", "M3/M4 formula, service gaps, seals and removable three-core stack."),
    ("D07", "Fan power and control", "rev_c_d07_fan_wiring.png", "Protected branches, control, fault indication and polarity."),
    ("D08", "Direct side-ear bracket interface", "rev_d_d08_existing_top_hole_mount.png", "Existing side posts, direct radiator ears, lower-saddle load path and B0/H0-L/H0-R/W0/P0/R0 controls."),
    ("D09", "Closed boxes on independent brackets", "rev_g_d09_independent_electrical_mounts.png", "Relay Rev D and MIDI Rev D on separate staggered mounts, cable exits and E2/E2a gates."),
    ("D10", "Procurement and physical-fit control", "rev_g_d10_procurement_fit_control.png", "Fixed nominal purchase dimensions, centred fan coordinates, depth budget and signed fit-release boundary."),
]


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("J40 Cooling System  |  Rev G  |  ")
    set_run(run, 9, False, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_landscape_drawing_appendix(document):
    add_heading(document, "Appendix A. Controlled Rev G drawing sheets", 1)
    add_paragraph(
        document,
        "These ten sheets show the Rev G Toyota-donor-led all-front architecture. Red values are fixed performance/release limits. "
        "Purple dimensions and interfaces are measured from the vehicle or actual component. Written values "
        "control; do not scale the images.",
    )
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.40)
    section.right_margin = Inches(0.40)
    section.header_distance = Inches(0.15)
    section.footer_distance = Inches(0.15)
    for index, (code, title, filename, _) in enumerate(DRAWING_SHEETS):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"{code} — {title}")
        set_run(run, 10, True, NAVY)
        document.add_picture(str(ASSET / filename), width=Inches(9.8))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.paragraphs[-1].paragraph_format.space_after = Pt(1)
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(0)
        run = caption.add_run("Written requirements govern; red = fixed/release limit; purple = field measure / owner approval.")
        set_run(run, 8, False, MUTED, True)
        if index < len(DRAWING_SHEETS) - 1:
            document.add_page_break()


def build_document_rev_e_reference():
    """Historical Rev E document builder; never use for current fabrication output."""
    save_dimensioned_drawing_set()
    document = Document()
    configure_document(document)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run("PAKISTAN FABRICATOR SPECIFICATION")
    set_run(r, 11, True, GOLD)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("J40 Integrated Radiator\nand Cooling System")
    set_run(r, 27, True, NAVY)
    p = document.add_paragraph()
    r = p.add_run("Toyota 2H • R134a A/C • 50°C continuous design ambient • separate charge-air path")
    set_run(r, 12.5, False, MUTED)
    p = document.add_paragraph()
    r = p.add_run("REV E  |  01 AUGUST 2026  |  ALL DIMENSIONS IN mm")
    set_run(r, 10, True, BLUE)
    add_callout(
        document,
        "DESIGN REQUIREMENT — NOT YET A PROVEN RATING",
        "Final manufacture remains HOLD. The installed vehicle may be described as 50°C-capable only after "
        "the component evidence, fit gates and full Rev E acceptance test pass.",
    )
    add_paragraph(
        document,
        "KARIGAR KE LIYE: 50°C grille-air, A/C ON aur approved full-load fuelling par coolant stable rehna "
        "chahiye. Sirf “4-row” ya free-air CFM claim pass nahin. Pehle complete dummy fit aur written approval.",
        italic=True,
    )
    add_table(
        document,
        ["DESIGN CEILING", "MAIN PACK", "CHARGE-AIR PACK"],
        [(
            "150 bhp / 112 kW crankshaft cooling envelope; not an approved engine output",
            "pusher(s) → condenser → radiator → full shroud → mechanical puller",
            "separate side fresh-air duct → charge cooler → dedicated fan → separate hot outlet",
        )],
        [2800, 3280, 3280],
        9.0,
    )

    document.add_page_break()
    add_heading(document, "1. Rev E packaging — Toyota parts, centred fans, fixed front width", 1)
    add_callout(
        document,
        "CONTROLLED PACKAGING RULE",
        "The radiator / condenser / twin-fan stack stays between the existing uprights. Both matching Toyota/Denso "
        "motor centres are level and symmetric about the measured active-fin centreline. The drier sits behind an "
        "upright; relays/fuses sit above or behind the carrier. The independent charge-air module starts aft/outboard "
        "with its smallest dimension lateral and a ≤160 mm projection target. M8 is a full-size-template gate.",
        "F3ECF8",
        PURPLE,
    )
    add_picture(document, ASSET / "rev_e_ph01_standard_toyota_compact_split.png", 6.20)
    add_paragraph(
        document,
        "PH01 — Split-out component visual: standard Toyota/Denso candidates, central pusher / condenser / radiator / "
        "puller-shroud layers, compact independent charge-air module, mounts and rear/upper electrical parts. This is a generated "
        "layout aid; D01–D09 and the accepted full-size templates control manufacture.",
        after=0,
        italic=True,
    )

    document.add_page_break()
    add_heading(document, "1.1 Rev E component assembly", 2)
    add_picture(document, ASSET / "rev_e_ph02_standard_toyota_compact_assembled.png", 4.20)
    add_paragraph(
        document,
        "PH02 — Fully assembled visual: two identical front fans are level and centred in one shallow sealed shroud; "
        "service parts add no side width, while the charge-air module mounts compactly aft/outboard.",
        after=0,
        italic=True,
    )

    document.add_page_break()
    add_heading(document, "1.2 Proposed installation on the existing top holes", 2)
    add_picture(document, ASSET / "rev_e_ph03_standard_toyota_compact_installed.png", 6.20)
    add_paragraph(
        document,
        "PH03E — Current proposed installed visual using the real chassis photograph as its geometric base. One "
        "visible vertical through-bolt is centred on each original top-return hole, with the handed carrier pickup "
        "directly below. The two equal Toyota/Denso 248 mm candidate front fans are symmetric about the condenser "
        "active-fin centreline in one custom sealed shroud. The drier is behind an upright and the relay/fuse plate "
        "is above/rear, so neither adds front-view width or moves a fan. "
        "The lower saddles carry the cooling-pack weight and the two top bolts only locate and retain. "
        "This remains a generated design visual, not a measured fit check or evidence of completed fabrication.",
        after=0,
        italic=True,
    )

    document.add_page_break()
    add_heading(document, "1.3 Existing-hole attachment close-up", 2)
    add_picture(document, ASSET / "rev_d_ph04_existing_top_hole_mount_closeup.png", 6.20)
    add_paragraph(
        document,
        "PH04 — Both required fastener stacks: bolt head and large washer above the retained original hole; "
        "sleeved EPDM isolator and 4 mm handed carrier pad immediately below; accessible washer and locking nut "
        "underneath. D08 and U1–U7—not this generated image—control the dimensions and physical dry fit.",
        after=0,
        italic=True,
    )

    document.add_page_break()
    add_heading(document, "1.4 Rev E decision and release status", 1)
    add_picture(document, ASSET / "rev_c_d01_complete_stack.png", 5.10)
    add_paragraph(
        document,
        "Rev E retains the separate intercooler air path introduced by Rev D and repackages the front module. The previous lower-front core masked "
        "about 39% of the nominal radiator face and made a 50°C/A/C claim indefensible. The side/wing charge "
        "cooler now has its own sealed fresh inlet, fan and hot-air exit."
    )
    add_callout(
        document,
        "SUPERSEDES OLD FAN INSTRUCTIONS",
        "Do not use the older one 12-inch, one 14-inch, optional 12–14-inch, two nominal 9-inch or side-service-tower instructions. "
        "Fan selection now follows installed airflow at static pressure and the complete mock-up.",
        "FFF8E8",
        GOLD,
    )
    add_heading(document, "Released now", 2)
    add_bullets(document, [
        "Vehicle measurement/evidence, supplier quotation and curve review.",
        "Full-size dummies and removable bracket/duct tack fitting.",
        "Instrument and controlled-test planning.",
    ])
    add_heading(document, "Final manufacture remains HOLD until", 2)
    add_numbers(document, [
        "M1–M8, U1–U7, F1–F7 and E1 pass on the actual vehicle.",
        "The actual parts/dummies fit with bonnet, grille, guard and accessories represented.",
        "Radiator and fan performance evidence is accepted.",
        "The owner signs the as-measured drawing.",
        "The completed vehicle passes the 50°C acceptance test.",
    ])

    add_heading(document, "2. Duty and turbo non-derate requirement", 1)
    add_table(document, ["ITEM", "CONTROLLED REQUIREMENT"], [
        ("Ambient", "50°C dry-bulb air measured at the grille/cooling-pack inlet; this is air, not coolant"),
        ("Vehicle", "bonnet closed; final grille/guard/bumper/winch/seals; loaded operating weight"),
        ("Engine", "final approved turbo/fuelling, up to 150 bhp crankshaft cooling-design ceiling"),
        ("A/C", "ON and stabilised with final condenser, charge and cabin blower"),
        ("Continuous", "≥60 min after stabilisation with no continuing coolant-temperature rise"),
        ("Radiator duty", "≥115 kW continuous; ≥130 kW for 10 min beginning from the stable condition"),
        ("Heat soak", "52°C controlled inlet for 10 min followed by hot restart"),
    ], [2100, 7260], 9.1)
    add_callout(
        document,
        "TURBO WORDING",
        "No cooling-related boost reduction or derate is permitted inside the approved 150 bhp crankshaft "
        "thermal-design envelope. This does not approve 150 bhp or 8–10 psi. Start at 5–7 psi; engine health, "
        "maps, EGT, drive pressure, oil, smoke, fuelling, clutch and driveline still govern.",
        "F3ECF8",
        PURPLE,
    )
    add_heading(document, "3. One-page shop specification", 1)
    add_table(document, ["PART", "REQUIRED"], [
        ("Radiator", "largest practical M1/M2 core; ≥0.250 m² net finned face preferred; smaller only with valid duty proof"),
        ("Radiator capacity", "≥115 kW continuous at stated point; ≥130 kW/10 min"),
        ("Condenser", "559 × 356 × 21 nominal R134a basis; actual part measured"),
        ("Front pusher(s)", "preferred twin Toyota/Denso 248 mm candidate in one centred custom sealed shroud; ≥3,000 m³/h installed at 75 Pa, 13.5 V"),
        ("Mechanical fan", "retain/rebuild + sealed full shroud; ≥9,000 m³/h installed at 125 Pa, 1,500 engine rpm"),
        ("Charge cooler", "separate side/wing path; ≥15 kW; IAT ≤80°C; complete route Δp ≤10 kPa"),
        ("Charge fan", "dedicated sealed unit; ≥2,500 m³/h installed at 75 Pa through complete duct/core/outlet"),
    ], [2200, 7160], 8.8)

    document.add_page_break()
    add_heading(document, "3.1 Standard parts first — custom metal only where necessary", 1)
    add_paragraph(
        document,
        "Use standard catalogue parts only where the measured envelope and written duty are met; availability and fit are not assumed. "
        "A Toyota part number identifies a candidate only: take the actual part to the vehicle, measure it and "
        "pass the relevant test before release."
    )
    add_table(document, ["ITEM", "STANDARD / TOYOTA CANDIDATE", "CUSTOM + RELEASE CONTROL"], [
        (
            "Radiator / coolant",
            "HJ47/2H pattern 16400-68030 or sound original Toyota tanks; 16571-68020 upper hose; 16572-68020 lower hose; cap candidate 16401-41021; Toyota-pattern clamps.",
            "Local core/adapter only if needed. Actual necks, cap seat and verified 2H pressure, M1/M2, coolant Δp and ≥115/130 kW proof control.",
        ),
        (
            "Front motors / blades",
            "Prado 120 / GX470 family 88590-60040/-60050/-60051/-60060; motor 88550-12160; blade 88453-60010. Published equivalent: 12 V, 96 W, 8 A, 2,400 rpm, 248 mm.",
            "One custom close sealed shroud and mounting ears. Complete donor shrouds not assumed to fit. F4 fit + F5 ≥3,000 m³/h at 75 Pa/13.5 V; ≥3,300 procurement target.",
        ),
        (
            "Condenser",
            "New common 14 × 22 in nominal R134a parallel-flow, 559 × 356 × 21 basis; common #8/#6 O-ring ports preferred.",
            "Four isolated tabs/short adapters only. Measure seams, manifolds, ports, ears and tools; pass 50°C A/C test.",
        ),
        (
            "Drier / A/C service",
            "New Toyota receiver/drier 88471-34010 if ports fit, otherwise new common #6 O-ring R134a drier with trinary provision; new barrier hose, HNBR seals and crimp fittings.",
            "Vertical rubber-lined clamp behind/in the shadow of an upright; outlet turns rearward. It may not add front-view width. Never reuse a drier; A/C technician verifies the system.",
        ),
        (
            "Electrical",
            "Toyota/Denso/Sumitomo plugs with new terminals/seals/cable; Toyota/Denso relay 90987-02027 candidate per front motor, otherwise sealed ISO 40 A; blade/MAXI fuses/holders.",
            "Vehicle-length loom and protected upper/rear crossrail plate—not a side tower. Verify relay rating; one fuse/relay branch per motor; size from measured hot run/start current.",
        ),
        (
            "Charge-air joints",
            "Common 57 mm / 2.25 in beaded aluminium tube, silicone straight/elbow couplers and T-bolt clamps.",
            "Vehicle-specific pipes/supports. M8 template, ≥15 kW proof and complete route ≤10 kPa control.",
        ),
        (
            "Fasteners / isolation",
            "Metric M6/M8 class 8.8 zinc hardware, large washers, prevailing-torque nuts, crush sleeves, EPDM isolators and lined clamps.",
            "Handed pickup pads, lower saddles and structural carrier. U1–U7 and hand-entering bolt checks control.",
        ),
        (
            "Side core / fan",
            "Prefer a Toyota/Denso catalogue fan only when its complete envelope and pressure-flow proof suit; otherwise use the smallest locally serviceable sealed high-static fan that passes.",
            "Fore–aft brackets, full duct/shroud and separate hot outlet; smallest dimension lateral, ≤160 mm projection target. M8 and complete-path ≥2,500 m³/h at 75 Pa control.",
        ),
    ], [1500, 4050, 3810], 7.8)
    add_callout(
        document,
        "CUSTOM-ONLY LIST",
        "Main carrier/top ears, lower saddles, front twin-fan shroud, mechanical shroud unless the sound Toyota "
        "part fits, side-pack brackets/duct/shroud, vehicle-length loom and measured adapters. Heat exchangers, "
        "motors, blades, drier, controls, connectors, hose, seals, clamps and fasteners remain service parts.",
        "F3ECF8",
        PURPLE,
    )

    document.add_page_break()
    add_heading(document, "4. Engine radiator", 1)
    add_picture(document, ASSET / "rev_c_d02_radiator_assembly.png", 4.20)
    add_paragraph(
        document,
        "Recore the original copper/brass unit first only if its tanks/interfaces pass inspection and the completed "
        "radiator meets the duty. Copper/brass or serviceable aluminium is acceptable. Row count and material "
        "alone are not acceptance."
    )
    add_table(document, ["PARAMETER", "REQUIREMENT"], [
        ("Core face", "≥0.250 m² net finned face preferred; maximise measured M1/M2"),
        ("Smaller core", "allowed only with valid ≥115 kW continuous and ≥130 kW/10 min installed-condition proof"),
        ("Interfaces", "copy physical sample for tank, 38 OD target necks, centres/angles, filler, overflow, drain and brackets"),
        ("Cap", "Toyota 16401-41021 candidate only after its seat and verified pressure match the 2H; a higher-pressure cap is not an upgrade"),
        ("Coolant circuit", "verify thermostat/bypass, high-point bleed, flow, pressure drop and lower-hose anti-collapse"),
        ("Dust margin", "normal grille/guard/screens fitted; ≥10% clean-core thermal/airflow margin"),
        ("Bench tests", "verified system pressure ≥5 min, no leak/loss; uniform flow; record result"),
    ], [2200, 7160], 8.8)
    add_heading(document, "4.1 Required radiator curve/report", 2)
    add_bullets(document, [
        "50°C grille ambient and actual radiator air-in temperature behind operating condenser.",
        "Coolant in/out temperatures, mixture, flow and coolant-side pressure loss.",
        "Installed airflow or face-velocity grid and air-side pressure loss.",
        "Fan/engine speed, voltage, clean-core condition and duration.",
        "≥60 min stable at 115 kW; then ≥130 kW for 10 min with no continuing rise.",
    ])
    add_callout(document, "NO CURVE?", "Do not invent a rating. Build an instrumented prototype and pass the complete vehicle test.")

    add_heading(document, "5. Main airflow, A/C and electrical", 1)
    add_picture(document, ASSET / "rev_c_d04_component_dimensions.png", 5.15)
    add_heading(document, "5.1 Mechanical puller", 2)
    add_bullets(document, [
        "Reject cracked, loose, distorted, contact-marked or weld-repaired blades; inspect hub, water pump and mounts.",
        "Record OD, swept circle, blade depth, centre X/Y and nearest point through rotation/movement.",
        "Rigid full-face perimeter-sealed shroud; 35–50% blade depth inside; ≥15 radial movement clearance.",
        "Radiator rear to fan ≥20 static, 25–30 preferred.",
        "≥9,000 m³/h installed at 125 Pa at 1,500 engine rpm; record a face-velocity grid.",
    ])
    add_heading(document, "5.2 Front electric pusher assembly", 2)
    add_bullets(document, [
        "One full-width assembly or multiple matched sealed pushers; use practical condenser face area.",
        "Centre the complete module on the measured condenser active finned-face centreline. With multiple fans, use equal left/right motor offsets and equal uncovered edge bands.",
        "Any offset needs a recorded physical obstruction plus an approved coverage sketch and installed-flow proof; wiring convenience or styling is not justification.",
        "Two common 12-inch or 14-inch fans do not fit side-by-side within a nominal 559-wide condenser before frames/clearance. Compact matched units or one full-width module still require F4 fit and F5 performance proof.",
        "Preferred standard candidate: matched Prado 120 / GX470 family motors 88550-12160 and blades 88453-60010, referenced by assemblies 88590-60040/-60050/-60051/-60060. Published equivalent data is 248 mm, 12 V, 96 W, 8 A and 2,400 rpm; it is not an airflow curve.",
        "Candidate geometry: two 248 mm sweeps; ≥258 mm ring openings; 266 mm motor C-C; centres x = W_active/2 ±133 from the left active-fin edge and at the same y; 524 mm nominal ring group. Use numerical x = 146.5 and 412.5 mm only if measured W_active is actually 559 mm. F4 controls.",
        "Use only the standard motors/blades/plugs in one custom full-width sealed shroud. Do not assume two complete donor shrouds fit.",
        "There is no side service tower. Put the drier behind/in the shadow of an upright, turn ports rearward, and put relays/fuses on the upper/rear carrier plate. They may not widen the front module, cover active fins or move either fan.",
        "≥3,000 m³/h installed at 75 Pa and 13.5 V through final restriction. Free-air CFM does not pass.",
        "Prefer supplier/procurement proof of ≥3,300 m³/h installed at the same point for 10% reserve; ≥3,000 remains the hard acceptance minimum.",
        "Independent close frame/shroud; no core ties and no fan load on condenser.",
        "Multiple motors use separately fused/relayed branches; A/C/trinary control and fail-safe pressure protection.",
        "One-fan-failed test must warn and fail safe; full 50°C A/C performance is not required with a fault.",
    ])
    add_heading(document, "5.3 Condenser and drier", 2)
    add_paragraph(
        document,
        "559 × 356 × 21 nominal parallel-flow condenser; measure all seams, manifolds, ports and ears. Four "
        "independent isolated tabs. Keep 15 clear to radiator preferred; 10 absolute only with proof. Fit the "
        "drier vertically behind/in the rear shadow of an existing upright, outside the sealed pusher shroud and active fin face. Prefer a "
        "new Toyota 88471-34010 receiver/drier if its ports match; otherwise use a new common #6 O-ring R134a "
        "drier with trinary provision. Never reuse a donor drier. Use new barrier hose, HNBR seals/crimps and have "
        "the A/C technician verify threads, switch set-points, evacuation, charge, pressures and 50°C performance."
    )

    add_heading(document, "6. Separate side/wing charge-air pack", 1)
    # This is a body-page overview; D06 is repeated full-size in Appendix A.
    # Keep the overview compact enough for the complete controlled requirement
    # table and its no-recirculation callout to remain on the same page.
    add_picture(document, ASSET / "rev_c_d06_side_geometry.png", 3.70)
    add_paragraph(
        document,
        "Select the side only after the actual turbo, steering, battery, air cleaner, A/C compressor, bonnet, "
        "wing, downpipe and tool paths are trial-fitted. The pack requires a protected fresh-air inlet, sealed "
        "duct, charge cooler, full-core shroud/fan, separate unrestricted hot-air exit, recirculation seals, "
        "drain and cleaning access. It begins aft and outboard of the selected upright, with the core approximately "
        "90° to the fixed-width central stack. Arrange the long axis fore–aft and the smallest complete-pack dimension "
        "laterally; target no more than 160 mm projection, subject to M8 and the actual fan depth."
    )
    add_table(document, ["PARAMETER", "REQUIREMENT"], [
        ("Heat rejection", "≥15 kW continuous after heat soak"),
        ("Test point", "50°C fresh-air inlet; 0.20 kg/s charge flow; 130°C nominal compressor discharge; record boost/pressure ratio"),
        ("Manifold IAT", "≤80°C during stabilised rated full load at 50°C; idle heat soak recorded separately"),
        ("Complete route Δp", "≤10 kPa / 1.45 psi compressor outlet to plenum; target ≤7 kPa"),
        ("Core Δp", "target ≤5 kPa at rated flow"),
        ("Fan", "≥2,500 m³/h installed at 75 Pa through complete inlet/guard/core/shroud/outlet"),
        ("Proof test", "≥2 × declared maximum boost and never less than 30 psi, using guarded safe method"),
        ("M8 template gate", "full-size core, fan/shroud, inlet/outlet ducts, clamps, mounts and service-tool envelope; long axis fore–aft, smallest dimension lateral, ≤160 mm projection target; trial-fit before manufacture"),
        ("Connections", "57 mm / 2.25 in beaded baseline; larger only by approved flow/fit review"),
    ], [2200, 7160], 8.3)
    add_callout(
        document,
        "NO RECIRCULATION",
        "Do not use engine-bay air as the normal inlet and do not discharge onto turbo/exhaust, battery, wiring, "
        "brakes or main radiator inlet.",
        "FFF8E8",
        GOLD,
    )

    add_heading(document, "7. Mounting, shroud and packaging", 1)
    add_picture(document, ASSET / "rev_c_d05_mounting_shroud.png", 6.35)
    add_bullets(document, [
        "Inspect both existing formed uprights, inward top returns, welds and original holes before reuse.",
        "Use removable rails and rubber-lined lower saddles. The saddles carry all pack weight; upper bolts locate/retain only.",
        "Every exchanger, fan, shroud, duct and drier removes separately without cutting.",
        "No drilling/welding a pressure part or core; no plastic core ties; no forced bolt alignment.",
        "Protect final cores from welding/grinding and all edges from tanks, hoses, wires and fins.",
    ])

    document.add_page_break()
    add_heading(document, "7.1 Existing top-hole carrier pickups", 1)
    add_picture(document, ASSET / "rev_d_d08_existing_top_hole_mount.png", 6.35)
    add_bullets(document, [
        "The two retained original top-return holes are the master upper datums: one positive vertical through-bolt per side.",
        "No new chassis hole, slot or reaming. Do not use a nearby hook, side clamp, tall post or substitute pickup.",
        "TP-L/TP-R: handed 4 mm mild-steel pads, nominal 50 × 50 top, 50-wide downleg, 30 × 30 × 3 gusset; trim to each measured return.",
        "Provisional M8 × 1.25 class 8.8 only if U1 accepts it. Stack: head/large washer; original return; 6 mm EPDM 60–65 Shore A with hard crush sleeve; pickup pad; washer/prevailing nut.",
        "Transfer both centres with one rigid 1:1 template. With the pack supported on its lower saddles, both bolts must enter freely by hand.",
    ])
    add_callout(
        document,
        "MANUFACTURE HOLD — U1–U7",
        "Photographs do not prove hole diameter, pitch, tab thickness, elevation or bonnet/tool clearance. Record "
        "U1–U7 on the vehicle and owner-approve the as-measured carrier before finish welding.",
        "FFF8E8",
        GOLD,
    )

    add_heading(document, "7.2 M3/M4 depth formula", 2)
    add_table(document, ["GATE", "FORMULA / PASS"], [
        ("M3", "10 build + actual front fan depth + ≥5 fan/condenser clear + actual condenser depth + 15 preferred gap"),
        ("M4", "M3 + actual radiator depth + 10 vehicle/fabrication tolerance"),
        ("M5", "radiator rear to nearest mechanical-fan point ≥20 static; 25–30 preferred"),
        ("M8", "full-size template proves complete side path aft/outboard of selected upright; long axis fore–aft, smallest dimension lateral, ≤160 projection target; no added front width"),
    ], [1700, 7660], 9.0)

    document.add_page_break()
    add_heading(document, "8. Electric fan power and controls", 1)
    add_picture(document, ASSET / "rev_c_d07_fan_wiring.png", 6.35)
    add_bullets(document, [
        "Each motor has a correctly sized fuse, sealed relay, weatherproof connector and equal-capacity ground.",
        "Preferred front-branch relay candidate is Toyota/Denso 90987-02027, one per motor, only after its moulded/contact rating or Toyota EWD is checked against measured hot run/start current; otherwise use sealed ISO 40 A relays.",
        "Mount the front relays and fuses above or behind the carrier crossrail, protected from heat and splash; do not create a side service tower.",
        "Fuse/wire from measured run/start current, temperature rating and voltage drop; do not copy a generic value.",
        "Side charge fan runs whenever the engine runs unless a documented fail-safe boost/IAT controller passes testing; manual override and fault indication required.",
        "Bench-test polarity: main grille → engine; side fresh inlet → core → separate hot outlet.",
        "Hot idle with A/C, blower, lights and all fans: motor voltage within 0.5 V of battery/alternator unless supplier tighter.",
        "Record alternator output, fan current, voltage and relay/connector temperature; protect wiring from heat/movement/water.",
    ])

    document.add_page_break()
    add_heading(document, "9. Mandatory measurement and performance gates", 1)
    gates = [
        ("M1", "main clear width", "largest practical core; ≥0.250 m² face preferred; tanks/ports/removal fit"),
        ("M2", "saddle to bonnet/latch", "radiator overall H + 10; filler/hose/service paths fit"),
        ("M3", "front obstruction to radiator front", "formula in section 7.2 using actual component depths"),
        ("M4", "front obstruction to radiator rear", "M3 + actual radiator D + 10 tolerance"),
        ("M5", "radiator rear to fan", "≥20 static; 25–30 preferred through movement"),
        ("M6", "lowest main-pack edge", "≥25 above protected line unless stronger guard approved"),
        ("M7", "main-pack service path", "each component removes separately without cutting"),
        ("M8", "compact side pack + ducts", "full-size template: aft/outboard of selected upright, long axis fore–aft and smallest dimension lateral; ≤160 projection target; no conflict, added front width or recirculation"),
        ("U1-L/R", "original hole Ø / condition", "both sound and round; select bolt after measurement; no chassis reaming"),
        ("U2", "actual top-hole pitch", "record with one rigid 1:1 template; do not infer from photograph"),
        ("U3-L/R", "hole to edge / tab width", "large washer, isolator and handed pickup pad fully supported"),
        ("U4", "left/right elevation delta", "record; form carrier/saddles to fit without bolt pull"),
        ("U5-L/R", "hole to carrier offset", "4 mm pickup directly below each original hole; no post/hook/clamp"),
        ("U6-L/R", "tab thickness + upper clearance", "bolt head/washer and tool clear bonnet, cap and hose"),
        ("U7", "top-hole plane to saddle datum", "saddles carry mass; both upper bolts enter freely by hand"),
        ("F1", "mechanical fan sweep", "record OD/full sweep"),
        ("F2", "blade insertion/clear", "35–50%; ≥15 radial through movement"),
        ("F3", "main fan centre", "record X/Y on as-built"),
        ("F4", "front fan envelope", "record active-face W/H/C-L; centres x = W_active/2 ±133 and same y; complete frame/guards/plugs/wires fit M1–M4; drier/electrics behind/above add no width"),
        ("F5", "front installed flow", "≥3,000 m³/h at 75 Pa, 13.5 V"),
        ("F6", "mechanical installed flow", "≥9,000 m³/h at 125 Pa, 1,500 engine rpm"),
        ("F7", "side installed flow", "≥2,500 m³/h at 75 Pa through complete path"),
        ("E1", "electric fan supply / branches", "record part nos., run/start A, fuses, relays, cable, sealed connectors and loaded V; no hot/drop fault; one branch fault cannot stop every electric fan"),
    ]
    add_table(document, ["ID", "MEASURE", "PASS"], gates, [700, 3100, 5560], 8.35)
    add_callout(
        document,
        "FIT REMAINS CONDITIONAL",
        "The current front-opening estimate is not a manufacture dimension. Put a tape/ruler in every evidence "
        "photo and issue an as-measured drawing before buying cores or fans.",
        "F3ECF8",
        PURPLE,
    )

    document.add_page_break()
    add_heading(document, "10. Required fabrication sequence", 1)
    add_numbers(document, [
        "Inspect/measure the vehicle, accessories, fan and both original top-return holes; complete M1–M8, U1–U7 and F1–F3.",
        "Obtain radiator, fan and charge-cooler performance curves.",
        "Select the largest practical radiator and side-pack location.",
        "Make complete full-size dummies including tanks, ports, fan motors, guards, plugs, wires, ducts and tools.",
        "Make one rigid 1:1 template from both original holes; tack TP-L/TP-R below them, remove carrier, then drill only the two round carrier-pad holes.",
        "Fit the main pack on its lower saddles; both top bolts must enter freely by hand. Fit the side charge pack independently.",
        "Close bonnet and represent final grille/guard/bumper/winch, A/C, battery, intake, steering, turbo and downpipe.",
        "Prove upper fastener/tool clearance, movement, protection, cleaning and separate removal paths.",
        "Photograph every M/U/F/E gate and both complete upper fastener stacks; issue the as-measured/as-selected evidence pack.",
        "Obtain written owner release; manufacture/recore, bench-test, install and instrument.",
        "Perform section 11. Do not call the system 50°C-rated until every acceptance item passes.",
    ])
    add_heading(document, "11. 50°C commissioning and acceptance", 1)
    add_paragraph(
        document,
        "A competent vehicle/turbo/A/C technician must supervise testing on a controlled dyno, proving ground or "
        "safe loaded route. Do not improvise a full-load public-road test."
    )
    add_heading(document, "11.1 Log at least once per second", 2)
    add_bullets(document, [
        "Grille ambient and radiator air-in behind the operating condenser.",
        "Engine-out/head coolant, radiator in/out, coolant pressure where practical.",
        "Oil temperature and hot oil pressure.",
        "Compressor-out, post-intercooler and manifold IAT; boost/plenum pressure.",
        "Pre-turbine EGT and turbine drive pressure.",
        "Fan voltage/current/speed, alternator voltage, A/C pressures/vent temperature.",
        "Engine/road speed, load, vehicle weight and smoke/fault observations.",
    ])
    add_heading(document, "11.2 Test sequence", 2)
    add_table(document, ["TEST", "CONDITION"], [
        ("Hot idle", "50°C, A/C/blower maximum, lights/normal loads, bonnet closed, 30 min"),
        ("Low-speed climb", "50°C, A/C ON, loaded vehicle, final approved tune, 20 min after stabilisation"),
        ("Sustained load", "50°C, A/C ON, repeatable high load including nominal 60 km/h pull where safe, 20 min"),
        ("Continuous proof", "≥60 min stable at 115 kW radiator duty or equivalent instrumented vehicle load"),
        ("Overload", "from stable state, 130 kW equivalent for 10 min"),
        ("Heat soak/restart", "52°C inlet for 10 min per plan, then hot restart and 5 min stabilisation"),
    ], [1900, 7460], 8.7)

    add_heading(document, "11.3 PASS / FAIL limits", 1)
    add_table(document, ["PARAMETER", "PASS"], [
        ("Coolant", "stable/no upward trend; ≤100°C continuous and ≤105°C transient, or verified lower engine limit"),
        ("Boiling margin", "≥10°C below pressure-adjusted boiling point; no loss/boil/hose collapse"),
        ("Manifold IAT", "≤80°C at stabilised rated full load, 50°C ambient"),
        ("Charge Δp", "≤10 kPa compressor outlet to plenum at rated flow"),
        ("Turbo/oil", "EGT, drive pressure, oil temperature/pressure within exact engine/turbo agreed limits"),
        ("A/C", "pressures inside component limits and agreed vent/cabin performance"),
        ("Power", "no cooling-related boost reduction/derate inside approved 150 bhp cooling envelope"),
        ("Electrical", "installed fan duties achieved; no abnormal current, drop or hot relay/connector"),
        ("Structure", "no leak, rubbing, fan contact, abnormal movement or fin damage"),
    ], [2100, 7260], 8.8)
    add_callout(
        document,
        "AUTOMATIC FAIL",
        "Coolant or manifold-air temperature continues to rise after the stated stabilisation period, even if "
        "coolant has not boiled.",
    )

    document.add_page_break()
    add_heading(document, "12. Handover record and signatures", 1)
    records = [
        ("[ ]", "M1–M8 / U1–U7 / F1–F7 / E1 as-measured sheet", ""),
        ("[ ]", "Radiator make/model/envelope", ""),
        ("[ ]", "115/130 kW report + water/air Δp", ""),
        ("[ ]", "Coolant flow/mix/cap/thermostat/bleed/hose", ""),
        ("[ ]", "Condenser model/envelope", ""),
        ("[ ]", "Front fans: installed flow/static/current", ""),
        ("[ ]", "Mechanical installed airflow", ""),
        ("[ ]", "Charge cooler: heat rejection/Δp/proof", ""),
        ("[ ]", "Side fan/duct installed flow/no-recirculation", ""),
        ("[ ]", "Alternator hot-idle output/voltage drop", ""),
        ("[ ]", "50°C + 52°C raw logs/graphs/report", ""),
        ("[ ]", "A/C pressures/vent performance", ""),
        ("[ ]", "As-built drawings/photos", ""),
    ]
    add_table(document, ["CHECK", "RECORD", "RESULT / EVIDENCE"], records, [800, 4300, 4260], 8.7)
    add_paragraph(document, "\nRadiator fabricator: __________________________________  Date: __________")
    add_paragraph(document, "A/C technician: _______________________________________  Date: __________")
    add_paragraph(document, "Turbo/engine tuner: ___________________________________  Date: __________")
    add_paragraph(document, "Owner final release: __________________________________  Date: __________")
    add_heading(document, "Controlled references", 2)
    add_bullets(document, [
        "docs/J40-integrated-cooling-pack-fabricator-specification-rev-c.md (controlled Rev E text)",
        "docs/2h-turbo-suitability-and-options-20260717.md",
        "SAE J1994 heat-exchanger heat-transfer/pressure-drop testing.",
        "SAE J1339 engine-cooling fan performance.",
        "SAE J1393 heavy-duty vehicle cooling test procedures.",
        "SAE J819 engine cooling field test.",
        "Pakistan Meteorological Department, Pakistan Climate 2024.",
    ])
    add_landscape_drawing_appendix(document)
    document.core_properties.title = "J40 Integrated Radiator and Cooling System — Pakistan Fabricator Specification — Rev E"
    document.core_properties.subject = "50°C cooling design with independent charge-air path and all required fans"
    document.core_properties.author = "J40 Project"
    document.core_properties.comments = f"Controlled text source: {SOURCE.name}"
    document.save(OUT)
    print(OUT)


# The active Rev G builder below supersedes the retained prior-revision
# reference above; only this function is invoked by the script entry point.
def build_document():
    save_dimensioned_drawing_set()
    document = Document()
    configure_document(document)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(42)
    r = p.add_run("PAKISTAN FABRICATOR SPECIFICATION")
    set_run(r, 11, True, GOLD)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("J40 Integrated Radiator\nand Cooling System")
    set_run(r, 27, True, NAVY)
    p = document.add_paragraph()
    r = p.add_run("Toyota 2H • R134a A/C • Rev G Toyota-donor-led all-front stack • best-effort / test-gated")
    set_run(r, 12.5, False, MUTED)
    p = document.add_paragraph()
    r = p.add_run("REV G  |  01 AUGUST 2026  |  ALL DIMENSIONS IN mm")
    set_run(r, 10, True, BLUE)
    add_callout(
        document,
        "DESIGN BASIS — NOT A GUARANTEED 50°C OR NO-DERATE RATING",
        "Rev G is a simple best-effort all-front architecture using standard service parts where they are locally obtainable or sourceable; the source status is listed immediately below. It may be described as 50°C-capable or no-derate only after measured fit, component evidence and the complete-vehicle acceptance test pass.",
    )
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("START HERE — WHAT TO BUY, RETAIN AND FABRICATE")
    set_run(r, 11, True, NAVY)
    add_paragraph(
        document,
        "This is not a Prado radiator conversion or a catalogue direct-fit kit. Prado 120/GX470 donor scope is limited to two matched Toyota/Denso front-fan motor, 248 mm blade and plug/pigtail sets.",
        after=3,
    )
    sourcing_table = add_table(document, ["SOURCE CLASS", "QTY / ITEM", "EXACT ITEM / SEARCH REFERENCE", "PAKISTAN ACTION"], [
        ("BUY USED TOYOTA DONOR", "2 matched front-fan sets", "Prado 120/GX470 Toyota/Denso: assemblies 88590-60040 / -60050 / -60051 / -60060; motor 88550-12160; 248 mm blade 88453-60010; plugs/pigtails", "Breaker/Bilal Ganj route. Match, check hand/depth/rotation and bench-test. No Prado radiator or complete donor shroud."),
        ("RETAIN / REBUILD EXISTING", "1 rear fan/hub + 2 boxes", "Existing Toyota 2H puller/hub; blade 16361-68030 / -68031. Existing covered Relay Rev D and closed MIDI Rev D boxes", "Already held. Inspect/rebuild and trial-fit the actual closed boxes; replace only if rejected."),
        ("SAMPLE, THEN LOCAL RECORE", "1 radiator", "HJ47/2H pattern 16400-68030 or sound original tanks. 530 × 435 × 64 is active-core basis only—not finished size", "Local stock unconfirmed. Sample, record complete R0, then release local recore only after hole fixture and dummy pass."),
        ("SOURCE NEW STANDARD — VERIFY STOCK", "1 each + service hardware", "Condenser 559 × 356 × 21; charge cooler 500 × 180 × 50, 57 mm outlets; drier 88471-34010 only if ports fit; hoses, relays/connectors, M6/M8 hardware, EPDM", "Use local A/C, turbo and hardware suppliers. Stock changes; complete envelopes and ports remain physical-sample first."),
        ("FABRICATE LOCALLY", "Measured interfaces only", "2 side rails, 2 direct 4 mm top ears, 2 lower saddles, centred twin-pusher shroud, rear shroud if needed, tabs, seals/screen, 2 separate electrical brackets/hoods and adapters", "Only after real parts, rigid top-hole template and full-size dummy pass. No crossbar, shared tray, new chassis hole, slot or forced alignment."),
    ], [1500, 1500, 3700, 2660], 7.1)
    for row in sourcing_table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    add_paragraph(
        document,
        "PURCHASING STATUS — Part numbers identify candidates; Pakistan sourcing describes the expected route, not guaranteed live stock or direct fit. Source samples and make the fixture now. Final parts, drilling and coating remain on HOLD until their physical-fit gates pass.",
        after=0,
        italic=True,
    )

    add_heading_new_page(document, "1. Rev G packaging — Toyota donor parts, bracket-first measured fit", 1)
    add_callout(document, "CONTROLLED ARCHITECTURE", "Grille → centred matched Prado 120/GX470-family Denso motor/blade pair in a custom shroud → 500 × 180 × 50 charge cooler → ≥10 mm gap → 559 × 356 × 21 condenser → 15 mm target gap (10 mm absolute minimum) → HJ47/2H-pattern recored/custom radiator using a 530 × 435 × 64 active-core thermal basis only, with complete envelope controlled by R0 → sealed full-face mechanical puller shroud. Rev G has exactly three fans: two front electric pushers and one retained rear engine-driven puller. Prado/GX sourcing is for matched fan motors, blades, connectors and pigtails only; no Prado radiator or direct-fit donor-shroud assumption.", "F3ECF8", PURPLE)
    add_picture(document, ASSET / "rev_g_ph01_toyota_donor_split.png", 6.20)
    add_paragraph(document, "PH01 — Rev G split visual: the serviceable parts and fabricated interfaces are separated. It is an AI layout aid only; D01–D10, signed measurements, actual components and the accepted full-size dummy govern manufacture.", italic=True)
    add_heading_new_page(document, "1.1 Rev G component assembly", 2)
    add_picture(document, ASSET / "rev_g_ph02_toyota_donor_assembled.png", 5.40)
    add_paragraph(document, "PH02 — Off-vehicle central cooling-module assembly. The two equal Toyota/Denso 248 mm electric pushers are centred and level; the retained engine-driven puller is visible behind the radiator. Exactly three fans total. The side rails, direct top ears and lower saddles belong to the removable radiator module. Chassis posts and electrical boxes are intentionally omitted: PH03/PH04 alone illustrate the original-hole attachment, while PH05/D09 and E2/E2a control the separate upper/rear electrical installation.", italic=True)
    add_heading_new_page(document, "1.2 Proposed installation on the existing top holes", 2)
    add_picture(document, ASSET / "rev_g_ph03_bracket_correct_installed.png", 6.20)
    add_paragraph(document, "PH03 — Corrected installation proposal on the authoritative bracket photo. Both original inward top-return plates and round holes remain visible. At each side, follow the vertical M8/washer stack through that original hole and sleeved EPDM isolator into the independent black 4 mm ear, which continues directly to the corresponding removable radiator side rail. The nearby inboard fastener joins the ear to the rail; it is not a substitute chassis pickup. A lower saddle beneath that same rail carries the weight. There is no transverse radiator carrier or intermediate bracket. This AI image is illustrative only; the rigid two-hole template, B0/H0-L/H0-R/W0/P0/R0 measurements and full-size dummy govern fit.", italic=True)
    add_heading_new_page(document, "1.3 Existing-hole attachment close-up", 2)
    add_picture(document, ASSET / "rev_g_ph04_direct_side_ear_closeup.png", 6.20)
    add_paragraph(document, "PH04 — Direct side-ear concept at one original top hole. D08 and the signed B0/H0-L/H0-R/W0/P0/R0 sheet retain full control of hole transfer, isolator stack, zero-pull bolt hand-entry and bonnet/tool clearance.", italic=True)
    add_heading_new_page(document, "1.4 Closed fuse/relay boxes on independent brackets", 2)
    add_picture(document, ASSET / "rev_g_ph05_independent_electrical_mounts.png", 6.20)
    add_paragraph(document, "PH05 — Tight electrical-placement detail only: the larger Relay box is upper/forward and the smaller MIDI box lower/rear, centred one behind the other above and behind the radiator top tank. Each closed box has its own local non-radiator structural bracket and rain hood; the visible gap separates the brackets. Both complete envelopes stay inside W0/upright width and clear of the active fin/fan aperture. No common plate is used. Every fuse, holder and live fuse stud stays inside MIDI. D09 and the actual parts at E2/E2a control dimensions, exact positions and service sweeps.", italic=True)

    add_heading(document, "2. Duty, claim and release logic", 1)
    add_table(document, ["ITEM", "REV G CONTROL"], [
        ("50°C / A/C basis", "A demanding best-effort design and test condition, not a present guarantee."),
        ("Engine basis", "Final approved turbo/fuelling up to the 150 bhp cooling-design ceiling; engine, turbo, EGT, oil, smoke, clutch and driveline still govern."),
        ("No-derate wording", "Permitted only after the complete vehicle passes the logged acceptance test at the agreed condition. No claim before then."),
        ("Charge air", "≥15 kW heat-rejection target; IAT ≤80°C and complete route Δp ≤10 kPa targets. All require representative installed test evidence."),
        ("Total fans", "Exactly 3: two centred front electric pushers plus one retained rear engine-driven mechanical puller. No third electric or fourth fan in Rev G."),
        ("Front pushers", "Two matched Toyota/Denso 248 mm candidates, centred in one sealed module; prove ≥3,000 m³/h installed through the complete three-core stack at 13.5 V."),
        ("Mechanical puller", "Retain/rebuild with sealed full-face shroud; prove ≥9,000 m³/h installed through the complete stack at 1,500 rpm."),
    ], [2400, 6960], 8.7)

    add_heading(document, "3. One-page shop specification", 1)
    add_table(document, ["PART", "REQUIRED"], [
        ("Radiator mounts", "Retain the two existing chassis posts: 48 mm face, 410 mm upright and 58 mm inward top return are reference dimensions only. Fit one removable side rail and one short 4 mm direct ear per radiator side beneath the original top hole; put one lower saddle under each same side rail. No transverse top carrier, new chassis hole, slot or reaming."),
        ("Twin pusher module", "Two matched Prado 120/GX470-family Denso 248 mm motor/blade/connector sets; ≥258 mm rings; exactly 266 mm centres; custom sealed symmetric shroud; independent fuse/relay branches."),
        ("Charge-air cooler", "New universal bar-and-plate body exactly 500 W × 180 H × 50 D with 57 mm beaded outlets; ≥15 kW, IAT ≤80°C and route Δp ≤10 kPa are installed-test targets."),
        ("Condenser", "New 14 × 22 nominal R134a parallel-flow body 559 W × 356 H × 21 D on isolated removable tabs; ≥10 mm ahead and 15 mm target / 10 mm absolute minimum behind."),
        ("Radiator / puller", "HJ47/2H-pattern recore/custom radiator; 530 W × 435 H × 64 D is the active-core thermal basis, not the finished radiator size. Record R0 as the complete W × H × D including tanks, seams, filler, drain, ears and lower locators. Build this core only if complete R0 fits signed W0/H0-L/H0-R/P0 and the bonnet/bridge; otherwise repackage the core within the measured R0 envelope and re-prove duty. Retain the Toyota mechanical puller and sealed full-face shroud. Do not buy a Prado radiator."),
        ("Electrical boxes", "Reuse the covered Relay Rev D box and closed gasketed MIDI Rev D fuse box. Give each box its own local structural accessory bracket and rain hood. Keep both complete envelopes inside W0 / the existing upright front-view silhouette; stagger vertically and/or in depth so the combined projected width is no greater than the Relay box's 360 mm base envelope. Never arrange them side-by-side across vehicle width. No common cross-car plate; no radiator, post-return or battery-stand load. Keep every fuse, holder and live stud inside MIDI; covers shut in operation."),
        ("Custom work", "Two radiator side rails/direct ears, two lower saddles, exchanger rails, shrouds, tabs, two local electrical brackets/hoods, seals/ducts and measured adapters only. Heat exchangers, motors, blades, drier, controls, hoses, seals, clamps and fasteners stay service parts."),
    ], [2200, 7160], 8.4)

    add_heading_new_page(document, "3.1 Purchase list — release in three phases", 2)
    add_callout(document, "BUY SAMPLES FIRST", "Do not buy the final radiator or cut radiator ears from catalogue dimensions. First make one rigid template of both original top holes and record B0/H0-L/H0-R/W0/P0. Source physical fan, radiator-pattern and electrical-box samples; record complete R0 and service envelopes; then prove the full-size dummy with both top bolts entering fully by hand and with zero pull.", "FFF8E8", GOLD)
    add_table(document, ["PHASE", "ITEMS TO PURCHASE", "RELEASE CHECK"], [
        ("A — sample / mock-up now", "2 matching Prado 120/GX470-family Denso motor + 248 mm blade + plug/pigtail sets; HJ47/2H radiator/tank pattern or original tanks; cap and hose samples; actual Relay Rev D and MIDI Rev D boxes; template and dummy material", "Record donor identity and complete outside envelopes. Confirm the two motors match. Record B0/H0-L/H0-R/W0/P0/R0. Do not assume a Prado radiator or complete donor shroud fits."),
        ("B — after bracket fixture", "HJ47-pattern recore/custom radiator using the 530 × 435 × 64 active-core basis only if complete R0 fits; 559 × 356 × 21 condenser; 500 × 180 × 50 charge cooler with 57 mm outlets; new port-matched drier; 57 mm pipe/couplers; barrier hose/HNBR seals; metric hardware/isolation", "Rigid hole template and lower-saddle fixture pass; both top bolts hand-enter with zero pull; complete R0 fits signed W0/H0-L/H0-R/P0; approved full-size dummy closes with bonnet/grille/bumper."),
        ("C — electrical / finish", "2 current-rated relays; branch fuses/holders/bus/cable/terminals wholly inside MIDI; sealed connectors/glands; aluminium for two independent local brackets and two hoods; stone screen; finish and test materials", "Measure hot run/start current. E1/E2/E2a prove each box's own bracket, lid/cover/tool sweep, ≥25 mm cable/cover clearance, downward/rearward exits and dry/clean weather check before drilling/coating."),
    ], [1450, 5100, 2810], 8.0)
    add_paragraph(document, "Full seller-facing list: docs/j40-rev-g-toyota-donor-cooling-pack-purchase-list-20260801.md. Machine-readable schedule: data/manual/fabrication/front_cooling_stack_rev_c/rev_g_toyota_donor_purchase_list.csv.", italic=True)
    add_picture(document, ASSET / "rev_g_d10_procurement_fit_control.png", 6.15)
    add_paragraph(document, "D10 — Fixed nominal purchase dimensions and fan coordinates. Actual complete parts, signed vehicle dimensions and the accepted dummy still control fabrication.", italic=True)

    add_heading_new_page(document, "4. Airflow, charge-air and A/C requirements", 1)
    add_picture(document, ASSET / "rev_c_d04_component_dimensions.png", 5.15)
    add_bullets(document, [
        "Rev G has exactly three fans: two centred electric pushers at the front and the retained engine-driven mechanical puller behind the radiator. No additional fan is released.",
        "Keep the full stack between the uprights; no part may add front-view width.",
        "Seal the pusher module to the charge-air cooler and seal core perimeters to prevent bypass. Do not use core ties or put fan load on an exchanger.",
        "Use ≥10 mm actual clearance from charge-air cooler to condenser and 15 mm target / 10 mm absolute minimum from condenser to radiator, after isolators and movement allowance.",
        "Toyota/Denso part numbers identify candidates only. Published free-air data is not an installed-flow pass.",
        "The F3 pusher acceptance is ≥3,000 m³/h at 13.5 V through the complete final grille, stone screen and three-core stack.",
        "The F4 mechanical-puller acceptance is ≥9,000 m³/h through the complete final stack at 1,500 engine rpm.",
        "The charge-air targets are ≥15 kW, IAT ≤80°C and complete compressor-outlet-to-plenum Δp ≤10 kPa at the agreed representative load; log the result before claiming them.",
    ])
    add_heading_new_page(document, "4.1 Drier and electrical service parts", 2)
    add_picture(document, ASSET / "rev_g_d09_independent_electrical_mounts.png", 5.85)
    add_paragraph(document, "Place the drier behind or in the shadow of one chassis upright. Mount the covered Relay Rev D box and the closed, gasketed MIDI Rev D fuse box on two independent local accessory brackets attached only to verified non-radiator structure. Keep each complete box, hood, bracket, cable bend and service sweep inside W0 and the existing uprights' front-view silhouette. Stagger the boxes vertically and/or in depth so the MIDI projected width lies wholly within the Relay box's 360 mm base envelope; never place the two boxes side-by-side across vehicle width. Neither bracket may cross the radiator, load a top return, cover an active fin or block module removal. Relay Rev D keeps its 360 × 245 × 3 mm base and 300 × 197 × 3 mm insulating sheet. MIDI Rev D keeps its 210 × 165 × 65 mm body, 230 × 185 mm lid envelope and 140 × 85 × 5 mm insulating subplate; every fuse, holder, bus and live fuse stud stays inside it. Fabricate each local bracket from 3 mm 5052-H32 aluminium with 20 mm folded returns after transferring the real box footprint. Give each box its own shallow 1.2–1.5 mm aluminium hood with 25 mm turned edges and a baffled lower drain. Keep covers shut in operation; allow the actual cover/lid/tool sweep and at least 25 mm around cable bends and covers. Route cables down/rear through correctly sized glands and drip loops. E2/E2a are separate per-box fit, service and weather gates and must establish the exact final positions. After hose-spray and dust exposure, both interiors must be dry and clean. Do not claim a formal IP rating unless the purchased boxes carry it or the completed installation is tested to it.")

    add_heading_new_page(document, "5. Mounting, shrouds and D08 direct radiator ears", 1)
    add_picture(document, ASSET / "rev_c_d05_mounting_shroud.png", 4.65)
    add_bullets(document, [
        "Every exchanger, both shrouds, the drier and both electrical boxes with their own local brackets remove separately without cutting; each exchanger is rubber-isolated.",
        "Use the retained two chassis posts as found: 48 mm face, 410 mm upright and 58 mm inward return are reference geometry, not radiator dimensions. Verify B0/H0-L/H0-R/W0/P0 on the vehicle.",
        "Lower saddles carry radiator mass directly beneath the same removable side rails that carry the short upper ears. Upper D08 bolts locate/restrain only; both enter fully by hand with zero pull and at least 5 mm rail-to-post hard clearance.",
        "No transverse radiator carrier and no new chassis hole, slot or reaming. Copy both original holes with one rigid 1:1 template; an independent transfer plate per side is allowed only if the physical radiator sample requires it.",
        "No drilling/welding a pressure part, header, tube, core or fin; protect finished cores from welding/grinding.",
    ])
    add_heading(document, "5.1 M3 complete-depth formula", 2)
    add_table(document, ["ITEM", "FORMULA / PASS"], [
        ("M3", "actual pusher depth + fan/charge-cooler clearance + actual charge-cooler depth + ≥10 gap + actual condenser depth + 15 target / 10 absolute-minimum gap + actual radiator depth"),
        ("Absolute component basis", "55 + 5 + 50 + 10 + 21 + 10 + 64 = 215 mm"),
        ("Preferred component basis", "55 + 10 + 50 + 10 + 21 + 15 + 64 = 225 mm"),
        ("Excluded extras", "stone screen, guards, plugs, seams, brackets, pipe bends, tools and vibration clearance; add their measured envelopes to M3"),
        ("M5", "radiator rear to nearest mechanical-fan point ≥20 mm static; 25–30 mm preferred, with ≥15 mm radial clearance through engine movement"),
    ], [1900, 7460], 8.1)

    add_heading_new_page(document, "6. Fan power, controls and installed-airflow tests", 1)
    add_picture(document, ASSET / "rev_c_d07_fan_wiring.png", 6.10)
    add_bullets(document, [
        "Rev G uses exactly two front electric motors plus the retained rear mechanical puller. Do not substitute one pusher or add a third electric/fourth fan without a new controlled revision.",
        "Each Toyota/Denso front motor has a correctly sized fuse, sealed relay, weatherproof connector and equal-capacity ground. Every fuse, holder and live fuse stud is inside the closed MIDI box; none may be exposed.",
        "Toyota/Denso 90987-02027 is a relay candidate only; verify its rating against measured hot run/start current or use sealed ISO 40 A relays.",
        "Bench-test polarity: grille → pushers → charge-air cooler → condenser → radiator → mechanical puller → engine.",
        "At hot idle with A/C, blower, lights and both pushers on, each motor must be within 0.5 V of battery/alternator unless supplier tighter; log voltage/current and relay/connector temperature.",
        "F3 and F4 test the completed, sealed three-core stack—not a bare core, a free-air fan or an assumed static-pressure point.",
        "The covered Relay Rev D box and closed MIDI Rev D box use the two independent local structural brackets shown by D09/E2/E2a. Keep lids shut, cable exits downward/rearward and the battery-side master cutoff feeding them through a protected main cable and accessible disconnect.",
    ])

    add_heading_new_page(document, "7. Mandatory measurement and performance gates", 1)
    gates = [
        ("B0", "original top-hole pattern", "record both hole diameters/condition/edge distance/tab thickness and centre-to-centre with one rigid 1:1 two-hole template; chassis holes remain unchanged"),
        ("H0-L / H0-R", "left/right hole plane to lower saddle seat", "measure each side after saddles are established; complete radiator side rails and ears fit without preload; 410 mm post upright is not radiator height"),
        ("W0", "minimum post clear width", "measure top/middle/bottom including weld beads; complete R0 radiator, rails and at least 5 mm hard clearance per side fit"),
        ("P0", "top-return/post plane to radiator rail", "record left/right depth and bonnet/bridge/cap/tool clearances; direct ears land flat beneath returns without a cross-car bar"),
        ("R0", "complete radiator envelope", "record W × H × D including tanks, seams, filler, drain, ears and lower locators; 530 × 435 × 64 is active-core basis only; final radiator remains HOLD until R0 fits W0/H0-L/H0-R/P0"),
        ("M1", "clear width", "≥569 mm for the 559 mm condenser plus 5 mm each side; every complete core/tank/bracket/plug/wire bend still fits between uprights"),
        ("M2", "clear height", "complete radiator including tanks, filler and brackets—not only the 435 mm active core—fits with bonnet latch, grille and lower protection represented"),
        ("M3", "complete front-to-rear depth", "full formula in section 5.1 plus guards, plugs, seams, brackets, pipe bends, tools and vibration clearance fits"),
        ("M4", "pipe and service routes", "charge pipes turn rearward/inboard at lower corners; A/C fittings and tools remain accessible without extra width or rubbing"),
        ("M5", "rear fan clearance", "≥20 mm static axial and ≥15 mm radial through engine movement; 25–30 mm axial preferred"),
        ("M6", "full-size vehicle dummy", "bonnet, grille and bumper/guard close; assembly is central/symmetric and every module removes separately"),
        ("BRKT", "direct-ear fixture/dummy", "lower saddles sit under the same radiator side rails; each short ear lands beneath its original hole; both provisional M8 bolts enter fully by hand with zero pull; no transverse carrier"),
        ("F1", "twin pusher envelope", "equal 248 mm candidates; ≥258 mm rings; exactly 266 mm C-C; on the 559 body centres x=146.5/412.5 and y=178; same height/equal 17.5 mm side bands; no contact"),
        ("F2", "mechanical fan/shroud", "full-face seal, 35–50% blade insertion and all movement clearances pass"),
        ("F3", "pusher installed flow", "≥3,000 m³/h at 13.5 V through final grille, screen and complete sealed three-core stack"),
        ("F4", "mechanical installed flow", "≥9,000 m³/h through complete sealed stack at 1,500 engine rpm"),
        ("E1", "electric supply", "correct polarity and independent branches; record run/start A, fuses, relays, cable, loaded voltage, temperatures and hot-idle alternator output"),
        ("E2", "two local electrical mounts", "actual Relay and MIDI boxes each have their own non-radiator structural bracket and hood; complete envelopes stay inside W0/upright silhouette, are vertically/depth staggered with MIDI projected width wholly within the Relay 360 mm envelope, never side-by-side across vehicle width; brackets do not cross the radiator or load a post return, and each box remains independently removable"),
        ("E2a", "per-box service / weather", "actual cover/lid/tool sweeps pass with ≥25 mm cable/cover clearance; all MIDI fuses/live studs are inside; lids shut; glands exit down/rear; after hose-spray/dust check both interiors are dry and clean; no unverified IP claim"),
        ("T1", "50°C hot idle", "A/C and cabin blower maximum for 30 min; coolant/A/C pressures stable; no recirculation, leak or electrical heating"),
        ("T2", "50°C loaded operation", "low-speed climb and sustained load pass; 115 kW equivalent stable for 60 min and 130 kW equivalent for 10 min"),
        ("T3", "charge-air proof", "≥15 kW target; manifold IAT ≤80°C; turbo-to-manifold Δp ≤10 kPa; no leak"),
    ]
    add_table(document, ["ID", "MEASURE", "PASS"], gates, [700, 3100, 5560], 8.2)
    add_callout(document, "FIT AND PERFORMANCE REMAIN CONDITIONAL", "Do not claim 50°C capability or no cooling-related derate from drawings, catalogue data, a dummy or a single isolated component test. The completed vehicle test is the release gate.", "F3ECF8", PURPLE)

    add_heading_new_page(document, "8. Fabrication, commissioning and handover", 1)
    add_numbers(document, [
        "Source only Phase A physical samples. Measure B0/H0-L/H0-R/W0/P0/R0 and M1–M6; prove the BRKT, F1–F2 and separate E2/E2a dummies before buying Phase B final cores.",
        "Make one rigid 1:1 template that captures both original top holes. Make a full-size radiator dummy including complete tanks, seams, filler, drain, ears, lower locators, cap/tool sweep and hose routes. Also mock each real electrical box separately with its own cover/lid sweep, cable bends, hood and removal path.",
        "Set the two lower saddles under the proposed radiator side rails. Tack one short ear per side directly beneath its original top return; prove both top bolts enter fully by hand with zero pull and at least 5 mm rail/post clearance. Finish-weld off the vehicle only after the dummy passes.",
        "Trial-fit the three exchangers with the specified ≥10 mm and 15 mm target / 10 mm minimum gaps, isolation, seals and actual mechanical-fan clearance; no side additions are allowed.",
        "Photograph E2/E2a for each real closed enclosure with its own bracket/hood, lid/cover sweeps, ≥25 mm cable/cover clearance, downward/rearward glands and tools. Only after each passes may its local bracket be drilled and finished.",
        "Fabricate only the radiator side rails/direct ears/lower saddles, exchanger rails, shrouds, tabs, seals/ducts, two local electrical brackets/hoods and measured adapters. Retain service components as replaceable parts. Put every fuse, holder, bus and live fuse stud inside MIDI and keep both covers shut in operation.",
        "Perform the hose-spray/dust check and inspect both boxes dry and clean inside. Complete F3/F4/E1/E2/E2a, then instrument the vehicle. Log grille/post-cooler/post-condenser air, coolant, oil, compressor-out/post-cooler/manifold IAT, boost, EGT, fan electrical data, A/C pressures and load once per second where practical.",
        "Complete T1–T3 at the stated ambient/load/A/C conditions. Any 50°C or no-derate statement is withheld unless the full acceptance pass is documented.",
    ])
    add_heading(document, "8.1 Handover record", 2)
    records = [
        ("[ ]", "B0 / H0-L / H0-R / W0 / P0 / R0 / M1–M6 / BRKT / F1–F4 / E1–E2a / T1–T3 as-measured sheet", ""),
        ("[ ]", "Three exchanger models/envelopes and isolation/removal evidence", ""),
        ("[ ]", "Pusher F3 and mechanical-puller F4 complete-stack test reports", ""),
        ("[ ]", "Charge-air heat-rejection / IAT / Δp test report", ""),
        ("[ ]", "Electrical hot-idle test and A/C report", ""),
        ("[ ]", "Two independent electrical-bracket E2/E2a fit, closed lids/covers, clearance, glands, protected feed/disconnect and separate-removal photographs", ""),
        ("[ ]", "Hose-spray/dust result with both electrical-box interiors photographed dry and clean", ""),
        ("[ ]", "Complete vehicle acceptance logs / graphs / owner release", ""),
    ]
    add_table(document, ["CHECK", "RECORD", "RESULT / EVIDENCE"], records, [800, 4300, 4260], 8.6)
    add_paragraph(document, "\nRadiator fabricator: __________________________________  Date: __________")
    add_paragraph(document, "A/C technician: _______________________________________  Date: __________")
    add_paragraph(document, "Turbo/engine tuner: ___________________________________  Date: __________")
    add_paragraph(document, "Owner final release: __________________________________  Date: __________")
    add_heading(document, "Controlled references", 2)
    add_bullets(document, [
        "docs/J40-integrated-cooling-pack-fabricator-specification-rev-c.md (controlled Rev G source specification)",
        "docs/j40-rev-g-toyota-donor-cooling-pack-purchase-list-20260801.md (seller-facing purchase and receiving list)",
        "SAE J1994 heat-exchanger heat-transfer/pressure-drop testing.",
        "SAE J1339 engine-cooling fan performance.",
        "SAE J1393 heavy-duty vehicle cooling test procedures.",
    ])
    add_landscape_drawing_appendix(document)
    document.core_properties.title = "J40 Integrated Radiator and Cooling System — Pakistan Fabricator Specification — Rev G"
    document.core_properties.subject = "Toyota-donor-led, best-effort, test-gated all-front three-core cooling stack"
    document.core_properties.author = "J40 Project"
    document.core_properties.comments = f"Controlled text source: {SOURCE.name}"
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
